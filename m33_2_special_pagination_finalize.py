from __future__ import annotations

"""Ajustes finales M33.2 para evitar colas editoriales huérfanas.

Esta capa no modifica texto, hechos, reglas, conclusiones, fuentes ni efectos jurídicos.
Únicamente reduce espacios verticales —sin bajar Book Antiqua 11 pt— en dos variantes
muy concretas de la familia especial M33.2:

* ``guide``: evita que un encabezado final y un párrafo breve queden aislados;
* ``communication``: evita páginas finales dedicadas únicamente a la firma cuando el
  cierre sustantivo puede convivir de forma legible con ella.

Los demás instrumentos especiales permanecen byte a byte fuera de esta capa.
"""

from functools import wraps
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from document_standard_v33 import audit_docx_legal_standard
from m33_2_special_reference_format import classify_m33_2_special_document

GUIDE_TITLE_AFTER_PT = 6
GUIDE_SUBTITLE_AFTER_PT = 7
GUIDE_SECTION_BEFORE_PT = 5
GUIDE_SECTION_AFTER_PT = 3
GUIDE_BODY_AFTER_PT = 3

COMM_TITLE_AFTER_PT = 7
COMM_SUBTITLE_AFTER_PT = 7
COMM_SECTION_BEFORE_PT = 4
COMM_SECTION_AFTER_PT = 3
COMM_SIGNATURE_BEFORE_PT = 2
COMM_SIGNATURE_AFTER_PT = 1
COMM_BODY_AFTER_PT = 4
COMM_NUMBERED_AFTER_PT = 2

_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+")


def _title_paragraph(document: Document, title: str):
    wanted = str(title or "").strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == wanted:
            return paragraph
    return next((p for p in document.paragraphs if p.style and p.style.name == "Title"), None)


def _paragraph_after(document: Document, paragraph):
    if paragraph is None:
        return None
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag.endswith("}p"):
            return next((p for p in document.paragraphs if p._p is sibling), None)
        sibling = sibling.getnext()
    return None


def _has_explicit_page_break(paragraph) -> bool:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    page_break_before = p_pr.find(qn("w:pageBreakBefore"))
    return bool(
        page_break_before is not None
        and page_break_before.get(qn("w:val"), "1") not in {"0", "false"}
    )


def _remove_top_level_blank_paragraphs(document: Document) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() or _has_explicit_page_break(paragraph):
            continue
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
            removed += 1
    return removed


def _is_signature_table(table) -> bool:
    tbl_pr = table._tbl.tblPr
    description = tbl_pr.find(qn("w:tblDescription")) if tbl_pr is not None else None
    return bool(
        description is not None
        and description.get(qn("w:val")) == "LegalAIZ-SignatureTable"
    )


def _compact_guide(document: Document, *, title: str) -> dict:
    title_p = _title_paragraph(document, title)
    subtitle = _paragraph_after(document, title_p)

    if title_p is not None:
        title_p.paragraph_format.space_after = Pt(GUIDE_TITLE_AFTER_PT)
    if subtitle is not None and subtitle.text.strip():
        style_name = subtitle.style.name if subtitle.style else ""
        if not style_name.lower().startswith("heading"):
            subtitle.paragraph_format.space_after = Pt(GUIDE_SUBTITLE_AFTER_PT)

    headings = 0
    body_paragraphs = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        if title_p is not None and paragraph._p is title_p._p:
            continue
        if subtitle is not None and paragraph._p is subtitle._p:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            paragraph.paragraph_format.space_before = Pt(GUIDE_SECTION_BEFORE_PT)
            paragraph.paragraph_format.space_after = Pt(GUIDE_SECTION_AFTER_PT)
            paragraph.paragraph_format.keep_with_next = True
            headings += 1
        else:
            paragraph.paragraph_format.space_after = Pt(GUIDE_BODY_AFTER_PT)
            body_paragraphs += 1

    return {
        "formatted_headings": headings,
        "formatted_body_paragraphs": body_paragraphs,
        "removed_blank_paragraphs": 0,
    }


def _compact_communication(document: Document, *, title: str) -> dict:
    title_p = _title_paragraph(document, title)
    removed_blank_paragraphs = _remove_top_level_blank_paragraphs(document)
    # La lista de párrafos cambia tras retirar blancos; se vuelve a resolver el título
    # y su subtítulo para no operar con referencias editoriales obsoletas.
    title_p = _title_paragraph(document, title)
    subtitle = _paragraph_after(document, title_p)

    if title_p is not None:
        title_p.paragraph_format.space_after = Pt(COMM_TITLE_AFTER_PT)
    if subtitle is not None and subtitle.text.strip():
        style_name = subtitle.style.name if subtitle.style else ""
        if not style_name.lower().startswith("heading"):
            subtitle.paragraph_format.space_after = Pt(COMM_SUBTITLE_AFTER_PT)

    headings = 0
    body_paragraphs = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if title_p is not None and paragraph._p is title_p._p:
            continue
        if subtitle is not None and paragraph._p is subtitle._p:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            signature = text.upper() in {"FIRMA", "FIRMAS"}
            paragraph.paragraph_format.space_before = Pt(
                COMM_SIGNATURE_BEFORE_PT if signature else COMM_SECTION_BEFORE_PT
            )
            paragraph.paragraph_format.space_after = Pt(
                COMM_SIGNATURE_AFTER_PT if signature else COMM_SECTION_AFTER_PT
            )
            paragraph.paragraph_format.keep_with_next = True
            headings += 1
        else:
            numbered = bool(_NUMBERED_RE.match(text))
            paragraph.paragraph_format.space_after = Pt(
                COMM_NUMBERED_AFTER_PT if numbered else COMM_BODY_AFTER_PT
            )
            paragraph.paragraph_format.line_spacing = 1.0
            body_paragraphs += 1

    signature_tables = 0
    for table in document.tables:
        if not _is_signature_table(table):
            continue
        signature_tables += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)

    return {
        "formatted_headings": headings,
        "formatted_body_paragraphs": body_paragraphs,
        "formatted_signature_tables": signature_tables,
        "removed_blank_paragraphs": removed_blank_paragraphs,
    }


def apply_m33_2_special_pagination_finalize(
    path: str | Path,
    *,
    product_code: str,
    title: str,
) -> dict:
    """Compacta solo guías y comunicaciones especiales, sin alterar contenido."""
    variant = classify_m33_2_special_document(product_code, title)
    if variant not in {"guide", "communication"}:
        return {
            "applied": False,
            "profile": "M33.2-special-pagination",
            "reason": "non_compactable_special_document",
        }

    target = Path(path)
    document = Document(target)
    if variant == "guide":
        details = _compact_guide(document, title=title)
    else:
        details = _compact_communication(document, title=title)

    document.save(target)
    report = audit_docx_legal_standard(target)
    if not report.get("valid"):
        raise ValueError(
            f"Documento especial M33.2 no supera auditoría tras ajuste de paginación: {report.get('findings')}"
        )

    return {
        "applied": True,
        "profile": f"M33.2-special-{variant}-pagination",
        "variant": variant,
        "font_size_preserved_pt": 11,
        **details,
    }


def install_m33_2_special_pagination_gate() -> bool:
    """Envuelve el constructor final y vuelve a gobernar el hash exacto resultante."""
    import docx_builder
    from legalai_platform.document_release_gate import enforce_document_release_gate, infer_product_code

    current = docx_builder.build_docx
    if getattr(current, "_legalaiz_m33_2_special_pagination", False):
        return True

    @wraps(current)
    def guarded_build_docx(*args, **kwargs):
        call_args = list(args)
        path = kwargs.get("path") or (call_args[0] if call_args else None)
        title = kwargs.get("title") or (call_args[1] if len(call_args) >= 2 else "")
        metadata = kwargs.get("metadata")
        if metadata is None and len(call_args) >= 4:
            metadata = call_args[3]
        product_code = str(
            kwargs.get("product_code")
            or infer_product_code(path or "document.docx", metadata)
            or ""
        ).upper()

        result = current(*args, **kwargs)
        final_path = Path(path or result)
        presentation = apply_m33_2_special_pagination_finalize(
            final_path,
            product_code=product_code,
            title=str(title or ""),
        )
        if presentation.get("applied"):
            enforce_document_release_gate(
                final_path,
                expected_product=product_code or None,
                metadata=metadata,
            )
        return result

    guarded_build_docx._legalaiz_m33_2_special_pagination = True
    guarded_build_docx._legalaiz_original = current
    docx_builder.build_docx = guarded_build_docx
    return True


__all__ = [
    "apply_m33_2_special_pagination_finalize",
    "install_m33_2_special_pagination_gate",
]
