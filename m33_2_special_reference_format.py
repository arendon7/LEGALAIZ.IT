from __future__ import annotations

"""Presentación M33.2 para instrumentos especiales, constancias y guías jurídicas.

Esta capa es exclusivamente editorial. No altera hechos, reglas, cálculos, fuentes,
obligaciones, solicitudes, conclusiones ni efectos jurídicos. Cierra las piezas que
no deben adoptar apariencia de contrato, escrito de radicación, informe analítico o
matriz operativa: títulos y acuerdos económicos, recibos/constancias, autorizaciones,
protocolos, resúmenes y guías de gestión.
"""

from functools import wraps
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from document_standard_v33 import audit_docx_legal_standard

FONT_NAME = "Book Antiqua"
BODY_PT = 11
TABLE_PT = 9.5
TITLE_PT = 12.5
NOTE_TITLE_PT = 13
SUBTITLE_PT = 10
PARAGRAPH_AFTER_PT = 5
SECTION_BEFORE_PT = 8
SECTION_AFTER_PT = 5
TITLE_AFTER_PT = 8
NAVY = "0D1324"
IVORY = "F7F5F1"
WHITE = "FFFFFF"
CHARCOAL = "1F1F1F"

SPECIAL_PRODUCT_CODES = frozenset({
    "CO-CD-001",
    "CO-CD-003",
    "CO-CD-004",
    "CO-TR-001",
    "CO-TR-002",
})

_SPECIAL_RULES = (
    ("CO-CD-001", "protocolo de actuación por posible suplantación de identidad", "guide"),
    ("CO-CD-003", "ejercicio del derecho de retracto", "communication"),
    ("CO-CD-003", "terminación por falta de entrega", "communication"),
    ("CO-CD-004", "estado de cuenta reconciliado", "statement"),
    ("CO-CD-004", "acuerdo de pago", "agreement"),
    ("CO-CD-004", "pagaré", "note"),
    ("CO-CD-004", "carta de instrucciones para diligenciamiento de pagaré", "instructions"),
    ("CO-CD-004", "recibo de pago y actualización de saldo", "receipt"),
    ("CO-CD-004", "paz y salvo o constancia de cierre", "certificate"),
    ("CO-TR-001", "autorización de gestión y consulta del expediente sast", "authorization"),
    ("CO-TR-001", "resumen consolidado de verificación sast", "guide"),
    ("CO-TR-002", "guía de radicación, alertas procesales y cierre del caso", "guide"),
)

_ORDINAL_HEADING_RE = re.compile(
    r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[ÉE]PTIMA|OCTAVA|NOVENA|"
    r"D[ÉE]CIMA(?:\s+(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[ÉE]PTIMA|OCTAVA|NOVENA))?|"
    r"VIG[ÉE]SIMA(?:\s+PRIMERA)?)\s*[:.]\s*(.+)$",
    re.IGNORECASE,
)
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+")


def classify_m33_2_special_document(product_code: str, title: str) -> str | None:
    code = str(product_code or "").strip().upper()
    if code not in SPECIAL_PRODUCT_CODES:
        return None
    lowered = str(title or "").strip().casefold()
    for expected_code, token, profile in _SPECIAL_RULES:
        if code == expected_code and token in lowered:
            return profile
    return None


def _set_run_font(
    run,
    *,
    size_pt: float = BODY_PT,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color: str = CHARCOAL,
) -> None:
    run.font.name = FONT_NAME
    rfonts = run._element.get_or_add_rPr().rFonts
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _title_paragraph(document: Document, title: str):
    wanted = str(title or "").strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == wanted:
            return paragraph
    return next((p for p in document.paragraphs if p.style and p.style.name == "Title"), None)


def _paragraph_after(document: Document, paragraph):
    if paragraph is None:
        return None
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            return next((p for p in document.paragraphs if p._p is sibling), None)
        sibling = sibling.getnext()
    return None


def _normalize_styles(document: Document) -> None:
    for style_name in ("Normal", "Title", "Heading 1"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_NAME
        style.font.size = Pt(BODY_PT)
        rfonts = style._element.get_or_add_rPr().rFonts
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT_NAME)


def _remove_body_branding(document: Document, title_paragraph) -> int:
    if title_paragraph is None:
        return 0
    removed = 0
    for paragraph in list(document.paragraphs):
        if paragraph._p is title_paragraph._p:
            break
        text = paragraph.text.strip().casefold()
        if not text or text in {"legalaiz.it", "más que respuestas, soluciones."}:
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def _format_title(document: Document, title: str, profile: str) -> None:
    title_p = _title_paragraph(document, title)
    if title_p is None:
        return
    underline = profile in {"agreement", "note", "instructions", "authorization", "communication"}
    title_size = NOTE_TITLE_PT if profile == "note" else TITLE_PT
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(TITLE_AFTER_PT)
    title_p.paragraph_format.keep_with_next = True
    for run in title_p.runs:
        _set_run_font(
            run,
            size_pt=title_size,
            bold=True,
            italic=False,
            underline=underline,
            color=NAVY,
        )

    subtitle = _paragraph_after(document, title_p)
    if subtitle is None or not subtitle.text.strip():
        return
    if subtitle.style and subtitle.style.name.lower().startswith("heading"):
        return
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(11)
    subtitle.paragraph_format.keep_with_next = True
    for run in subtitle.runs:
        _set_run_font(run, size_pt=SUBTITLE_PT, italic=True, color=CHARCOAL)


def _merge_ordinal_headings(document: Document) -> int:
    changed = 0
    for paragraph in list(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if not style_name.lower().startswith("heading"):
            continue
        match = _ORDINAL_HEADING_RE.match(paragraph.text.strip())
        if not match:
            continue
        following = _paragraph_after(document, paragraph)
        if following is None or not following.text.strip():
            continue
        following_style = following.style.name if following.style else ""
        if following_style.lower().startswith("heading"):
            continue

        body_text = following.text.strip()
        label = f"{match.group(1).upper()}. {match.group(2).upper()}:"
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        label_run = paragraph.add_run(label + " ")
        _set_run_font(label_run, size_pt=BODY_PT, bold=True, color=CHARCOAL)
        body_run = paragraph.add_run(body_text)
        _set_run_font(body_run, size_pt=BODY_PT, bold=False, color=CHARCOAL)
        paragraph.style = document.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_with_next = False
        _remove_paragraph(following)
        changed += 1
    return changed


def _format_paragraphs(document: Document, title: str, profile: str) -> int:
    title_p = _title_paragraph(document, title)
    subtitle = _paragraph_after(document, title_p)
    centered = {
        "CONSIDERACIONES",
        "FIRMA",
        "FIRMAS",
        "RECIBO DE PAGO",
        "PAZ Y SALVO Y CONSTANCIA DE CIERRE",
        "PAGARÉ — DATOS ESENCIALES",
    }
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if title_p is not None and paragraph._p is title_p._p:
            continue
        if subtitle is not None and paragraph._p is subtitle._p:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.upper() in centered else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(SECTION_BEFORE_PT)
            paragraph.paragraph_format.space_after = Pt(SECTION_AFTER_PT)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                _set_run_font(run, size_pt=BODY_PT, bold=True, italic=False, color=NAVY)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(PARAGRAPH_AFTER_PT)
            paragraph.paragraph_format.line_spacing = 1.0
            if _NUMBERED_RE.match(text):
                paragraph.paragraph_format.left_indent = Pt(22)
                paragraph.paragraph_format.first_line_indent = Pt(-11)
            for run in paragraph.runs:
                _set_run_font(run, size_pt=BODY_PT, color=CHARCOAL)
        changed += 1
    return changed


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _is_signature_table(table) -> bool:
    tbl_pr = table._tbl.tblPr
    description = tbl_pr.find(qn("w:tblDescription")) if tbl_pr is not None else None
    return bool(
        description is not None
        and description.get(qn("w:val")) == "LegalAIZ-SignatureTable"
    )


def _mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _format_tables(document: Document, profile: str) -> int:
    changed = 0
    formal_communication = profile == "communication"
    for table in document.tables:
        signature = _is_signature_table(table)
        for row_index, row in enumerate(table.rows):
            _prevent_row_split(row)
            if row_index == 0 and not signature:
                _mark_header_row(row)
            for cell_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if not signature:
                    if formal_communication:
                        fill = WHITE
                    else:
                        fill = NAVY if row_index == 0 else (IVORY if cell_index == 0 else WHITE)
                    _set_cell_shading(cell, fill)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2 if signature else 1)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        header_bold = not signature and (row_index == 0 or cell_index == 0)
                        _set_run_font(
                            run,
                            size_pt=TABLE_PT,
                            bold=(bool(run.bold) or header_bold),
                            color=(WHITE if (not signature and not formal_communication and row_index == 0) else CHARCOAL),
                        )
        changed += 1
    return changed


def apply_m33_2_special_format(
    path: str | Path,
    *,
    product_code: str,
    title: str,
) -> dict:
    profile = classify_m33_2_special_document(product_code, title)
    if profile is None:
        return {
            "applied": False,
            "profile": "M33.2-special",
            "reason": "non_special_document",
        }

    target = Path(path)
    document = Document(target)
    _normalize_styles(document)
    title_p = _title_paragraph(document, title)
    removed_branding = _remove_body_branding(document, title_p)
    merged_clauses = _merge_ordinal_headings(document) if profile in {"agreement", "note"} else 0
    _format_title(document, title, profile)
    paragraphs = _format_paragraphs(document, title, profile)
    tables = _format_tables(document, profile)
    document.save(target)

    report = audit_docx_legal_standard(target)
    if not report.get("valid"):
        raise ValueError(
            f"Documento especial M33.2 no supera auditoría documental: {report.get('findings')}"
        )
    return {
        "applied": True,
        "profile": f"M33.2-special-{profile}",
        "variant": profile,
        "font": FONT_NAME,
        "removed_body_branding": removed_branding,
        "merged_ordinal_clauses": merged_clauses,
        "formatted_paragraphs": paragraphs,
        "formatted_tables": tables,
        "paragraph_after_pt": PARAGRAPH_AFTER_PT,
    }


def install_m33_2_special_format_gate() -> bool:
    """Envuelve `build_docx` tras las demás familias y revalida el hash final."""
    import docx_builder
    from legalai_platform.document_release_gate import enforce_document_release_gate, infer_product_code

    current = docx_builder.build_docx
    if getattr(current, "_legalaiz_m33_2_special", False):
        return True

    @wraps(current)
    def guarded_build_docx(*args, **kwargs):
        call_args = list(args)
        path = kwargs.get("path") or (call_args[0] if call_args else None)
        title = kwargs.get("title") or (call_args[1] if len(call_args) >= 2 else "")
        metadata = kwargs.get("metadata")
        if metadata is None and len(call_args) >= 4:
            metadata = call_args[3]
        product_code = str(
            kwargs.get("product_code")
            or infer_product_code(path or "document.docx", metadata)
            or ""
        ).upper()

        result = current(*args, **kwargs)
        final_path = Path(path or result)
        presentation = apply_m33_2_special_format(
            final_path,
            product_code=product_code,
            title=str(title or ""),
        )
        if presentation.get("applied"):
            enforce_document_release_gate(
                final_path,
                expected_product=product_code or None,
                metadata=metadata,
            )
        return result

    guarded_build_docx._legalaiz_m33_2_special = True
    guarded_build_docx._legalaiz_original = current
    docx_builder.build_docx = guarded_build_docx
    return True


__all__ = [
    "SPECIAL_PRODUCT_CODES",
    "classify_m33_2_special_document",
    "apply_m33_2_special_format",
    "install_m33_2_special_format_gate",
]
