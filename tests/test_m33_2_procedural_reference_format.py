from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_builder import build_docx
from m33_2_procedural_reference_format import apply_m33_2_procedural_format


class ProceduralReferenceFormatM332Tests(unittest.TestCase):
    def _build(
        self,
        root: Path,
        *,
        title: str = "Reclamación directa de garantía legal",
        product_code: str = "CO-CD-003",
    ) -> Path:
        path = root / f"{product_code}_formal_m33_2.docx"
        build_docx(
            path,
            title,
            "Garantía legal · reclamación directa",
            [],
            [
                {
                    "heading": "ASUNTO Y OBJETO DE LA RECLAMACIÓN",
                    "paragraphs": [
                        "La persona consumidora formula reclamación directa y solicita una respuesta de fondo sustentada en los soportes del expediente.",
                        "La comunicación conserva la trazabilidad de los hechos y de los anexos aportados.",
                    ],
                },
                {
                    "heading": "I. HECHOS",
                    "numbered": [
                        "El producto fue entregado y posteriormente presentó una falla reportada por la persona consumidora.",
                        "La reclamación se acompaña de los soportes disponibles para su verificación.",
                    ],
                },
                {
                    "heading": "II. SOLICITUDES",
                    "numbered": [
                        "Emitir una respuesta individual, completa y verificable.",
                        "Informar la solución aplicable y conservar trazabilidad de su ejecución.",
                    ],
                },
                {
                    "heading": "FIRMA",
                    "_type": "signature",
                    "parties": [{"label": "PERSONA CONSUMIDORA", "name": "ANA DEMO", "id": "43.000.001"}],
                },
                {
                    "heading": "CONTROL DE USO, FUENTES Y REVISIÓN",
                    "_type": "control",
                    "text": "Documento sujeto a revisión jurídica y QA sobre la misma revisión.",
                },
            ],
            product_code=product_code,
            enforce_legal_standard=True,
        )
        return path

    def test_formal_writing_removes_body_branding_and_uses_legal_hierarchy(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = self._build(Path(tmp))
            result = apply_m33_2_procedural_format(
                path,
                product_code="CO-CD-003",
                title="Reclamación directa de garantía legal",
            )
            self.assertTrue(result["applied"])
            self.assertEqual(result["font"], "Book Antiqua")
            self.assertEqual(result["paragraph_after_pt"], 4)
            self.assertEqual(result["numbered_after_pt"], 2)
            self.assertEqual(result["linked_signature_context"], 0)

            document = Document(path)
            nonempty = [p for p in document.paragraphs if p.text.strip()]
            self.assertEqual(nonempty[0].text, "Reclamación directa de garantía legal")
            self.assertNotIn("LegalAIZ.it", [p.text.strip() for p in nonempty[:3]])

            title = nonempty[0]
            self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(any(run.bold for run in title.runs))
            self.assertTrue(any(run.underline for run in title.runs))
            self.assertTrue(all((run.font.name or "") == "Book Antiqua" for run in title.runs if run.text))
            self.assertAlmostEqual(title.paragraph_format.space_after.pt, 7.0, places=1)

            subtitle = next(p for p in document.paragraphs if p.text.strip() == "Garantía legal · reclamación directa")
            self.assertEqual(subtitle.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(any(run.italic for run in subtitle.runs))
            self.assertAlmostEqual(subtitle.paragraph_format.space_after.pt, 7.0, places=1)

            body = next(p for p in document.paragraphs if p.text.startswith("La persona consumidora"))
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertAlmostEqual(body.paragraph_format.space_after.pt, 4.0, places=1)

            heading = next(p for p in document.paragraphs if p.text.strip() == "I. HECHOS")
            self.assertEqual(heading.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(all(run.bold for run in heading.runs if run.text))
            self.assertAlmostEqual(heading.paragraph_format.space_before.pt, 4.0, places=1)
            self.assertAlmostEqual(heading.paragraph_format.space_after.pt, 3.0, places=1)

            numbered = next(p for p in document.paragraphs if p.text.startswith("1. El producto"))
            self.assertAlmostEqual(numbered.paragraph_format.space_after.pt, 2.0, places=1)

            signature_heading = next(p for p in document.paragraphs if p.text.strip() == "FIRMA")
            self.assertAlmostEqual(signature_heading.paragraph_format.space_before.pt, 2.0, places=1)
            self.assertAlmostEqual(signature_heading.paragraph_format.space_after.pt, 1.0, places=1)

    def test_dense_signature_closure_keeps_three_substantive_paragraphs_with_signature(self):
        with tempfile.TemporaryDirectory() as tmp:
            title = "Reclamación de hábeas data financiero — corrección, actualización y retiro condicionado"
            path = self._build(Path(tmp), title=title, product_code="CO-CD-001")
            result = apply_m33_2_procedural_format(path, product_code="CO-CD-001", title=title)
            self.assertTrue(result["applied"])
            self.assertEqual(result["linked_signature_context"], 3)

            document = Document(path)
            paragraphs = document.paragraphs
            signature_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "FIRMA")
            context = [p for p in paragraphs[:signature_index] if p.text.strip()][-3:]
            self.assertEqual(len(context), 3)
            self.assertTrue(all(p.paragraph_format.keep_with_next is True for p in context))
            self.assertTrue(all(any((run.font.name or "") == "Book Antiqua" for run in p.runs if run.text) for p in context))

    def test_top_level_blank_paragraphs_are_removed_without_touching_explicit_page_breaks(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = self._build(Path(tmp))
            document = Document(path)
            body_anchor = next(p for p in document.paragraphs if p.text.startswith("La comunicación conserva"))
            body_anchor.insert_paragraph_before("")
            explicit = body_anchor.insert_paragraph_before("")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            explicit.add_run()._r.append(br)
            document.save(path)

            result = apply_m33_2_procedural_format(
                path,
                product_code="CO-CD-003",
                title="Reclamación directa de garantía legal",
            )
            self.assertGreater(result["removed_blank_paragraphs"], 0)
            document = Document(path)
            plain_blanks = [
                p
                for p in document.paragraphs
                if not p.text.strip() and qn("w:br") not in {node.tag for node in p._p.iter()}
            ]
            self.assertEqual(plain_blanks, [])
            self.assertTrue(
                any(
                    any(br.get(qn("w:type")) == "page" for br in p._p.iter(qn("w:br")))
                    for p in document.paragraphs
                )
            )

    def test_non_formal_family_keeps_base_presentation(self):
        with tempfile.TemporaryDirectory() as tmp:
            title = "Diagnóstico jurídico y probatorio de garantía"
            path = self._build(Path(tmp), title=title)
            before = path.read_bytes()
            result = apply_m33_2_procedural_format(path, product_code="CO-CD-003", title=title)
            self.assertFalse(result["applied"])
            self.assertEqual(result["reason"], "non_formal_writing")
            self.assertEqual(path.read_bytes(), before)


if __name__ == "__main__":
    unittest.main()
