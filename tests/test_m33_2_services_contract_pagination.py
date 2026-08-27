from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from m33_2_contract_pagination_finalize import (
    SERVICES_CLAUSE_BEFORE_PT,
    SERVICES_PARAGRAPH_AFTER_PT,
    SERVICES_SIGNATURE_BEFORE_PT,
    finalize_contract_pagination,
)


class ServicesContractPaginationTests(unittest.TestCase):
    def _sample(self, path: Path) -> None:
        document = Document()
        title = document.add_paragraph("CONTRATO DE PRESTACIÓN DE SERVICIOS INDEPENDIENTES")
        title.style = document.styles["Title"]

        clause = document.add_paragraph()
        label = clause.add_run("PRIMERA. OBJETO: ")
        label.bold = True
        label.font.size = Pt(11)
        body = clause.add_run(
            "EL CONTRATISTA ejecutará con autonomía técnica los servicios y entregables pactados."
        )
        body.font.size = Pt(11)
        clause.paragraph_format.space_before = Pt(6)
        clause.paragraph_format.space_after = Pt(6)

        paragraph = document.add_paragraph(
            "La compactación editorial no modifica obligaciones, hechos, riesgos ni criterios de aceptación."
        )
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.size = Pt(11)

        signatures = document.add_paragraph("FIRMAS")
        signatures.paragraph_format.space_before = Pt(8)
        signatures.paragraph_format.space_after = Pt(6)
        for run in signatures.runs:
            run.font.size = Pt(11)
        document.save(path)

    def test_services_contract_compacts_spacing_without_reducing_font_or_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "services.docx"
            self._sample(path)
            before_text = "\n".join(p.text for p in Document(path).paragraphs)

            result = finalize_contract_pagination(path, product_code="CO-EM-003")
            self.assertTrue(result["applied"])
            self.assertEqual("M33.2-services-pagination", result["profile"])
            self.assertEqual(11, result["font_size_preserved_pt"])

            rendered = Document(path)
            after_text = "\n".join(p.text for p in rendered.paragraphs)
            self.assertEqual(before_text, after_text)

            clause = next(p for p in rendered.paragraphs if p.text.startswith("PRIMERA."))
            body = next(p for p in rendered.paragraphs if p.text.startswith("La compactación"))
            signatures = next(p for p in rendered.paragraphs if p.text == "FIRMAS")
            self.assertAlmostEqual(clause.paragraph_format.space_before.pt, SERVICES_CLAUSE_BEFORE_PT)
            self.assertAlmostEqual(clause.paragraph_format.space_after.pt, SERVICES_PARAGRAPH_AFTER_PT)
            self.assertAlmostEqual(body.paragraph_format.space_after.pt, SERVICES_PARAGRAPH_AFTER_PT)
            self.assertAlmostEqual(signatures.paragraph_format.space_before.pt, SERVICES_SIGNATURE_BEFORE_PT)

            sizes = [run.font.size.pt for p in rendered.paragraphs for run in p.runs if run.font.size]
            self.assertTrue(sizes)
            self.assertTrue(all(abs(size - 11) < 0.01 for size in sizes))
            self.assertGreater(result["paragraphs_compacted"], 0)
            self.assertGreater(result["clauses_compacted"], 0)
            self.assertEqual(1, result["signature_headings"])


if __name__ == "__main__":
    unittest.main()
