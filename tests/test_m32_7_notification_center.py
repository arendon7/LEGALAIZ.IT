from __future__ import annotations

from datetime import datetime, timedelta
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase
import sqlite3

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace, PermissionDenied
from legalai_platform.approval_notification_center import (
    ApprovalNotificationCenter,
    BusinessCalendar,
    NotificationIntegrityError,
)


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ApprovalNotificationCenterM327Tests(TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = self.root / "m327.db"
        self.fixed_now = datetime(2026, 8, 6, 9, 0, tzinfo=BOGOTA)
        con = self.db()
        con.executescript(
            """
            CREATE TABLE users(
              id TEXT PRIMARY KEY,
              name TEXT NOT NULL,
              email TEXT,
              role TEXT NOT NULL,
              specialty TEXT,
              active INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,
              owner_id TEXT,
              specialist_id TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,
              case_id TEXT NOT NULL,
              product_code TEXT NOT NULL,
              name TEXT NOT NULL,
              kind TEXT NOT NULL,
              mime_type TEXT NOT NULL,
              file_path TEXT,
              updated_at TEXT NOT NULL
            );
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana QA", "ana@example.test", "admin", "QA documental", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-OTHER", "Otro abogado", "otro@example.test", "specialist", "Litigios", 1),
                ("USR-CLIENT", "Cliente", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = self.root / f"{code}_{index:02d}.docx"
            self.write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), self.fixed_now.isoformat()),
            )
        con.commit(); con.close()
        self.admin = {"id":"USR-ADMIN","role":"admin","name":"Ana QA"}
        self.legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}
        self.other = {"id":"USR-OTHER","role":"specialist","name":"Otro abogado"}
        self.client = {"id":"USR-CLIENT","role":"client","name":"Cliente"}
        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            document_lookup=self.document_lookup,
            access_check=self.access_check,
            upload_validator=self.upload_validator,
        )
        self.operations = ApprovalDeskOperations(
            self.root / "approval-desk",
            workspace=self.workspace,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.center = ApprovalNotificationCenter(
            self.root / "approval-desk",
            operations=self.operations,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )

    def tearDown(self):
        self.temporary.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, code: str):
        document = Document()
        document.add_heading(f"DOCUMENTO DEMOSTRATIVO {code}", 0)
        document.add_paragraph("Contenido sintético para validar el centro operativo M32.7.")
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
        if not row: return False
        if user["role"] == "admin": return True
        if user["role"] == "client": return row["owner_id"] == user["id"]
        return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

    def document_lookup(self, user, document_id):
        con = self.db()
        row = con.execute(
            "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
            (document_id,),
        ).fetchone()
        con.close()
        if not row or not self.access_check(user, row["case_id"]): return None
        return dict(row)

    @staticmethod
    def upload_validator(filename, data):
        return DOCX_MIME, sha256(data).hexdigest(), "clean:test"

    @staticmethod
    def first_case():
        return "DSK-DOC-01"

    def prepare_first(self, *, overdue: bool = True):
        self.operations.sync_portfolio(self.admin)
        self.operations.update_assignment(self.admin, self.first_case(), "USR-LEGAL", "USR-ADMIN")
        if overdue:
            self.operations.update_deadline(
                self.admin,
                self.first_case(),
                (self.fixed_now - timedelta(hours=2)).isoformat(),
                24,
            )

    def test_calendario_habil_salta_cierre_explicito_y_fin_de_semana(self):
        calendar = BusinessCalendar({
            "calendar_id": "test-co",
            "name": "Prueba",
            "timezone": "America/Bogota",
            "weekdays": [0,1,2,3,4],
            "open_time": "08:00",
            "close_time": "17:00",
            "holidays": ["2026-08-07"],
        })
        result = calendar.add_business_hours(datetime(2026,8,6,16,0,tzinfo=BOGOTA), 2)
        self.assertEqual(result, datetime(2026,8,10,9,0,tzinfo=BOGOTA))
        self.assertEqual(calendar.business_hours_between(datetime(2026,8,6,16,0,tzinfo=BOGOTA), result), 2)

    def test_evaluacion_es_idempotente_y_dirige_bandejas(self):
        self.prepare_first()
        first = self.center.evaluate(self.admin, self.first_case())
        self.assertGreater(len(first["created_notifications"]), 0)
        second = self.center.evaluate(self.admin, self.first_case())
        self.assertEqual(second["created_notifications"], [])
        self.assertGreater(second["suppressed_duplicates"], 0)
        legal_inbox = self.center.inbox(self.legal)
        self.assertGreater(legal_inbox["metrics"]["total"], 0)
        self.assertTrue(all(item["recipient_id"] == "USR-LEGAL" for item in legal_inbox["notifications"]))
        all_inbox = self.center.inbox(self.admin, include_all=True)
        self.assertGreaterEqual(all_inbox["metrics"]["total"], legal_inbox["metrics"]["total"])

    def test_cola_externa_no_declara_entrega(self):
        self.prepare_first()
        self.center.update_policy(self.admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
        })
        result = self.center.evaluate(self.admin, self.first_case())
        self.assertFalse(result["external_delivery_performed"])
        self.assertGreater(len(result["queued_messages"]), 0)
        outbox = self.center.outbox(self.admin)
        self.assertEqual(outbox["metrics"]["delivered"], 0)
        self.assertTrue(all(item["status"] == "queued" for item in outbox["messages"]))
        self.assertTrue(all(item["provider"] is None for item in outbox["messages"]))
        self.assertTrue(all(item["contains_document_content"] is False for item in outbox["messages"]))

    def test_lectura_aplazamiento_y_reconocimiento_son_append_only(self):
        self.prepare_first()
        self.center.evaluate(self.admin, self.first_case())
        item = self.center.inbox(self.legal)["notifications"][0]
        read = self.center.mark_read(self.legal, item["notification_id"])
        self.assertTrue(next(row for row in read["notifications"] if row["notification_id"] == item["notification_id"])["read"])
        snoozed = self.center.snooze(
            self.legal,
            item["notification_id"],
            (self.fixed_now + timedelta(hours=24)).isoformat(),
        )
        self.assertTrue(next(row for row in snoozed["notifications"] if row["notification_id"] == item["notification_id"])["snoozed"])
        acknowledged = self.center.acknowledge(self.legal, item["notification_id"], "Recibida para gestión")
        self.assertTrue(next(row for row in acknowledged["notifications"] if row["notification_id"] == item["notification_id"])["acknowledged"])
        self.assertTrue(self.center.verify_chain()["valid"])
        event_types = [event["event_type"] for event in self.center._read_events()]
        self.assertIn("notification.read", event_types)
        self.assertIn("notification.snoozed", event_types)
        self.assertIn("notification.acknowledged", event_types)

    def test_otro_especialista_no_puede_operar_notificacion_ajena(self):
        self.prepare_first()
        self.center.evaluate(self.admin, self.first_case())
        item = self.center.inbox(self.legal)["notifications"][0]
        with self.assertRaises(PermissionDenied):
            self.center.mark_read(self.other, item["notification_id"])

    def test_programacion_habil_conserva_declaracion_no_legal(self):
        self.prepare_first(overdue=False)
        self.center.update_calendar(self.admin, {
            "name": "Jornada interna",
            "weekdays": [0,1,2,3,4],
            "open_time": "08:00",
            "close_time": "17:00",
            "holidays": ["2026-08-07"],
        })
        result = self.center.schedule_case(
            self.admin,
            self.first_case(),
            10,
            datetime(2026,8,6,16,0,tzinfo=BOGOTA).isoformat(),
        )
        self.assertEqual(result["schedule"]["due_at"], datetime(2026,8,10,17,0,tzinfo=BOGOTA).isoformat(timespec="seconds"))
        self.assertFalse(result["schedule"]["legal_deadline"])
        self.assertEqual(result["business_sla"]["calendar"]["name"], "Jornada interna")

    def test_carga_de_trabajo_es_operativa_y_acotada_por_rol(self):
        self.prepare_first()
        admin = self.center.workload(self.admin)
        legal = next(row for row in admin["professionals"] if row["professional"]["id"] == "USR-LEGAL")
        self.assertEqual(legal["legal_assignments"], 1)
        self.assertEqual(legal["overdue"], 1)
        personal = self.center.workload(self.legal)
        self.assertEqual(len(personal["professionals"]), 1)
        self.assertEqual(personal["professionals"][0]["professional"]["id"], "USR-LEGAL")

    def test_manipulacion_bloquea_evaluacion(self):
        self.prepare_first()
        self.center.evaluate(self.admin, self.first_case())
        path = self.root / "approval-desk" / "notification-center" / "events.jsonl"
        path.write_text(path.read_text(encoding="utf-8").replace('"notification.created"', '"notification.altered"', 1), encoding="utf-8")
        self.assertFalse(self.center.verify_chain()["valid"])
        with self.assertRaises(NotificationIntegrityError):
            self.center.evaluate(self.admin, self.first_case())


class ApprovalNotificationCenterStaticM327Tests(TestCase):
    def test_handler_activos_y_marcadores_estan_conectados(self):
        root = Path(__file__).resolve().parents[1]
        index = (root / "app" / "index.html").read_text(encoding="utf-8")
        script = (root / "app" / "modules" / "notification_center_m32_7.js").read_text(encoding="utf-8")
        styles = (root / "app" / "modules" / "notification_center_m32_7.css").read_text(encoding="utf-8")
        run = (root / "run.py").read_text(encoding="utf-8")
        self.assertIn("notification_center_m32_7.css", index)
        self.assertIn("notification_center_m32_7.js", index)
        self.assertIn("Notificaciones, escalamiento y carga", script)
        self.assertIn("Entrega real deshabilitada", script)
        self.assertIn("m327-center", styles)
        self.assertIn("http_handler_m32_7 import Handler", run)
        self.assertIn("http_handler_m32_6 import Handler", run)
