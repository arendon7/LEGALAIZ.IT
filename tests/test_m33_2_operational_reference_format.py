from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx_builder import build_docx
from m33_2_operational_reference_format import apply_m33_2_operational_format


class OperationalReferenceFormatM332Tests(unittest.TestCase):
    def _build(self, root: Path, *, title: str, table: list[tuple[str, ...]], product_code: str = "CO-SA-001") -> Path:
        path = root / "operational_m33_2.docx"
        build_docx(
            path,
            title,
            "Soportes, términos y trazabilidad sujetos a verificación",
            [],
            [
                {
                    "heading": "OBJETO",
                    "paragraphs": [
                        "La pieza organiza datos jurídicos y probatorios para facilitar su cotejo sin sustituir la revisión profesional del expediente.",
                    ],
                },
                {
                    "heading": "I. MATRIZ DE CONTROL",
                    "table": table,
                },
                {
                    "heading": "II. REGLAS DE ACTUALIZACIÓN",
                    "numbered": [
                        "Conservar la versión anterior cuando cambie un dato relevante.",
                        "Vincular cada actualización con el soporte que permita verificarla.",
                    ],
                },
            ],
            product_code=product_code,
            enforce_legal_standard=True,
            append_default_control=False,
        )
        return path

    def test_five_column_matrix_uses_landscape_and_operational_table_controls(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Matriz probatoria y guía de radicación en salud"
            path = self._build(
                root,
                title=title,
                table=[
                    ("ID", "Soporte", "Qué acredita", "Sensibilidad", "Control"),
                    ("SA-01", "Orden o fórmula", "Prestación prescrita", "Alta", "Conservar copia exacta"),
                    ("SA-02", "Radicado y acuse", "Recepción, fecha y hora", "Baja", "Crítico para cómputo"),
                ],
            )
            result = apply_m33_2_operational_format(path, product_code="CO-SA-001", title=title)
            self.assertTrue(result["applied"])
            self.assertTrue(result["landscape"])
            self.assertEqual(result["font"], "Book Antiqua")

            document = Document(path)
            self.assertEqual(document.sections[0].orientation, WD_ORIENT.LANDSCAPE)
            title_p = next(p for p in document.paragraphs if p.text.strip() == title)
            self.assertEqual(title_p.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(all((run.font.name or "") == "Book Antiqua" for run in title_p.runs if run.text))

            table = document.tables[0]
            first_row_pr = table.rows[0]._tr.get_or_add_trPr()
            self.assertIsNotNone(first_row_pr.find(qn("w:tblHeader")))
            for row in table.rows:
                self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))
            shd = table.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertIsNotNone(shd)
            self.assertEqual(shd.get(qn("w:fill")), "0D1324")

    def test_four_column_calendar_stays_portrait_and_uses_fixed_table_layout(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Calendario jurídico de protección al consumidor"
            path = self._build(
                root,
                title=title,
                product_code="CO-CD-003",
                table=[
                    ("Hito", "Regla", "Fecha preliminar", "Control"),
                    ("Reclamación directa", "Respuesta de fondo sujeta al régimen aplicable", "28 de agosto de 2026", "Cotejar festivos y recepción"),
                ],
            )
            result = apply_m33_2_operational_format(path, product_code="CO-CD-003", title=title)
            self.assertTrue(result["applied"])
            self.assertFalse(result["landscape"])

            document = Document(path)
            self.assertNotEqual(document.sections[0].orientation, WD_ORIENT.LANDSCAPE)
            layout = document.tables[0]._tbl.tblPr.find(qn("w:tblLayout"))
            self.assertIsNotNone(layout)
            self.assertEqual(layout.get(qn("w:type")), "fixed")
            header_width = document.tables[0].cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:tcW"))
            rule_width = document.tables[0].cell(0, 1)._tc.get_or_add_tcPr().find(qn("w:tcW"))
            self.assertLess(int(header_width.get(qn("w:w"))), int(rule_width.get(qn("w:w"))))

    def test_analytical_document_is_not_reformatted_as_operational(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Diagnóstico jurídico y operativo de barrera en salud"
            path = self._build(
                root,
                title=title,
                table=[("Variable", "Resultado"), ("Riesgo", "Prioritario")],
            )
            before = path.read_bytes()
            result = apply_m33_2_operational_format(path, product_code="CO-SA-001", title=title)
            self.assertFalse(result["applied"])
            self.assertEqual(result["reason"], "non_operational_document")
            self.assertEqual(path.read_bytes(), before)

    def test_formal_claim_is_not_reformatted_as_operational(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Reclamación directa de garantía legal"
            path = self._build(
                root,
                title=title,
                product_code="CO-CD-003",
                table=[("Solicitud", "Soporte"), ("Reparación", "Factura")],
            )
            before = path.read_bytes()
            result = apply_m33_2_operational_format(path, product_code="CO-CD-003", title=title)
            self.assertFalse(result["applied"])
            self.assertEqual(path.read_bytes(), before)


if __name__ == "__main__":
    unittest.main()
