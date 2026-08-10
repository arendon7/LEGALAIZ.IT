from __future__ import annotations

import re
import sys
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
RUNTIME_MODULES = ROOT / "legalai_runtime_modules"
if str(RUNTIME_MODULES) not in sys.path:
    sys.path.insert(0, str(RUNTIME_MODULES))

from co_em_003_document_factory_v245 import CoEm003DocumentFactoryV245
from m33_document_presentation import APPROVAL_CANDIDATE_MODE, audit_m33_presentation


class ControlledEvaluator:
    def __init__(self, documents: list[str], blocks: list[str] | None = None):
        self.documents = list(documents)
        self.blocks = list(blocks or [])

    def evaluate(self, answers):
        return {
            "blocked": False,
            "missing_fields": [],
            "documents": self.documents,
            "readiness": "ready_for_human_review",
            "status": "ready_for_human_review",
            "professional_review_required": True,
            "professional_reviews": ["Revisión jurídica sustantiva", "QA visual humano"],
            "review_requirements": ["Revisión jurídica sustantiva", "QA visual humano"],
            "findings": [], "blockers": [], "warnings": [], "blocks": self.blocks,
        }


def services_answers() -> dict:
    return {
        "client": {
            "identification": {"type": "legal_person", "name": "Soluciones Andinas S.A.S.", "identification_number": "901234567-8", "domicile": "Medellín, Antioquia", "email": "juridica@demo.legalaiz.it"},
            "signatory": {"name": "María Fernanda Gómez Ruiz", "identification_number": "43678901", "capacity": "representante legal"},
        },
        "contractor": {
            "identification": {"type": "legal_person", "name": "Consultoría Documental Segura S.A.S.", "identification_number": "900765432-1", "domicile": "Bogotá D.C.", "email": "contratos@demo.legalaiz.it"},
            "signatory": {"name": "Juan David Torres Mejía", "identification_number": "79543210", "capacity": "representante legal"},
        },
        "service": {
            "object": "prestar servicios independientes de diagnóstico, diseño y mejora de procesos documentales y tecnológicos",
            "expected_result": "entregar una arquitectura documentada, matrices de control, configuración funcional y evidencia de pruebas",
        },
        "scope": {
            "included": ["Levantar y documentar requerimientos y restricciones.", "Diseñar matrices, flujos y controles de trazabilidad.", "Configurar un entorno demostrativo y ejecutar pruebas acordadas.", "Entregar documentación técnica y sesiones de transferencia."],
            "excluded": ["Representación judicial o administrativa.", "Operación permanente de sistemas del contratante.", "Adquisiciones o licencias no aprobadas por escrito."],
            "deliverables": [{"id": "E1", "name": "Informe de diagnóstico"}, {"id": "E2", "name": "Diseño funcional y matriz de controles"}, {"id": "E3", "name": "Cierre y transferencia"}],
        },
        "schedule": {"start_date": "2026-08-15", "end_date": "2026-10-15", "duration": "dos meses", "milestones": "diagnóstico, diseño, configuración, pruebas y cierre"},
        "execution": {"arrangement": "ejecución autónoma por resultados, con reuniones de coordinación y sin sujeción a jornada laboral", "subcontracting": "solo para componentes especializados, con autorización previa y obligaciones equivalentes"},
        "fees": {"model": "fixed", "currency": "COP", "amount": 48000000, "payment_term": "treinta días calendario después de aceptación y factura válida"},
        "independence": {"no_exclusivity": "no existe exclusividad salvo conflicto específico informado y aceptado"},
        "data": {"personal": True, "roles": "responsable y encargado según actividad", "security": "mínimo privilegio y trazabilidad"},
        "ip": {"preexisting": "cada parte conserva sus activos previos", "results": "resultados según anexo", "third_party": "licencias de terceros"},
        "risk": {"allocation": "cada parte asume riesgos bajo su control", "liability": "daños directos ciertos y probados"},
        "termination": {"rules": "incumplimiento grave o acuerdo", "cure_period": "diez días hábiles"},
        "closure": {"transition": "entrega ordenada", "return_destroy": "devolución o eliminación segura"},
        "dispute": {"mechanism": "negotiation_conciliation_courts", "city": "Medellín"},
    }


def _primary_path(factory, manifest, document_id: str) -> tuple[Path, dict]:
    item = next(item for item in manifest["documents"] if item["id"] == document_id)
    candidates = sorted((factory.output_dir / manifest["generation_id"]).rglob(item["filename"]))
    if len(candidates) != 1:
        raise AssertionError(candidates)
    return candidates[0], item


class ServicesReferenceM330Tests(unittest.TestCase):
    def test_primary_services_contract_is_recomposed_under_m33(self):
        with tempfile.TemporaryDirectory() as tmp:
            factory = CoEm003DocumentFactoryV245(Path(tmp), ControlledEvaluator(["DOC-EM-CONTRACT-001"], ["EM-BASE-001", "EM-SCOPE-001", "EM-FEES-001"]))
            manifest = factory.generate(services_answers(), actor={"id": "qa-m33", "role": "qa"})
            contract, primary = _primary_path(factory, manifest, "DOC-EM-CONTRACT-001")
            report = audit_m33_presentation(contract, APPROVAL_CANDIDATE_MODE)
            self.assertTrue(report["valid"], report["findings"])
            self.assertEqual(primary.get("document_standard"), "M33.0")
            self.assertEqual(primary.get("presentation_mode"), APPROVAL_CANDIDATE_MODE)
            self.assertEqual(factory.VERSION, "2.45")

            document = Document(contract)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            title = "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES INDEPENDIENTES"
            title_paragraph = next(p for p in document.paragraphs if p.text.strip() == title)
            clause_paragraphs = [p for p in paragraphs if re.match(r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|VIGÉSIMA|TRIGÉSIMA|CUADRAGÉSIMA|QUINCUAGÉSIMA)", p)]

            self.assertIn(title, text)
            self.assertEqual(title_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(any(run.bold and run.underline for run in title_paragraph.runs))
            self.assertNotIn("Soluciones Andinas S.A.S. · Consultoría Documental Segura S.A.S.", text)
            self.assertNotIn("Medellín · 15 de agosto de 2026", text)
            self.assertIn("Entre Soluciones Andinas S.A.S.", text)
            self.assertIn("NIT 901.234.567-8", text)
            self.assertIn("NIT 900.765.432-1", text)
            self.assertIn("María Fernanda Gómez Ruiz", text)
            self.assertIn("Juan David Torres Mejía", text)
            self.assertNotIn("identificado en la ficha contractual", text)
            self.assertNotIn("identificado en la misma ficha", text)
            self.assertIn("PRIMERA: OBJETO", text)
            self.assertIn("ANEXO NO. 1", text.upper())
            self.assertNotIn("CONTROL DE USO, FUENTES Y REVISIÓN", text)
            self.assertNotIn("BORRADOR CONTROLADO", text)
            self.assertNotIn("NO FIRMAR", text)
            self.assertNotIn("________", text)
            self.assertNotIn("La plataforma conservará", text)
            self.assertIn("cuarenta y ocho millones de pesos moneda corriente", text)
            self.assertIn("no establece exclusividad general ni obligación de no competencia", text)
            self.assertEqual(document.core_properties.subject, "CO-EM-003")
            self.assertGreaterEqual(len((primary.get("review_evidence") or {}).get("legal_sources") or []), 6)
            self.assertGreater(len(text.split()), 3_100)
            self.assertGreaterEqual(len(clause_paragraphs), 40)
            self.assertIn("representante legal de Soluciones Andinas S.A.S. · NIT 901.234.567-8", table_text)
            self.assertIn("representante legal de Consultoría Documental Segura S.A.S. · NIT 900.765.432-1", table_text)

            with ZipFile(contract) as archive:
                styles = archive.read("word/styles.xml").decode("utf-8")
                body = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("Book Antiqua", styles)
            self.assertNotIn("Times New Roman", styles)
            self.assertNotIn("Arial", styles)
            self.assertIn('w:tblDescription w:val="LegalAIZ-SignatureTable"', body)

    def test_manifest_remains_pending_for_both_human_approvals(self):
        with tempfile.TemporaryDirectory() as tmp:
            factory = CoEm003DocumentFactoryV245(Path(tmp), ControlledEvaluator(["DOC-EM-CONTRACT-001"], ["EM-BASE-001"]))
            manifest = factory.generate(services_answers(), actor={"id": "qa-m33", "role": "qa"})
            legal = manifest.get("legal_approval")
            qa = manifest.get("qa_approval")
            legal_status = legal.get("status") if isinstance(legal, dict) else legal
            qa_status = qa.get("status") if isinstance(qa, dict) else qa
            self.assertEqual(str(legal_status).casefold(), "pending")
            self.assertEqual(str(qa_status).casefold(), "pending")
            self.assertFalse(bool(manifest.get("released", False)))


if __name__ == "__main__":
    unittest.main()
