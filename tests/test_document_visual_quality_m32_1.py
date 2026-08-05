from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

ROOT = Path(__file__).resolve().parents[1]
RUNTIME_MODULES = ROOT / "legalai_runtime_modules"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(RUNTIME_MODULES) not in sys.path:
    sys.path.insert(0, str(RUNTIME_MODULES))

from legalai_platform.document_visual_quality import validate_visual_structure
from co_em_004_document_factory_v246 import CoEm004DocumentFactoryV246
from co_em_004_document_factory_v247 import CoEm004DocumentFactoryV247
from co_la_001_document_factory_v252 import CoLa001DocumentFactoryV252
from co_la_001_document_factory_v253 import CoLa001DocumentFactoryV253
from co_la_002_document_factory_v239 import CoLa002DocumentFactoryV239


class LaborEvaluator:
    def evaluate(self, answers):
        return {
            "blocked": False,
            "missing_fields": [],
            "documents": ["DOC-LA-CONTRACT-001", "ANX-LA-FUN-001"],
            "readiness": "ready_for_human_review",
            "review_requirements": ["Revisión jurídica laboral"],
            "warnings": [],
            "blocks": ["LABOR_BASE", "FUNCTIONS_ANNEX"],
        }


def add_page_field(document: Document, product: str):
    footer = document.sections[0].footer.paragraphs[0]
    footer.add_run(f"LegalAIZ.it | {product} | Página ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def create_quality_fixture(path: Path, product: str):
    document = Document()
    document.core_properties.subject = product
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"DOCUMENTO JURÍDICO {product}")
    run.bold = True
    for ordinal in ("PRIMERA", "SEGUNDA", "TERCERA", "CUARTA", "QUINTA"):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"CLÁUSULA {ordinal}. CONTROL: ").bold = True
        paragraph.add_run("Contenido jurídico suficientemente extenso para verificar estructura, trazabilidad, consistencia y revisión humana posterior.")
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    table.cell(0, 0).text = "Parte"
    table.cell(0, 1).text = "Firma"
    table.cell(1, 0).text = "EL EMPLEADOR"
    table.cell(2, 0).text = "EL TRABAJADOR"
    table.cell(3, 0).text = "LAS PARTES"
    add_page_field(document, product)
    document.save(path)


class DocumentVisualQualityTests(unittest.TestCase):
    def test_preflight_detects_legal_structure_and_page_field(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "fixture.docx"
            create_quality_fixture(path, "CO-LA-002")
            report = validate_visual_structure(path, expected_product="CO-LA-002")
            self.assertTrue(report["valid"], report["errors"])
            self.assertGreaterEqual(report["metrics"]["clauses"], 5)
            self.assertTrue(report["metrics"]["has_page_field"])
            self.assertEqual(report["metrics"]["tables_with_repeating_header"], 1)
            self.assertGreaterEqual(report["metrics"]["rows_protected_from_split"], 4)
            self.assertTrue(report["requires_human_visual_review"])

    def test_preflight_blocks_unusable_page_geometry(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "bad-layout.docx"
            document = Document()
            section = document.sections[0]
            section.page_width = Cm(8)
            section.page_height = Cm(10)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
            document.add_paragraph("Contenido jurídico de prueba con una geometría de página manifiestamente inutilizable para un documento profesional.")
            document.save(path)
            report = validate_visual_structure(path)
            self.assertFalse(report["valid"])
            self.assertTrue(any("tamaño de página" in error or "ancho imprimible" in error for error in report["errors"]))

    def test_co_la_002_manifest_contains_quality_and_visual_evidence(self):
        answers = {
            "employer": {"type": "legal_person", "legalName": "Empresa Demo S.A.S.", "identificationNumber": "900123456-7"},
            "employerSignatory": {"fullName": "Ana Representante", "positionOrCapacity": "representante legal"},
            "worker": {"fullName": "Carlos Trabajador", "identificationNumber": "1012345678"},
            "role": {
                "jobTitle": "Analista jurídico",
                "purpose": "apoyar la gestión contractual y documental",
                "functionsPlacement": "full_in_contract",
                "essentialFunctions": ["Revisar contratos", "Mantener trazabilidad documental"],
            },
            "work": {"mainWorkplace": "Medellín", "modality": "onsite", "actualStartDate": "2026-08-10"},
            "schedule": {"weeklyHours": 42, "type": "fixed"},
            "compensation": {"baseSalary": 3500000, "salaryType": "ordinary"},
        }
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest = CoLa002DocumentFactoryV239(root, LaborEvaluator()).generate(answers)
            self.assertTrue(manifest["requires_human_visual_review"])
            self.assertEqual(manifest["legal_approval"], "pending")
            self.assertEqual(manifest["qa_approval"], "pending")
            self.assertEqual(len(manifest["documents"]), 2)
            for item in manifest["documents"]:
                self.assertTrue(item["quality"]["valid"])
                self.assertTrue(item["visual_preflight"]["valid"])
                self.assertTrue(item["visual_preflight"]["requires_human_visual_review"])
            contract = root / "data" / "generated" / "co-la-002-v239" / manifest["generation_id"] / "CO-LA-002_Contrato_Indefinido.docx"
            document = Document(contract)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            footer_text = "\n".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)
            self.assertIn("CLÁUSULA PRIMERA. OBJETO:", text)
            self.assertIn("LegalAIZ.it | CO-LA-002", footer_text)

    def test_co_la_001_wrapper_attaches_both_quality_reports(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = Path(tmp)
            filename = "CO-LA-001_Informe_Tecnico_Liquidacion.docx"
            create_quality_fixture(target / filename, "CO-LA-001")
            generated = [{"id": "DOC-LA1-REPORT-001", "filename": filename}]
            calculation = {"economic_reconciliation": {"valid": True}}
            factory = CoLa001DocumentFactoryV253.__new__(CoLa001DocumentFactoryV253)
            with patch.object(
                CoLa001DocumentFactoryV252,
                "render_documents",
                return_value=({"blocks": []}, generated, {}, calculation),
            ):
                _, result, hashes, _ = factory.render_documents({}, target)
            self.assertTrue(result[0]["quality"]["valid"])
            self.assertTrue(result[0]["visual_preflight"]["valid"])
            self.assertEqual(len(hashes[filename]), 64)

    def test_co_em_004_wrapper_attaches_both_quality_reports(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = Path(tmp)
            filename = "CO-EM-004_Acuerdo_Confidencialidad_PI.docx"
            create_quality_fixture(target / filename, "CO-EM-004")
            generated = [{"id": "DOC-EM4-NDA-001", "filename": filename}]
            factory = CoEm004DocumentFactoryV247.__new__(CoEm004DocumentFactoryV247)
            with patch.object(
                CoEm004DocumentFactoryV246,
                "render_documents",
                return_value=({"blocks": []}, generated, {}),
            ):
                _, result, hashes = factory.render_documents({}, target)
            self.assertTrue(result[0]["quality"]["valid"])
            self.assertTrue(result[0]["visual_preflight"]["valid"])
            self.assertEqual(len(hashes[filename]), 64)


if __name__ == "__main__":
    unittest.main()
