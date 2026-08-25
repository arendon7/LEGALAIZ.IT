from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.shared import Pt

from docx_builder import build_docx
from m33_2_special_pagination_finalize import apply_m33_2_special_pagination_finalize
from m33_2_special_reference_format import apply_m33_2_special_format


class SpecialPaginationFinalizeM332Tests(unittest.TestCase):
    def _build(self, title: str, *, product_code: str = "CO-TR-002", with_signature: bool = False) -> Path:
        self._tmp = tempfile.TemporaryDirectory()
        target = Path(self._tmp.name) / "sample.docx"
        sections = [
            {"heading": "RADICACIÓN Y TRAZABILIDAD", "paragraphs": ["Conservar escrito, anexos, fecha y radicado."]},
            {"heading": "I. ALERTA", "paragraphs": ["Todo acto adverso exige revisión del expediente y de sus términos."]},
            {"heading": "II. SEMÁFORO", "table": [["Estado", "Acción"], ["Defensa abierta", "Audiencia y pruebas"]]},
            {
                "heading": "III. CIERRE",
                "numbered": [
                    "Registrar la actuación y conservar el soporte de recepción.",
                    "Confirmar el cierre únicamente cuando exista evidencia suficiente.",
                ],
            },
        ]
        if with_signature:
            sections.append(
                {
                    "heading": "FIRMA",
                    "_type": "signature",
                    "parties": [
                        {
                            "label": "PERSONA CONSUMIDORA",
                            "name": "VALENTINA MARÍA SUÁREZ GÓMEZ",
                            "id": "43.000.801",
                        }
                    ],
                }
            )
        build_docx(
            target,
            title,
            "Radicado · acto adverso · recurso · cobro · revisión profesional",
            [],
            sections,
            product_code=product_code,
            enforce_legal_standard=True,
            append_default_control=False,
        )
        return target

    def tearDown(self):
        tmp = getattr(self, "_tmp", None)
        if tmp is not None:
            tmp.cleanup()

    def test_guide_compacts_vertical_spacing_without_lowering_body_font(self):
        title = "Guía de radicación, alertas procesales y cierre del caso"
        target = self._build(title)
        apply_m33_2_special_format(target, product_code="CO-TR-002", title=title)
        result = apply_m33_2_special_pagination_finalize(target, product_code="CO-TR-002", title=title)
        self.assertTrue(result["applied"])
        self.assertEqual(result["variant"], "guide")
        self.assertEqual(result["font_size_preserved_pt"], 11)

        document = Document(target)
        headings = [p for p in document.paragraphs if p.style and p.style.name.lower().startswith("heading")]
        self.assertTrue(headings)
        self.assertTrue(all(round(p.paragraph_format.space_before.pt, 1) == 5.0 for p in headings))
        self.assertTrue(all(round(p.paragraph_format.space_after.pt, 1) == 3.0 for p in headings))
        body_runs = [r for p in document.paragraphs for r in p.runs if p.text.strip() and r.text.strip()]
        self.assertTrue(any(r.font.size == Pt(11) for r in body_runs))

    def test_communication_compacts_signature_closure_without_lowering_font(self):
        title = "Terminación por falta de entrega"
        target = self._build(title, product_code="CO-CD-003", with_signature=True)
        apply_m33_2_special_format(target, product_code="CO-CD-003", title=title)
        result = apply_m33_2_special_pagination_finalize(target, product_code="CO-CD-003", title=title)
        self.assertTrue(result["applied"])
        self.assertEqual(result["variant"], "communication")
        self.assertEqual(result["font_size_preserved_pt"], 11)
        self.assertGreaterEqual(result["formatted_signature_tables"], 1)

        document = Document(target)
        title_p = next(p for p in document.paragraphs if p.text.strip() == title)
        self.assertAlmostEqual(title_p.paragraph_format.space_after.pt, 7.0, places=1)
        numbered = next(p for p in document.paragraphs if p.text.startswith("1. Registrar"))
        self.assertAlmostEqual(numbered.paragraph_format.space_after.pt, 2.0, places=1)
        signature = next(p for p in document.paragraphs if p.text.strip() == "FIRMA")
        self.assertAlmostEqual(signature.paragraph_format.space_before.pt, 2.0, places=1)
        self.assertAlmostEqual(signature.paragraph_format.space_after.pt, 1.0, places=1)
        body_runs = [r for p in document.paragraphs for r in p.runs if p.text.strip() and r.text.strip()]
        self.assertTrue(any(r.font.size == Pt(11) for r in body_runs))

    def test_unrelated_document_is_byte_identical(self):
        title = "Diagnóstico jurídico de fotodetección"
        target = self._build(title)
        before = target.read_bytes()
        result = apply_m33_2_special_pagination_finalize(target, product_code="CO-TR-002", title=title)
        self.assertFalse(result["applied"])
        self.assertEqual(result["reason"], "non_compactable_special_document")
        self.assertEqual(before, target.read_bytes())

    def test_other_special_variant_is_byte_identical(self):
        title = "Acuerdo de pago"
        target = self._build(title, product_code="CO-CD-004")
        apply_m33_2_special_format(target, product_code="CO-CD-004", title=title)
        before = target.read_bytes()
        result = apply_m33_2_special_pagination_finalize(target, product_code="CO-CD-004", title=title)
        self.assertFalse(result["applied"])
        self.assertEqual(result["reason"], "non_compactable_special_document")
        self.assertEqual(before, target.read_bytes())


if __name__ == "__main__":
    unittest.main()
