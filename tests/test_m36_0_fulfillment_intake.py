from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import json
import sqlite3
import unittest

from docx import Document

from legalai_platform.approval_desk_operations import ApprovalDeskOperations
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace, PermissionDenied
from legalai_platform.fulfillment_intake_m36_0 import FulfillmentIntakeCenter, FulfillmentIntakeError


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class FakeActivation:
    def __init__(self):
        self.status = "ACTIVE"
        self.document_count = 2
        self.order_id = "ORD-M360-1"
        self.receipt = "RCPT-SBX-M360"

    def build(self, con, owner_id, case_id):
        if owner_id != "USR-CLIENT" or case_id != "LZ-M360-1":
            raise AssertionError("M36.0 intentó verificar otra identidad o expediente")
        return {
            "activation_status": self.status,
            "case": {"id": case_id, "product_code": "CO-CD-003"},
            "purchase_confirmation": {
                "order_id": self.order_id,
                "payment_intent_id": "PAY-M360-1",
                "receipt_number": self.receipt,
                "amount": 99800,
                "currency": "COP",
                "service_level": "solucion_revisada",
                "review_included": True,
            },
            "documents": {"count": self.document_count, "ready": self.status == "ACTIVE"},
            "journey": {"current_state": "GENERADO"},
        }


class FakeJourney:
    def __init__(self):
        self.state = "GENERADO"
        self.transitions = []

    def detail(self, con, case_id, actor):
        return {"case_id": case_id, "current_state": self.state}

    def transition(self, con, case_id, target, reason, evidence, confirmation, actor):
        if self.state != "GENERADO" or target != "EN_REVISION_JURIDICA":
            raise ValueError("Transición inesperada")
        self.transitions.append({"case_id": case_id, "target": target, "reason": reason, "evidence": evidence})
        self.state = target
        return {"case_id": case_id, "current_state": target}


class M360FulfillmentIntakeTests(unittest.TestCase):
    def setUp(self):
        self.tmp = TemporaryDirectory()
        self.root = Path(self.tmp.name)
        self.database = self.root / "m360.db"
        con = self.db()
        con.executescript(
            """
            CREATE TABLE users(
              id TEXT PRIMARY KEY,name TEXT NOT NULL,role TEXT NOT NULL,specialty TEXT,active INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,product_code TEXT NOT NULL,owner_id TEXT,specialist_id TEXT,status TEXT,risk TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,case_id TEXT NOT NULL,product_code TEXT NOT NULL,kind TEXT NOT NULL,name TEXT NOT NULL,
              mime_type TEXT NOT NULL,file_path TEXT,version TEXT,status TEXT,updated_at TEXT
            );
            CREATE TABLE m35_commerce_case_links(
              id TEXT PRIMARY KEY,user_id TEXT NOT NULL,case_id TEXT,product_code TEXT,state TEXT,order_id TEXT,created_at TEXT
            );
            CREATE TABLE audit_log(
              id INTEGER PRIMARY KEY AUTOINCREMENT,actor TEXT,entity_type TEXT,entity_id TEXT,action TEXT,detail TEXT,created_at TEXT
            );
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana Admin", "admin", "Gobernanza", 1),
                ("USR-LEGAL", "María Legal", "specialist", "Consumo", 1),
                ("USR-CLIENT", "Cliente", "client", None, 1),
            ],
        )
        con.execute(
            "INSERT INTO cases VALUES('LZ-M360-1','CO-CD-003','USR-CLIENT',NULL,'Expediente abierto','green')"
        )
        con.execute(
            "INSERT INTO cases VALUES('LZ-OTHER','CO-CD-003','USR-CLIENT',NULL,'Expediente abierto','green')"
        )
        self.paths = []
        for index, kind in enumerate(("main", "support"), 1):
            path = self.root / f"document-{index}.docx"
            self.write_docx(path, f"Documento M36.0 {index}")
            self.paths.append(path)
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?,?,?)",
                (f"DOC-M360-{index}", "LZ-M360-1", "CO-CD-003", kind, path.name, DOCX_MIME, str(path), "1.0", "Borrador", "2026-08-24T03:00:00-05:00"),
            )
        other = self.root / "other.docx"
        self.write_docx(other, "Documento no relacionado")
        con.execute(
            "INSERT INTO documents VALUES('DOC-OTHER','LZ-OTHER','CO-CD-003','main',?, ?, ?, '1.0','Borrador','2026-08-24T03:00:00-05:00')",
            (other.name, DOCX_MIME, str(other)),
        )
        con.execute(
            "INSERT INTO m35_commerce_case_links VALUES('CCL-M360-1','USR-CLIENT','LZ-M360-1','CO-CD-003','CASE_CREATED','ORD-M360-1','2026-08-24T03:00:00-05:00')"
        )
        con.commit(); con.close()

        self.admin = {"id": "USR-ADMIN", "role": "admin", "name": "Ana Admin"}
        self.client = {"id": "USR-CLIENT", "role": "client", "name": "Cliente"}
        self.activation = FakeActivation()
        self.journey = FakeJourney()
        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            access_check=self.access_check,
        )
        self.operations = ApprovalDeskOperations(
            self.root / "approval-desk",
            workspace=self.workspace,
            db_factory=self.db,
        )
        self.center = FulfillmentIntakeCenter(
            self.activation,
            self.workspace,
            self.operations,
            self.journey,
            db_factory=self.db,
        )

    def tearDown(self):
        self.tmp.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, title: str):
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph("Contenido sintético para intake profesional M36.0.")
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db()
        row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone()
        con.close()
        if not row:
            return False
        if user.get("role") == "admin":
            return True
        if user.get("role") == "client":
            return row["owner_id"] == user.get("id")
        return user.get("role") == "specialist" and row["specialist_id"] == user.get("id")

    def assert_code(self, expected, fn):
        with self.assertRaises(FulfillmentIntakeError) as ctx:
            fn()
        self.assertEqual(ctx.exception.code, expected)
        return ctx.exception

    def test_exact_activation_registers_two_desks_initializes_operations_and_advances_m24(self):
        result = self.center.activate(self.admin, "LZ-M360-1")
        self.assertFalse(result["idempotent"])
        self.assertEqual(result["state"], "EN_REVISION_JURIDICA")
        self.assertEqual(result["journey_state"], "EN_REVISION_JURIDICA")
        self.assertEqual(result["document_count"], 2)
        self.assertEqual(result["desk_case_ids"], ["DSK-DOC-M360-1", "DSK-DOC-M360-2"])
        self.assertEqual(len(self.journey.transitions), 1)
        for desk_id in result["desk_case_ids"]:
            detail = self.workspace.detail(self.admin, desk_id)
            self.assertEqual(detail["source_case_id"], "LZ-M360-1")
            self.assertEqual(detail["workflow_status"], "legal_pending")
            self.assertIsNone(detail["revisions"][0]["approvals"]["legal"])
            self.assertIsNone(detail["revisions"][0]["approvals"]["qa"])
            self.assertIsNone(detail.get("release"))
            audit = self.operations.verify_chain(desk_id)
            self.assertTrue(audit["valid"])
            self.assertEqual(audit["events"], 1)
            state = self.operations.state(self.admin, desk_id)
            self.assertEqual(state["operations"]["priority"], "normal")
            self.assertIsNone(state["operations"]["assigned_specialist"])
            self.assertIsNone(state["operations"]["assigned_qa"])
        self.assertFalse((self.root / "approval-desk" / "DSK-DOC-OTHER" / "case.json").exists())

    def test_retry_is_idempotent_without_new_revisions_events_or_journey_transition(self):
        first = self.center.activate(self.admin, "LZ-M360-1")
        revision_counts = {desk: len(self.workspace.detail(self.admin, desk)["revisions"]) for desk in first["desk_case_ids"]}
        event_counts = {desk: self.operations.verify_chain(desk)["events"] for desk in first["desk_case_ids"]}
        second = self.center.activate(self.admin, "LZ-M360-1")
        self.assertTrue(second["idempotent"])
        self.assertEqual(second["fulfillment_intake_id"], first["fulfillment_intake_id"])
        self.assertEqual(len(self.journey.transitions), 1)
        for desk in first["desk_case_ids"]:
            self.assertEqual(len(self.workspace.detail(self.admin, desk)["revisions"]), revision_counts[desk])
            self.assertEqual(self.operations.verify_chain(desk)["events"], event_counts[desk])
        con = self.db()
        self.assertEqual(con.execute("SELECT COUNT(*) FROM m36_fulfillment_intake").fetchone()[0], 1)
        con.close()

    def test_non_admin_cannot_activate_or_read_queue(self):
        with self.assertRaises(PermissionDenied):
            self.center.activate(self.client, "LZ-M360-1")
        with self.assertRaises(PermissionDenied):
            self.center.queue(self.client)

    def test_pending_m35_activation_blocks_before_creating_any_desk(self):
        self.activation.status = "DOCUMENTS_PENDING"
        self.assert_code("DOCUMENTS_NOT_READY", lambda: self.center.activate(self.admin, "LZ-M360-1"))
        self.assertFalse((self.root / "approval-desk" / "DSK-DOC-M360-1" / "case.json").exists())

    def test_document_count_drift_blocks_before_review_intake(self):
        self.activation.document_count = 1
        self.assert_code("DOCUMENT_COUNT_DRIFT", lambda: self.center.activate(self.admin, "LZ-M360-1"))

    def test_non_docx_document_blocks_entire_intake(self):
        con = self.db()
        con.execute("UPDATE documents SET mime_type='application/pdf' WHERE id='DOC-M360-2'")
        con.commit(); con.close()
        self.assert_code("DOCUMENT_NOT_REVIEWABLE", lambda: self.center.activate(self.admin, "LZ-M360-1"))
        self.assertFalse((self.root / "approval-desk" / "DSK-DOC-M360-1" / "case.json").exists())

    def test_document_mutation_after_intake_is_detected(self):
        self.center.activate(self.admin, "LZ-M360-1")
        self.write_docx(self.paths[0], "Documento M36.0 alterado")
        self.assert_code("FULFILLMENT_INTAKE_DRIFT", lambda: self.center.activate(self.admin, "LZ-M360-1"))

    def test_tampered_operations_chain_blocks_idempotent_read(self):
        first = self.center.activate(self.admin, "LZ-M360-1")
        path = self.root / "approval-desk" / first["desk_case_ids"][0] / "operations.jsonl"
        path.write_text(path.read_text(encoding="utf-8").replace('"priority.updated"', '"priority.tampered"'), encoding="utf-8")
        self.assert_code("OPERATIONS_TRACE_BROKEN", lambda: self.center.activate(self.admin, "LZ-M360-1"))

    def test_unready_m24_state_blocks_after_document_registration_without_false_ledger(self):
        self.journey.state = "INICIADO"
        self.assert_code("JOURNEY_NOT_READY_FOR_REVIEW", lambda: self.center.activate(self.admin, "LZ-M360-1"))
        con = self.db()
        self.assertEqual(con.execute("SELECT COUNT(*) FROM m36_fulfillment_intake").fetchone()[0], 0)
        con.close()
        self.assertTrue((self.root / "approval-desk" / "DSK-DOC-M360-1" / "case.json").is_file())

    def test_existing_review_phase_is_reconciled_without_duplicate_transition(self):
        self.journey.state = "EN_REVISION_JURIDICA"
        result = self.center.activate(self.admin, "LZ-M360-1")
        self.assertEqual(result["journey_state"], "EN_REVISION_JURIDICA")
        self.assertEqual(self.journey.transitions, [])

    def test_queue_uses_existing_m32_states_and_exposes_no_owner_or_hashes(self):
        self.center.activate(self.admin, "LZ-M360-1")
        queue = self.center.queue(self.admin)
        self.assertEqual(queue["metrics"]["cases"], 1)
        self.assertEqual(queue["metrics"]["documents"], 2)
        raw = json.dumps(queue, ensure_ascii=False)
        for forbidden in ("owner_id", "activation_sha256", "document_snapshot_sha256", "receipt_number", "payment_intent_id"):
            self.assertNotIn(forbidden, raw)

    def test_public_intake_declares_no_automatic_assignment_approval_or_release(self):
        result = self.center.activate(self.admin, "LZ-M360-1")
        governance = result["governance"]
        self.assertFalse(governance["automatic_assignment"])
        self.assertFalse(governance["automatic_legal_approval"])
        self.assertFalse(governance["automatic_qa_approval"])
        self.assertFalse(governance["automatic_release"])
        self.assertTrue(governance["dual_approval_preserved"])
        self.assertTrue(governance["m32_review_machinery_reused"])


if __name__ == "__main__":
    unittest.main()
