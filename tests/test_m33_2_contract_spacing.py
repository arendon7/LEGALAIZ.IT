from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from m33_2_contract_style_finalize import (
    CLAUSE_BEFORE_PT,
    PARAGRAPH_AFTER_PT,
    finalize_contract_style,
)


class ContractSpacingM332Tests(unittest.TestCase):
    def test_body_paragraphs_receive_visible_spacing_without_blank_paragraphs(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spacing.docx"
            document = Document()
            title = document.add_paragraph("CONTRATO DEMOSTRATIVO")
            title.style = document.styles["Title"]

            clause = document.add_paragraph()
            label = clause.add_run("PRIMERA: OBJETO: ")
            label.bold = True
            clause.add_run("La parte ejecutará el objeto con diligencia profesional.")
            document.add_paragraph(
                "La ejecución conservará soportes suficientes y trazabilidad de las entregas."
            )
            document.add_paragraph("CONSIDERACIONES")
            document.add_paragraph("FIRMAS")
            document.save(path)

            result = finalize_contract_style(path)
            self.assertEqual(result["paragraph_spacing_after_pt"], PARAGRAPH_AFTER_PT)
            self.assertEqual(result["clause_spacing_before_pt"], CLAUSE_BEFORE_PT)

            rendered = Document(path)
            nonempty = [p for p in rendered.paragraphs if p.text.strip()]
            self.assertEqual(len(nonempty), 5)
            self.assertEqual(len(rendered.paragraphs), 5)

            clause_p = next(p for p in rendered.paragraphs if p.text.startswith("PRIMERA."))
            body_p = next(p for p in rendered.paragraphs if p.text.startswith("La ejecución"))
            self.assertAlmostEqual(clause_p.paragraph_format.space_before.pt, CLAUSE_BEFORE_PT)
            self.assertAlmostEqual(clause_p.paragraph_format.space_after.pt, PARAGRAPH_AFTER_PT)
            self.assertAlmostEqual(body_p.paragraph_format.space_after.pt, PARAGRAPH_AFTER_PT)

    def test_title_keeps_its_own_spacing_profile(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "title-spacing.docx"
            document = Document()
            title = document.add_paragraph("CONTRATO DEMOSTRATIVO")
            title.style = document.styles["Title"]
            title.paragraph_format.space_after = None
            document.add_paragraph("Texto introductorio.")
            document.save(path)

            finalize_contract_style(path)
            rendered = Document(path)
            title_after = rendered.paragraphs[0].paragraph_format.space_after
            self.assertIsNone(title_after)
            self.assertAlmostEqual(
                rendered.paragraphs[1].paragraph_format.space_after.pt,
                PARAGRAPH_AFTER_PT,
            )


if __name__ == "__main__":
    unittest.main()
