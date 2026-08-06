from __future__ import annotations

from datetime import datetime, timedelta
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase
import sqlite3

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace, PermissionDenied
from legalai_platform.approval_notification_center import ApprovalNotificationCenter
from legalai_platform.contact_governance import (
    ContactGovernance,
    ContactGovernanceIntegrityError,
    GovernedTransactionalCommunications,
)


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ContactGovernanceM329Tests(TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = self.root / "m329.db"
        self.fixed_now = datetime(2026, 8, 6, 11, 0, tzinfo=BOGOTA)
        con = self.db()
        con.executescript(
            """
            CREATE TABLE users(
              id TEXT PRIMARY KEY,
              name TEXT NOT NULL,
              email TEXT,
              role TEXT NOT NULL,
              specialty TEXT,
              active INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,
              owner_id TEXT,
              specialist_id TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,
              case_id TEXT NOT NULL,
              product_code TEXT NOT NULL,
              name TEXT NOT NULL,
              kind TEXT NOT NULL,
              mime_type TEXT NOT NULL,
              file_path TEXT,
              updated_at TEXT NOT NULL
            );
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana Administradora", "ana@example.test", "admin", "Operación", 1),
                ("USR-QA", "Quinn QA", "qa@example.test", "qa", "QA", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-OTHER", "Otro abogado", "otro@example.test", "specialist", "Litigios", 1),
                ("USR-CLIENT", "Cliente", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = self.root / f"{code}_{index:02d}.docx"
            self.write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), self.fixed_now.isoformat()),
            )
        con.commit(); con.close()

        self.admin = {"id":"USR-ADMIN","role":"admin","name":"Ana Administradora"}
        self.qa = {"id":"USR-QA","role":"qa","name":"Quinn QA"}
        self.legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}
        self.other = {"id":"USR-OTHER","role":"specialist","name":"Otro abogado"}
        self.client = {"id":"USR-CLIENT","role":"client","name":"Cliente"}

        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            document_lookup=self.document_lookup,
            access_check=self.access_check,
            upload_validator=self.upload_validator,
        )
        self.operations = ApprovalDeskOperations(
            self.root / "approval-desk",
            workspace=self.workspace,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.center = ApprovalNotificationCenter(
            self.root / "approval-desk",
            operations=self.operations,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.governance = ContactGovernance(
            self.root / "approval-desk",
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.communications = GovernedTransactionalCommunications(
            self.root / "approval-desk",
            notification_center=self.center,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
            governance=self.governance,
        )

    def tearDown(self):
        self.temporary.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, code: str):
        document = Document()
        document.add_heading(f"DOCUMENTO DEMOSTRATIVO {code}", 0)
        document.add_paragraph("Contenido sintético para validar gobierno de contacto M32.9.")
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
        if not row: return False
        if user["role"] in {"admin", "qa"}: return True
        if user["role"] == "client": return row["owner_id"] == user["id"]
        return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

    def document_lookup(self, user, document_id):
        con = self.db()
        row = con.execute(
            "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
            (document_id,),
        ).fetchone()
        con.close()
        if not row or not self.access_check(user, row["case_id"]): return None
        return dict(row)

    @staticmethod
    def upload_validator(filename, data):
        return DOCX_MIME, sha256(data).hexdigest(), "clean:test"

    @staticmethod
    def first_case():
        return "DSK-DOC-01"

    def prepare_outbox(self):
        self.operations.sync_portfolio(self.admin)
        self.operations.update_assignment(self.admin, self.first_case(), "USR-LEGAL", "USR-ADMIN")
        self.operations.update_deadline(
            self.admin,
            self.first_case(),
            (self.fixed_now - timedelta(hours=2)).isoformat(),
            24,
        )
        self.center.update_policy(self.admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
        })
        evaluation = self.center.evaluate(self.admin, self.first_case())
        self.assertGreater(len(evaluation["queued_messages"]), 0)
        synced = self.communications.sync_outbox(self.admin)
        self.assertGreater(len(synced["imported_dispatches"]), 0)

    def register_client_relationship(self):
        return self.governance.record_relationship(self.admin, {
            "subject_id": "USR-CLIENT",
            "relationship_type": "client",
            "lawful_basis": "contract",
            "status": "active",
            "evidence_reference": "CONTRACT-CLIENT-001",
        })

    def grant_client(self, purpose="commercial_marketing", channel="email"):
        return self.governance.record_preference(self.admin, {
            "subject_id": "USR-CLIENT",
            "purpose": purpose,
            "channel": channel,
            "state": "granted",
            "basis": "consent",
            "evidence_reference": f"CONSENT-{purpose}-{channel}-001",
            "reason": "Autorización sintética verificable",
        })

    def test_alerta_profesional_es_permitida_y_supresion_la_bloquea(self):
        allowed = self.governance.evaluate(
            self.admin,
            subject_id="USR-LEGAL",
            purpose="professional_operational",
            channel="email",
        )["decision"]
        self.assertTrue(allowed["allowed"])
        self.assertEqual(allowed["declared_basis"], "internal_operational_policy")

        self.governance.add_suppression(self.qa, {
            "subject_id": "USR-LEGAL",
            "scope": "global",
            "reason": "Solicitud verificada de no contacto",
            "source": "verified_request",
        })
        blocked = self.governance.evaluate(
            self.admin,
            subject_id="USR-LEGAL",
            purpose="professional_operational",
            channel="email",
        )["decision"]
        self.assertFalse(blocked["allowed"])
        self.assertIn("active_suppression", blocked["reasons"])

    def test_marketing_exige_relacion_consentimiento_y_horario(self):
        missing = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
        )["decision"]
        self.assertFalse(missing["allowed"])
        self.assertIn("active_relationship_missing", missing["reasons"])
        self.assertIn("explicit_marketing_consent_missing", missing["reasons"])

        self.register_client_relationship(); preference = self.grant_client()
        self.assertFalse(preference["preference"]["evidence_reference_stored"])
        self.assertEqual(len(preference["preference"]["evidence_sha256"]), 64)
        allowed = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
        )["decision"]
        self.assertTrue(allowed["allowed"])
        self.assertEqual(allowed["declared_basis"], "consent")

        sunday = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
            scheduled_at=datetime(2026, 8, 9, 11, 0, tzinfo=BOGOTA),
        )["decision"]
        self.assertFalse(sunday["allowed"])
        self.assertIn("sunday", sunday["reasons"])

    def test_preferencia_negada_prevalece(self):
        self.register_client_relationship(); self.grant_client()
        self.governance.record_preference(self.qa, {
            "subject_id": "USR-CLIENT",
            "purpose": "commercial_marketing",
            "channel": "email",
            "state": "denied",
            "basis": "consent",
            "reason": "Revocatoria sintética",
        })
        decision = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
        )["decision"]
        self.assertFalse(decision["allowed"])
        self.assertIn("preference_denied", decision["reasons"])

    def test_frecuencia_bloquea_segundo_contacto_del_dia(self):
        self.register_client_relationship(); self.grant_client("collections", "email")
        first = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="collections",
            channel="email",
        )["decision"]
        self.assertTrue(first["allowed"])
        self.governance.record_contact(
            self.admin,
            decision_id=first["decision_id"],
            subject_id="USR-CLIENT",
            purpose="collections",
            channel="email",
            dispatch_id="DSP-COLLECTION-001",
            synthetic=True,
        )
        second = self.governance.evaluate(
            self.admin,
            subject_id="USR-CLIENT",
            purpose="collections",
            channel="email",
        )["decision"]
        self.assertFalse(second["allowed"])
        self.assertIn("daily_frequency_limit", second["reasons"])

    def test_aviso_requiere_activacion_independiente(self):
        created = self.governance.create_notice_version(self.admin, {
            "notice_id": "contact-governance",
            "name": "Aviso revisado",
            "text": "Aviso sintético con finalidades, canales, derechos y mecanismos de revocatoria.",
        })
        self.assertEqual(created["notice"]["version"], 2)
        with self.assertRaises(PermissionDenied):
            self.governance.activate_notice(self.admin, "contact-governance", 2)
        activated = self.governance.activate_notice(self.qa, "contact-governance", 2)
        item = next(row for row in activated["notices"] if row["version"] == 2)
        self.assertTrue(item["active"])
        self.assertEqual(item["approved_by"]["id"], "USR-QA")

    def test_m32_8_registra_contacto_sintetico_bajo_compuerta(self):
        self.prepare_outbox()
        result = self.communications.process(self.admin)
        self.assertGreater(len(result["accepted_sandbox"]), 0)
        self.assertEqual(result["governance_blocked"], [])
        self.assertFalse(result["real_contact_performed"])
        dashboard = self.governance.dashboard(self.admin)
        self.assertGreater(dashboard["metrics"]["synthetic_contacts"], 0)
        self.assertTrue(result["governance_audit"]["valid"])

    def test_supresion_bloquea_despacho_m32_8(self):
        self.prepare_outbox()
        self.governance.add_suppression(self.qa, {
            "subject_id": "USR-LEGAL",
            "scope": "purpose_channel",
            "purpose": "professional_operational",
            "channel": "email",
            "reason": "Bloqueo sintético verificado",
        })
        result = self.communications.process(self.admin)
        self.assertGreater(len(result["governance_blocked"]), 0)
        self.assertEqual(result["accepted_sandbox"], [])
        queue = self.communications.queue(self.admin)
        self.assertGreater(queue["metrics"]["dead_letter"], 0)

    def test_especialista_no_administra_ni_consulta_terceros(self):
        with self.assertRaises(PermissionDenied):
            self.governance.record_relationship(self.legal, {
                "subject_id": "USR-CLIENT",
                "relationship_type": "client",
                "lawful_basis": "contract",
                "status": "active",
                "evidence_reference": "X",
            })
        with self.assertRaises(PermissionDenied):
            self.governance.subject(self.legal, "USR-OTHER")
        own = self.governance.subject(self.legal, "USR-LEGAL")
        self.assertEqual(own["subject_id"], "USR-LEGAL")

    def test_manipulacion_bloquea_nuevas_decisiones(self):
        self.governance.evaluate(
            self.admin,
            subject_id="USR-LEGAL",
            purpose="professional_operational",
            channel="email",
        )
        path = self.root / "approval-desk" / "contact-governance" / "events.jsonl"
        path.write_text(path.read_text(encoding="utf-8").replace('"decision.recorded"', '"decision.altered"', 1), encoding="utf-8")
        self.assertFalse(self.governance.verify_chain()["valid"])
        with self.assertRaises(ContactGovernanceIntegrityError):
            self.governance.evaluate(
                self.admin,
                subject_id="USR-LEGAL",
                purpose="professional_operational",
                channel="email",
            )


class ContactGovernanceStaticM329Tests(TestCase):
    def test_handler_y_activos_estan_conectados(self):
        root = Path(__file__).resolve().parents[1]
        index = (root / "app" / "index.html").read_text(encoding="utf-8")
        script = (root / "app" / "modules" / "contact_governance_m32_9.js").read_text(encoding="utf-8")
        styles = (root / "app" / "modules" / "contact_governance_m32_9.css").read_text(encoding="utf-8")
        run = (root / "run.py").read_text(encoding="utf-8")
        routes = (root / "legalai_platform" / "routes" / "m32_8_transactional_communications_routes.py").read_text(encoding="utf-8")
        self.assertIn("contact_governance_m32_9.css", index)
        self.assertIn("contact_governance_m32_9.js", index)
        self.assertIn("Consentimientos, preferencias y supresiones", script)
        self.assertIn("m329-center", styles)
        self.assertIn("http_handler_m32_9 import Handler", run)
        self.assertIn("http_handler_m32_8 import Handler", run)
        self.assertIn("GovernedTransactionalCommunications", routes)
