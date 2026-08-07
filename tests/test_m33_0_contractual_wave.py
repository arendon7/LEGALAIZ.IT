from __future__ import annotations

import re
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
RUNTIME_MODULES = ROOT / "legalai_runtime_modules"
if str(RUNTIME_MODULES) not in sys.path:
    sys.path.insert(0, str(RUNTIME_MODULES))

from co_ar_001_document_factory_v251 import CoAr001DocumentFactoryV251
from co_ar_001_test_fixtures_v249 import complete_answers as lease_answers
from co_em_004_document_factory_v248 import CoEm004DocumentFactoryV248
from co_la_002_document_factory_v240 import CoLa002DocumentFactoryV240
from document_standard_v33 import audit_docx_legal_standard


class ControlledEvaluator:
    def __init__(self, documents: list[str], blocks: list[str] | None = None):
        self.documents = list(documents)
        self.blocks = list(blocks or [])

    def evaluate(self, answers):
        return {
            "blocked": False,
            "missing_fields": [],
            "documents": self.documents,
            "readiness": "ready_for_human_review",
            "status": "ready_for_human_review",
            "professional_review_required": True,
            "professional_reviews": ["Revisión jurídica sustantiva", "QA visual humano"],
            "review_requirements": ["Revisión jurídica sustantiva", "QA visual humano"],
            "findings": [],
            "blockers": [],
            "warnings": [],
            "blocks": self.blocks,
        }


class LeaseEvaluator(ControlledEvaluator):
    """La fábrica histórica de arrendamiento consulta también nombres de documentos."""

    def __init__(self):
        super().__init__(["DOC-AR-CONTRACT-001"], ["AR-BASE", "AR-PROPERTY", "AR-ECONOMICS"])
        self.documents = [
            {"id": "DOC-AR-CONTRACT-001", "name": "Contrato de arrendamiento"},
            {"id": "ANX-AR-INVENTORY-001", "name": "Inventario y estado del inmueble"},
            {"id": "ANX-AR-DELIVERY-001", "name": "Acta de entrega"},
            {"id": "ANX-AR-RETURN-001", "name": "Acta de restitución"},
        ]

    def evaluate(self, answers):
        result = super().evaluate(answers)
        result["documents"] = [item["id"] for item in self.documents]
        return result


def employment_answers() -> dict:
    return {
        "employer": {"type": "legal_person", "legalName": "Soluciones Andinas S.A.S.", "identificationNumber": "901234567-8"},
        "employerSignatory": {"fullName": "María Fernanda Gómez Ruiz", "positionOrCapacity": "representante legal"},
        "worker": {"fullName": "Carlos Andrés Pérez López", "identificationNumber": "1030123456"},
        "role": {
            "jobTitle": "Analista jurídico y documental",
            "purpose": "gestionar contratos, expedientes y controles de trazabilidad jurídica de la organización",
            "functionsPlacement": "full_in_contract",
            "essentialFunctions": [
                "Revisar y preparar contratos, comunicaciones y actas conforme a los procedimientos internos.",
                "Mantener actualizados los expedientes, matrices de obligaciones y registros de versiones.",
                "Reportar riesgos jurídicos, vencimientos e inconsistencias documentales.",
                "Coordinar la entrega de soportes y preservar la confidencialidad.",
            ],
        },
        "work": {"mainWorkplace": "Medellín, Antioquia", "modality": "onsite", "actualStartDate": "2026-08-10"},
        "schedule": {"weeklyHours": 42, "type": "fixed"},
        "compensation": {"baseSalary": 4200000, "salaryType": "ordinary"},
    }


def nda_answers() -> dict:
    return {
        "party_a": {
            "identification": {"type": "legal_person", "name": "Soluciones Andinas S.A.S.", "id_number": "901234567-8", "address": "Medellín, Antioquia", "email": "juridica@demo.legalaiz.it"},
            "signatory": {"name": "María Fernanda Gómez Ruiz", "id_number": "43678901", "capacity": "representante legal", "authority_source": "certificado de existencia y representación legal"},
        },
        "party_b": {
            "identification": {"type": "legal_person", "name": "Tecnología Segura S.A.S.", "id_number": "900765432-1", "address": "Bogotá D.C.", "email": "contratos@demo.legalaiz.it"},
            "signatory": {"name": "Juan David Torres", "id_number": "79543210", "capacity": "representante legal", "authority_source": "certificado de existencia y representación legal"},
        },
        "agreement": {"type": "mutual", "reciprocal": True, "purpose": "evaluar y ejecutar una integración tecnológica y documental entre las partes", "reference": "Proyecto Integración Segura 2026"},
        "information": {"categories": "arquitectura, código, documentación técnica, modelos jurídicos, precios, estrategias y datos operativos", "formats_sources": "documentos, reuniones, repositorios autorizados, demostraciones y accesos controlados", "exclusions": "información pública, conocida legítimamente o desarrollada de forma independiente"},
        "access": {"authorized_recipients": "personal directivo, jurídico y técnico expresamente asignado", "representatives": "asesores y subcontratistas autorizados y sometidos a obligaciones equivalentes", "need_to_know": "mínimo privilegio", "permitted_use": "evaluación, integración, pruebas y ejecución del proyecto", "compelled_disclosure": "notificación previa cuando sea posible y revelación mínima"},
        "security": {"controls": {"level": "enhanced", "technical": "cifrado, MFA, registro de accesos, segregación y copias de seguridad", "organizational": "mínimo privilegio, capacitación, gestión de terceros y respuesta a incidentes", "physical": "control de ingreso y custodia de soportes"}, "incident_protocol": "notificación, contención, investigación y preservación de evidencia"},
        "data": {"personal": False, "roles": {}, "lifecycle": "conservación durante la finalidad y eliminación al cierre", "crossborder": False},
        "ip": {"results_allocation": "case_by_case", "preexisting_materials": "herramientas, bibliotecas, plantillas y conocimientos identificados por cada parte", "source_code_reverse_engineering": "prohibición salvo autorización o excepción legal"},
        "ai": {"used": True, "training_outputs": "uso controlado sin entrenamiento ni retención con información protegida"},
        "term_remedies": {"agreement_years": 2, "ordinary_confidentiality_years": 5, "trade_secret_rule": "while_secret", "penalty_or_liability": "responsabilidad por daños directos, probados y causalmente vinculados"},
        "closure_confirmation": {"return_destroy": "devolución o eliminación segura", "retained_copies": "conservación limitada por obligación legal o defensa de derechos", "dispute_mechanism": "negotiation_conciliation"},
    }


def primary_path(factory, manifest: dict, document_id: str) -> tuple[Path, dict]:
    item = next(item for item in manifest["documents"] if item["id"] == document_id)
    candidates = sorted((factory.output_dir / manifest["generation_id"]).rglob(item["filename"]))
    if len(candidates) != 1:
        raise AssertionError(candidates)
    return candidates[0], item


def clause_count(document: Document) -> int:
    pattern = re.compile(r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|VIGÉSIMA|TRIGÉSIMA|CUADRAGÉSIMA)")
    return sum(bool(pattern.match((paragraph.text or "").strip())) for paragraph in document.paragraphs)


def approval_status(manifest: dict, key: str) -> str:
    value = manifest.get(key)
    return str(value.get("status") if isinstance(value, dict) else value).casefold()


class ContractualWaveM330Tests(unittest.TestCase):
    def _assert_primary(self, factory, answers, document_id: str, expected_title: str, expected_clause: str, minimum_words: int, minimum_clauses: int):
        manifest = factory.generate(answers, actor={"id": "qa-m33", "role": "qa"})
        path, item = primary_path(factory, manifest, document_id)
        report = audit_docx_legal_standard(path)
        self.assertTrue(report["valid"], report["findings"])
        self.assertEqual(item.get("document_standard"), "M33.0")
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        self.assertIn(expected_title, text)
        self.assertIn(expected_clause, text)
        self.assertIn("CONTROL DE USO, FUENTES Y REVISIÓN", text)
        self.assertNotIn("________", text)
        self.assertGreater(len(text.split()), minimum_words)
        self.assertGreaterEqual(clause_count(document), minimum_clauses)
        self.assertEqual(approval_status(manifest, "legal_approval"), "pending")
        self.assertEqual(approval_status(manifest, "qa_approval"), "pending")
        self.assertFalse(bool(manifest.get("released", False)))
        return manifest, path

    def test_employment_m33(self):
        with tempfile.TemporaryDirectory() as tmp:
            factory = CoLa002DocumentFactoryV240(Path(tmp), ControlledEvaluator(["DOC-LA-CONTRACT-001", "ANX-LA-FUN-001"], ["LABOR_BASE", "FUNCTIONS_ANNEX"]))
            self._assert_primary(factory, employment_answers(), "DOC-LA-CONTRACT-001", "CONTRATO INDIVIDUAL DE TRABAJO A TÉRMINO INDEFINIDO", "PRIMERA: OBJETO", 1_800, 25)
            self.assertEqual(factory.VERSION, "2.40")

    def test_nda_m33(self):
        with tempfile.TemporaryDirectory() as tmp:
            factory = CoEm004DocumentFactoryV248(Path(tmp), ControlledEvaluator(["DOC-EM4-NDA-001"], ["NDA_BASE", "SECURITY", "IP"]))
            self._assert_primary(factory, nda_answers(), "DOC-EM4-NDA-001", "ACUERDO DE CONFIDENCIALIDAD", "PRIMERA: OBJETO", 1_800, 20)
            self.assertEqual(factory.VERSION, "2.48")

    def test_lease_m33(self):
        with tempfile.TemporaryDirectory() as tmp:
            factory = CoAr001DocumentFactoryV251(Path(tmp), LeaseEvaluator())
            self._assert_primary(factory, lease_answers(), "DOC-AR-CONTRACT-001", "CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA", "PRIMERA: OBJETO", 1_800, 25)
            self.assertEqual(factory.VERSION, "2.51")


if __name__ == "__main__":
    unittest.main()
