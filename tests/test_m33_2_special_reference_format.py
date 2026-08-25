from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx_builder import build_docx
from m33_2_special_reference_format import (
    apply_m33_2_special_format,
    classify_m33_2_special_document,
)


class SpecialReferenceFormatM332Tests(unittest.TestCase):
    def _build(
        self,
        root: Path,
        *,
        title: str,
        product_code: str,
        heading: str = "PRIMERA: OBJETO",
    ) -> Path:
        path = root / "special_m33_2.docx"
        build_docx(
            path,
            title,
            "Documento sujeto a verificación y revisión profesional",
            [],
            [
                {
                    "heading": heading,
                    "paragraphs": [
                        "El documento conserva su contenido jurídico y únicamente cambia su composición editorial para facilitar lectura, firma, cotejo y trazabilidad.",
                    ],
                },
                {
                    "heading": "II. DATOS DE CONTROL",
                    "table": [
                        ("Elemento", "Información"),
                        ("Parte", "Persona jurídica demostrativa"),
                        ("Estado", "Por verificar"),
                    ],
                },
            ],
            product_code=product_code,
            enforce_legal_standard=True,
            append_default_control=False,
        )
        return path

    def test_residual_inventory_is_fully_classified(self):
        expected = {
            ("CO-CD-001", "Protocolo de actuación por posible suplantación de identidad"): "guide",
            ("CO-CD-003", "Ejercicio del derecho de retracto"): "communication",
            ("CO-CD-003", "Terminación por falta de entrega"): "communication",
            ("CO-CD-004", "Estado de cuenta reconciliado y liquidación de referencia"): "statement",
            ("CO-CD-004", "Acuerdo de pago"): "agreement",
            ("CO-CD-004", "Pagaré"): "note",
            ("CO-CD-004", "Carta de instrucciones para diligenciamiento de pagaré"): "instructions",
            ("CO-CD-004", "Recibo de pago y actualización de saldo"): "receipt",
            ("CO-CD-004", "Paz y salvo o constancia de cierre"): "certificate",
            ("CO-TR-001", "Autorización de gestión y consulta del expediente SAST"): "authorization",
            ("CO-TR-001", "Resumen consolidado de verificación SAST y siguientes actuaciones"): "guide",
            ("CO-TR-002", "Guía de radicación, alertas procesales y cierre del caso"): "guide",
        }
        for (product_code, title), profile in expected.items():
            with self.subTest(product_code=product_code, title=title):
                self.assertEqual(
                    classify_m33_2_special_document(product_code, title),
                    profile,
                )

    def test_agreement_merges_ordinal_heading_and_uses_instrument_title(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Acuerdo de pago"
            path = self._build(root, title=title, product_code="CO-CD-004")
            result = apply_m33_2_special_format(
                path,
                product_code="CO-CD-004",
                title=title,
            )
            self.assertTrue(result["applied"])
            self.assertEqual(result["variant"], "agreement")
            self.assertEqual(result["merged_ordinal_clauses"], 1)

            document = Document(path)
            title_p = next(p for p in document.paragraphs if p.text.strip() == title)
            self.assertEqual(title_p.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(all(run.underline for run in title_p.runs if run.text))
            self.assertTrue(
                any(p.text.startswith("PRIMERA. OBJETO:") for p in document.paragraphs)
            )
            visible = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            self.assertNotIn("LegalAIZ.it", visible[:2])

    def test_guide_uses_non_underlined_title_and_structured_table(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Guía de radicación, alertas procesales y cierre del caso"
            path = self._build(
                root,
                title=title,
                product_code="CO-TR-002",
                heading="I. RUTA DE GESTIÓN",
            )
            result = apply_m33_2_special_format(
                path,
                product_code="CO-TR-002",
                title=title,
            )
            self.assertTrue(result["applied"])
            self.assertEqual(result["variant"], "guide")

            document = Document(path)
            title_p = next(p for p in document.paragraphs if p.text.strip() == title)
            self.assertTrue(all(not run.underline for run in title_p.runs if run.text))
            table = document.tables[0]
            shading = table.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertIsNotNone(shading)
            self.assertEqual(shading.get(qn("w:fill")), "0D1324")
            self.assertIsNotNone(
                table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
            )

    def test_unrelated_analytical_document_is_not_reformatted(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            title = "Diagnóstico jurídico y operativo de barrera en salud"
            path = self._build(root, title=title, product_code="CO-CD-004")
            before = path.read_bytes()
            result = apply_m33_2_special_format(
                path,
                product_code="CO-CD-004",
                title=title,
            )
            self.assertFalse(result["applied"])
            self.assertEqual(result["reason"], "non_special_document")
            self.assertEqual(path.read_bytes(), before)


if __name__ == "__main__":
    unittest.main()
