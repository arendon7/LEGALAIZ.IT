from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from types import ModuleType

from docx import Document

from docx_builder import build_docx
from legalai_platform.runtime_m33_overrides import _install_m33_docx_presentation_policy
from m33_wave3_runtime import document_specs_m33_all
from tests.test_m33_0_procedural_wave import PRODUCTS, labor_fixture


def labor_specs():
    answers, result = labor_fixture()
    specs = document_specs_m33_all(
        "CASE-M33-LABOR-FINAL",
        "CO-LA-001",
        answers,
        result,
        PRODUCTS["CO-LA-001"],
        "2026-08-08T08:00:00-05:00",
        [],
    )
    return answers, result, specs


def spec_of(specs: list[dict], kind: str) -> dict:
    return next(spec for spec in specs if spec.get("kind") == kind)


def visible_docx_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


class LaborProceduralFinalizeM330Tests(unittest.TestCase):
    def test_calculation_has_one_reconciled_data_layer(self):
        _, _, specs = labor_specs()
        calculation = spec_of(specs, "calculation")
        text = json.dumps(calculation, ensure_ascii=False)
        self.assertIn("COP $6.173.734", text)
        self.assertIn("Laura Isabel Gómez Pérez", text)
        self.assertIn("210", text)
        self.assertIn("ANEXO No. 1 — TRAZA REPRODUCIBLE DE ESTA LIQUIDACIÓN", text)
        self.assertNotIn("ANEXO No. 1 — MATRICES DEL MOTOR DE LIQUIDACIÓN", text)
        self.assertNotIn("No informado", text)
        self.assertNotIn("Dato pendiente de verificación", text)
        self.assertNotIn('["Duración 30/360", "0 días"]', text)

    def test_calculation_reconciles_gross_prior_payments_and_net(self):
        _, result, specs = labor_specs()
        calculation = spec_of(specs, "calculation")
        text = json.dumps(calculation, ensure_ascii=False)
        c = result["calculation"]
        self.assertEqual(c["gross_total"] - c["prior_paid_total"], c["total"])
        self.assertIn("COP $6.173.734 - COP $0 = COP $6.173.734", text)
        self.assertIn("Número de líneas calculadas", text)
        self.assertIn('"5"', text)

    def test_claim_individualizes_amounts_and_prescription_control(self):
        _, _, specs = labor_specs()
        claim = spec_of(specs, "claim")
        text = json.dumps(claim, ensure_ascii=False)
        self.assertIn("COP $6.173.734", text)
        self.assertIn("artículo 488", text)
        self.assertIn("artículo 489", text)
        self.assertIn("prueba verificable de contenido, fecha de envío", text)
        self.assertIn("La sola preparación o envío de este documento no permite afirmar que el efecto interruptivo se produjo", text)
        self.assertIn("pagar oportunamente los valores que sean reconocidos como ciertos y debidos", text.casefold())

    def test_visible_subtitles_are_client_facing(self):
        _, _, specs = labor_specs()
        for kind in ("calculation", "claim"):
            subtitle = str(spec_of(specs, kind).get("subtitle") or "")
            self.assertNotIn("Composición jurídica profunda", subtitle)
            self.assertNotIn("Modelo madurado", subtitle)
            self.assertNotIn("M33.0", subtitle)

    def test_internal_control_keeps_legal_sources_but_is_not_public(self):
        _, _, specs = labor_specs()
        for kind in ("calculation", "claim"):
            spec = spec_of(specs, kind)
            public_text = json.dumps(spec.get("sections") or [], ensure_ascii=False)
            controls = spec.get("internal_review_sections") or []
            self.assertTrue(spec.get("internal_controls_externalized"))
            self.assertFalse(any(section.get("_type") == "control" for section in spec.get("sections") or []))
            self.assertTrue(any(section.get("_suppress_default_control") for section in spec.get("sections") or []))
            self.assertGreaterEqual(len(controls), 1)
            self.assertNotIn("aprobación jurídica y QA", public_text)
            self.assertNotIn("mismo hash", public_text)
            self.assertNotIn("M33-test", public_text)
            self.assertNotIn("Versión del motor", public_text)
        internal = json.dumps(spec_of(specs, "calculation").get("internal_review_sections") or [], ensure_ascii=False)
        self.assertIn("Código Sustantivo del Trabajo", internal)
        self.assertIn("Ley 52 de 1975", internal)
        self.assertIn("Ley 2466 de 2025", internal)

    def test_runtime_policy_only_suppresses_marked_sections(self):
        core_module = ModuleType("fake_core_m33")
        calls: list[dict] = []

        def fake_builder(*args, **kwargs):
            calls.append(dict(kwargs))
            return kwargs

        core_module.build_docx = fake_builder
        wrapped = _install_m33_docx_presentation_policy(core_module)
        marked = [{"heading": "Documento cliente", "_suppress_default_control": True}]
        ordinary = [{"heading": "Documento ordinario"}]

        wrapped(Path("marked.docx"), "T", "S", [], marked)
        self.assertFalse(calls[-1].get("append_default_control"))
        wrapped(Path("ordinary.docx"), "T", "S", [], ordinary)
        self.assertNotIn("append_default_control", calls[-1])
        self.assertIs(wrapped, _install_m33_docx_presentation_policy(core_module))

    def test_calculation_and_claim_render_strictly_without_internal_box(self):
        _, _, specs = labor_specs()
        with tempfile.TemporaryDirectory() as tmp:
            for kind in ("calculation", "claim"):
                spec = spec_of(specs, kind)
                path = Path(tmp) / f"{kind}.docx"
                build_docx(
                    path,
                    spec["title"],
                    spec.get("subtitle", ""),
                    [],
                    spec["sections"],
                    product_code="CO-LA-001",
                    enforce_legal_standard=True,
                    append_default_control=not bool(spec.get("internal_controls_externalized")),
                )
                self.assertTrue(path.is_file())
                self.assertGreater(path.stat().st_size, 5_000)
                document = Document(path)
                self.assertGreater(len(document.paragraphs), 5)
                visible = visible_docx_text(document)
                self.assertNotIn("CONTROL DE USO", visible)
                self.assertNotIn("aprobación jurídica y QA", visible)
                self.assertNotIn("misma revisión y hash", visible)
                self.assertNotIn("mismo hash", visible)
                self.assertNotIn("M33-test", visible)
                self.assertNotIn("Versión del motor", visible)


if __name__ == "__main__":
    unittest.main()
