from __future__ import annotations

from datetime import datetime, timedelta
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase
import sqlite3

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace, PermissionDenied
from legalai_platform.approval_notification_center import ApprovalNotificationCenter
from legalai_platform.transactional_communications import (
    CommunicationsIntegrityError,
    TransactionalCommunications,
)


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class TransactionalCommunicationsM328Tests(TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = self.root / "m328.db"
        self.fixed_now = datetime(2026, 8, 6, 9, 30, tzinfo=BOGOTA)
        con = self.db()
        con.executescript(
            """
            CREATE TABLE users(
              id TEXT PRIMARY KEY,
              name TEXT NOT NULL,
              email TEXT,
              role TEXT NOT NULL,
              specialty TEXT,
              active INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,
              owner_id TEXT,
              specialist_id TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,
              case_id TEXT NOT NULL,
              product_code TEXT NOT NULL,
              name TEXT NOT NULL,
              kind TEXT NOT NULL,
              mime_type TEXT NOT NULL,
              file_path TEXT,
              updated_at TEXT NOT NULL
            );
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana Administradora", "ana@example.test", "admin", "Operación", 1),
                ("USR-QA", "Quinn QA", "qa@example.test", "qa", "QA", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-OTHER", "Otro abogado", "otro@example.test", "specialist", "Litigios", 1),
                ("USR-CLIENT", "Cliente", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = self.root / f"{code}_{index:02d}.docx"
            self.write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), self.fixed_now.isoformat()),
            )
        con.commit(); con.close()
        self.admin = {"id":"USR-ADMIN","role":"admin","name":"Ana Administradora"}
        self.qa = {"id":"USR-QA","role":"qa","name":"Quinn QA"}
        self.legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}
        self.other = {"id":"USR-OTHER","role":"specialist","name":"Otro abogado"}
        self.client = {"id":"USR-CLIENT","role":"client","name":"Cliente"}
        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            document_lookup=self.document_lookup,
            access_check=self.access_check,
            upload_validator=self.upload_validator,
        )
        self.operations = ApprovalDeskOperations(
            self.root / "approval-desk",
            workspace=self.workspace,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.center = ApprovalNotificationCenter(
            self.root / "approval-desk",
            operations=self.operations,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )
        self.communications = TransactionalCommunications(
            self.root / "approval-desk",
            notification_center=self.center,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )

    def tearDown(self):
        self.temporary.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, code: str):
        document = Document()
        document.add_heading(f"DOCUMENTO DEMOSTRATIVO {code}", 0)
        document.add_paragraph("Contenido sintético para validar comunicaciones M32.8.")
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
        if not row: return False
        if user["role"] in {"admin", "qa"}: return True
        if user["role"] == "client": return row["owner_id"] == user["id"]
        return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

    def document_lookup(self, user, document_id):
        con = self.db()
        row = con.execute(
            "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
            (document_id,),
        ).fetchone()
        con.close()
        if not row or not self.access_check(user, row["case_id"]): return None
        return dict(row)

    @staticmethod
    def upload_validator(filename, data):
        return DOCX_MIME, sha256(data).hexdigest(), "clean:test"

    @staticmethod
    def first_case():
        return "DSK-DOC-01"

    def prepare_outbox(self):
        self.operations.sync_portfolio(self.admin)
        self.operations.update_assignment(self.admin, self.first_case(), "USR-LEGAL", "USR-ADMIN")
        self.operations.update_deadline(
            self.admin,
            self.first_case(),
            (self.fixed_now - timedelta(hours=2)).isoformat(),
            24,
        )
        self.center.update_policy(self.admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
        })
        evaluation = self.center.evaluate(self.admin, self.first_case())
        self.assertGreater(len(evaluation["queued_messages"]), 0)

    def test_sincronizacion_es_idempotente_y_minimiza_datos(self):
        self.prepare_outbox()
        first = self.communications.sync_outbox(self.admin)
        self.assertGreater(len(first["imported_dispatches"]), 0)
        second = self.communications.sync_outbox(self.admin)
        self.assertEqual(second["imported_dispatches"], [])
        queue = self.communications.queue(self.admin)
        self.assertFalse(queue["recipient_addresses_stored"])
        self.assertFalse(queue["contains_document_content"])
        self.assertTrue(all("email" not in item for item in queue["dispatches"]))
        self.assertTrue(all(item["attachments"] == [] for item in queue["dispatches"]))
        self.assertTrue(all(len(item["template_sha256"]) == 64 for item in queue["dispatches"]))

    def test_procesamiento_sandbox_no_declara_entrega_real(self):
        self.prepare_outbox(); self.communications.sync_outbox(self.admin)
        result = self.communications.process(self.admin)
        self.assertGreater(len(result["accepted_sandbox"]), 0)
        self.assertFalse(result["real_delivery_performed"])
        queue = self.communications.queue(self.admin)
        self.assertGreater(queue["metrics"]["accepted_sandbox"], 0)
        accepted = [item for item in queue["dispatches"] if item["status"] == "accepted_sandbox"]
        self.assertTrue(all(item["provider"] == "sandbox" for item in accepted))
        self.assertTrue(all(item["real_delivery"] is False for item in accepted))
        self.assertTrue(all("recipient_address" not in item for item in accepted))

    def test_recibo_sintetico_es_idempotente_y_no_es_prueba_real(self):
        self.prepare_outbox(); self.communications.sync_outbox(self.admin); self.communications.process(self.admin)
        dispatch = next(item for item in self.communications.queue(self.admin)["dispatches"] if item["status"] == "accepted_sandbox")
        first = self.communications.record_receipt(
            self.qa,
            dispatch["dispatch_id"],
            provider_status="delivered",
            provider_event_id="EVT-SBX-001",
            synthetic=True,
        )
        self.assertGreater(first["metrics"]["delivered_sandbox"], 0)
        events_before = self.communications.verify_chain()["events"]
        self.communications.record_receipt(
            self.qa,
            dispatch["dispatch_id"],
            provider_status="delivered",
            provider_event_id="EVT-SBX-001",
            synthetic=True,
        )
        self.assertEqual(self.communications.verify_chain()["events"], events_before)
        delivered = next(item for item in self.communications.queue(self.admin)["dispatches"] if item["dispatch_id"] == dispatch["dispatch_id"])
        self.assertEqual(delivered["status"], "delivered_sandbox")
        self.assertTrue(delivered["receipts"][0]["synthetic"])
        self.assertFalse(delivered["real_delivery"])

    def test_plantilla_requiere_activacion_independiente(self):
        created = self.communications.create_template_version(self.admin, {
            "template_id": "professional-alert",
            "name": "Alerta revisada",
            "subject": "LegalAIZ.it · {{product_code}} · {{title}}",
            "body": "Hola {{recipient_name}}. Revise el expediente {{case_id}} antes de {{due_at}}.",
        })
        self.assertEqual(created["template"]["version"], 2)
        with self.assertRaises(PermissionDenied):
            self.communications.activate_template(self.admin, "professional-alert", 2)
        activated = self.communications.activate_template(self.qa, "professional-alert", 2)
        row = next(item for item in activated["templates"] if item["template_id"] == "professional-alert" and item["version"] == 2)
        self.assertTrue(row["active"])
        self.assertEqual(row["approved_by"]["id"], "USR-QA")
        self.assertFalse(row["contains_document_content"])
        self.assertFalse(row["attachments_allowed"])

    def test_correo_inexistente_va_a_cola_muerta_sin_reintento(self):
        self.prepare_outbox(); self.communications.sync_outbox(self.admin)
        con = self.db(); con.execute("UPDATE users SET email=NULL WHERE id='USR-LEGAL'"); con.commit(); con.close()
        result = self.communications.process(self.admin)
        self.assertGreater(len(result["dead_lettered"]), 0)
        queue = self.communications.queue(self.admin)
        self.assertGreater(queue["metrics"]["dead_letter"], 0)

    def test_rbac_limita_procesamiento_y_bandeja_personal(self):
        self.prepare_outbox(); self.communications.sync_outbox(self.admin)
        with self.assertRaises(PermissionDenied):
            self.communications.process(self.legal)
        with self.assertRaises(PermissionDenied):
            self.communications.sync_outbox(self.legal)
        personal = self.communications.queue(self.legal)
        self.assertTrue(all(item["recipient_id"] == "USR-LEGAL" for item in personal["dispatches"]))
        other = self.communications.queue(self.other)
        self.assertEqual(other["dispatches"], [])

    def test_manipulacion_bloquea_procesamiento(self):
        self.prepare_outbox(); self.communications.sync_outbox(self.admin)
        path = self.root / "approval-desk" / "transactional-communications" / "events.jsonl"
        path.write_text(path.read_text(encoding="utf-8").replace('"dispatch.imported"', '"dispatch.altered"', 1), encoding="utf-8")
        self.assertFalse(self.communications.verify_chain()["valid"])
        with self.assertRaises(CommunicationsIntegrityError):
            self.communications.process(self.admin)


class TransactionalCommunicationsStaticM328Tests(TestCase):
    def test_handler_activos_y_marcadores_estan_conectados(self):
        root = Path(__file__).resolve().parents[1]
        index = (root / "app" / "index.html").read_text(encoding="utf-8")
        script = (root / "app" / "modules" / "transactional_communications_m32_8.js").read_text(encoding="utf-8")
        styles = (root / "app" / "modules" / "transactional_communications_m32_8.css").read_text(encoding="utf-8")
        run = (root / "run.py").read_text(encoding="utf-8")
        self.assertIn("transactional_communications_m32_8.css", index)
        self.assertIn("transactional_communications_m32_8.js", index)
        self.assertIn("Plantillas, despachos y evidencia", script)
        self.assertIn("Solo sandbox", script)
        self.assertIn("m328-center", styles)
        self.assertIn("http_handler_m32_8 import Handler", run)
        self.assertIn("http_handler_m32_7 import Handler", run)
