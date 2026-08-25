from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from legalai_platform.document_approval_desk import DocumentApprovalDesk, ReleaseBlocked
from m33_document_presentation import (
    APPROVAL_CANDIDATE_MODE,
    REVIEW_MODE,
    build_m33_presentation,
)
from m33_services_legal_finalize import compose_services_m33_final
from tests.test_m33_0_services_reference import services_answers


AUTHOR = {"id": "author-m33", "role": "author", "name": "Autor M33"}
LEGAL = {"id": "legal-m33", "role": "specialist", "name": "Especialista M33"}
QA = {"id": "qa-m33", "role": "qa", "name": "QA M33"}
ADMIN = {"id": "admin-m33", "role": "admin", "name": "Administración M33"}


def _sha(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


class M330PresentationModesTests(unittest.TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        answers = services_answers()
        answers.setdefault("service", {})["professional"] = True
        self.composition = compose_services_m33_final(answers)
        self.metadata = [
            ("Producto", "CO-EM-003"),
            ("Estándar documental", "M33.0"),
            ("Estado", "Candidato sujeto a revisión jurídica y QA"),
        ]

    def tearDown(self):
        self.temporary.cleanup()

    def test_review_copy_keeps_internal_status_and_sources(self):
        target = self.root / "review.docx"
        evidence = build_m33_presentation(
            path=target,
            title=self.composition["title"],
            subtitle=self.composition.get("subtitle") or "",
            metadata=self.metadata,
            sections=self.composition["sections"],
            product_code="CO-EM-003",
            presentation_mode=REVIEW_MODE,
        )
        text = _docx_text(target)
        self.assertIn("BORRADOR CONTROLADO", text)
        self.assertIn("CONTROL DE USO, FUENTES Y REVISIÓN", text)
        self.assertIn("Estándar documental", text)
        self.assertGreaterEqual(len(evidence.get("legal_sources") or []), 6)
        self.assertEqual(evidence["presentation_mode"], REVIEW_MODE)

    def test_approval_candidate_is_clean_but_evidence_remains_external(self):
        target = self.root / "approval-candidate.docx"
        evidence = build_m33_presentation(
            path=target,
            title=self.composition["title"],
            subtitle=self.composition.get("subtitle") or "",
            metadata=self.metadata,
            sections=self.composition["sections"],
            product_code="CO-EM-003",
            presentation_mode=APPROVAL_CANDIDATE_MODE,
        )
        text = _docx_text(target)
        self.assertNotIn("BORRADOR CONTROLADO", text)
        self.assertNotIn("NO FIRMAR", text)
        self.assertNotIn("CONTROL DE USO, FUENTES Y REVISIÓN", text)
        self.assertNotIn("Estándar documental", text)
        self.assertNotIn("Candidato jurídico M33.0", text)
        self.assertGreaterEqual(len(evidence.get("legal_sources") or []), 6)
        self.assertEqual(evidence["presentation_mode"], APPROVAL_CANDIDATE_MODE)
        self.assertIn("mismo que posteriormente puede liberarse", evidence["release_rule"])

    def test_release_copies_exact_approved_candidate_hash(self):
        candidate = self.root / "CO-EM-003_aprobacion.docx"
        build_m33_presentation(
            path=candidate,
            title=self.composition["title"],
            subtitle=self.composition.get("subtitle") or "",
            metadata=self.metadata,
            sections=self.composition["sections"],
            product_code="CO-EM-003",
            presentation_mode=APPROVAL_CANDIDATE_MODE,
        )
        candidate_hash = _sha(candidate)

        desk = DocumentApprovalDesk(self.root / "desk")
        case = desk.create_case(
            case_id="CASE-M33-PRESENTATION",
            product_code="CO-EM-003",
            document_id="DOC-EM-CONTRACT-001",
            title=self.composition["title"],
            actor=AUTHOR,
            source_generation_id="GEN-M33-PRESENTATION",
        )
        revision = desk.add_revision(
            case_id=case["case_id"],
            source_file=candidate,
            actor=AUTHOR,
            note="Instrumento limpio sometido a aprobación M33.0.",
        )
        self.assertEqual(revision["sha256"], candidate_hash)

        with self.assertRaises(ReleaseBlocked):
            desk.release(
                case_id=case["case_id"],
                revision_id=revision["revision_id"],
                actor=ADMIN,
                expected_sha256=candidate_hash,
            )

        desk.approve(
            case_id=case["case_id"],
            revision_id=revision["revision_id"],
            approval_type="legal",
            decision="approve",
            actor=LEGAL,
            comment="Revisión jurídica del instrumento limpio.",
            expected_sha256=candidate_hash,
        )
        desk.approve(
            case_id=case["case_id"],
            revision_id=revision["revision_id"],
            approval_type="qa",
            decision="approve",
            actor=QA,
            comment="QA del mismo instrumento y hash.",
            expected_sha256=candidate_hash,
        )
        release = desk.release(
            case_id=case["case_id"],
            revision_id=revision["revision_id"],
            actor=ADMIN,
            expected_sha256=candidate_hash,
        )
        self.assertEqual(release["sha256"], candidate_hash)
        self.assertEqual(release["status"], "released_exact_hash")


if __name__ == "__main__":
    unittest.main()
