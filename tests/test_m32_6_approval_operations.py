from __future__ import annotations

from datetime import datetime, timedelta
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase
from zipfile import ZipFile
import json
import sqlite3

from docx import Document

from legalai_platform.approval_desk_operations import (
    ApprovalDeskOperations,
    OperationsIntegrityError,
    PORTFOLIO_CODES,
)
from legalai_platform.approval_desk_workspace import (
    ApprovalDeskError,
    ApprovalDeskWorkspace,
    PermissionDenied,
)
from zoneinfo import ZoneInfo


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ApprovalDeskOperationsM326Tests(TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = self.root / "m326.db"
        self.fixed_now = datetime(2026, 8, 6, 9, 0, tzinfo=BOGOTA)
        con = self.db()
        con.executescript(
            """
            CREATE TABLE users(
              id TEXT PRIMARY KEY,
              name TEXT NOT NULL,
              role TEXT NOT NULL,
              specialty TEXT,
              active INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,
              owner_id TEXT,
              specialist_id TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,
              case_id TEXT NOT NULL,
              product_code TEXT NOT NULL,
              name TEXT NOT NULL,
              kind TEXT NOT NULL,
              mime_type TEXT NOT NULL,
              file_path TEXT,
              updated_at TEXT NOT NULL
            );
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana QA", "admin", "QA documental", 1),
                ("USR-LEGAL", "María Jurídica", "specialist", "Contratos", 1),
                ("USR-OTHER", "Otro abogado", "specialist", "Litigios", 1),
                ("USR-CLIENT", "Cliente", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = self.root / f"{code}_{index:02d}.docx"
            self.write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), self.fixed_now.isoformat()),
            )
        con.commit(); con.close()
        self.admin = {"id":"USR-ADMIN","role":"admin","name":"Ana QA"}
        self.legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}
        self.other = {"id":"USR-OTHER","role":"specialist","name":"Otro abogado"}
        self.client = {"id":"USR-CLIENT","role":"client","name":"Cliente"}
        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            document_lookup=self.document_lookup,
            access_check=self.access_check,
            upload_validator=self.upload_validator,
        )
        self.operations = ApprovalDeskOperations(
            self.root / "approval-desk",
            workspace=self.workspace,
            db_factory=self.db,
            now_factory=lambda: self.fixed_now,
        )

    def tearDown(self):
        self.temporary.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, code: str):
        document = Document()
        document.add_heading(f"DOCUMENTO DEMOSTRATIVO {code}", 0)
        document.add_paragraph("Contenido sintético para validar operación y trazabilidad M32.6.")
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
        if not row: return False
        if user["role"] == "admin": return True
        if user["role"] == "client": return row["owner_id"] == user["id"]
        return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

    def document_lookup(self, user, document_id):
        con = self.db()
        row = con.execute(
            "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
            (document_id,),
        ).fetchone()
        con.close()
        if not row or not self.access_check(user, row["case_id"]): return None
        return dict(row)

    @staticmethod
    def upload_validator(filename, data):
        return DOCX_MIME, sha256(data).hexdigest(), "clean:test"

    def sync(self):
        result = self.operations.sync_portfolio(self.admin)
        self.assertEqual(result["bootstrap"]["created_count"], 11)
        return result

    def first_case(self):
        return "DSK-DOC-01"

    def assign_first(self):
        return self.operations.update_assignment(self.admin, self.first_case(), "USR-LEGAL", "USR-ADMIN")

    def test_sincronizacion_cubre_los_once_productos_sin_aprobarlos(self):
        result = self.sync()
        self.assertEqual(result["portfolio"]["expected_products"], 11)
        self.assertEqual(result["portfolio"]["covered_products"], 11)
        self.assertEqual(result["portfolio"]["coverage_percent"], 100)
        portfolio = self.operations.portfolio(self.admin)
        self.assertEqual(len(portfolio["cases"]), 11)
        self.assertEqual(portfolio["metrics"]["released"], 0)
        self.assertEqual(portfolio["metrics"]["unassigned"], 11)

    def test_asignacion_es_exclusiva_de_admin_y_separa_funciones(self):
        self.sync()
        with self.assertRaises(PermissionDenied):
            self.operations.update_assignment(self.legal, self.first_case(), "USR-LEGAL", "USR-ADMIN")
        with self.assertRaises(ApprovalDeskError):
            self.operations.update_assignment(self.admin, self.first_case(), "USR-LEGAL", "USR-LEGAL")
        result = self.assign_first()
        self.assertEqual(result["state"]["operations"]["assigned_specialist"]["id"], "USR-LEGAL")
        con = self.db(); assigned = con.execute("SELECT specialist_id FROM cases WHERE id='CASE-01'").fetchone()[0]; con.close()
        self.assertEqual(assigned, "USR-LEGAL")
        self.assertEqual(self.workspace.detail(self.legal, self.first_case())["source_case_id"], "CASE-01")
        with self.assertRaises(PermissionDenied):
            self.workspace.detail(self.other, self.first_case())

    def test_prioridad_y_vencimiento_generan_estado_sla_y_alerta(self):
        self.sync(); self.assign_first()
        high = self.operations.update_priority(self.admin, self.first_case(), "high")["state"]
        self.assertEqual(high["operations"]["sla_hours"], 24)
        self.assertEqual(high["sla"]["status"], "in_time")
        overdue_at = (self.fixed_now - timedelta(hours=2)).isoformat()
        overdue = self.operations.update_deadline(self.admin, self.first_case(), overdue_at, 24)["state"]
        self.assertEqual(overdue["sla"]["status"], "overdue")
        self.assertIn("sla_overdue", {item["code"] for item in overdue["alerts"]})
        acknowledged = self.operations.acknowledge_alert(self.legal, self.first_case(), "sla_overdue", "Escalado a administración")["state"]
        alert = next(item for item in acknowledged["alerts"] if item["code"] == "sla_overdue")
        self.assertTrue(alert["acknowledged"])

    def test_notas_y_actividad_consolidan_operacion_y_aprobacion(self):
        self.sync(); self.assign_first()
        self.operations.add_note(self.legal, self.first_case(), "Se inicia cotejo de variables y anexos.")
        detail = self.operations.case_detail(self.legal, self.first_case())
        sources = {item["source"] for item in detail["activity"]}
        self.assertIn("operations", sources)
        self.assertIn("approval", sources)
        self.assertEqual(detail["operations"]["notes"][-1]["text"], "Se inicia cotejo de variables y anexos.")

    def test_cadena_operativa_alterada_bloquea_nuevos_eventos_y_exportacion(self):
        self.sync(); self.assign_first()
        self.operations.add_note(self.legal, self.first_case(), "Nota inicial íntegra.")
        path = self.root / "approval-desk" / self.first_case() / "operations.jsonl"
        content = path.read_text(encoding="utf-8").replace('"assignment.updated"', '"assignment.altered"')
        path.write_text(content, encoding="utf-8")
        self.assertFalse(self.operations.verify_chain(self.first_case())["valid"])
        with self.assertRaises(OperationsIntegrityError):
            self.operations.add_note(self.legal, self.first_case(), "No debe persistir.")
        with self.assertRaises(OperationsIntegrityError):
            self.operations.export_dossier(self.admin, self.first_case())

    def test_expediente_pendiente_declara_revision_humana_requerida(self):
        self.sync(); self.assign_first()
        target, _ = self.operations.export_dossier(self.admin, self.first_case())
        with ZipFile(target) as archive:
            self.assertIn("expediente_aprobacion.json", archive.namelist())
            self.assertIn("revision_vigente.docx", archive.namelist())
            dossier = json.loads(archive.read("expediente_aprobacion.json"))
        self.assertTrue(dossier["human_review_required"])
        self.assertFalse(dossier["professional_approval_complete"])
        self.assertEqual(dossier["case"]["product_code"], "CO-EM-003")

    def test_expediente_liberado_refleja_aprobacion_dual_del_hash(self):
        self.sync(); self.assign_first()
        detail = self.workspace.detail(self.legal, self.first_case())
        current = detail["revisions"][0]
        self.workspace.approve(self.legal, self.first_case(), {
            "revision_id": current["revision_id"], "approval_type": "legal", "decision": "approve",
            "comment": "Aprobación sintética.", "expected_sha256": current["sha256"],
        })
        self.workspace.approve(self.admin, self.first_case(), {
            "revision_id": current["revision_id"], "approval_type": "qa", "decision": "approve",
            "comment": "QA sintético.", "expected_sha256": current["sha256"],
        })
        self.workspace.release(self.admin, self.first_case(), {
            "revision_id": current["revision_id"], "expected_sha256": current["sha256"],
        })
        dossier = self.operations.build_dossier(self.admin, self.first_case())
        self.assertTrue(dossier["professional_approval_complete"])
        self.assertFalse(dossier["human_review_required"])
        self.assertEqual(dossier["release"]["sha256"], current["sha256"])


class ApprovalDeskOperationsStaticM326Tests(TestCase):
    def test_activos_handler_y_documentacion_estan_conectados(self):
        root = Path(__file__).resolve().parents[1]
        index = (root / "app" / "index.html").read_text(encoding="utf-8")
        script = (root / "app" / "modules" / "approval_operations_m32_6.js").read_text(encoding="utf-8")
        styles = (root / "app" / "modules" / "approval_operations_m32_6.css").read_text(encoding="utf-8")
        run = (root / "run.py").read_text(encoding="utf-8")
        self.assertIn("approval_operations_m32_6.css", index)
        self.assertIn("approval_operations_m32_6.js", index)
        self.assertIn("Cobertura, responsables y SLA", script)
        self.assertIn("Descargar expediente de aprobación", script)
        self.assertIn("m326-portfolio", styles)
        self.assertIn("http_handler_m32_6 import Handler", run)
