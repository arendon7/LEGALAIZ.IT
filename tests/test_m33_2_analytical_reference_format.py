from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx_builder import build_docx
from m33_2_analytical_reference_format import apply_m33_2_analytical_format


class AnalyticalReferenceFormatM332Tests(unittest.TestCase):
    def _build(self, root: Path, *, title: str = "Diagnóstico jurídico de obligación y ruta de cobro") -> Path:
        path = root / "CO-CD-004_analytical_m33_2.docx"
        build_docx(
            path,
            title,
            "Origen, exigibilidad, saldo, intereses, títulos y riesgos del expediente",
            [],
            [
                {
                    "heading": "OBJETO Y ALCANCE DEL DIAGNÓSTICO",
                    "paragraphs": [
                        "Este diagnóstico reconstruye la obligación antes de producir un requerimiento o un acuerdo y separa los hechos acreditados de las conclusiones preliminares.",
                        "La generación documental no transforma una cifra informada por una parte en saldo incontrovertible ni sustituye la revisión profesional.",
                    ],
                },
                {
                    "heading": "I. IDENTIFICACIÓN DEL EXPEDIENTE",
                    "table": [
                        ("Elemento", "Información del expediente"),
                        ("Parte acreedora", "Insumos Empresariales Demo S.A.S."),
                        ("Parte deudora", "Comercializadora Horizonte S.A.S."),
                    ],
                },
                {
                    "heading": "II. RESULTADO PRELIMINAR",
                    "paragraphs": [
                        "Clasificación preliminar: la obligación requiere cotejo documental antes de formalizar cualquier instrumento ejecutivo.",
                    ],
                    "numbered": [
                        "Verificar el negocio causal y la exigibilidad.",
                        "Conciliar pagos, ajustes e intereses contra los soportes del expediente.",
                    ],
                },
                {
                    "heading": "CONTROL DE USO, FUENTES Y REVISIÓN",
                    "_type": "control",
                    "text": "Documento sujeto a revisión jurídica y QA sobre la misma revisión.",
                },
            ],
            product_code="CO-CD-004",
            enforce_legal_standard=True,
        )
        return path

    def test_analytical_document_uses_report_hierarchy_and_table_profile(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = self._build(Path(tmp))
            result = apply_m33_2_analytical_format(
                path,
                product_code="CO-CD-004",
                title="Diagnóstico jurídico de obligación y ruta de cobro",
            )
            self.assertTrue(result["applied"])
            self.assertEqual(result["profile"], "M33.2-analytical")
            self.assertEqual(result["font"], "Book Antiqua")
            self.assertEqual(result["table_header_fill"], "0D1324")

            document = Document(path)
            nonempty = [p for p in document.paragraphs if p.text.strip()]
            self.assertEqual(nonempty[0].text, "Diagnóstico jurídico de obligación y ruta de cobro")
            self.assertNotIn("LegalAIZ.it", [p.text.strip() for p in nonempty[:3]])

            title = nonempty[0]
            self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(all(run.bold for run in title.runs if run.text))
            self.assertTrue(all(run.font.all_caps for run in title.runs if run.text))
            self.assertFalse(any(run.underline for run in title.runs if run.text))

            subtitle = next(p for p in document.paragraphs if p.text.startswith("Origen, exigibilidad"))
            self.assertEqual(subtitle.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(any(run.italic for run in subtitle.runs))

            body = next(p for p in document.paragraphs if p.text.startswith("Este diagnóstico reconstruye"))
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertAlmostEqual(body.paragraph_format.space_after.pt, 6.0, places=1)

            heading = next(p for p in document.paragraphs if p.text.strip() == "I. IDENTIFICACIÓN DEL EXPEDIENTE")
            self.assertEqual(heading.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(all(run.bold for run in heading.runs if run.text))

            table = document.tables[0]
            first_header_cell = table.rows[0].cells[0]
            header_shading = first_header_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertIsNotNone(header_shading)
            self.assertEqual(header_shading.get(qn("w:fill")), "0D1324")
            first_data_cell = table.rows[1].cells[0]
            first_col_shading = first_data_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertIsNotNone(first_col_shading)
            self.assertEqual(first_col_shading.get(qn("w:fill")), "F7F5F1")
            self.assertIsNotNone(table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")))
            self.assertTrue(all(table.rows[i]._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None for i in range(len(table.rows))))

            lead = next(p for p in document.paragraphs if p.text.startswith("Clasificación preliminar:"))
            self.assertTrue(lead.runs[0].bold)
            self.assertFalse(lead.runs[-1].bold)

    def test_formal_claim_is_not_reformatted_as_analytical_report(self):
        with tempfile.TemporaryDirectory() as tmp:
            title = "Reclamación directa de garantía legal"
            path = self._build(Path(tmp), title=title)
            before = path.read_bytes()
            result = apply_m33_2_analytical_format(path, product_code="CO-CD-004", title=title)
            self.assertFalse(result["applied"])
            self.assertEqual(result["reason"], "non_analytical_document")
            self.assertEqual(path.read_bytes(), before)

    def test_matrix_family_stays_outside_analytical_profile(self):
        with tempfile.TemporaryDirectory() as tmp:
            title = "Matriz de trazabilidad jurídica, técnica y temporal del SAST"
            path = self._build(Path(tmp), title=title)
            before = path.read_bytes()
            result = apply_m33_2_analytical_format(path, product_code="CO-TR-001", title=title)
            self.assertFalse(result["applied"])
            self.assertEqual(path.read_bytes(), before)


if __name__ == "__main__":
    unittest.main()
