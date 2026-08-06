from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase, mock
import sqlite3

from docx import Document

from legalai_platform.approval_desk_workspace import (
    ApprovalDeskError,
    ApprovalDeskWorkspace,
    ImmutableRecordError,
    PermissionDenied,
    ReleaseBlocked,
)


class ApprovalDeskWorkspaceM325Tests(TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.database = self.root / "test.db"
        con = self.db()
        con.executescript(
            """
            CREATE TABLE cases(
              id TEXT PRIMARY KEY,
              owner_id TEXT,
              specialist_id TEXT
            );
            CREATE TABLE documents(
              id TEXT PRIMARY KEY,
              case_id TEXT NOT NULL,
              product_code TEXT NOT NULL,
              name TEXT NOT NULL,
              kind TEXT NOT NULL,
              mime_type TEXT NOT NULL,
              file_path TEXT,
              updated_at TEXT NOT NULL
            );
            """
        )
        con.execute("INSERT INTO cases VALUES(?,?,?)", ("LZ-CASE-001", "USR-CLIENT", "USR-LEGAL"))
        self.source = self.root / "contrato.docx"
        self.write_docx(self.source, "Versión uno del contrato.")
        con.execute(
            "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
            (
                "DOC-001",
                "LZ-CASE-001",
                "CO-EM-003",
                "Contrato de prestación de servicios.docx",
                "main",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                str(self.source),
                "2026-08-05T22:00:00-05:00",
            ),
        )
        con.commit(); con.close()
        self.admin = {"id":"USR-ADMIN","role":"admin","name":"Ana QA"}
        self.legal = {"id":"USR-LEGAL","role":"specialist","name":"María Abogada"}
        self.other_legal = {"id":"USR-OTHER","role":"specialist","name":"Otro especialista"}
        self.client = {"id":"USR-CLIENT","role":"client","name":"Cliente"}
        self.workspace = ApprovalDeskWorkspace(
            self.root / "approval-desk",
            db_factory=self.db,
            document_lookup=self.document_lookup,
            access_check=self.access_check,
            upload_validator=self.upload_validator,
        )

    def tearDown(self):
        self.temporary.cleanup()

    def db(self):
        con = sqlite3.connect(self.database)
        con.row_factory = sqlite3.Row
        return con

    @staticmethod
    def write_docx(path: Path, text: str):
        document = Document()
        document.add_heading("CONTRATO", 0)
        document.add_paragraph(text)
        document.save(path)

    def access_check(self, user, case_id):
        con = self.db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
        if not row: return False
        if user["role"] == "admin": return True
        if user["role"] == "client": return row["owner_id"] == user["id"]
        return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

    def document_lookup(self, user, document_id):
        con = self.db()
        row = con.execute(
            "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
            (document_id,),
        ).fetchone()
        con.close()
        if not row or not self.access_check(user, row["case_id"]): return None
        return dict(row)

    @staticmethod
    def upload_validator(filename, data):
        if not filename.lower().endswith(".docx") or not data.startswith(b"PK"):
            raise ValueError("Archivo inválido")
        return (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            sha256(data).hexdigest(),
            "clean:test",
        )

    def bootstrap(self):
        result = self.workspace.bootstrap(self.admin)
        self.assertEqual(result["created_count"], 1)
        return result["created"][0]

    def test_solo_administracion_prepara_bandeja(self):
        with self.assertRaises(PermissionDenied):
            self.workspace.bootstrap(self.legal)
        case_id = self.bootstrap()
        self.assertEqual(case_id, "DSK-DOC-001")
        detail = self.workspace.detail(self.admin, case_id)
        self.assertEqual(detail["case"]["source_generation_id"], "LZ-CASE-001")
        self.assertEqual(detail["case"]["revision_count"], 1)

    def test_cliente_no_ingresa_a_mesa_pero_puede_recibir_liberado(self):
        case_id = self.bootstrap()
        with self.assertRaises(PermissionDenied):
            self.workspace.list_for_user(self.client)
        with self.assertRaises(PermissionDenied):
            self.workspace.detail(self.client, case_id)
        with self.assertRaises(ReleaseBlocked):
            self.workspace.released_path(self.client, case_id)

    def test_aislamiento_impide_acceso_a_especialista_no_asignado(self):
        case_id = self.bootstrap()
        with self.assertRaises(PermissionDenied):
            self.workspace.detail(self.other_legal, case_id)
        self.assertEqual(self.workspace.list_for_user(self.other_legal)["metrics"]["total"], 0)

    def test_flujo_operativo_libera_hash_aprobado(self):
        case_id = self.bootstrap()
        detail = self.workspace.detail(self.legal, case_id)
        current = detail["revisions"][0]
        finding = self.workspace.add_finding(self.legal, case_id, {
            "revision_id": current["revision_id"],
            "severity": "major",
            "description": "Debe precisarse el alcance de la obligación principal.",
            "page": 1,
            "clause": "PRIMERA",
        })
        with self.assertRaises(ReleaseBlocked):
            self.workspace.approve(self.legal, case_id, {
                "revision_id": current["revision_id"],
                "approval_type": "legal",
                "decision": "approve",
                "comment": "Aprobado",
                "expected_sha256": current["sha256"],
            })
        self.workspace.resolve_finding(self.legal, case_id, finding["finding_id"], {
            "resolution": "La cláusula fue cotejada y el alcance quedó confirmado.",
            "state": "resolved",
        })
        legal = self.workspace.approve(self.legal, case_id, {
            "revision_id": current["revision_id"],
            "approval_type": "legal",
            "decision": "approve",
            "comment": "Contenido jurídico cotejado.",
            "expected_sha256": current["sha256"],
        })
        qa = self.workspace.approve(self.admin, case_id, {
            "revision_id": current["revision_id"],
            "approval_type": "qa",
            "decision": "approve",
            "comment": "Integridad, formato y trazabilidad verificados.",
            "expected_sha256": current["sha256"],
        })
        self.assertNotEqual(legal["actor"]["id"], qa["actor"]["id"])
        release = self.workspace.release(self.admin, case_id, {
            "revision_id": current["revision_id"],
            "expected_sha256": current["sha256"],
        })
        target, delivered = self.workspace.released_path(self.client, case_id)
        self.assertEqual(release["sha256"], delivered["sha256"])
        self.assertEqual(sha256(target.read_bytes()).hexdigest(), current["sha256"])
        summary = self.workspace.list_for_user(self.admin)
        self.assertEqual(summary["cases"][0]["status"], "released")
        self.assertTrue(summary["cases"][0]["audit_valid"])

    def test_nueva_revision_requiere_archivo_diferente_y_obsoleta_decisiones(self):
        case_id = self.bootstrap()
        with self.assertRaises(ImmutableRecordError):
            self.workspace.register_current_document(self.legal, case_id, "Sin cambios")
        self.write_docx(self.source, "Versión dos con alcance corregido.")
        revision = self.workspace.register_current_document(self.legal, case_id, "Corrección de alcance")
        self.assertEqual(revision["revision_id"], "REV-0002")
        detail = self.workspace.detail(self.legal, case_id)
        self.assertEqual(detail["case"]["current_revision_id"], "REV-0002")
        self.assertIsNone(detail["revisions"][-1]["approvals"]["legal"])
        comparison = self.workspace.compare(self.legal, case_id, "REV-0001", "REV-0002")
        self.assertTrue(comparison["changed"])
        self.assertGreaterEqual(comparison["summary"]["added_lines"], 1)

    def test_carga_docx_crea_revision_hija(self):
        case_id = self.bootstrap()
        uploaded = self.root / "ajuste.docx"
        self.write_docx(uploaded, "Revisión cargada por el especialista.")
        revision = self.workspace.upload_revision(
            self.legal,
            case_id,
            "ajuste.docx",
            uploaded.read_bytes(),
            "Ajuste solicitado en mesa.",
        )
        self.assertEqual(revision["revision_id"], "REV-0002")
        self.assertEqual(revision["parent_revision_id"], "REV-0001")
        self.assertEqual(revision["upload_sha256"], revision["sha256"])
        self.assertEqual(revision["filename"], "ajuste.docx")
        detail = self.workspace.detail(self.legal, case_id)
        persisted_note = detail["revisions"][-1]["note"]
        self.assertIn("Validación M32.5", persisted_note)
        self.assertIn("clean:test", persisted_note)
        self.assertIn("ajuste.docx", persisted_note)
        self.assertIn(revision["sha256"], persisted_note)

    def test_vista_estructural_declara_ausencia_de_paginacion(self):
        case_id = self.bootstrap()
        with mock.patch("legalai_platform.approval_desk_workspace.shutil.which", return_value=None):
            preview = self.workspace.preview(self.legal, case_id, "REV-0001")
        self.assertFalse(preview["rendered"])
        self.assertIsNone(preview["page_count"])
        self.assertTrue(preview["structural_preview"])
        self.assertIn("no acredita", preview["warning"])

    def test_segmentos_con_recorrido_de_ruta_son_rechazados(self):
        with self.assertRaises(ApprovalDeskError):
            self.workspace.detail(self.admin, "../secreto")


    def test_cadena_invalida_bloquea_liberacion(self):
        case_id = self.bootstrap()
        detail = self.workspace.detail(self.legal, case_id)
        current = detail["revisions"][0]
        self.workspace.approve(self.legal, case_id, {
            "revision_id": current["revision_id"],
            "approval_type": "legal",
            "decision": "approve",
            "comment": "Aprobación jurídica de prueba.",
            "expected_sha256": current["sha256"],
        })
        self.workspace.approve(self.admin, case_id, {
            "revision_id": current["revision_id"],
            "approval_type": "qa",
            "decision": "approve",
            "comment": "Aprobación QA de prueba.",
            "expected_sha256": current["sha256"],
        })
        events = self.workspace.root / case_id / "events.jsonl"
        lines = events.read_text(encoding="utf-8").splitlines()
        lines[0] = lines[0].replace('"case.created"', '"case.altered"')
        events.write_text("\n".join(lines) + "\n", encoding="utf-8")
        detail = self.workspace.detail(self.admin, case_id)
        self.assertFalse(detail["audit"]["valid"])
        self.assertEqual(detail["workflow_status"], "audit_invalid")
        with self.assertRaisesRegex(ReleaseBlocked, "cadena de auditoría"):
            self.workspace.release(self.admin, case_id, {
                "revision_id": current["revision_id"],
                "expected_sha256": current["sha256"],
            })

    def test_cadena_alterada_bloquea_descarga_ya_liberada(self):
        case_id = self.bootstrap()
        detail = self.workspace.detail(self.legal, case_id)
        current = detail["revisions"][0]
        for user, approval_type in ((self.legal, "legal"), (self.admin, "qa")):
            self.workspace.approve(user, case_id, {
                "revision_id": current["revision_id"],
                "approval_type": approval_type,
                "decision": "approve",
                "comment": f"Aprobación {approval_type} de prueba.",
                "expected_sha256": current["sha256"],
            })
        self.workspace.release(self.admin, case_id, {
            "revision_id": current["revision_id"],
            "expected_sha256": current["sha256"],
        })
        events = self.workspace.root / case_id / "events.jsonl"
        events.write_text(events.read_text(encoding="utf-8").replace('"document.released"', '"document.altered"'), encoding="utf-8")
        with self.assertRaisesRegex(ReleaseBlocked, "cadena de auditoría"):
            self.workspace.released_path(self.client, case_id)

class ApprovalDeskStaticInterfaceM325Tests(TestCase):
    def test_activos_y_handler_estan_conectados(self):
        root = Path(__file__).resolve().parents[1]
        index = (root / "app" / "index.html").read_text(encoding="utf-8")
        script = (root / "app" / "modules" / "approval_desk_m32_5.js").read_text(encoding="utf-8")
        styles = (root / "app" / "modules" / "approval_desk_m32_5.css").read_text(encoding="utf-8")
        run = (root / "run.py").read_text(encoding="utf-8")
        self.assertIn("approval_desk_m32_5.css", index)
        self.assertIn("approval_desk_m32_5.js", index)
        self.assertIn("/mesa-juridica", script)
        self.assertIn("Registrar hallazgo inmutable", script)
        self.assertIn("Liberar hash exacto", script)
        self.assertIn("Visor por revisión", script)
        self.assertIn("m325-page-sheet", styles)
        self.assertIn("http_handler_m32_5 import Handler", run)
