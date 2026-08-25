from __future__ import annotations

import hashlib
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from m33_2_contract_pagination_finalize import (
    EMPLOYMENT_CLAUSE_BEFORE_PT,
    EMPLOYMENT_PARAGRAPH_AFTER_PT,
    EMPLOYMENT_SIGNATURE_BEFORE_PT,
    finalize_contract_pagination,
)


class EmploymentContractPaginationTests(unittest.TestCase):
    def _sample(self, path: Path) -> None:
        document = Document()
        title = document.add_paragraph("CONTRATO DE TRABAJO A TÉRMINO INDEFINIDO")
        title.style = document.styles["Title"]

        clause = document.add_paragraph()
        label = clause.add_run("PRIMERA. OBJETO: ")
        label.bold = True
        label.font.size = Pt(11)
        body = clause.add_run("LA PERSONA TRABAJADORA prestará sus servicios con diligencia.")
        body.font.size = Pt(11)
        clause.paragraph_format.space_before = Pt(6)
        clause.paragraph_format.space_after = Pt(6)

        paragraph = document.add_paragraph("Las obligaciones conservan íntegro su contenido jurídico.")
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.size = Pt(11)

        signatures = document.add_paragraph("FIRMAS")
        signatures.paragraph_format.space_before = Pt(8)
        signatures.paragraph_format.space_after = Pt(6)
        for run in signatures.runs:
            run.font.size = Pt(11)
        document.save(path)

    def test_employment_contract_compacts_spacing_without_reducing_font_or_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "employment.docx"
            self._sample(path)
            before_text = "\n".join(p.text for p in Document(path).paragraphs)

            result = finalize_contract_pagination(path, product_code="CO-LA-002")
            self.assertTrue(result["applied"])
            self.assertEqual(11, result["font_size_preserved_pt"])

            rendered = Document(path)
            after_text = "\n".join(p.text for p in rendered.paragraphs)
            self.assertEqual(before_text, after_text)

            clause = next(p for p in rendered.paragraphs if p.text.startswith("PRIMERA."))
            body = next(p for p in rendered.paragraphs if p.text.startswith("Las obligaciones"))
            signatures = next(p for p in rendered.paragraphs if p.text == "FIRMAS")
            self.assertAlmostEqual(clause.paragraph_format.space_before.pt, EMPLOYMENT_CLAUSE_BEFORE_PT)
            self.assertAlmostEqual(clause.paragraph_format.space_after.pt, EMPLOYMENT_PARAGRAPH_AFTER_PT)
            self.assertAlmostEqual(body.paragraph_format.space_after.pt, EMPLOYMENT_PARAGRAPH_AFTER_PT)
            self.assertAlmostEqual(signatures.paragraph_format.space_before.pt, EMPLOYMENT_SIGNATURE_BEFORE_PT)
            sizes = [run.font.size.pt for p in rendered.paragraphs for run in p.runs if run.font.size]
            self.assertTrue(sizes)
            self.assertTrue(all(abs(size - 11) < 0.01 for size in sizes))

    def test_other_contracts_are_byte_identical_and_not_compacted(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "services.docx"
            self._sample(path)
            before = hashlib.sha256(path.read_bytes()).hexdigest()
            result = finalize_contract_pagination(path, product_code="CO-EM-003")
            after = hashlib.sha256(path.read_bytes()).hexdigest()
            self.assertFalse(result["applied"])
            self.assertEqual(before, after)


if __name__ == "__main__":
    unittest.main()
