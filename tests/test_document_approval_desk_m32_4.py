from __future__ import annotations

import json
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from legalai_platform.document_approval_desk import (
    DocumentApprovalDesk,
    ImmutableRecordError,
    PermissionDenied,
    ReleaseBlocked,
)


LEGAL = {"id": "legal-001", "role": "specialist", "name": "Especialista Jurídico Demo"}
QA = {"id": "qa-001", "role": "qa", "name": "QA Demo"}
ADMIN = {"id": "admin-001", "role": "admin", "name": "Administrador Demo"}
AUTHOR = {"id": "author-001", "role": "author", "name": "Autor Demo"}


def make_docx(path: Path, title: str, body: str) -> Path:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(body)
    document.save(path)
    return path


class DocumentApprovalDeskM324Tests(unittest.TestCase):
    def setUp(self):
        self.temporary = TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.desk = DocumentApprovalDesk(self.root / "desk")
        self.source = make_docx(self.root / "source.docx", "Contrato demo", "Cláusula inicial.")
        self.case = self.desk.create_case(
            case_id="CASE-DEMO-001",
            product_code="CO-EM-003",
            document_id="DOC-EM-CONTRACT-001",
            title="Contrato de prestación de servicios",
            actor=AUTHOR,
            source_generation_id="GEN-DEMO-001",
        )
        self.revision = self.desk.add_revision(
            case_id=self.case["case_id"],
            source_file=self.source,
            actor=AUTHOR,
            note="Primera revisión",
        )

    def tearDown(self):
        self.temporary.cleanup()

    def approve_both(self):
        self.desk.approve(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            approval_type="legal",
            decision="approve",
            actor=LEGAL,
            comment="Aprobación jurídica demo.",
            expected_sha256=self.revision["sha256"],
        )
        self.desk.approve(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            approval_type="qa",
            decision="approve",
            actor=QA,
            comment="Aprobación QA demo.",
            expected_sha256=self.revision["sha256"],
        )

    def test_flujo_completo_libera_exactamente_el_hash_aprobado(self):
        finding = self.desk.add_finding(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            actor=LEGAL,
            severity="minor",
            description="Ajustar puntuación de una cláusula.",
            page=1,
            clause="PRIMERA",
            block_id="CLAUSE-001",
        )
        self.desk.resolve_finding(
            case_id=self.case["case_id"],
            finding_id=finding["finding_id"],
            actor=LEGAL,
            resolution="Puntuación verificada.",
        )
        self.approve_both()
        release = self.desk.release(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            actor=ADMIN,
            expected_sha256=self.revision["sha256"],
        )
        self.assertEqual(release["sha256"], self.revision["sha256"])
        self.assertEqual(release["status"], "released_exact_hash")
        detail = self.desk.detail(self.case["case_id"])
        self.assertEqual(detail["case"]["status"], "released")
        self.assertEqual(detail["release"]["revision_id"], self.revision["revision_id"])
        self.assertTrue(detail["audit"]["valid"])
        self.assertGreaterEqual(detail["audit"]["events"], 7)

    def test_hash_equivocado_bloquea_aprobacion(self):
        with self.assertRaises(ReleaseBlocked):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="legal",
                decision="approve",
                actor=LEGAL,
                comment="No debe aprobar.",
                expected_sha256="0" * 64,
            )

    def test_manipulacion_del_archivo_invalida_aprobacion_y_liberacion(self):
        stored = (
            self.root
            / "desk"
            / self.case["case_id"]
            / "revisions"
            / self.revision["revision_id"]
            / "document.docx"
        )
        stored.write_bytes(stored.read_bytes() + b"tamper")
        with self.assertRaises(ReleaseBlocked):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="legal",
                decision="approve",
                actor=LEGAL,
                comment="No debe aprobar.",
                expected_sha256=self.revision["sha256"],
            )

    def test_hallazgo_bloqueante_impide_aprobacion(self):
        self.desk.add_finding(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            actor=LEGAL,
            severity="blocking",
            description="Nombre de parte inconsistente.",
            page=1,
            clause="COMPARECENCIA",
        )
        with self.assertRaises(ReleaseBlocked):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="legal",
                decision="approve",
                actor=LEGAL,
                comment="No debe aprobar.",
                expected_sha256=self.revision["sha256"],
            )

    def test_qa_no_puede_aprobar_antes_de_legal(self):
        with self.assertRaises(ReleaseBlocked):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="qa",
                decision="approve",
                actor=QA,
                comment="Orden incorrecto.",
                expected_sha256=self.revision["sha256"],
            )

    def test_misma_persona_no_puede_cubrir_ambas_aprobaciones(self):
        dual_actor_legal = {"id": "dual-001", "role": "specialist"}
        dual_actor_qa = {"id": "dual-001", "role": "qa"}
        self.desk.approve(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            approval_type="legal",
            decision="approve",
            actor=dual_actor_legal,
            comment="Primera decisión.",
            expected_sha256=self.revision["sha256"],
        )
        with self.assertRaises(ReleaseBlocked):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="qa",
                decision="approve",
                actor=dual_actor_qa,
                comment="No debe cubrir QA.",
                expected_sha256=self.revision["sha256"],
            )

    def test_decision_inmutable_no_se_sobrescribe(self):
        self.desk.approve(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            approval_type="legal",
            decision="reject",
            actor=LEGAL,
            comment="Requiere ajustes.",
            expected_sha256=self.revision["sha256"],
        )
        with self.assertRaises(ImmutableRecordError):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="legal",
                decision="approve",
                actor=LEGAL,
                comment="Intento de sobrescritura.",
                expected_sha256=self.revision["sha256"],
            )

    def test_nueva_revision_vuelve_obsoletas_las_aprobaciones_anteriores(self):
        self.approve_both()
        updated = make_docx(self.root / "updated.docx", "Contrato demo", "Cláusula inicial corregida.")
        revision_two = self.desk.add_revision(
            case_id=self.case["case_id"],
            source_file=updated,
            actor=AUTHOR,
            note="Segunda revisión",
            parent_revision_id=self.revision["revision_id"],
        )
        with self.assertRaises(ReleaseBlocked):
            self.desk.release(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                actor=ADMIN,
                expected_sha256=self.revision["sha256"],
            )
        detail = self.desk.detail(self.case["case_id"])
        self.assertEqual(detail["case"]["current_revision_id"], revision_two["revision_id"])
        self.assertEqual(detail["revisions"][-1]["approval_state"], {"legal": "pending", "qa": "pending"})

    def test_revision_sin_padre_vigente_es_rechazada(self):
        updated = make_docx(self.root / "updated.docx", "Contrato demo", "Cambio.")
        with self.assertRaises(ImmutableRecordError):
            self.desk.add_revision(
                case_id=self.case["case_id"],
                source_file=updated,
                actor=AUTHOR,
                note="Rama paralela no permitida",
                parent_revision_id="REV-9999",
            )

    def test_comparacion_identifica_lineas_agregadas_y_retiradas(self):
        updated = make_docx(self.root / "updated.docx", "Contrato demo", "Cláusula inicial corregida y ampliada.")
        revision_two = self.desk.add_revision(
            case_id=self.case["case_id"],
            source_file=updated,
            actor=AUTHOR,
            note="Segunda revisión",
            parent_revision_id=self.revision["revision_id"],
        )
        comparison = self.desk.compare(
            case_id=self.case["case_id"],
            from_revision_id=self.revision["revision_id"],
            to_revision_id=revision_two["revision_id"],
        )
        self.assertTrue(comparison["changed"])
        self.assertGreater(comparison["summary"]["added_lines"], 0)
        self.assertGreater(comparison["summary"]["removed_lines"], 0)
        self.assertTrue(any("corregida y ampliada" in line for line in comparison["diff_lines"]))

    def test_roles_no_autorizados_son_rechazados(self):
        with self.assertRaises(PermissionDenied):
            self.desk.add_finding(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                actor=AUTHOR,
                severity="minor",
                description="No autorizado.",
            )
        with self.assertRaises(PermissionDenied):
            self.desk.approve(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                approval_type="legal",
                decision="approve",
                actor=QA,
                comment="Rol equivocado.",
                expected_sha256=self.revision["sha256"],
            )

    def test_cadena_de_auditoria_detecta_manipulacion(self):
        audit_path = self.root / "desk" / self.case["case_id"] / "events.jsonl"
        lines = audit_path.read_text(encoding="utf-8").splitlines()
        event = json.loads(lines[0])
        event["payload"]["title"] = "Texto manipulado"
        lines[0] = json.dumps(event, ensure_ascii=False, separators=(",", ":"))
        audit_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        verification = self.desk.verify_audit_chain(self.case["case_id"])
        self.assertFalse(verification["valid"])
        self.assertTrue(any("Hash inválido" in error for error in verification["errors"]))

    def test_no_se_puede_liberar_con_hallazgo_menor_abierto(self):
        self.desk.add_finding(
            case_id=self.case["case_id"],
            revision_id=self.revision["revision_id"],
            actor=QA,
            severity="minor",
            description="Ajustar espaciado final.",
            page=1,
        )
        # La aprobación solo bloquea hallazgos mayores/bloqueantes, pero la liberación exige cierre total.
        self.approve_both()
        with self.assertRaises(ReleaseBlocked):
            self.desk.release(
                case_id=self.case["case_id"],
                revision_id=self.revision["revision_id"],
                actor=ADMIN,
                expected_sha256=self.revision["sha256"],
            )


if __name__ == "__main__":
    unittest.main()
