from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


UNRESOLVED_PATTERN = re.compile(r"\{\{[^{}]+\}\}|\b(?:NULL|undefined|N/A)\b", re.I)
GEN_RE = re.compile(r"COEM003-[A-F0-9]{12}")

ONES = ["cero", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
TEENS = {10:"diez",11:"once",12:"doce",13:"trece",14:"catorce",15:"quince",16:"dieciséis",17:"diecisiete",18:"dieciocho",19:"diecinueve",20:"veinte",21:"veintiuno",22:"veintidós",23:"veintitrés",24:"veinticuatro",25:"veinticinco",26:"veintiséis",27:"veintisiete",28:"veintiocho",29:"veintinueve"}
TENS = {30:"treinta",40:"cuarenta",50:"cincuenta",60:"sesenta",70:"setenta",80:"ochenta",90:"noventa"}
HUNDREDS = {100:"cien",200:"doscientos",300:"trescientos",400:"cuatrocientos",500:"quinientos",600:"seiscientos",700:"setecientos",800:"ochocientos",900:"novecientos"}
ORDINALS = [
    "PRIMERA", "SEGUNDA", "TERCERA", "CUARTA", "QUINTA", "SEXTA", "SÉPTIMA", "OCTAVA", "NOVENA", "DÉCIMA",
    "DÉCIMA PRIMERA", "DÉCIMA SEGUNDA", "DÉCIMA TERCERA", "DÉCIMA CUARTA", "DÉCIMA QUINTA", "DÉCIMA SEXTA",
    "DÉCIMA SÉPTIMA", "DÉCIMA OCTAVA", "DÉCIMA NOVENA", "VIGÉSIMA", "VIGÉSIMA PRIMERA", "VIGÉSIMA SEGUNDA",
    "VIGÉSIMA TERCERA", "VIGÉSIMA CUARTA", "VIGÉSIMA QUINTA", "VIGÉSIMA SEXTA", "VIGÉSIMA SÉPTIMA",
    "VIGÉSIMA OCTAVA", "VIGÉSIMA NOVENA", "TRIGÉSIMA", "TRIGÉSIMA PRIMERA", "TRIGÉSIMA SEGUNDA",
]


def number_to_words_es(value: int) -> str:
    value = int(value)
    if value < 0:
        return "menos " + number_to_words_es(-value)
    if value < 10:
        return ONES[value]
    if value < 30:
        return TEENS[value]
    if value < 100:
        tens, rem = divmod(value, 10)
        return TENS[tens * 10] + (" y " + ONES[rem] if rem else "")
    if value < 1000:
        hundreds, rem = divmod(value, 100)
        head = "ciento" if hundreds == 1 and rem else HUNDREDS[hundreds * 100]
        return head + (" " + number_to_words_es(rem) if rem else "")
    if value < 1_000_000:
        thousands, rem = divmod(value, 1000)
        head = "mil" if thousands == 1 else number_to_words_es(thousands) + " mil"
        return head + (" " + number_to_words_es(rem) if rem else "")
    if value < 1_000_000_000:
        millions, rem = divmod(value, 1_000_000)
        head = "un millón" if millions == 1 else number_to_words_es(millions) + " millones"
        return head + (" " + number_to_words_es(rem) if rem else "")
    billions, rem = divmod(value, 1_000_000_000)
    head = "mil millones" if billions == 1 else number_to_words_es(billions) + " mil millones"
    return head + (" " + number_to_words_es(rem) if rem else "")


class CoEm003DocumentFactoryV243:
    VERSION = "2.43"
    """Deep DOCX factory for the canonical private-services product CO-EM-003."""

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.output_dir = self.root / "data" / "generated" / "co-em-003-v243"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _v(data, path, default=""):
        cur = data
        for part in path.split("."):
            if not isinstance(cur, dict) or part not in cur:
                return default
            cur = cur[part]
        return default if cur is None else cur

    @staticmethod
    def _as_list(value):
        if value in (None, "", [], {}):
            return []
        if isinstance(value, list):
            return value
        if isinstance(value, tuple):
            return list(value)
        if isinstance(value, str):
            return [x.strip(" -•\t") for x in value.replace(";", "\n").splitlines() if x.strip(" -•\t")]
        return [value]

    @staticmethod
    def _as_dict(value):
        return value if isinstance(value, dict) else {}

    @staticmethod
    def _key_label(key):
        labels = {
            "id": "Código", "name": "Nombre", "title": "Título", "deliverable": "Entregable",
            "due_date": "Fecha de entrega", "milestone": "Hito", "acceptance_criteria": "Criterio de aceptación",
            "criteria": "Criterio", "format": "Formato", "location": "Lugar", "tools": "Herramientas",
            "amount": "Valor", "value": "Valor", "total": "Total", "currency": "Moneda",
            "payment_terms": "Condiciones de pago", "payment_schedule": "Calendario de pago",
            "invoice_requirements": "Requisitos de facturación", "billing": "Facturación", "taxes": "Impuestos y retenciones",
            "withholdings": "Retenciones", "cap": "Límite", "limit": "Límite", "exclusions": "Exclusiones",
            "insurance_required": "Seguro requerido", "insurance": "Seguro", "third_party_claims": "Reclamaciones de terceros",
            "client": "Contratante", "contractor": "Contratista", "categories": "Categorías",
            "new_results": "Resultados nuevos", "background": "Activos preexistentes", "open_source": "Código abierto",
            "notice": "Preaviso", "cure_period": "Período de subsanación", "serious_breach": "Incumplimiento grave",
            "transition": "Transición", "support": "Soporte", "authorized": "Autorización",
            "start_date": "Fecha de inicio", "end_date": "Fecha de terminación", "renewal": "Renovación",
            "milestones": "Hitos", "dependencies": "Dependencias", "client_delays": "Retrasos atribuibles al contratante",
            "third_party": "Componentes de terceros", "advance": "Anticipo", "approval": "Aprobación",
            "authorization": "Autorización", "evidence": "Soportes", "receipts": "Comprobantes",
            "access_controls": "Controles de acceso", "coordination": "Coordinación",
        }
        return labels.get(str(key), str(key).replace("_", " ").capitalize())

    @staticmethod
    def _field_value(key, value):
        key_text = str(key or "").lower()
        if key_text == "date" or key_text.endswith("_date"):
            return CoEm003DocumentFactoryV243._date_es(value)
        return CoEm003DocumentFactoryV243._plain(value)

    @staticmethod
    def _plain(value, default=""):
        if value in (None, ""):
            return default
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, bool):
            return "Sí" if value else "No"
        if isinstance(value, (int, float)):
            return str(value)
        if isinstance(value, list):
            parts = []
            for item in value:
                if isinstance(item, dict):
                    parts.append("; ".join(f"{CoEm003DocumentFactoryV243._key_label(k)}: {CoEm003DocumentFactoryV243._field_value(k, v)}" for k, v in item.items() if v not in (None, "", [], {})))
                else:
                    parts.append(str(item))
            return "; ".join(x for x in parts if x)
        if isinstance(value, dict):
            return "; ".join(f"{CoEm003DocumentFactoryV243._key_label(k)}: {CoEm003DocumentFactoryV243._field_value(k, v)}" for k, v in value.items() if v not in (None, "", [], {}))
        return str(value)

    @staticmethod
    def _ensure_period(value):
        text = str(value or "").strip()
        if not text:
            return text
        return text if text[-1] in ".;:!?" else text + "."

    @staticmethod
    def _join_sentences(*parts):
        return " ".join(
            CoEm003DocumentFactoryV243._ensure_period(part)
            for part in parts
            if str(part or "").strip()
        ).strip()

    @staticmethod
    def _date_es(value):
        if not value:
            return ""
        try:
            parsed = date.fromisoformat(str(value)[:10])
            months = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            return f"{parsed.day} de {months[parsed.month - 1]} de {parsed.year}"
        except (TypeError, ValueError):
            return str(value)

    @staticmethod
    def _money(value, currency="COP"):
        try:
            amount = int(float(value))
            formatted = "$" + f"{amount:,}".replace(",", ".")
            words = number_to_words_es(amount).upper()
            suffix = "PESOS M/CTE" if currency in ("COP", "", None) else currency
            return f"{words} {suffix} ({formatted} {currency or 'COP'})"
        except (TypeError, ValueError):
            return str(value or "")

    @staticmethod
    def _set_doc_styles(doc: Document):
        sec = doc.sections[0]
        sec.top_margin = Cm(2.3)
        sec.bottom_margin = Cm(2.1)
        sec.left_margin = Cm(2.6)
        sec.right_margin = Cm(2.6)
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.2)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for style_name in ("Title", "Heading 1", "Heading 2"):
            style = doc.styles[style_name]
            style.font.name = "Aptos Display"
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Heading 1"].font.size = Pt(12)
        doc.styles["Heading 2"].font.size = Pt(11)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("LegalAIZ.it | CO-EM-003 | Documento sujeto a revisión jurídica y QA | Página ")
        run.font.size = Pt(8)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer._p.append(fld)

    @staticmethod
    def _compact_styles(doc):
        sec = doc.sections[0]
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        normal = doc.styles["Normal"]
        normal.font.size = Pt(9.5)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.0
        doc.styles["Title"].font.size = Pt(14)
        doc.styles["Heading 1"].font.size = Pt(11)
        doc.styles["Heading 2"].font.size = Pt(10.5)

    @staticmethod
    def _title(doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(12)
        return p

    @staticmethod
    def _section(doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11.5)
        return p

    @staticmethod
    def _paragraph(doc, text, bold_lead=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
        p.add_run(str(text).strip())
        return p

    @staticmethod
    def _clause(doc, number, heading, body):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.keep_together = False
        ordinal = ORDINALS[number - 1] if 0 < number <= len(ORDINALS) else str(number)
        r = p.add_run(f"CLÁUSULA {ordinal}. {heading.upper()}: ")
        r.bold = True
        p.add_run(body.strip())
        return p

    @staticmethod
    def _subparagraph(doc, label, body):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(label.upper() + ": ")
        r.bold = True
        p.add_run(body.strip())
        return p

    @staticmethod
    def _bullets(doc, items):
        for item in items:
            if isinstance(item, dict):
                text = "; ".join(f"{CoEm003DocumentFactoryV243._key_label(k)}: {CoEm003DocumentFactoryV243._field_value(k, v)}" for k, v in item.items() if v not in (None, "", [], {}))
            else:
                text = str(item)
            if not text.strip():
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(text)

    @staticmethod
    def _table(doc, headers, rows, widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        trPr = hdr._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        trPr.append(tblHeader)
        for cell, label in zip(hdr.cells, headers):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(label))
            r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(CoEm003DocumentFactoryV243._plain(value))
        return table

    def _party(self, answers, prefix, role_label):
        ptype = self._v(answers, f"{prefix}.type")
        ident = self._v(answers, f"{prefix}.identification")
        ident_d = self._as_dict(ident)
        name = ident_d.get("legal_name") or ident_d.get("full_name") or ident_d.get("name")
        identification = ident_d.get("nit") or ident_d.get("id_number") or ident_d.get("identification_number") or ident_d.get("identification")
        domicile = ident_d.get("domicile") or ident_d.get("address") or ident_d.get("city") or ""
        email = ident_d.get("email") or ""
        if not name and isinstance(ident, str):
            name = ident
        signatory = self._v(answers, f"{prefix}.signatory")
        sign_d = self._as_dict(signatory)
        sign_name = sign_d.get("full_name") or sign_d.get("name") or ""
        sign_id = sign_d.get("identification") or sign_d.get("id_number") or sign_d.get("identification_number") or ""
        capacity = sign_d.get("capacity") or sign_d.get("position") or ""
        authority = sign_d.get("authority_source") or sign_d.get("authority") or ""
        if isinstance(signatory, str) and not sign_name:
            sign_name = signatory
        party_text = f"{name or role_label}, identificado(a) con {identification or 'identificación consignada en el expediente'}"
        if domicile:
            party_text += f", con domicilio en {domicile}"
        if ptype == "legal_person":
            party_text += f", representado(a) para este acto por {sign_name or 'el firmante registrado'}, identificado(a) con {sign_id or 'identificación registrada'}, quien actúa como {capacity or 'representante autorizado'}"
            if authority:
                party_text += f" con fundamento en {authority}"
        return {"name": name or role_label, "id": identification or "", "domicile": domicile, "email": email, "signatory": sign_name or name or role_label, "capacity": capacity or role_label, "text": party_text}

    def _term_text(self, answers):
        term = self._v(answers, "term")
        if isinstance(term, dict):
            start = self._date_es(term.get("start_date")) or self._plain(term.get("start"))
            end = self._date_es(term.get("end_date")) or self._plain(term.get("end"))
            duration = self._plain(term.get("duration"))
            renewal = self._plain(term.get("renewal"))
            parts = []
            if start: parts.append(f"inicia el {start}")
            if end: parts.append(f"termina el {end}")
            elif duration: parts.append(f"tendrá una duración de {duration}")
            if renewal: parts.append(f"la renovación se regirá por {renewal}")
            return "; ".join(parts) + "."
        return self._plain(term, "El plazo se encuentra definido en el expediente contractual.")

    def _fees_text(self, answers):
        terms = self._v(answers, "fees.financial_terms")
        model = self._v(answers, "fees.model")
        labels = {"fixed":"precio fijo", "monthly":"honorarios mensuales", "hourly":"tarifa por hora", "milestone":"pago por hitos", "success":"remuneración por resultado o éxito", "mixed":"esquema mixto"}
        if isinstance(terms, dict):
            amount = terms.get("amount") or terms.get("value") or terms.get("total")
            currency = terms.get("currency") or "COP"
            amount_text = self._money(amount, currency) if amount not in (None, "") else self._plain(terms.get("amount_text"))
            payment = self._plain(terms.get("payment_terms") or terms.get("payment_schedule") or terms.get("due_date"))
            invoice = self._plain(terms.get("invoice_requirements") or terms.get("billing"))
            taxes = self._plain(terms.get("taxes") or terms.get("withholdings"))
            return self._join_sentences(
                f"La contraprestación se pacta bajo un modelo de {labels.get(model, model)} por {amount_text or 'el valor definido en el anexo económico'}",
                payment or "El pago se realizará contra factura o cuenta de cobro válida y aceptación del hito o período correspondiente",
                invoice,
                taxes or "Se practicarán las retenciones legalmente procedentes",
            )
        return f"La contraprestación se pacta bajo un modelo de {labels.get(model, model)}. {self._plain(terms)}".strip()

    def _contract(self, answers, evaluation, target):
        client = self._party(answers, "client", "EL CONTRATANTE")
        contractor = self._party(answers, "contractor", "EL CONTRATISTA")
        service_object = self._plain(self._v(answers, "service.object"))
        expected = self._plain(self._v(answers, "service.expected_result"))
        included = self._as_list(self._v(answers, "scope.included"))
        excluded = self._as_list(self._v(answers, "scope.excluded"))
        deliverables = self._as_list(self._v(answers, "scope.deliverables"))
        schedule = self._plain(self._v(answers, "schedule"), "el cronograma aprobado en el Anexo de alcance")
        arrangement = self._plain(self._v(answers, "execution.arrangement"), "los lugares y medios definidos por EL CONTRATISTA, coordinados con EL CONTRATANTE")
        termination = self._plain(self._v(answers, "termination"), "las causales y períodos de subsanación previstos en esta cláusula")
        closure = self._plain(self._v(answers, "closure"), "entrega ordenada, devolución de activos, cierre de accesos y acta final")
        liability = self._plain(self._v(answers, "liability"), "responsabilidad por daños directos, previsibles y debidamente probados, con las exclusiones y límites acordados")
        risk_allocation = self._plain(self._v(answers, "risk_allocation"), "cada parte asumirá los riesgos que se encuentren bajo su control razonable")
        expenses = self._ensure_period(self._plain(self._v(answers, "fees.expenses"), "Los gastos reembolsables requerirán autorización previa y soportes idóneos."))
        dispute = self._v(answers, "disputes.mechanism")

        doc = Document()
        self._set_doc_styles(doc)
        self._title(doc, "Contrato de prestación de servicios independientes")
        appearance = (
            f"Entre los suscritos, de una parte, {client['text']}, quien para efectos del presente contrato se denominará EL CONTRATANTE; "
            f"y de otra parte, {contractor['text']}, quien se denominará EL CONTRATISTA; conjuntamente LAS PARTES, se celebra el presente contrato privado de prestación de servicios independientes, regido por sus estipulaciones y por las normas civiles, comerciales y demás disposiciones colombianas aplicables."
        )
        self._paragraph(doc, appearance)

        self._section(doc, "Consideraciones")
        considerations = [
            ("PRIMERA", f"Que EL CONTRATANTE requiere un servicio independiente relacionado con {service_object}, orientado a obtener {expected}."),
            ("SEGUNDA", "Que EL CONTRATISTA declaró contar con la capacidad jurídica, técnica, administrativa, financiera y, cuando sea exigible, profesional o regulatoria necesaria para asumir el encargo."),
            ("TERCERA", "Que LAS PARTES desean delimitar el alcance, los entregables, los criterios de aceptación, la contraprestación, la distribución de riesgos y los mecanismos de cambio y cierre, evitando que instrucciones de coordinación se conviertan en subordinación laboral."),
            ("CUARTA", "Que los anexos, actas y órdenes de cambio identificados en el contrato forman parte del expediente únicamente cuando hayan sido generados, aceptados y versionados de manera trazable."),
            ("QUINTA", "Que LAS PARTES obrarán de buena fe, cooperarán para mitigar riesgos y documentarán oportunamente las decisiones que afecten precio, plazo, alcance, seguridad, datos personales o propiedad intelectual."),
        ]
        for label, body in considerations:
            self._paragraph(doc, body, f"{label}: ")

        self._section(doc, "Cláusulas")
        clauses = []
        clauses.append(("Objeto", f"EL CONTRATISTA se obliga, con autonomía técnica y administrativa, a ejecutar {service_object}, con el propósito de entregar {expected}. El servicio comprende las actividades expresamente incluidas y no se extenderá por analogía a labores, resultados, capacidades o responsabilidades ajenas al alcance aprobado. El objeto debe interpretarse junto con el Anexo de alcance, entregables y cronograma."))
        clauses.append(("Alcance, inclusiones y exclusiones", "El alcance incluye: " + ("; ".join(self._plain(x) for x in included) if included else "las actividades consignadas en el Anexo de alcance") + ". Quedan excluidos: " + ("; ".join(self._plain(x) for x in excluded) if excluded else "los trabajos que no hayan sido expresamente incluidos") + ". Las exclusiones no podrán suplirse mediante instrucciones informales; cualquier ampliación material requerirá control de cambios."))
        clauses.append(("Entregables, hitos y documentación", "EL CONTRATISTA entregará productos verificables, completos y utilizables, de acuerdo con los hitos y formatos aprobados. Cada entregable deberá identificar versión, fecha, responsable, dependencias, supuestos, fuentes y limitaciones relevantes. La entrega no implica aceptación automática ni transfiere riesgos que correspondan a EL CONTRATANTE."))
        clauses.append(("Aceptación, observaciones y subsanación", "EL CONTRATANTE contará con el período definido en el Anexo de alcance para revisar cada entregable frente a criterios objetivos. Podrá aceptarlo, aceptarlo con reservas o formular observaciones específicas. EL CONTRATISTA atenderá observaciones atribuibles al alcance contratado dentro del plazo acordado. El silencio solo producirá aceptación cuando LAS PARTES lo hayan establecido expresamente; no podrán rechazarse entregables por requisitos nuevos o subjetivos no comunicados."))
        clauses.append(("Control de cambios", "Toda variación material de alcance, metodología comprometida, entregables, cronograma, recursos, precio, riesgos, tratamiento de datos o titularidad de resultados deberá documentarse mediante ACT-EM-CHANGE-001 u otro instrumento equivalente. La solicitud deberá describir el cambio, su justificación y el impacto. Ninguna parte estará obligada a ejecutar un cambio material antes de su aprobación, salvo medidas urgentes para proteger personas, información o bienes."))
        clauses.append(("Autonomía técnica y administrativa", f"EL CONTRATISTA organizará los métodos, secuencia, personal, medios y lugar de ejecución, sujeto únicamente al resultado, los estándares pactados, la coordinación necesaria y las obligaciones legales. La ejecución se desarrollará en {arrangement}. Las reuniones, reportes, validaciones, accesos o controles de seguridad no constituyen por sí solos subordinación, siempre que no se utilicen para imponer de manera continuada el modo, tiempo o cantidad de trabajo como en una relación laboral."))
        clauses.append(("Inexistencia de vínculo laboral y primacía de la realidad", "El contrato no crea relación laboral, sociedad, agencia, mandato ni representación, y cada parte conserva su autonomía. Sin embargo, esta declaración no podrá utilizarse para desconocer una relación laboral si en la ejecución real concurren prestación personal, remuneración y subordinación. LAS PARTES se obligan a corregir prácticas incompatibles con la independencia y a escalar jurídicamente cualquier cambio operativo que incremente el riesgo laboral."))
        clauses.append(("Personal, sustitución y subcontratación", "EL CONTRATISTA podrá apoyarse en personal propio o subcontratistas idóneos cuando ello sea compatible con la naturaleza del servicio y haya sido informado o autorizado en los casos definidos. Continuará siendo responsable por la dirección, pago, afiliaciones, confidencialidad, seguridad y calidad de dicho personal. La autorización de terceros no crea vínculo con EL CONTRATANTE. Cuando el encargo sea estrictamente personal por habilitación o confianza especial, la sustitución requerirá aprobación previa."))
        clauses.append(("Lugar, medios y accesos", "EL CONTRATISTA utilizará sus propios medios salvo aquellos que EL CONTRATANTE suministre por necesidad técnica, seguridad o interoperabilidad. El uso de instalaciones, correo, plataformas, credenciales o equipos del contratante se limitará al servicio y no integrará al contratista como empleado. Los accesos se otorgarán bajo mínimo privilegio, serán auditables y deberán cerrarse al terminar o cuando dejen de ser necesarios."))
        clauses.append(("Plazo, cronograma y dependencias", f"El contrato {self._term_text(answers)} Su ejecución seguirá {schedule}. Los plazos se suspenderán o ajustarán cuando el retraso provenga de información, aprobaciones, accesos, materiales o decisiones a cargo de EL CONTRATANTE, o de eventos no imputables a EL CONTRATISTA, siempre que sean informados oportunamente y se adopten medidas de mitigación."))
        clauses.append(("Honorarios, facturación y pago", self._fees_text(answers) + " La aceptación de una factura no implica renuncia a objeciones documentadas ni autoriza compensaciones o descuentos no pactados. Los pagos causados y pendientes conservarán su exigibilidad después de la terminación."))
        clauses.append(("Impuestos, aportes y seguridad social", "Cada parte cumplirá las obligaciones tributarias y parafiscales que legalmente le correspondan. EL CONTRATISTA asumirá sus aportes como independiente cuando sean aplicables y suministrará los soportes que razonablemente pueda verificar EL CONTRATANTE. La verificación no convierte al contratante en empleador ni lo libera de las obligaciones que la normativa de riesgos laborales le atribuya por clase de riesgo o forma de contratación."))
        clauses.append(("Anticipos, gastos y reembolsos", expenses + " Los anticipos conservarán destinación específica, estarán sujetos a legalización y no se confundirán con honorarios causados. Los costos ordinarios propios de la organización empresarial del contratista se entienden incluidos salvo pacto expreso. Ninguna parte asumirá gastos extraordinarios de la otra sin autorización verificable."))
        clauses.append(("Obligaciones del contratante", "EL CONTRATANTE deberá: suministrar información completa y oportuna; designar interlocutores con capacidad suficiente; facilitar decisiones, accesos y dependencias a su cargo; revisar entregables dentro del plazo; pagar lo debido; informar riesgos de sus instalaciones y sistemas; proteger al personal del contratista frente a peligros bajo su control; abstenerse de impartir órdenes laborales; y documentar los cambios y observaciones de manera trazable."))
        clauses.append(("Obligaciones del contratista", "EL CONTRATISTA deberá ejecutar el servicio con diligencia profesional; mantener capacidad y habilitaciones; cumplir el alcance y cronograma; informar riesgos, retrasos y conflictos de interés; conservar registros; proteger información, datos y activos; utilizar personal idóneo; cumplir las medidas de seguridad aplicables; respetar derechos de terceros; atender observaciones procedentes; y entregar resultados, archivos y soportes en formatos acordados."))
        clauses.append(("Cumplimiento regulatorio, ética y conflictos de interés", "Cada parte cumplirá las normas aplicables a su actividad, incluyendo reglas profesionales, ambientales, técnicas, de libre competencia, anticorrupción y prevención de operaciones ilícitas. EL CONTRATISTA declarará conflictos de interés reales o potenciales que puedan afectar objetividad, confidencialidad o independencia. No se pacta exclusividad general, salvo delimitación expresa, necesaria y proporcional en un anexo."))
        clauses.append(("Seguridad y salud en el trabajo", "Cuando la ejecución se realice en instalaciones o procesos del contratante, LAS PARTES coordinarán afiliación, inducción, información de peligros, permisos, elementos de protección, investigación de incidentes y demás controles exigibles. EL CONTRATISTA dirigirá a su personal, pero cumplirá las reglas de ingreso y seguridad que protejan a las personas sin constituir subordinación laboral. Una actividad crítica no deberá iniciarse si faltan controles indispensables."))
        clauses.append(("Confidencialidad y secretos empresariales", "Cuando se active ANX-EM-CONF-001, la información protegida solo podrá utilizarse para el contrato y compartirse con personas que necesiten conocerla y estén sometidas a reserva. Se aplicarán excepciones para información pública, conocida legítimamente, obtenida de tercero autorizado, desarrollada independientemente o revelada por obligación legal, defensa de derechos o denuncia legítima. Los secretos empresariales se protegerán mientras conserven tal carácter."))
        clauses.append(("Tratamiento de datos personales", "Cuando una parte trate datos por cuenta de la otra, ANX-EM-DATA-001 definirá roles, finalidades, instrucciones, titulares, categorías, seguridad, subencargados, transferencias, incidentes, ejercicio de derechos, conservación y eliminación. Ninguna instrucción autoriza tratamientos contrarios a la ley. Cada parte responderá por decisiones que tome como responsable y por incumplimientos bajo su control."))
        clauses.append(("Propiedad intelectual y resultados", "ANX-EM-IP-001 distinguirá materiales preexistentes, resultados creados, componentes de terceros, código abierto, datos, invenciones, secretos y conocimiento general. La transferencia o licencia deberá identificar modalidades de explotación, territorio, duración, contraprestación y limitaciones. Los derechos morales permanecerán en sus titulares. No se transfieren de forma indeterminada todas las creaciones futuras ni activos ajenos al objeto."))
        clauses.append(("Uso de inteligencia artificial", "Si se autoriza inteligencia artificial, ANX-EM-AI-001 clasificará usos permitidos, controlados y prohibidos; sistemas autorizados; información que no podrá ingresarse; deberes de revisión humana; trazabilidad; verificación de sesgos, exactitud y licencias; identificación de contenidos sintéticos; y responsabilidad sobre decisiones finales. El uso de IA no reduce el estándar de diligencia ni traslada automáticamente titularidad sobre entradas o resultados."))
        clauses.append(("Responsabilidad", f"La responsabilidad se regirá por {liability}. Como regla, solo serán indemnizables los daños directos, ciertos, previsibles y probados que guarden relación causal con un incumplimiento imputable. Cualquier límite deberá ser razonable frente al valor y riesgo del contrato y no cobijará dolo, culpa grave cuando no sea limitable, violación deliberada de confidencialidad, infracción de propiedad intelectual, tratamiento ilícito de datos ni obligaciones expresamente excluidas."))
        clauses.append(("Indemnidad y reclamaciones de terceros", f"La distribución de riesgos será la siguiente: {risk_allocation}. La parte que reciba una reclamación cubierta deberá notificarla oportunamente, permitir la participación de la parte responsable, preservar evidencia y mitigar daños. No podrá reconocer responsabilidad, transigir o comprometer a la otra parte sin autorización cuando ello afecte su defensa. La indemnidad no operará por actos propios de la parte protegida."))
        if "EM-INS-001" in evaluation.get("blocks", []):
            clauses.append(("Seguros", "EL CONTRATISTA mantendrá las pólizas o coberturas expresamente definidas en el expediente, con amparos, límites, deducibles y vigencias proporcionales al riesgo. La existencia de seguro no amplía la responsabilidad pactada ni sustituye controles preventivos. Los certificados podrán ser verificados antes de iniciar actividades críticas."))
        clauses.append(("Fuerza mayor y deber de mitigación", "La parte afectada por un evento imprevisible e irresistible informará su ocurrencia, alcance y duración estimada, protegerá los resultados y adoptará medidas razonables de mitigación. Las obligaciones imposibles se suspenderán durante el evento; las obligaciones dinerarias causadas no se extinguen por esta sola circunstancia. Si el impedimento hace inútil el contrato o supera el período acordado, cualquiera podrá solicitar terminación y liquidación de lo ejecutado."))
        clauses.append(("Suspensión", "LAS PARTES podrán suspender total o parcialmente el servicio por acuerdo, fuerza mayor, falta de dependencias esenciales, riesgo grave, incumplimiento subsanable o decisión regulatoria. El acta de suspensión identificará causa, fecha, custodia de información y activos, costos inevitables, actividades de preservación y condiciones de reinicio. La suspensión no autoriza mantener disponibilidad personal indefinida sin contraprestación."))
        clauses.append(("Terminación anticipada", f"La terminación se regirá por {termination}. Salvo incumplimiento grave, violación no subsanable, riesgo urgente o prohibición legal, la parte incumplida dispondrá de un período razonable para corregir. La terminación sin incumplimiento deberá respetar el preaviso pactado y reconocer servicios ejecutados, entregables aceptables, gastos autorizados no recuperables y componentes variables ya causados."))
        clauses.append(("Efectos de la terminación, transición y cierre", f"Al finalizar se ejecutará {closure}. EL CONTRATISTA entregará avances utilizables, archivos fuente acordados, documentación, activos y accesos; EL CONTRATANTE pagará lo causado y decidirá sobre entregables parciales. Las obligaciones de confidencialidad, datos, propiedad intelectual, responsabilidad, pagos, auditoría y controversias sobrevivirán en la medida necesaria. ACT-EM-CLOSE-001 no constituirá paz y salvo general salvo estipulación expresa y específica."))
        clauses.append(("Registros, auditoría y trazabilidad", "Cada parte conservará los registros necesarios para acreditar entregas, aprobaciones, cambios, facturación, seguridad, datos, licencias y cumplimiento. Las auditorías deberán limitarse al contrato, proteger información de terceros, realizarse con aviso razonable y no interferir indebidamente con la operación. Los hallazgos se comunicarán y permitirán contradicción y plan de mejora."))
        clauses.append(("Notificaciones", f"Las comunicaciones contractuales se enviarán a {client['email'] or 'el correo registrado por EL CONTRATANTE'} y {contractor['email'] or 'el correo registrado por EL CONTRATISTA'}, o a los canales posteriormente informados. Las notificaciones de incumplimiento, terminación, reclamación o cambio deberán permitir acreditar envío, contenido y recepción. El correo ordinario no sustituirá formalidades especiales legalmente exigibles."))
        if dispute == "arbitration":
            clauses.append(("Solución de controversias", "LAS PARTES intentarán una negociación directa. Si no hay acuerdo, la controversia será sometida a arbitraje en derecho conforme a la cláusula compromisoria aprobada, que deberá identificar centro, sede, número de árbitros, reglas y distribución inicial de costos. La cláusula arbitral es autónoma y no impide solicitar medidas cautelares cuando proceda."))
        elif dispute == "amicable_composition":
            clauses.append(("Solución de controversias", "LAS PARTES agotarán negociación directa y someterán las diferencias técnicas o contractuales definidas a amigable composición, bajo el centro, procedimiento y alcance indicados en el expediente. Las materias no cubiertas podrán acudir a la jurisdicción competente. La elección deberá considerar costos y cuantía."))
        elif dispute == "negotiation_conciliation_courts":
            clauses.append(("Solución de controversias", "LAS PARTES agotarán negociación directa y, cuando resulte procedente, conciliación extrajudicial en derecho antes de acudir a los jueces competentes de Colombia. La solicitud de medidas cautelares o los casos exceptuados por la ley no se entenderán impedidos."))
        else:
            clauses.append(("Solución de controversias", "LAS PARTES intentarán resolver directamente las diferencias de buena fe. Si no alcanzan acuerdo dentro del plazo definido, podrán acudir a los jueces competentes de Colombia, sin perjuicio de los requisitos de procedibilidad que correspondan."))
        clauses.append(("Integridad, prelación, modificaciones y firma", "El contrato, sus otrosíes, anexos, órdenes de cambio y actas vigentes contienen el acuerdo. En caso de contradicción prevalecerán las normas imperativas, el otrosí específico, el contrato, el anexo de alcance y luego los demás anexos. Las modificaciones materiales deberán constar por escrito y generar nueva revisión. El contrato podrá firmarse física o electrónicamente mediante un mecanismo que permita identificar al firmante, demostrar su intención y preservar la integridad y disponibilidad del documento."))

        for idx, (heading, body) in enumerate(clauses, 1):
            self._clause(doc, idx, heading, body)
            if heading == "Entregables, hitos y documentación" and deliverables:
                self._bullets(doc, deliverables)

        self._section(doc, "Documentos relacionados")
        related = evaluation.get("documents", [])
        self._bullets(doc, [f"{doc_id}: documento generado y sujeto a la misma revisión del expediente." for doc_id in related if doc_id != "DOC-EM-CONTRACT-001"])
        self._paragraph(doc, "Los anexos de confidencialidad, datos, propiedad intelectual e inteligencia artificial prevalecerán sobre disposiciones generales del contrato únicamente en su materia y dentro de la revisión vigente.")

        # Mantener la sección de firmas completa en una página independiente evita
        # que la tabla se fracture entre páginas en contratos extensos.
        doc.add_page_break()
        self._section(doc, "Firmas")
        self._paragraph(doc, f"Se firma en la fecha registrada en el expediente contractual. LAS PARTES declaran haber leído el documento, comprendido su alcance y recibido copia íntegra de sus anexos aplicables.")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        signature_header_props = table.rows[0]._tr.get_or_add_trPr()
        signature_header = OxmlElement("w:tblHeader")
        signature_header.set(qn("w:val"), "true")
        signature_header_props.append(signature_header)
        values = [
            ("EL CONTRATANTE", "EL CONTRATISTA"),
            (client["signatory"], contractor["signatory"]),
            (client["capacity"], contractor["capacity"]),
            ("Firma: ______________________________", "Firma: ______________________________"),
        ]
        for r_idx, row in enumerate(table.rows):
            for cell, value in zip(row.cells, values[r_idx]):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(value)
                run.bold = r_idx == 0
        doc.core_properties.title = "Contrato de prestación de servicios independientes"
        doc.core_properties.subject = f"CO-EM-003 v{self.VERSION}"
        doc.core_properties.comments = "Generado por LegalAIZ.it. Requiere aprobación jurídica y QA antes de liberación."
        doc.save(target)

    def _annex_header(self, doc, answers, doc_id, title):
        client = self._party(answers, "client", "EL CONTRATANTE")
        contractor = self._party(answers, "contractor", "EL CONTRATISTA")
        self._title(doc, title)
        relation_text = f"Documento {doc_id}, vinculado al contrato de prestación de servicios independientes celebrado entre {client['name']} y {contractor['name']}"
        self._paragraph(doc, self._ensure_period(relation_text) + " Su contenido corresponde a la revisión vigente del expediente y no podrá aplicarse separadamente a otra relación contractual.")
        return client, contractor

    def _scope_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc); self._compact_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-SCOPE-001", "Anexo de alcance, entregables y cronograma")
        self._section(doc, "Objetivo y resultado")
        self._paragraph(doc, self._plain(self._v(answers, "service.object")), "OBJETO: ")
        self._paragraph(doc, self._plain(self._v(answers, "service.expected_result")), "RESULTADO ESPERADO: ")
        self._section(doc, "Actividades incluidas")
        self._bullets(doc, self._as_list(self._v(answers, "scope.included")))
        self._section(doc, "Actividades excluidas")
        excluded = self._as_list(self._v(answers, "scope.excluded"))
        self._bullets(doc, excluded or ["Toda actividad, resultado o servicio que no se encuentre expresamente incluido."])
        self._section(doc, "Entregables, hitos y aceptación")
        ds = self._as_list(self._v(answers, "scope.deliverables"))
        rows = []
        for i, item in enumerate(ds, 1):
            if isinstance(item, dict):
                due_or_milestone = item.get("due_date") or item.get("milestone") or "Según cronograma"
                if item.get("due_date"):
                    due_or_milestone = self._date_es(item.get("due_date"))
                rows.append([item.get("id") or i, item.get("name") or item.get("deliverable") or item.get("title") or f"Entregable {i}", due_or_milestone, item.get("acceptance_criteria") or item.get("criteria") or "Conformidad con alcance", item.get("format") or "Digital"])
            else:
                rows.append([i, item, "Según cronograma", "Conformidad con alcance", "Digital"])
        self._table(doc, ["No.", "Entregable", "Fecha o hito", "Criterio de aceptación", "Formato"], rows)
        self._section(doc, "Cronograma, dependencias y supuestos")
        self._paragraph(doc, self._plain(self._v(answers, "schedule")))
        arrangement = self._plain(self._v(answers, "execution.arrangement"))
        if arrangement:
            self._paragraph(doc, arrangement, "CONDICIONES DE EJECUCIÓN: ")
        self._subparagraph(doc, "Dependencias del contratante", "La entrega oportuna de información, accesos, decisiones, muestras, aprobaciones y recursos identificados constituye una dependencia del cronograma. Los retrasos imputables a estas dependencias ajustarán razonablemente las fechas.")
        self._subparagraph(doc, "Supuestos", "Los resultados se evaluarán bajo los supuestos técnicos, comerciales o informativos expresamente registrados. Un cambio material de supuesto se gestionará mediante control de cambios.")
        self._subparagraph(doc, "Aceptación", "EL CONTRATANTE deberá formular observaciones específicas dentro del período acordado. EL CONTRATISTA atenderá las que correspondan al alcance. Los requisitos nuevos requerirán orden de cambio.")
        doc.save(target)

    def _fees_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-FEES-001", "Anexo económico y de honorarios")
        model = self._v(answers, "fees.model")
        self._clause(doc, 1, "Modelo económico", self._fees_text(answers))
        self._clause(doc, 2, "Causación", "Cada pago se causará por el período, hito, hora, resultado o combinación expresamente identificada. La causación no dependerá de una aceptación discrecional ni podrá modificarse retroactivamente.")
        terms = self._as_dict(self._v(answers, "fees.financial_terms"))
        rows = []
        currency = terms.get("currency") or "COP"
        for k, v in terms.items():
            if v in (None, "", [], {}):
                continue
            rendered = self._money(v, currency) if k in {"amount", "value", "total"} else self._plain(v)
            rows.append([self._key_label(k), rendered])
        if rows:
            self._table(doc, ["Concepto", "Condición"], rows)
        self._clause(doc, 3, "Facturación y soportes", "La factura o cuenta de cobro deberá cumplir los requisitos aplicables y acompañarse únicamente de soportes razonables. La objeción deberá indicar el concepto discutido y no autoriza retener valores no controvertidos.")
        self._clause(doc, 4, "Gastos, anticipos y reembolsos", self._plain(self._v(answers, "fees.expenses"), "Los gastos reembolsables requerirán autorización previa y soportes; los anticipos deberán legalizarse y los saldos reintegrarse."))
        self._clause(doc, 5, "Terminación", "La terminación no extingue pagos ya causados. Los entregables parciales utilizables, gastos autorizados no recuperables y componentes por éxito ya consolidados se liquidarán conforme a la fórmula acordada.")
        doc.save(target)

    def _conf_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-CONF-001", "Anexo de confidencialidad y secretos empresariales")
        clauses = [
            ("Información protegida", "Comprende información técnica, científica, comercial, financiera, jurídica, operativa, estratégica, de seguridad, personal y de terceros que sea identificada como reservada o que razonablemente deba tratarse como tal por su naturaleza y contexto."),
            ("Uso permitido", "La parte receptora utilizará la información únicamente para ejecutar el contrato, aplicará acceso por necesidad de conocer y exigirá obligaciones equivalentes a su personal, asesores y subcontratistas."),
            ("Exclusiones", "No será confidencial la información pública sin incumplimiento; conocida legítimamente; recibida de tercero autorizado; desarrollada independientemente; o cuya divulgación haya sido aprobada por escrito."),
            ("Divulgaciones legítimas", "La reserva no impide cumplir una orden legal, ejercer defensa, consultar asesores sometidos a confidencialidad, reportar irregularidades o acudir a autoridades. Cuando sea jurídicamente posible se informará previamente y se limitará la revelación."),
            ("Seguridad e incidentes", "La parte receptora adoptará medidas técnicas, humanas y administrativas proporcionales, evitará copias innecesarias y reportará accesos, pérdidas o divulgaciones no autorizadas sin demora injustificada."),
            ("Devolución y conservación", "Al terminar se devolverá o eliminará la información y sus copias, salvo aquella que deba conservarse por ley, respaldo inalterable, defensa de derechos o auditoría, la cual permanecerá protegida y no se utilizará para otros fines."),
            ("Duración", "La obligación se mantendrá durante el contrato y por el plazo indicado en el expediente. Los secretos empresariales se protegerán mientras conserven esa calidad y se mantengan medidas razonables de reserva."),
            ("Remedios", "La parte afectada podrá solicitar medidas de protección, cesación, recuperación de información y reparación de daños probados. No se presume automáticamente un perjuicio ni se excluye el debido análisis de causalidad y responsabilidad."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.save(target)

    def _ip_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-IP-001", "Anexo de propiedad intelectual y activos preexistentes")
        allocation = self._plain(self._v(answers, "ip.allocation") or self._v(answers, "ip.details"))
        clauses = [
            ("Clasificación de activos", "LAS PARTES distinguirán materiales preexistentes, resultados desarrollados, componentes de terceros, software de código abierto, bases de datos, información no protegible, secretos empresariales, invenciones y conocimiento general."),
            ("Materiales preexistentes", "Cada parte conserva la titularidad de sus activos previos. Su incorporación deberá declararse y dará únicamente la licencia necesaria para usar el entregable conforme al contrato, salvo pacto escrito diferente."),
            ("Resultados del encargo", allocation or "La titularidad o licencia de los resultados se definirá por entregable, teniendo en cuenta autoría, encargo, aportes, contraprestación y finalidad. No se presume una cesión universal sobre activos no identificados."),
            ("Derechos patrimoniales", "Toda cesión o licencia identificará modalidades de explotación, territorio, duración, exclusividad y facultad de transformación o sublicencia. Los actos de transferencia constarán por escrito. La falta de precisión se interpretará restrictivamente conforme al régimen aplicable."),
            ("Derechos morales", "Los derechos morales permanecen en cabeza de los autores y no se entienden renunciados. LAS PARTES coordinarán créditos, modificaciones y acciones de protección sin impedir la explotación legítimamente autorizada."),
            ("Terceros y código abierto", "EL CONTRATISTA documentará licencias, atribuciones, restricciones de copyleft, componentes de terceros y obligaciones de divulgación. No incorporará materiales incompatibles con el uso esperado sin autorización informada."),
            ("Invenciones y propiedad industrial", "Los resultados potencialmente patentables o registrables se documentarán antes de divulgar. LAS PARTES definirán inventores, titulares, costos, territorios, explotación y cooperación en trámites según sus contribuciones y el acuerdo aplicable."),
            ("Archivos fuente y continuidad", "Se entregarán archivos editables, código fuente, documentación, credenciales institucionales, modelos, parámetros y demás elementos expresamente incluidos. No deberán entregarse secretos de terceros ni herramientas generales excluidas."),
            ("Inteligencia artificial", "Los resultados generados con IA deberán identificarse cuando sea material. No se garantiza protección exclusiva cuando falte autoría humana o existan restricciones del proveedor. Las entradas confidenciales y datos personales se regirán por los anexos correspondientes."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.save(target)

    def _data_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-DATA-001", "Acuerdo de encargo o tratamiento de datos personales")
        details = self._plain(self._v(answers, "data_processing.roles") or self._v(answers, "data_processing.details"))
        clauses = [
            ("Roles y objeto", details or "LAS PARTES identificarán quién decide sobre finalidades y medios como responsable y quién trata datos por cuenta de aquel como encargado, sin perjuicio de roles independientes para tratamientos propios."),
            ("Instrucciones documentadas", "El encargado tratará únicamente conforme a instrucciones lícitas y documentadas. Informará si una instrucción puede infringir la normativa y no utilizará los datos para finalidades propias incompatibles."),
            ("Categorías y titulares", "El expediente deberá identificar titulares, datos, operaciones, finalidades, sistemas, ubicaciones y plazo. Los datos sensibles o de menores requieren controles reforzados y habilitación jurídica suficiente."),
            ("Confidencialidad y acceso", "El acceso se limitará a personal autorizado, capacitado y sometido a confidencialidad. Se aplicará mínimo privilegio, autenticación, registro y revisión periódica de permisos."),
            ("Seguridad", "Se adoptarán medidas técnicas, humanas y administrativas proporcionales, incluyendo gestión de vulnerabilidades, cifrado cuando proceda, respaldo, continuidad, segregación, control de cambios y eliminación segura."),
            ("Subencargados y transferencias", "No se incorporarán subencargados ni transferencias nacionales o internacionales fuera de las condiciones autorizadas. El encargado impondrá obligaciones equivalentes y continuará respondiendo por su selección y supervisión."),
            ("Incidentes", "El encargado reportará sin demora injustificada cualquier evento que comprometa confidencialidad, integridad o disponibilidad, indicando hechos conocidos, categorías afectadas, medidas adoptadas y plan de contención, preservando evidencias."),
            ("Derechos de titulares", "Las solicitudes se remitirán o atenderán conforme a los roles definidos. El encargado colaborará con consultas, reclamos, actualización, rectificación, supresión y prueba de autorización dentro de los plazos aplicables."),
            ("Devolución, supresión y conservación", "Al finalizar, los datos se devolverán o eliminarán según instrucción, salvo obligación legal de conservación. Las copias de respaldo quedarán bloqueadas y se suprimirán en su ciclo normal."),
            ("Auditoría y responsabilidad", "El responsable podrá verificar razonablemente el cumplimiento sin acceder a información ajena. Cada parte responderá por sus decisiones, instrucciones y controles, y cooperará en investigaciones o requerimientos de autoridad."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.save(target)

    def _ai_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM-AI-001", "Anexo de uso controlado de inteligencia artificial")
        self._section(doc, "Clasificación de usos")
        self._table(doc, ["Clasificación", "Usos"], [
            ["Permitidos", "Ideación con información pública, corrección lingüística y apoyo administrativo sin datos reservados."],
            ["Controlados", "Redacción, análisis, código, imágenes, modelos y síntesis sujetos a sistema autorizado, revisión humana y trazabilidad."],
            ["Prohibidos", "Ingreso de datos sensibles en sistemas no aprobados, decisiones finales automáticas, suplantación, evasión de controles, clonación no autorizada y uso ilícito de contenidos."],
        ])
        clauses = [
            ("Sistemas autorizados", "Solo podrán utilizarse proveedores, modelos, cuentas y configuraciones aprobados. Los términos del proveedor deberán ser compatibles con confidencialidad, datos, licencias y ubicación del tratamiento."),
            ("Entradas", "No se ingresarán secretos, datos personales, información privilegiada, credenciales, código restringido ni materiales de terceros sin autorización y controles suficientes."),
            ("Resultados y revisión humana", "Todo resultado relevante deberá verificarse por una persona competente antes de usarlo. Se revisarán exactitud, sesgo, alucinaciones, seguridad, licencias, atribución y adecuación al propósito."),
            ("Trazabilidad", "Cuando sea material se registrarán herramienta, fecha, finalidad, versión, persona revisora y modificaciones. Los registros no deberán reproducir datos confidenciales innecesarios."),
            ("Propiedad intelectual", "El uso de IA no garantiza titularidad ni exclusividad. Se respetarán derechos sobre entradas, datasets, estilos, marcas y componentes; se informará la presencia material de contenido sintético cuando corresponda."),
            ("Decisiones y responsabilidad", "La herramienta no adoptará decisiones finales que afecten derechos o comprometan a LAS PARTES sin revisión humana autorizada. EL CONTRATISTA conserva responsabilidad por el entregable y no podrá excusarse en el proveedor de IA."),
            ("Incidentes", "Se reportarán fugas, respuestas inseguras, generación ilícita, sesgos materiales o acceso indebido. Podrá suspenderse la herramienta hasta completar evaluación y medidas correctivas."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.save(target)

    def _acceptance_act(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ACT-EM-ACCEPT-001", "Acta de aceptación de entregables")
        rows = []
        for i, item in enumerate(self._as_list(self._v(answers, "scope.deliverables")), 1):
            name = item.get("name") or item.get("deliverable") or item.get("title") if isinstance(item, dict) else item
            criteria = item.get("acceptance_criteria") or item.get("criteria") if isinstance(item, dict) else "Conformidad con el alcance"
            rows.append([i, name, criteria, "☐ Aceptado  ☐ Con reservas  ☐ Rechazado", ""])
        self._table(doc, ["No.", "Entregable", "Criterio", "Decisión", "Observaciones"], rows)
        self._paragraph(doc, "Las observaciones deberán ser específicas, trazables y referidas al alcance. La aceptación con reservas identificará ajustes y plazo. La firma no implica renuncia a garantías ocultas ni autoriza exigir requisitos nuevos sin control de cambios.")
        doc.save(target)

    def _change_act(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ACT-EM-CHANGE-001", "Orden o acta de control de cambios")
        self._table(doc, ["Campo", "Información"], [
            ["Solicitud", ""], ["Solicitante", ""], ["Justificación", ""], ["Alcance anterior", ""], ["Alcance propuesto", ""],
            ["Impacto en entregables", ""], ["Impacto en plazo", ""], ["Impacto económico", ""], ["Impacto en riesgos, datos o PI", ""], ["Fecha efectiva", ""],
        ])
        self._paragraph(doc, "El cambio solo será exigible después de aprobación por representantes autorizados. Las actividades urgentes de preservación deberán documentarse y no implican aceptación automática del impacto económico o temporal.")
        doc.save(target)

    def _closure_act(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ACT-EM-CLOSE-001", "Acta de terminación, entrega y cierre")
        self._table(doc, ["Componente", "Estado / detalle", "Responsable", "Fecha"], [
            ["Entregables finales y parciales", "", "", ""], ["Archivos fuente y documentación", "", "", ""],
            ["Activos y materiales", "", "", ""], ["Accesos y credenciales institucionales", "", "", ""],
            ["Datos personales: devolución o eliminación", "", "", ""], ["Información confidencial", "", "", ""],
            ["Propiedad intelectual y licencias", "", "", ""], ["Facturas y pagos pendientes", "", "", ""],
            ["Reclamaciones o salvedades", "", "", ""], ["Transición y soporte posterior", "", "", ""],
        ])
        self._paragraph(doc, "No deberán consignarse contraseñas en texto plano. La firma acredita la información expresamente registrada, pero no constituye paz y salvo general, novación, renuncia de derechos ni aceptación de obligaciones no determinadas, salvo declaración específica y separada.")
        doc.save(target)

    def _annex(self, doc_id, answers, target):
        dispatch = {
            "ANX-EM-SCOPE-001": self._scope_annex,
            "ANX-EM-FEES-001": self._fees_annex,
            "ANX-EM-CONF-001": self._conf_annex,
            "ANX-EM-IP-001": self._ip_annex,
            "ANX-EM-DATA-001": self._data_annex,
            "ANX-EM-AI-001": self._ai_annex,
            "ACT-EM-ACCEPT-001": self._acceptance_act,
            "ACT-EM-CHANGE-001": self._change_act,
            "ACT-EM-CLOSE-001": self._closure_act,
        }
        if doc_id not in dispatch:
            raise ValueError(f"Documento no soportado: {doc_id}")
        dispatch[doc_id](answers, target)

    @staticmethod
    def _extract_text(docx_path):
        doc = Document(docx_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)

    def render_documents(self, answers, target_folder):
        evaluation = self.evaluator.evaluate(answers)
        if evaluation["blocked"]:
            messages = "; ".join(x["message"] for x in evaluation["findings"] if x["severity"] == "blocker")
            raise ValueError("El expediente contiene bloqueos jurídicos: " + messages)
        if evaluation["missing_fields"]:
            missing = ", ".join(x["label"] for x in evaluation["missing_fields"])
            raise ValueError("Faltan datos esenciales: " + missing)
        target_folder = Path(target_folder)
        target_folder.mkdir(parents=True, exist_ok=False)
        contract = target_folder / "CO-EM-003_Contrato_Servicios_Independientes.docx"
        self._contract(answers, evaluation, contract)
        generated = [{"id": "DOC-EM-CONTRACT-001", "filename": contract.name}]
        for doc_id in evaluation["documents"]:
            if doc_id == "DOC-EM-CONTRACT-001":
                continue
            target = target_folder / f"{doc_id}.docx"
            self._annex(doc_id, answers, target)
            generated.append({"id": doc_id, "filename": target.name})
        hashes, unresolved = {}, []
        for item in generated:
            path = target_folder / item["filename"]
            text = self._extract_text(path)
            if UNRESOLVED_PATTERN.search(text):
                unresolved.append(item["id"])
            hashes[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()
        if unresolved:
            raise ValueError("Se detectaron variables o valores centinela sin resolver: " + ", ".join(unresolved))
        return evaluation, generated, hashes

    def generate(self, answers, actor=None):
        generation_id = "COEM003-" + uuid.uuid4().hex[:12].upper()
        folder = self.output_dir / generation_id
        folder.mkdir(parents=True, exist_ok=False)
        documents_dir = folder / "documents" / "revision-0001"
        evaluation, generated, hashes = self.render_documents(answers, documents_dir)
        manifest = {
            "generation_id": generation_id,
            "product_id": "CO-EM-003",
            "version": self.VERSION,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "actor": {"id": (actor or {}).get("id"), "role": (actor or {}).get("role")},
            "readiness": evaluation["readiness"],
            "requires_professional_review": bool(evaluation["review_requirements"] or evaluation["warnings"]),
            "review_requirements": evaluation["review_requirements"],
            "documents": generated,
            "document_folder": "documents/revision-0001",
            "hashes": hashes,
            "selected_blocks": evaluation["blocks"],
            "unresolved_variables": 0,
            "status": "draft_generated",
            "current_revision": 1,
            "legal_approval": {"status": "pending"},
            "qa_approval": {"status": "pending"},
            "released": False,
        }
        (folder / "answers.json").write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self._package(folder, self.output_dir / f"{generation_id}.zip")
        manifest["package_filename"] = package.name
        manifest["package_sha256"] = hashlib.sha256(package.read_bytes()).hexdigest()
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return manifest

    @staticmethod
    def _package(folder, target):
        with ZipFile(target, "w", ZIP_DEFLATED) as zf:
            for path in sorted(Path(folder).rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        return target

    def package_path(self, generation_id):
        if not GEN_RE.fullmatch(generation_id or ""):
            return None
        path = self.output_dir / f"{generation_id}.zip"
        return path if path.is_file() else None
