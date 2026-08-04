from __future__ import annotations

import hashlib
import json
import re
import secrets
from datetime import date, datetime, timezone
from pathlib import Path
import os
from typing import Any, Iterable, Optional, Tuple
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SENTINEL_RE = re.compile(
    r"(?:\[\[|\]\]|\{\{|\}\}|\bNULL\b|\bundefined\b|\bN/?A\b|<[^>]+_PENDIENTE>)",
    re.IGNORECASE,
)
GEN_RE = re.compile(r"COTR001-[A-F0-9]{12}")



def _runtime_root(project_root: Path) -> Path:
    raw = os.environ.get("LEGAL_RUNTIME_DIR", "").strip()
    path = Path(raw).expanduser() if raw else Path(project_root) / "runtime"
    if not path.is_absolute():
        path = Path(project_root) / path
    return path.resolve()

class DocumentGenerationError(ValueError):
    pass


class CoTr001DocumentFactoryV258:
    VERSION = "2.58"
    PRODUCT_ID = "CO-TR-001"
    PACKAGE_NAME = "00_PAQUETE_DOCUMENTAL_CO-TR-001_V258.docx"

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.product_dir = self.root / "app" / "assets" / "advanced-legal-library" / self.PRODUCT_ID
        self.output_dir = _runtime_root(self.root) / "generated" / "co-tr-001"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.document_specs = self._load_json("DOCUMENTOS_V258.json")
        self.blocks = self._load_json("BLOQUES_V258.json")
        self.sources = self._load_json("FUENTES_V258.json")

    def _load_json(self, name: str) -> dict[str, Any]:
        path = self.product_dir / name
        if not path.is_file():
            raise FileNotFoundError("No se encontró el activo requerido: %s" % path)
        return json.loads(path.read_text(encoding="utf-8"))

    @staticmethod
    def _now() -> str:
        return datetime.now(timezone.utc).isoformat()

    @staticmethod
    def _hash_file(path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    @staticmethod
    def _hash_obj(value: Any) -> str:
        raw = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        return hashlib.sha256(raw).hexdigest()

    @staticmethod
    def get_value(data: dict[str, Any], path: str, default: Any = None) -> Any:
        value: Any = data
        for part in path.split("."):
            if not isinstance(value, dict) or part not in value:
                return default
            value = value[part]
        return value

    @staticmethod
    def _human(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "Sí" if value else "No"
        if isinstance(value, (list, tuple)):
            return "; ".join(CoTr001DocumentFactoryV258._human(item) for item in value if item not in (None, ""))
        rendered = str(value).strip()
        translations = {
            "precheck": "chequeo preliminar",
            "registration": "inscripción verificada",
            "ready_with_review": "listo con revisión profesional",
            "ready": "listo",
            "incomplete": "incompleto",
            "blocked": "bloqueado",
            "blocked_professional_review": "bloqueado para revisión profesional",
            "pending_legal_review": "pendiente de revisión jurídica",
            "pending_qa_review": "pendiente de revisión QA",
            "released": "liberado",
            "yellow": "amarillo",
            "green": "verde",
            "red": "rojo",
            "unpaid": "no pagado",
            "paid": "pagado",
            "agreement": "acuerdo de pago",
            "none": "sin etapa informada",
            "collection": "cobro coactivo",
            "payment_order": "mandamiento de pago",
            "embargo": "embargo",
            "judicial": "proceso judicial",
        }
        return translations.get(rendered.lower(), rendered)

    @staticmethod
    def _safe_filename(value: str) -> str:
        value = re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("._")
        return value or "documento"

    @staticmethod
    def _set_repeat_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        element = OxmlElement("w:tblHeader")
        element.set(qn("w:val"), "true")
        tr_pr.append(element)

    @staticmethod
    def _shade(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    @staticmethod
    def _page_field(paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Página ")
        run.font.size = Pt(8)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instr, end])

    def _configure_document(self, doc: Document, generation_id: str, package: bool = False) -> None:
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(1.9)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.3)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for name, size, color in (
            ("Title", 17, RGBColor(13, 19, 36)),
            ("Heading 1", 13, RGBColor(13, 19, 36)),
            ("Heading 2", 11.5, RGBColor(37, 99, 235)),
        ):
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.size = Pt(size)
            style.font.color.rgb = color

        if "LegalAIZ Notice" not in styles:
            notice = styles.add_style("LegalAIZ Notice", WD_STYLE_TYPE.PARAGRAPH)
            notice.font.name = "Aptos"
            notice.font.size = Pt(8.7)
            notice.font.color.rgb = RGBColor(75, 75, 75)
            notice.paragraph_format.space_after = Pt(6)

        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0]
            p.text = "LegalAIZ.it · %s · v%s · %s" % (self.PRODUCT_ID, self.VERSION, generation_id)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(80, 80, 80)

            footer = section.footer
            fp = footer.paragraphs[0]
            fp.text = (
                "Documento generado con tecnología. Resultado preliminar sujeto a verificación de fuentes, hechos, "
                "decisión firme y aplicabilidad individual."
            )
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in fp.runs:
                run.font.size = Pt(7.2)
                run.font.color.rgb = RGBColor(100, 100, 100)
            self._page_field(footer.add_paragraph())

        if package:
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("PAQUETE DOCUMENTAL CO-TR-001").bold = True
            p2 = doc.add_paragraph(style="LegalAIZ Notice")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run("Chequeo SAST + inscripción verificada · Versión 2.58")
            p3 = doc.add_paragraph(style="LegalAIZ Notice")
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p3.add_run("BORRADOR CONTROLADO — NO ES DECISIÓN ADMINISTRATIVA NI GARANTIZA RESULTADO")
            r.bold = True
            r.font.color.rgb = RGBColor(160, 70, 20)
            doc.add_page_break()

    def _add_title(self, doc: Document, title: str) -> None:
        p = doc.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title.upper()).bold = True
        n = doc.add_paragraph(style="LegalAIZ Notice")
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        n.add_run("Producto %s · Colombia · Versión %s" % (self.PRODUCT_ID, self.VERSION))
        d = doc.add_paragraph(style="LegalAIZ Notice")
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = d.add_run("BORRADOR CONTROLADO — VERIFICAR ANTES DE USAR")
        rr.bold = True
        rr.font.color.rgb = RGBColor(160, 70, 20)

    def _add_control_notice(self, doc: Document, evaluation: dict[str, Any]) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        self._set_repeat_header(table.rows[0])
        cell = table.cell(0, 0)
        self._shade(cell, "F7F5F1")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.add_run("CONTROL JURÍDICO PREVIO. ").bold = True
        p.add_run(
            "La coincidencia se limita al cruce de organismo y fecha contra una instantánea de fuentes oficiales. "
            "No acredita por sí sola que el dispositivo concreto esté afectado ni que exista una decisión firme aplicable al caso."
        )
        if evaluation.get("review_requirements"):
            p2 = cell.add_paragraph()
            p2.add_run("Revisión profesional: ").bold = True
            p2.add_run("; ".join(evaluation["review_requirements"]))
        if evaluation.get("release_blockers"):
            p3 = cell.add_paragraph()
            p3.add_run("Bloqueos de liberación: ").bold = True
            p3.add_run("; ".join(evaluation["release_blockers"]))

    def _add_key_value_table(self, doc: Document, rows: Iterable[Tuple[str, Any]], title: Optional[str] = None) -> None:
        if title:
            doc.add_paragraph(title, style="Heading 1")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0]
        self._set_repeat_header(hdr)
        hdr.cells[0].text = "Campo"
        hdr.cells[1].text = "Valor"
        for cell in hdr.cells:
            self._shade(cell, "0D1324")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = self._human(value) or "No informado"
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_match_table(self, doc: Document, matches: list[dict[str, Any]]) -> None:
        doc.add_paragraph("Coincidencias de la matriz", style="Heading 1")
        if not matches:
            p = doc.add_paragraph()
            p.add_run("No se encontró coincidencia en los rangos cargados. ").bold = True
            p.add_run(
                "Este resultado no demuestra la legalidad del comparendo, no descarta otras irregularidades y no sustituye la revisión del expediente."
            )
            return
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0]
        self._set_repeat_header(hdr)
        labels = ["ID", "Grupo", "Organismo", "Inicio", "Fin", "Fuente"]
        for cell, label in zip(hdr.cells, labels):
            cell.text = label
            self._shade(cell, "0D1324")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(8)
        for item in matches:
            cells = table.add_row().cells
            values = [
                item.get("id"), item.get("group"), item.get("authority_name"), item.get("start_date"), item.get("end_date"), item.get("source_id")
            ]
            for cell, value in zip(cells, values):
                cell.text = self._human(value)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
        for item in matches:
            p = doc.add_paragraph()
            p.add_run("%s — causa publicada: " % self._human(item.get("id"))).bold = True
            p.add_run(self._human(item.get("cause")) or "Sin descripción adicional.")
            if item.get("notes"):
                p.add_run(" Observación: ").bold = True
                p.add_run(self._human(item.get("notes")))

    def _add_findings(self, doc: Document, evaluation: dict[str, Any]) -> None:
        doc.add_paragraph("Hallazgos y controles", style="Heading 1")
        findings = evaluation.get("findings") or []
        if not findings:
            doc.add_paragraph("No se registraron hallazgos adicionales en la evaluación.")
            return
        for item in findings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("%s [%s]: " % (self._human(item.get("id")), self._human(item.get("risk")).upper())).bold = True
            p.add_run(self._human(item.get("message")))

    def _add_limits(self, doc: Document) -> None:
        doc.add_paragraph("Límites de interpretación", style="Heading 1")
        limits = [
            "Una investigación publicada no equivale automáticamente a una decisión sancionatoria firme.",
            "El artículo 158A exige verificar la firmeza de la decisión de la Superintendencia de Transporte y su correspondencia con autoridad, periodo, dispositivo y comparendo individual.",
            "El chequeo no suspende términos, no reemplaza recursos o acciones y no constituye representación judicial.",
            "Un pago o acuerdo de pago exige análisis individual; este producto no promete devolución.",
            "La instantánea de fuentes debe revalidarse antes de cualquier uso externo o decisión profesional.",
        ]
        for item in limits:
            doc.add_paragraph(item, style="List Bullet")

    def _add_next_steps(self, doc: Document, evaluation: dict[str, Any]) -> None:
        doc.add_paragraph("Siguientes pasos", style="Heading 1")
        steps = []
        if evaluation.get("matches"):
            steps.extend([
                "Obtener el comparendo, soportes de validación y evidencia del dispositivo SAST concreto.",
                "Verificar si existe decisión firme de la Superintendencia de Transporte aplicable al organismo y periodo.",
                "Contrastar el expediente individual, estado procesal, pagos, acuerdos y términos vigentes.",
            ])
        else:
            steps.extend([
                "Revisar el expediente individual y la notificación, aunque no exista coincidencia en esta matriz.",
                "Verificar actualizaciones oficiales posteriores a la fecha de la instantánea.",
            ])
        if evaluation.get("professional_review_required"):
            steps.append("Someter el caso y sus anexos a revisión jurídica profesional antes de actuar.")
        for idx, item in enumerate(steps, start=1):
            p = doc.add_paragraph()
            p.add_run("%d. " % idx).bold = True
            p.add_run(item)

    def _source_by_id(self) -> dict[str, dict[str, Any]]:
        return {item["id"]: item for item in self.sources.get("sources", [])}

    def _render_preliminary(self, doc: Document, answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        self._add_control_notice(doc, evaluation)
        authority_resolution = evaluation.get("authority_resolution") or {}
        self._add_key_value_table(
            doc,
            [
                ("Organismo ingresado", self.get_value(answers, "authority.name")),
                ("Organismo normalizado", authority_resolution.get("authority_key") or authority_resolution.get("normalized")),
                ("Fecha de la presunta infracción", evaluation.get("infraction_date")),
                ("Número de comparendo", self.get_value(answers, "infraction.comparendo_number")),
                ("Ubicación o dispositivo", self.get_value(answers, "device.location") or self.get_value(answers, "device.id")),
                ("Modo", evaluation.get("mode")),
                ("Estado", evaluation.get("status")),
                ("Riesgo", evaluation.get("risk")),
                ("Coincidencias", evaluation.get("match_count")),
                ("Instantánea de fuentes", evaluation.get("source_snapshot")),
            ],
            "Entradas y resultado",
        )
        self._add_match_table(doc, evaluation.get("matches") or [])
        self._add_findings(doc, evaluation)
        self._add_limits(doc)
        self._add_next_steps(doc, evaluation)

    def _render_trace(self, doc: Document, answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        self._add_control_notice(doc, evaluation)
        self._add_key_value_table(
            doc,
            [
                ("Organismo consultado", self.get_value(answers, "authority.name")),
                ("Fecha consultada", evaluation.get("infraction_date")),
                ("Instantánea de matriz", evaluation.get("source_snapshot")),
                ("Hash de entradas", self._hash_obj({"authority": self.get_value(answers, "authority.name"), "date": evaluation.get("infraction_date")})),
                ("Número de coincidencias", evaluation.get("match_count")),
            ],
            "Trazabilidad del chequeo",
        )
        self._add_match_table(doc, evaluation.get("matches") or [])
        doc.add_paragraph("Fuentes oficiales registradas", style="Heading 1")
        source_map = self._source_by_id()
        used_ids = {item.get("source_id") for item in evaluation.get("matches") or []}
        used_ids.update({"ST-2026-37-WEB", "ST-2026-12-WEB", "LEY-2251-2022-ART18", "LEY-1843-2017"})
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0]
        self._set_repeat_header(hdr)
        for cell, label in zip(hdr.cells, ["ID", "Autoridad", "Fecha", "Título / URL"]):
            cell.text = label
            self._shade(cell, "0D1324")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(8)
        for source_id in sorted(item for item in used_ids if item in source_map):
            item = source_map[source_id]
            cells = table.add_row().cells
            values = [item.get("id"), item.get("authority"), item.get("date"), "%s\n%s" % (item.get("title"), item.get("url"))]
            for cell, value in zip(cells, values):
                cell.text = self._human(value)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(7.5)
        p = doc.add_paragraph(style="LegalAIZ Notice")
        p.add_run("Política de actualización: ").bold = True
        p.add_run("revisión mensual y actualización extraordinaria ante decisiones firmes, correcciones o nuevas matrices oficiales.")

    def _render_case_file(self, doc: Document, answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        self._add_control_notice(doc, evaluation)
        self._add_key_value_table(
            doc,
            [
                ("Nombre", self.get_value(answers, "identity.full_name")),
                ("Tipo y número de documento", "%s %s" % (self._human(self.get_value(answers, "identity.document_type")), self._human(self.get_value(answers, "identity.document_number")))),
                ("Correo", self.get_value(answers, "identity.email")),
                ("Ciudad", self.get_value(answers, "identity.city")),
                ("Organismo", self.get_value(answers, "authority.name")),
                ("Fecha de la infracción", evaluation.get("infraction_date")),
                ("Comparendo", self.get_value(answers, "infraction.comparendo_number")),
                ("Dispositivo o ubicación", self.get_value(answers, "device.id") or self.get_value(answers, "device.location")),
                ("Estado de pago", self.get_value(answers, "payment.status")),
                ("Etapa informada", self.get_value(answers, "procedure.stage")),
            ],
            "Identificación mínima del expediente",
        )
        doc.add_paragraph("Registro de consentimientos", style="Heading 1")
        consents = [
            ("Carácter preliminar", self.get_value(answers, "consents.preliminary"), "Comprende que el chequeo no decide el caso."),
            ("Tratamiento de datos", self.get_value(answers, "consents.data_processing"), "Autoriza crear y conservar el expediente para la finalidad informada."),
        ]
        for label, value, purpose in consents:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("%s: %s. " % (label, self._human(value))).bold = True
            p.add_run(purpose)
        p = doc.add_paragraph(style="LegalAIZ Notice")
        p.add_run("Minimización y acceso restringido. ").bold = True
        p.add_run(
            "Solo deben conservarse datos pertinentes para verificar el caso. El expediente debe mantenerse con controles de acceso, trazabilidad y posibilidad de consulta, corrección o supresión conforme a la política aplicable."
        )
        self._add_match_table(doc, evaluation.get("matches") or [])
        self._add_findings(doc, evaluation)
        doc.add_paragraph("Lista de verificación de anexos", style="Heading 1")
        items = [
            "Copia legible del comparendo y sus soportes.",
            "Prueba de la ubicación o identificación del dispositivo SAST.",
            "Consulta del estado del comparendo y del procedimiento.",
            "Constancias de pago o acuerdo, si existen.",
            "Actos, comunicaciones, mandamientos o decisiones recibidas.",
            "Evidencia de términos o audiencias próximas.",
        ]
        for item in items:
            doc.add_paragraph("☐ " + item)

    def _render_review(self, doc: Document, answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        self._add_control_notice(doc, evaluation)
        self._add_key_value_table(
            doc,
            [
                ("Organismo", self.get_value(answers, "authority.name")),
                ("Fecha de la presunta infracción", evaluation.get("infraction_date")),
                ("Comparendo", self.get_value(answers, "infraction.comparendo_number")),
                ("Coincidencias", evaluation.get("match_count")),
                ("Estado de pago", self.get_value(answers, "payment.status")),
                ("Etapa procesal", self.get_value(answers, "procedure.stage")),
                ("Término urgente", self.get_value(answers, "procedure.imminent_deadline")),
                ("Posible fraude", self.get_value(answers, "security.possible_fraud")),
                ("Riesgo", evaluation.get("risk")),
            ],
            "Resumen para revisión",
        )
        self._add_match_table(doc, evaluation.get("matches") or [])
        self._add_findings(doc, evaluation)
        doc.add_paragraph("Cuestiones que debe resolver el especialista", style="Heading 1")
        questions = [
            "¿La actuación oficial publicada se convirtió en decisión firme y cuál es su alcance exacto?",
            "¿El dispositivo concreto y la fecha del caso están dentro del periodo jurídicamente afectado?",
            "¿Qué actos administrativos existen y cuál es la etapa procesal real?",
            "¿Hay términos, recursos, cobro coactivo, embargo o proceso judicial que exijan actuación inmediata?",
            "¿Existe pago o acuerdo y qué consecuencias jurídicas y probatorias produce?",
            "¿Qué documentos adicionales deben solicitarse antes de definir una estrategia?",
        ]
        for item in questions:
            doc.add_paragraph(item, style="List Bullet")
        self._add_limits(doc)

    def _render_content(self, doc: Document, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        self._add_title(doc, spec["title"])
        kind = spec.get("kind")
        if kind == "report":
            self._render_preliminary(doc, answers, evaluation)
        elif kind == "traceability":
            self._render_trace(doc, answers, evaluation)
        elif kind == "case_file":
            self._render_case_file(doc, answers, evaluation)
        elif kind == "review":
            self._render_review(doc, answers, evaluation)
        else:
            raise DocumentGenerationError("Tipo documental no soportado: %s" % kind)

    def _validate_required(self, answers: dict[str, Any], spec: dict[str, Any]) -> None:
        missing = []
        for path in spec.get("required", []):
            value = self.get_value(answers, path)
            if value in (None, "", False):
                missing.append(path)
        if missing:
            raise DocumentGenerationError("Faltan variables requeridas para %s: %s" % (spec["id"], ", ".join(missing)))

    def _selected_specs(self, evaluation: dict[str, Any], mode: str) -> list[dict[str, Any]]:
        ids = set(evaluation.get("documents") or [])
        specs = []
        for spec in self.document_specs.get("documents", []):
            if spec["id"] in ids and mode in spec.get("modes", [mode]):
                specs.append(spec)
        return specs

    def _assert_clean_docx(self, path: Path) -> None:
        if not path.is_file() or path.stat().st_size < 5000:
            raise DocumentGenerationError("El documento generado es inexistente o anormalmente pequeño: %s" % path.name)
        with ZipFile(path) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        if SENTINEL_RE.search(document_xml):
            raise DocumentGenerationError("El documento contiene variables, centinelas o marcadores sin resolver: %s" % path.name)

    def render_documents(
        self,
        answers: dict[str, Any],
        target_dir: Path,
        generation_id: str,
        mode: str = "precheck",
    ) -> Tuple[dict[str, Any], list[str], dict[str, str]]:
        evaluation = self.evaluator.evaluate(answers, mode=mode)
        if evaluation.get("missing_fields"):
            fields = ", ".join(item.get("path", "") for item in evaluation["missing_fields"])
            raise DocumentGenerationError("No es posible generar con datos esenciales pendientes: %s" % fields)
        specs = self._selected_specs(evaluation, mode)
        if not specs:
            raise DocumentGenerationError("La evaluación no seleccionó documentos para el modo solicitado.")

        target_dir = Path(target_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        documents: list[str] = []
        hashes: dict[str, str] = {}

        for spec in specs:
            self._validate_required(answers, spec)
            filename = self._safe_filename(spec["filename"]) + ".docx"
            path = target_dir / filename
            doc = Document()
            self._configure_document(doc, generation_id)
            self._render_content(doc, spec, answers, evaluation)
            doc.core_properties.title = spec["title"]
            doc.core_properties.subject = self.PRODUCT_ID
            doc.core_properties.author = "LegalAIZ.it"
            doc.core_properties.comments = "Borrador controlado generado con tecnología; requiere revisión antes de uso externo."
            doc.save(path)
            self._assert_clean_docx(path)
            documents.append(filename)
            hashes[filename] = self._hash_file(path)

        package_path = target_dir / self.PACKAGE_NAME
        package = Document()
        self._configure_document(package, generation_id, package=True)
        for index, spec in enumerate(specs):
            if index:
                package.add_section(WD_SECTION.NEW_PAGE)
            self._render_content(package, spec, answers, evaluation)
        package.core_properties.title = "Paquete documental CO-TR-001 v2.58"
        package.core_properties.author = "LegalAIZ.it"
        package.core_properties.comments = "Paquete documental controlado; no constituye decisión administrativa."
        package.save(package_path)
        self._assert_clean_docx(package_path)
        documents.insert(0, self.PACKAGE_NAME)
        hashes[self.PACKAGE_NAME] = self._hash_file(package_path)
        return evaluation, documents, hashes

    def generate(self, answers: dict[str, Any], actor: dict[str, Any], mode: str = "precheck") -> dict[str, Any]:
        generation_id = "COTR001-%s" % secrets.token_hex(6).upper()
        if not GEN_RE.fullmatch(generation_id):
            raise DocumentGenerationError("No fue posible crear un identificador de generación válido.")
        folder = self.output_dir / generation_id
        if folder.exists():
            raise DocumentGenerationError("Colisión de identificador de generación.")
        document_dir = folder / "documents" / "revision-0001"
        evaluation, documents, hashes = self.render_documents(answers, document_dir, generation_id, mode=mode)

        manifest = {
            "generation_id": generation_id,
            "product_id": self.PRODUCT_ID,
            "version": self.VERSION,
            "mode": mode,
            "created_at": self._now(),
            "created_by": {"id": actor.get("id"), "role": actor.get("role")},
            "source_snapshot": evaluation.get("source_snapshot"),
            "source_revalidation_required": True,
            "evaluation_hash": self._hash_obj(evaluation),
            "answers_hash": self._hash_obj(answers),
            "selected_blocks": evaluation.get("blocks", []),
            "review_requirements": evaluation.get("review_requirements", []),
            "release_blockers": evaluation.get("release_blockers", []),
            "document_folder": "documents/revision-0001",
            "documents": documents,
            "hashes": hashes,
            "status": "draft",
            "notice": evaluation.get("document_disclaimer"),
        }
        folder.mkdir(parents=True, exist_ok=True)
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "answers.json").write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "evaluation.json").write_text(json.dumps(evaluation, ensure_ascii=False, indent=2), encoding="utf-8")

        package_path = self.output_dir / (generation_id + ".zip")
        with ZipFile(package_path, "w", ZIP_DEFLATED) as zf:
            for path in sorted(folder.rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        return {
            "generation_id": generation_id,
            "folder": str(folder),
            "package": str(package_path),
            "package_sha256": self._hash_file(package_path),
            "documents": documents,
            "hashes": hashes,
            "evaluation": evaluation,
            "mode": mode,
        }
