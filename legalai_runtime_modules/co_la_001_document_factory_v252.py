from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from co_em_003_document_factory_v243 import (
    CoEm003DocumentFactoryV243,
    UNRESOLVED_PATTERN,
    number_to_words_es,
)


GEN_RE = re.compile(r"COLA001-[A-F0-9]{12}")


class CoLa001DocumentFactoryV252(CoEm003DocumentFactoryV243):
    """Fábrica determinística y documental para CO-LA-001.

    Genera una estimación trazable por concepto y documentos jurídicos de apoyo.
    No convierte en deuda cierta los conceptos controvertidos ni incorpora de forma
    automática sanciones cuya procedencia depende de valoración probatoria.
    """

    VERSION = "2.52"
    SMMLV_2026 = 1_750_905.0
    TRANSPORT_AID_2026 = 249_095.0
    TRANSPORT_THRESHOLD = 2.0
    CESANTIA_INTEREST_RATE = 0.12

    ENUM_LABELS = {
        "worker": "trabajador", "employer": "empleador", "advisor": "abogado o asesor", "other": "otro",
        "indefinite": "término indefinido", "fixed": "término fijo", "work": "obra o labor", "verbal": "verbal o no documentado", "unknown": "por verificar",
        "resignation": "renuncia", "mutual": "mutuo acuerdo", "without_cause": "terminación sin justa causa", "just_cause": "terminación con justa causa",
        "fixed_or_work_end": "finalización del término u obra", "active": "relación vigente", "ended": "terminada", "suspended": "suspendida",
        "ordinary": "ordinario", "integral": "integral", "mixed": "mixto o discutido",
        "yes": "sí", "no": "no", "review": "requiere revisión", "complete": "completos", "partial": "parciales", "none": "sin soportes",
        "conciliation": "conciliación", "lawsuit": "proceso judicial", "settlement": "transacción o acuerdo previo",
    }

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.output_dir = self.root / "data" / "generated" / "co-la-001-v252"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _set_doc_styles(doc: Document):
        sec = doc.sections[0]
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.55)
        sec.right_margin = Cm(2.55)
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.1)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for style_name in ("Title", "Heading 1", "Heading 2"):
            style = doc.styles[style_name]
            style.font.name = "Aptos Display"
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Heading 1"].font.size = Pt(12)
        doc.styles["Heading 2"].font.size = Pt(11)
        footer = sec.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("LegalAIZ.it | CO-LA-001 | Documento sujeto a revisión jurídica, contable y QA | Página ")
        run.font.size = Pt(8)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer._p.append(fld)

    @classmethod
    def _label(cls, value, default="No informado"):
        if isinstance(value, bool):
            return "Sí" if value else "No"
        if value in (None, ""):
            return default
        return cls.ENUM_LABELS.get(str(value).strip(), str(value).strip())

    @staticmethod
    def _num(value):
        try:
            return 0.0 if value in (None, "") else float(value)
        except (TypeError, ValueError):
            return 0.0

    @staticmethod
    def _dt(value):
        if not value:
            return None
        try:
            return date.fromisoformat(str(value))
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _commercial_days(start: date, end: date, inclusive=True):
        if not start or not end or end < start:
            return 0
        d1 = min(start.day, 30)
        d2 = min(end.day, 30)
        days = (end.year - start.year) * 360 + (end.month - start.month) * 30 + (d2 - d1)
        return max(days + (1 if inclusive else 0), 0)

    @staticmethod
    def _calendar_year_segments(start: date, end: date):
        cursor = start
        while cursor <= end:
            segment_end = min(end, date(cursor.year, 12, 30))
            yield cursor, segment_end
            cursor = date(cursor.year + 1, 1, 1)

    @staticmethod
    def _money(value, currency="COP"):
        try:
            amount = int(round(float(value)))
            formatted = "$" + f"{amount:,}".replace(",", ".")
            return f"{formatted} {currency or 'COP'}"
        except (TypeError, ValueError):
            return str(value or "")

    @classmethod
    def _money_words(cls, value):
        amount = max(0, int(round(cls._num(value))))
        return f"{number_to_words_es(amount).upper()} PESOS M/CTE ({cls._money(amount)})"

    @classmethod
    def _table(cls, doc, headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr = table.rows[0]
        tr_pr = hdr._tr.get_or_add_trPr()
        h = OxmlElement("w:tblHeader"); h.set(qn("w:val"), "true"); tr_pr.append(h)
        tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, label in zip(hdr.cells, headers):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(label)); r.bold = True; r.font.size = Pt(8.2)
        for row in rows:
            ro = table.add_row(); ro._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for cell, value in zip(ro.cells, row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(value)); r.font.size = Pt(8.2)
        return table

    @staticmethod
    def _signature_table(doc, worker, employer):
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        values = [
            ("LA PERSONA TRABAJADORA", "EL EMPLEADOR"),
            (worker["name"], employer["name"]),
            (worker["id"], employer["id"]),
            ("Firma: __________________________", "Firma: __________________________"),
        ]
        for i, row in enumerate(table.rows):
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            if i == 0:
                h = OxmlElement("w:tblHeader"); h.set(qn("w:val"), "true"); row._tr.get_or_add_trPr().append(h)
            for j, cell in enumerate(row.cells):
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(values[i][j]); r.font.size = Pt(9); r.bold = i == 0
        return table

    def _party(self, answers, prefix, fallback):
        ident = self._as_dict(self._v(answers, f"{prefix}.identity"))
        signatory = self._as_dict(self._v(answers, "employer.signatory")) if prefix == "employer" else {}
        name = self._plain(ident.get("name"), fallback)
        ident_no = self._plain(ident.get("id_number"), "identificación no informada")
        email = self._plain(ident.get("email"), "correo no informado")
        address = self._plain(ident.get("address"), "domicilio no informado")
        signatory_name = self._plain(signatory.get("name"), name)
        capacity = self._plain(signatory.get("capacity"), "responsable autorizado")
        authority = self._plain(signatory.get("authority_source"))
        text = f"{name}, identificado(a) con {ident_no}, con domicilio o dirección en {address} y correo {email}"
        if prefix == "employer" and signatory:
            text += f", actuando por conducto de {signatory_name}, en calidad de {capacity}"
            if authority:
                text += f", con fundamento en {authority}"
        return {"name": name, "id": ident_no, "email": email, "address": address, "text": text, "signatory": signatory_name, "capacity": capacity}

    def calculate(self, answers):
        start = self._dt(self._v(answers, "relationship.start_date"))
        end = self._dt(self._v(answers, "relationship.end_date"))
        cutoff = self._dt(self._v(answers, "claim.cutoff_date")) or end
        if not start or not end or end < start:
            raise ValueError("Las fechas del vínculo no permiten ejecutar el cálculo.")

        ces_start = self._dt(self._v(answers, "periods.cesantias_start_date")) or start
        prima_start = self._dt(self._v(answers, "periods.prima_start_date")) or start
        ces_start = max(ces_start, start)
        prima_start = max(prima_start, start)

        fixed_salary = self._num(self._v(answers, "compensation.base_salary"))
        variable = self._num(self._v(answers, "compensation.variable.monthly_average")) if self._v(answers, "compensation.variable.exists") is True else 0.0
        other = self._as_dict(self._v(answers, "compensation.other_salary"))
        other_salary = sum(self._num(other.get(k)) for k in ("overtime", "surcharges", "in_kind"))
        salary = fixed_salary + variable + other_salary
        if salary <= 0:
            raise ValueError("La base salarial debe ser positiva.")

        aid_answer = self._v(answers, "compensation.transport_aid")
        aid_applied = aid_answer == "yes" and salary <= self.SMMLV_2026 * self.TRANSPORT_THRESHOLD
        transport = self.TRANSPORT_AID_2026 if aid_applied else 0.0
        benefit_base = salary + transport

        link_days = self._commercial_days(start, end)
        ces_days = self._commercial_days(ces_start, end) if ces_start <= end else 0
        prima_days = self._commercial_days(prima_start, end) if prima_start <= end else 0
        salary_due_days = max(0.0, self._num(self._v(answers, "periods.salary_due_days")))
        vacation_pending_days = max(0.0, self._num(self._v(answers, "periods.vacation_pending_days")))
        vacation_accrued_ceiling = round(link_days * 15 / 360, 4)

        gross = {
            "salary_due": salary / 30 * salary_due_days,
            "cesantias": benefit_base * ces_days / 360,
            "prima": benefit_base * prima_days / 360,
            "vacation": salary / 30 * vacation_pending_days,
        }
        interest_segments = []
        interest_total = 0.0
        if ces_days:
            for seg_start, seg_end in self._calendar_year_segments(ces_start, end):
                seg_days = self._commercial_days(seg_start, seg_end)
                seg_ces = benefit_base * seg_days / 360
                seg_interest = seg_ces * self.CESANTIA_INTEREST_RATE * seg_days / 360
                interest_total += seg_interest
                interest_segments.append({
                    "year": seg_start.year,
                    "start": seg_start.isoformat(),
                    "end": seg_end.isoformat(),
                    "days": seg_days,
                    "cesantias_base": round(seg_ces, 2),
                    "interest": round(seg_interest, 2),
                })
        gross["cesantias_interest"] = interest_total

        contract_type = self._v(answers, "relationship.contract_type")
        termination_type = self._v(answers, "relationship.termination_type")
        without_claim = self._v(answers, "termination.without_cause_claim") in {"yes", "unknown"} or termination_type == "without_cause"
        indemnity_days = 0.0
        indemnity_formula = "No aplica con la información suministrada."
        if without_claim:
            if contract_type == "indefinite":
                excess_days = max(link_days - 360, 0)
                if salary < 10 * self.SMMLV_2026:
                    indemnity_days = 30 + excess_days * 20 / 360
                    indemnity_formula = "30 días por el primer año y 20 días proporcionales por cada año adicional, para salario inferior a 10 SMLMV."
                else:
                    indemnity_days = 20 + excess_days * 15 / 360
                    indemnity_formula = "20 días por el primer año y 15 días proporcionales por cada año adicional, para salario igual o superior a 10 SMLMV."
            elif contract_type == "fixed":
                fixed_end = self._dt(self._v(answers, "relationship.special_term.fixed_end_date"))
                if not fixed_end or fixed_end <= end:
                    raise ValueError("La indemnización de término fijo requiere una fecha final posterior a la terminación.")
                indemnity_days = (fixed_end - end).days
                indemnity_formula = "Salarios correspondientes al tiempo calendario faltante para cumplir el término pactado."
            elif contract_type == "work":
                remaining = self._num(self._v(answers, "relationship.special_term.work_remaining_days"))
                if remaining <= 0:
                    raise ValueError("La indemnización de obra o labor requiere días restantes verificables.")
                indemnity_days = max(remaining, 15)
                indemnity_formula = "Tiempo estimado restante de la obra o labor, con mínimo de quince días."
        gross["indemnity"] = salary / 30 * indemnity_days

        prior = self._as_dict(self._v(answers, "payments.prior"))
        indemnity_prior_group = self._num(prior.get("indemnity"))
        indemnity_prior_direct = self._num(self._v(answers, "termination.indemnity_already_paid"))
        indemnity_prior = max(indemnity_prior_group, indemnity_prior_direct)
        prior_map = {
            "salary_due": self._num(prior.get("salary")),
            "cesantias": self._num(prior.get("cesantias")),
            "cesantias_interest": self._num(prior.get("interests")),
            "prima": self._num(prior.get("prima")),
            "vacation": self._num(prior.get("vacation")),
            "indemnity": indemnity_prior,
        }
        labels = {
            "salary_due": "Salario pendiente",
            "cesantias": "Cesantías",
            "cesantias_interest": "Intereses a las cesantías",
            "prima": "Prima de servicios",
            "vacation": "Vacaciones compensables",
            "indemnity": "Indemnización por terminación",
        }
        formulas = {
            "salary_due": "salario mensual ÷ 30 × días pendientes",
            "cesantias": "base prestacional × días del período ÷ 360",
            "cesantias_interest": "cesantías de cada segmento × 12 % × días del segmento ÷ 360",
            "prima": "base prestacional × días del período ÷ 360",
            "vacation": "salario mensual ÷ 30 × días pendientes confirmados",
            "indemnity": indemnity_formula,
        }
        source_ids = {
            "salary_due": ["LA1-S1"], "cesantias": ["LA1-S1", "LA1-S2"], "cesantias_interest": ["LA1-S3"],
            "prima": ["LA1-S1"], "vacation": ["LA1-S1"], "indemnity": ["LA1-S1", "LA1-S4"],
        }
        line_items = []
        notes = []
        for key in ("salary_due", "cesantias", "cesantias_interest", "prima", "vacation", "indemnity"):
            gross_value = max(0.0, gross.get(key, 0.0))
            paid = max(0.0, prior_map.get(key, 0.0))
            net = max(0.0, gross_value - paid)
            if paid > gross_value + 0.01:
                notes.append(f"El pago previo informado para {labels[key].lower()} supera el valor bruto; el saldo se llevó a cero y debe conciliarse.")
            line_items.append({
                "key": key, "label": labels[key], "gross": round(gross_value, 2), "prior_paid": round(paid, 2),
                "net": round(net, 2), "formula": formulas[key], "source_ids": source_ids[key],
            })

        deductions = self._as_dict(self._v(answers, "payments.deductions")) if self._v(answers, "payments.deductions.exists") is True else {}
        disputed_deductions = max(0.0, self._num(deductions.get("amount")))
        if disputed_deductions:
            notes.append("Las deducciones discutidas no fueron restadas del total; su procedencia requiere soporte, autorización o fundamento legal.")
        if aid_answer == "yes" and not aid_applied:
            notes.append("El auxilio de transporte no se aplicó porque el ingreso salarial informado supera el umbral parametrizado de dos SMLMV.")
        if aid_answer == "unknown":
            notes.append("El auxilio de transporte no se aplicó porque su procedencia no fue confirmada.")
        if vacation_pending_days > vacation_accrued_ceiling + 0.01:
            notes.append(f"Los días de vacaciones pendientes ({vacation_pending_days:g}) superan el máximo teórico causado ({vacation_accrued_ceiling:g}); debe revisarse el historial de disfrute y acumulación.")
        if indemnity_prior_group and indemnity_prior_direct and abs(indemnity_prior_group - indemnity_prior_direct) > 0.01:
            notes.append("Existen dos cifras distintas de indemnización pagada; se usó la mayor para evitar doble descuento y se requiere conciliación.")

        subtotal = sum(x["gross"] for x in line_items)
        prior_total = sum(x["prior_paid"] for x in line_items)
        total = sum(x["net"] for x in line_items)
        delay_days = max(0.0, self._num(self._v(answers, "termination.delay_days")))
        moratory_reference = salary / 30 * min(delay_days, 720) if self._v(answers, "termination.moratory_claim") in {"yes", "review"} else 0.0

        return {
            "engine_version": self.VERSION,
            "parameter_version": "CO-LA-001-2026.1",
            "verified_at": "2026-07-24",
            "periods": {
                "employment": {"start": start.isoformat(), "end": end.isoformat(), "days_30_360": link_days},
                "cesantias": {"start": ces_start.isoformat(), "end": end.isoformat(), "days_30_360": ces_days},
                "prima": {"start": prima_start.isoformat(), "end": end.isoformat(), "days_30_360": prima_days},
                "vacation": {"pending_days": round(vacation_pending_days, 4), "accrued_ceiling_days": vacation_accrued_ceiling},
            },
            "salary": round(salary, 2), "fixed_salary": round(fixed_salary, 2), "variable_average": round(variable, 2),
            "other_salary_average": round(other_salary, 2), "transport_aid_applied": aid_applied,
            "transport_aid_value": round(transport, 2), "benefit_base": round(benefit_base, 2),
            "indemnity_days": round(indemnity_days, 4), "indemnity_formula": indemnity_formula,
            "interest_segments": interest_segments, "line_items": line_items,
            "gross_total": round(subtotal, 2), "prior_payments_total": round(prior_total, 2), "net_total": round(total, 2),
            "disputed_deductions_excluded": round(disputed_deductions, 2),
            "moratory_reference_not_added": round(moratory_reference, 2),
            "moratory_reference_days": round(min(delay_days, 720), 2),
            "notes": notes,
            "assumptions": [
                "La convención 30/360 se utiliza para prestaciones proporcionales; la indemnización de término fijo usa días calendario faltantes.",
                "La remuneración variable y otros factores se incluyen únicamente en la cifra promedio suministrada y sujeta a soporte.",
                "La cifra moratoria es una referencia ilustrativa y no forma parte del total estimado.",
            ],
            "exclusions": [
                "Sanciones moratorias, indexación, intereses judiciales, costas y perjuicios no se suman automáticamente.",
                "Horas extra, recargos o comisiones no cuantificados y soportados quedan fuera del cálculo.",
                "Estabilidad reforzada, contrato realidad, fueros y regímenes colectivos requieren análisis especializado.",
                "Los descuentos controvertidos no se restan sin fundamento verificable.",
            ],
            "cutoff_date": cutoff.isoformat() if cutoff else end.isoformat(),
        }

    def _new_doc(self):
        doc = Document()
        self._set_doc_styles(doc)
        return doc

    def _header(self, doc, title, subtitle=None):
        self._title(doc, title)
        if subtitle:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle); r.italic = True; r.font.size = Pt(9)

    def _calc_rows(self, calc):
        return [[x["label"], x["formula"], self._money(x["gross"]), self._money(x["prior_paid"]), self._money(x["net"])] for x in calc["line_items"]]

    def _report(self, answers, evaluation, calc, target):
        doc = self._new_doc()
        self._header(doc, "INFORME TÉCNICO DE LIQUIDACIÓN LABORAL", "CO-LA-001 · Estimación por concepto, período y soporte")
        worker = self._party(answers, "worker", "Persona trabajadora")
        employer = self._party(answers, "employer", "Empleador")
        self._section(doc, "1. Identificación y alcance")
        self._paragraph(doc, f"Persona trabajadora: {worker['text']}.")
        self._paragraph(doc, f"Empleador: {employer['text']}.")
        self._paragraph(doc, f"Fecha de corte: {self._date_es(calc['cutoff_date'])}. Solicitante: {self._label(self._v(answers, 'claimant.role'))}.")
        self._paragraph(doc, "Este informe presenta una estimación técnica y trazable de acreencias laborales con base exclusiva en los datos y soportes declarados. No constituye sentencia, reconocimiento definitivo de deuda ni reemplaza la revisión profesional del expediente.")

        self._section(doc, "2. Cronología y parámetros")
        self._table(doc, ["Dato", "Valor"], [
            ["Inicio real", self._date_es(calc["periods"]["employment"]["start"])],
            ["Terminación o corte", self._date_es(calc["periods"]["employment"]["end"])],
            ["Modalidad contractual", self._label(self._v(answers, "relationship.contract_type"))],
            ["Situación al corte", self._label(self._v(answers, "relationship.termination_type"))],
            ["Días del vínculo (30/360)", calc["periods"]["employment"]["days_30_360"]],
            ["Parámetros", f"SMLMV 2026 {self._money(self.SMMLV_2026)}; auxilio de transporte {self._money(self.TRANSPORT_AID_2026)}; intereses a cesantías 12 % anual"],
        ])

        self._section(doc, "3. Base salarial")
        self._table(doc, ["Componente", "Valor mensual"], [
            ["Salario básico", self._money(calc["fixed_salary"])],
            ["Promedio variable", self._money(calc["variable_average"])],
            ["Otros factores salariales informados", self._money(calc["other_salary_average"])],
            ["Salario base del cálculo", self._money(calc["salary"])],
            ["Auxilio aplicado", self._money(calc["transport_aid_value"]) if calc["transport_aid_applied"] else "No aplicado"],
            ["Base prestacional", self._money(calc["benefit_base"])],
        ])

        self._section(doc, "4. Liquidación por concepto")
        self._table(doc, ["Concepto", "Fórmula", "Bruto", "Pago previo", "Saldo"], self._calc_rows(calc))
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"TOTAL NETO ESTIMADO: {self._money_words(calc['net_total'])}"); r.bold = True
        self._paragraph(doc, f"Total bruto matemático: {self._money(calc['gross_total'])}. Pagos previos imputados al concepto correspondiente: {self._money(calc['prior_payments_total'])}.")

        self._section(doc, "5. Intereses a las cesantías")
        if calc["interest_segments"]:
            self._table(doc, ["Año", "Desde", "Hasta", "Días", "Cesantías del segmento", "Interés"], [
                [x["year"], self._date_es(x["start"]), self._date_es(x["end"]), x["days"], self._money(x["cesantias_base"]), self._money(x["interest"])]
                for x in calc["interest_segments"]
            ])
        else:
            self._paragraph(doc, "No se registró un período liquidable de cesantías.")

        self._section(doc, "6. Indemnización y mora")
        self._paragraph(doc, f"Indemnización estándar: {calc['indemnity_days']:.4f} días; fórmula aplicada: {calc['indemnity_formula']}")
        self._paragraph(doc, f"Referencia moratoria no sumada al total: {self._money(calc['moratory_reference_not_added'])}, calculada únicamente como escenario ilustrativo sobre {calc['moratory_reference_days']:.0f} días. Su procedencia exige analizar deuda, oportunidad de pago y conducta de buena o mala fe.")

        self._section(doc, "7. Hallazgos, supuestos y exclusiones")
        for item in evaluation.get("findings", []):
            self._paragraph(doc, f"{item['id']} — {item['message']}")
        self._subparagraph(doc, "Notas de conciliación", " ".join(calc["notes"]) if calc["notes"] else "No se identificaron notas adicionales de conciliación matemática.")
        self._bullets(doc, calc["assumptions"])
        self._bullets(doc, calc["exclusions"])

        if self._v(answers, "compensation.salary_type") != "integral":
            doc.add_page_break()
        self._section(doc, "8. Conclusión y advertencia profesional")
        self._paragraph(doc, "La cifra neta estimada debe cotejarse con contrato, desprendibles, certificaciones, extractos, fondos de cesantías, carta de terminación, pagos y actuaciones previas. Cualquier cambio en fechas, salario, pagos, modalidad o soportes exige una nueva revisión inmutable y reinicia las aprobaciones.")
        self._signature_table(doc, worker, employer)
        doc.save(target)

    def _concepts_annex(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "ANEXO DE CONCEPTOS, PERÍODOS Y FÓRMULAS", "ANX-LA1-CONCEPTS-001")
        self._section(doc, "Matriz de liquidación")
        self._table(doc, ["Concepto", "Período/base", "Fórmula", "Bruto", "Pagado", "Saldo", "Fuentes"], [
            [x["label"], self._period_for_key(calc, x["key"]), x["formula"], self._money(x["gross"]), self._money(x["prior_paid"]), self._money(x["net"]), ", ".join(x["source_ids"])]
            for x in calc["line_items"]
        ])
        self._section(doc, "Parámetros aplicados")
        self._table(doc, ["Parámetro", "Valor", "Tratamiento"], [
            ["SMLMV 2026", self._money(self.SMMLV_2026), "Umbral para auxilio e indemnización del contrato indefinido"],
            ["Auxilio de transporte 2026", self._money(self.TRANSPORT_AID_2026), "Se incluye en cesantías y prima cuando fue confirmado y procede"],
            ["Interés a cesantías", "12 % anual", "Proporcional por segmentos anuales"],
            ["Convención prestacional", "30/360", "Aplicada a cesantías, prima y duración laboral estimada"],
            ["Vacaciones", "Salario ÷ 30 × días pendientes", "No incluye auxilio de transporte"],
        ])
        self._section(doc, "Trazabilidad")
        self._paragraph(doc, "Cada renglón separa valor bruto, pago previo confirmado y saldo. Un pago solo se imputa al concepto correspondiente; los excesos se llevan a saldo cero y se marcan para conciliación, sin trasladarlos automáticamente a otra partida.")
        doc.save(target)

    @staticmethod
    def _period_for_key(calc, key):
        p = calc["periods"]
        if key in {"cesantias", "cesantias_interest"}:
            return f"{p['cesantias']['start']} a {p['cesantias']['end']} · {p['cesantias']['days_30_360']} días"
        if key == "prima":
            return f"{p['prima']['start']} a {p['prima']['end']} · {p['prima']['days_30_360']} días"
        if key == "vacation":
            return f"{p['vacation']['pending_days']} días pendientes"
        return f"{p['employment']['start']} a {p['employment']['end']}"

    def _evidence_annex(self, answers, evaluation, target):
        doc = self._new_doc(); self._header(doc, "MATRIZ DE SOPORTES, PAGOS Y DIFERENCIAS", "ANX-LA1-EVIDENCE-001")
        evidence = self._as_dict(self._v(answers, "evidence.items"))
        labels = {"contract": "Contrato", "payroll": "Nómina", "bank": "Extractos o comprobantes", "cesantias": "Fondo de cesantías", "termination_letter": "Carta de terminación"}
        self._table(doc, ["Soporte", "Estado", "Observación"], [[labels[k], self._label(evidence.get(k)), "Debe cotejarse con fechas, bases, pagos y firmas"] for k in labels])
        self._section(doc, "Nivel global de soporte")
        self._paragraph(doc, f"Nivel informado: {self._label(self._v(answers, 'evidence.support_level'))}. Cálculo del empleador: {self._label(self._v(answers, 'evidence.employer_calculation'))}.")
        self._section(doc, "Hallazgos pendientes")
        pending = [x["message"] for x in evaluation.get("findings", []) if x.get("severity") in {"review", "blocker"}]
        self._bullets(doc, pending or ["No se reportaron hallazgos adicionales, sin perjuicio de la revisión de autenticidad e integridad de los soportes."])
        self._section(doc, "Pagos previos y descuentos")
        prior = self._as_dict(self._v(answers, "payments.prior")); deductions = self._as_dict(self._v(answers, "payments.deductions"))
        payment_labels = {
            "salary": "Salario",
            "cesantias": "Cesantías",
            "interests": "Intereses a las cesantías",
            "prima": "Prima de servicios",
            "vacation": "Vacaciones",
            "indemnity": "Indemnización por terminación",
        }
        self._table(doc, ["Concepto", "Valor informado"], [[payment_labels.get(k, k.replace("_", " ").capitalize()), self._money(v)] for k, v in prior.items()] or [["Pagos", "No informados"]])
        if deductions.get('exists') is True:
            self._paragraph(doc, f"Deducciones discutidas: {self._money(deductions.get('amount'))}; motivo: {self._plain(deductions.get('reason'), 'no informado')}; autorización o fundamento: {self._plain(deductions.get('authorization'), 'no informado')}.")
        else:
            self._paragraph(doc, "No se informaron deducciones controvertidas para este expediente.")
        doc.save(target)

    def _claim_document(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "RECLAMACIÓN DIRECTA DE ACREENCIAS LABORALES", "DOC-LA1-CLAIM-001")
        worker = self._party(answers, "worker", "Persona trabajadora"); employer = self._party(answers, "employer", "Empleador")
        for line in (self._date_es(calc["cutoff_date"]), "Señores", employer["name"], employer["email"]):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(1)
            p.add_run(line)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("Asunto: "); r.bold = True
        p.add_run("Reclamación directa, solicitud de pago y preservación de derechos.")
        self._paragraph(doc, f"Yo, {worker['name']}, identificado(a) con {worker['id']}, presento reclamación escrita respecto de las acreencias derivadas de la relación laboral desarrollada entre {self._date_es(calc['periods']['employment']['start'])} y {self._date_es(calc['periods']['employment']['end'])}.")
        self._section(doc, "HECHOS")
        facts = [
            f"La modalidad informada fue {self._label(self._v(answers, 'relationship.contract_type'))} y la situación de terminación o corte fue {self._label(self._v(answers, 'relationship.termination_type'))}.",
            f"La base salarial mensual estimada asciende a {self._money(calc['salary'])}, sujeta a conciliación con soportes.",
            f"Los pagos previos confirmados e imputados por concepto suman {self._money(calc['prior_payments_total'])}.",
            f"El saldo neto técnico estimado asciende a {self._money_words(calc['net_total'])}.",
        ]
        self._bullets(doc, facts)
        self._section(doc, "PRETENSIONES")
        claims = [
            "Entregar una liquidación detallada por concepto, período, base y pago previo.",
            f"Pagar o conciliar los saldos que resulten acreditados, cuyo valor técnico preliminar es {self._money(calc['net_total'])}.",
            "Remitir los soportes de nómina, consignaciones, seguridad social, cesantías y terminación que sustenten cualquier diferencia.",
        ]
        if self._v(answers, 'payments.deductions.exists') is True:
            claims.append("Explicar y soportar las deducciones aplicadas, absteniéndose de imputar valores no autorizados o legalmente procedentes.")
        claims.append("Dar respuesta escrita, completa y verificable por el canal contractual informado.")
        self._bullets(doc, claims)
        self._section(doc, "CUADRO DE VALORES")
        self._table(doc, ["Concepto", "Bruto", "Pago previo", "Saldo"], [[x["label"], self._money(x["gross"]), self._money(x["prior_paid"]), self._money(x["net"])] for x in calc["line_items"]])
        self._paragraph(doc, "La presente reclamación no implica renuncia, transacción ni aceptación de una cifra definitiva. Se formula con base en la información disponible y podrá precisarse al recibir los soportes faltantes.")
        for line in ("Atentamente,", worker["name"], f"{worker['id']} · {worker['email']}"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(0)
            p.add_run(line)
        doc.save(target)

    def _support_request(self, answers, target):
        doc = self._new_doc(); self._header(doc, "SOLICITUD DE SOPORTES LABORALES", "DOC-LA1-SUPPORT-REQUEST-001")
        worker = self._party(answers, "worker", "Persona trabajadora"); employer = self._party(answers, "employer", "Empleador")
        self._paragraph(doc, f"{worker['name']} solicita a {employer['name']} copia legible, íntegra y verificable de los documentos necesarios para conciliar la liquidación laboral.")
        self._bullets(doc, [
            "Contrato, otrosíes, anexos, certificaciones y comunicaciones sobre inicio o terminación.",
            "Desprendibles de nómina, comprobantes de pago, extractos o soportes de transferencias.",
            "Detalle de salario básico, comisiones, recargos, pagos en especie y demás factores salariales.",
            "Certificaciones y extractos de cesantías, intereses, prima y vacaciones.",
            "Planillas de seguridad social y soportes de novedades relevantes.",
            "Liquidación preparada por el empleador, bases, fórmulas, descuentos y autorizaciones.",
            "Carta de terminación, descargos, decisiones o soportes de la causa invocada, cuando existan.",
        ])
        self._paragraph(doc, "La solicitud se limita a información relacionada con la persona trabajadora y no exige revelar datos de terceros que puedan ser protegidos mediante supresión, anonimización o entrega parcial.")
        doc.save(target)

    def _prescription_report(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "INFORME DE EXIGIBILIDAD Y PRESCRIPCIÓN", "DOC-LA1-PRESCRIPTION-001")
        earliest = self._v(answers, "prescription.rights_dates.earliest"); latest = self._v(answers, "prescription.rights_dates.latest")
        claim = self._v(answers, "prescription.last_written_claim")
        self._table(doc, ["Hito", "Fecha", "Observación"], [
            ["Derecho más antiguo informado", self._date_es(earliest), "Debe identificarse por concepto y fecha de exigibilidad"],
            ["Derecho más reciente informado", self._date_es(latest), "No todos los conceptos se hacen exigibles al mismo tiempo"],
            ["Último reclamo escrito", self._date_es(claim) if claim else "No informado", "Debe acreditarse su recepción por el empleador"],
            ["Corte del análisis", self._date_es(calc["cutoff_date"]), "Fecha de referencia del expediente"],
        ])
        self._section(doc, "Metodología")
        self._paragraph(doc, "La prescripción debe estudiarse individualmente para salarios, prestaciones, vacaciones, indemnizaciones y demás derechos, a partir de su exigibilidad. Un reclamo escrito puede producir efectos interruptivos en las condiciones legales aplicables, por lo que deben conservarse texto, fecha, destinatario y prueba de recepción.")
        self._section(doc, "Alertas")
        self._bullets(doc, [
            "No debe utilizarse una sola fecha para todos los conceptos sin reconstruir su causación y exigibilidad.",
            "Una actuación conciliatoria, transacción o proceso puede modificar el análisis y debe cotejarse antes de reclamar.",
            "Este informe no declara prescrito un derecho; identifica riesgos temporales para revisión profesional inmediata.",
        ])
        doc.save(target)

    def _moratory_report(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "ANÁLISIS PRELIMINAR DE INDEMNIZACIÓN MORATORIA", "DOC-LA1-MORATORY-001")
        self._paragraph(doc, "La indemnización moratoria no se impone de manera automática por cualquier diferencia. Su análisis exige determinar la existencia de salarios o prestaciones exigibles, la oportunidad del pago y la conducta del empleador, incluyendo razones serias y atendibles que puedan acreditar buena fe.")
        self._table(doc, ["Elemento", "Información"], [
            ["Días de retardo informados", str(self._v(answers, "termination.delay_days") or 0)],
            ["Hechos sobre buena fe", self._plain(self._v(answers, "termination.good_faith_facts"), "No informados")],
            ["Referencia ilustrativa no sumada", self._money(calc["moratory_reference_not_added"])],
            ["Saldo técnico base", self._money(calc["net_total"])],
        ])
        self._section(doc, "Factores de evaluación")
        self._bullets(doc, [
            "Claridad y exigibilidad de la deuda al terminar el vínculo.",
            "Calidad de los soportes y consistencia de la liquidación del empleador.",
            "Existencia de controversias reales, pagos parciales, consignaciones o retenciones autorizadas.",
            "Conducta desplegada para verificar, pagar, conciliar o explicar las diferencias.",
            "Duración del retardo y régimen aplicable después de los primeros veinticuatro meses.",
        ])
        self._paragraph(doc, "La referencia numérica se presenta únicamente para dimensionar el riesgo y no se incorpora al total reclamado sin aprobación jurídica expresa.")
        doc.save(target)

    def _conciliation(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "PROPUESTA DE CONCILIACIÓN LABORAL", "DOC-LA1-CONCILIATION-001")
        worker = self._party(answers, "worker", "Persona trabajadora"); employer = self._party(answers, "employer", "Empleador")
        self._paragraph(doc, f"{worker['name']} y {employer['name']} manifiestan su intención de explorar una solución verificable respecto de los conceptos relacionados en el expediente, sin que este borrador constituya por sí solo conciliación, transacción o renuncia.")
        self._table(doc, ["Concepto", "Saldo técnico"], [[x["label"], self._money(x["net"])] for x in calc["line_items"]])
        self._paragraph(doc, f"Valor técnico de referencia: {self._money_words(calc['net_total'])}.")
        self._bullets(doc, [
            "Precisar qué conceptos y períodos son objeto del acuerdo.",
            "Distinguir derechos ciertos e indiscutibles de materias realmente controvertidas.",
            "Definir monto, oportunidad, impuestos, retenciones y forma de acreditación del pago.",
            "Someter el acuerdo al mecanismo y autoridad competente cuando ello sea necesario para su eficacia.",
            "Mantener a salvo los derechos no incluidos de forma expresa y válida.",
        ])
        doc.save(target)

    def _payment_agreement(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "ACUERDO DE PAGO DE ACREENCIAS LABORALES", "AGR-LA1-PAYMENT-001")
        worker = self._party(answers, "worker", "Persona trabajadora"); employer = self._party(answers, "employer", "Empleador")
        terms = self._as_dict(self._v(answers, "settlement.payment_terms"))
        installments = max(1, int(self._num(terms.get("installments")) or 1))
        initial = max(0.0, self._num(terms.get("initial_payment")))
        balance = max(0.0, calc["net_total"] - initial)
        installment_value = balance / installments if installments else balance
        self._paragraph(doc, f"Entre {worker['name']} y {employer['name']} se documenta un plan de pago sobre un valor de referencia de {self._money_words(calc['net_total'])}, sujeto a verificación y aprobación jurídica.")
        self._table(doc, ["Condición", "Valor"], [
            ["Pago inicial", self._money(initial)],
            ["Saldo financiado", self._money(balance)],
            ["Número de cuotas", str(installments)],
            ["Valor orientativo por cuota", self._money(installment_value)],
            ["Canal de pago", self._plain(terms.get("payment_channel"), "por definir antes de firma")],
        ])
        self._section(doc, "Cláusulas esenciales")
        clauses = [
            ("Objeto y reconocimiento", "El acuerdo identifica únicamente los conceptos y valores expresamente incluidos, sin transformar en renunciables los derechos ciertos e indiscutibles."),
            ("Forma de pago", "Los pagos deberán efectuarse en las fechas, canales y referencias verificables que se incorporen antes de firma."),
            ("Imputación", "Cada pago se imputará al concepto acordado y deberá generar constancia de saldo."),
            ("Incumplimiento", "El retardo o incumplimiento producirá las consecuencias pactadas y legalmente admisibles, sin multas desproporcionadas ni renuncias generales."),
            ("Paz y salvo limitado", "Cualquier paz y salvo operará solo respecto de conceptos efectivamente pagados y válidamente incluidos."),
            ("No renuncia", "El documento no implica renuncia a derechos mínimos ni a conceptos omitidos, desconocidos o no conciliables."),
            ("Revisión", "La versión final requiere aprobación jurídica y QA antes de liberación."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        self._signature_table(doc, worker, employer)
        doc.save(target)

    def _closure_act(self, answers, calc, target):
        doc = self._new_doc(); self._header(doc, "ACTA DE CUMPLIMIENTO Y CIERRE", "ACT-LA1-CLOSE-001")
        worker = self._party(answers, "worker", "Persona trabajadora"); employer = self._party(answers, "employer", "Empleador")
        self._paragraph(doc, f"Las partes dejan constancia del estado de ejecución del acuerdo relacionado con el expediente de {worker['name']} y {employer['name']}.")
        self._table(doc, ["Verificación", "Resultado"], [
            ["Valor técnico inicial", self._money(calc["net_total"])],
            ["Pagos acreditados", "Debe diligenciarse con comprobantes antes de firma"],
            ["Saldo final", "Debe calcularse con base en pagos efectivamente acreditados"],
            ["Soportes anexos", "Comprobantes, recibos y conciliación por concepto"],
        ])
        self._paragraph(doc, "La firma del acta no constituye paz y salvo general ni renuncia a derechos no incluidos, salvo el alcance específico y jurídicamente válido que resulte de un acuerdo previo cumplido y aprobado.")
        self._signature_table(doc, worker, employer)
        doc.save(target)

    def _annex(self, doc_id, answers, evaluation, calc, target):
        mapping = {
            "ANX-LA1-CONCEPTS-001": self._concepts_annex,
            "ANX-LA1-EVIDENCE-001": lambda a, c, t: self._evidence_annex(a, evaluation, t),
            "DOC-LA1-CLAIM-001": self._claim_document,
            "DOC-LA1-SUPPORT-REQUEST-001": lambda a, c, t: self._support_request(a, t),
            "DOC-LA1-PRESCRIPTION-001": self._prescription_report,
            "DOC-LA1-MORATORY-001": self._moratory_report,
            "DOC-LA1-CONCILIATION-001": self._conciliation,
            "AGR-LA1-PAYMENT-001": self._payment_agreement,
            "ACT-LA1-CLOSE-001": self._closure_act,
        }
        if doc_id not in mapping:
            raise ValueError(f"Documento no soportado: {doc_id}")
        mapping[doc_id](answers, calc, target)

    def render_documents(self, answers, target_folder):
        evaluation = self.evaluator.evaluate(answers)
        if evaluation.get("blocked"):
            messages = "; ".join(x["message"] for x in evaluation.get("blockers", []))
            raise ValueError("El expediente contiene bloqueos jurídicos: " + messages)
        if evaluation.get("missing_fields"):
            labels = ", ".join(x["label"] for x in evaluation["missing_fields"][:15])
            raise ValueError("Faltan datos esenciales: " + labels)
        calc = self.calculate(answers)
        target_folder = Path(target_folder)
        target_folder.mkdir(parents=True, exist_ok=False)
        report = target_folder / "CO-LA-001_Informe_Tecnico_Liquidacion_Laboral.docx"
        self._report(answers, evaluation, calc, report)
        generated = [{"id": "DOC-LA1-CALCULATION-001", "filename": report.name}]
        for doc_id in evaluation.get("documents", []):
            if doc_id == "DOC-LA1-CALCULATION-001":
                continue
            target = target_folder / f"{doc_id}.docx"
            self._annex(doc_id, answers, evaluation, calc, target)
            generated.append({"id": doc_id, "filename": target.name})
        hashes = {}
        unresolved = []
        for item in generated:
            path = target_folder / item["filename"]
            text = self._extract_text(path)
            if UNRESOLVED_PATTERN.search(text):
                unresolved.append(item["id"])
            hashes[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()
        if unresolved:
            raise ValueError("Se detectaron variables o valores centinela sin resolver: " + ", ".join(unresolved))
        (target_folder / "calculation.json").write_text(json.dumps(calc, ensure_ascii=False, indent=2), encoding="utf-8")
        hashes["calculation.json"] = hashlib.sha256((target_folder / "calculation.json").read_bytes()).hexdigest()
        return evaluation, generated, hashes, calc

    def generate(self, answers, actor=None):
        generation_id = "COLA001-" + uuid.uuid4().hex[:12].upper()
        folder = self.output_dir / generation_id
        folder.mkdir(parents=True, exist_ok=False)
        documents_dir = folder / "documents" / "revision-0001"
        evaluation, generated, hashes, calc = self.render_documents(answers, documents_dir)
        manifest = {
            "generation_id": generation_id,
            "product_id": "CO-LA-001",
            "version": self.VERSION,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "actor": {"id": (actor or {}).get("id"), "role": (actor or {}).get("role")},
            "status": "draft_generated",
            "workflow_status": "pending_legal_review",
            "current_revision": 1,
            "documents": generated,
            "document_folder": "documents/revision-0001",
            "hashes": hashes,
            "selected_blocks": evaluation.get("blocks", []),
            "professional_reviews": evaluation.get("professional_reviews", []),
            "calculation_summary": {
                "gross_total": calc["gross_total"],
                "gross_total_formatted": self._money(calc["gross_total"]),
                "prior_payments_total": calc["prior_payments_total"],
                "prior_payments_formatted": self._money(calc["prior_payments_total"]),
                "net_total": calc["net_total"],
                "net_total_formatted": self._money(calc["net_total"]),
                "moratory_reference_not_added": calc["moratory_reference_not_added"],
                "moratory_reference_formatted": self._money(calc["moratory_reference_not_added"]),
            },
            "unresolved_variables": 0,
            "legal_approval": {"status": "pending"},
            "qa_approval": {"status": "pending"},
            "released": False,
        }
        (folder / "answers.json").write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self._package(folder, self.output_dir / f"{generation_id}.zip")
        manifest["package_filename"] = package.name
        manifest["package_sha256"] = hashlib.sha256(package.read_bytes()).hexdigest()
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return manifest

    @staticmethod
    def _package(folder, target):
        with ZipFile(target, "w", ZIP_DEFLATED) as zf:
            for path in sorted(Path(folder).rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        return target

    def package_path(self, generation_id):
        if not GEN_RE.fullmatch(generation_id or ""):
            return None
        path = self.output_dir / f"{generation_id}.zip"
        return path if path.is_file() else None
