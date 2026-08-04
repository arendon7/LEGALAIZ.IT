from __future__ import annotations

from pathlib import Path
from typing import Any, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from co_tr_001_document_factory_v258 import (
    CoTr001DocumentFactoryV258,
    DocumentGenerationError,
    GEN_RE,
    SENTINEL_RE,
)


class CoTr001DocumentFactoryV259(CoTr001DocumentFactoryV258):
    """Fábrica documental cerrada de CO-TR-001 v2.59."""

    VERSION = "2.59"
    PRODUCT_ID = "CO-TR-001"
    PACKAGE_NAME = "00_PAQUETE_DOCUMENTAL_CO-TR-001_V259.docx"

    def __init__(self, root: Path, evaluator):
        super().__init__(root, evaluator)
        self.document_specs = self._load_json("DOCUMENTOS_V259.json")
        self.blocks = self._load_json("BLOQUES_V259.json")
        self.sources = self._load_json("FUENTES_V259.json")
        self.source_verification = self._load_json("SOURCE_VERIFICATION_V259.json")

    def _configure_document(self, doc: Document, generation_id: str, package: bool = False) -> None:
        super()._configure_document(doc, generation_id, package=package)
        # La plantilla heredada de v2.58 conserva un rótulo literal en la portada.
        # Se corrige en cada ejecución sin alterar la base ni el resto del formato.
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "Versión 2.58" in run.text:
                    run.text = run.text.replace("Versión 2.58", "Versión 2.59")

    @staticmethod
    def _prevent_row_split(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    def _render_content(self, doc: Document, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any]) -> None:
        super()._render_content(doc, spec, answers, evaluation)
        # Evita que una misma fila continúe en otra página, especialmente en la
        # tabla de fuentes con URLs extensas. Los encabezados siguen repitiéndose.
        for table in doc.tables:
            for row in table.rows:
                self._prevent_row_split(row)

    def render_documents(
        self,
        answers: dict[str, Any],
        target_dir: Path,
        generation_id: str,
        mode: str = "precheck",
    ) -> Tuple[dict[str, Any], list[str], dict[str, str]]:
        evaluation, documents, hashes = super().render_documents(
            answers,
            target_dir,
            generation_id,
            mode=mode,
        )
        package_path = Path(target_dir) / self.PACKAGE_NAME
        package = Document(package_path)
        package.core_properties.title = "Paquete documental CO-TR-001 v2.59"
        package.core_properties.subject = "CO-TR-001 - cierre controlado"
        package.core_properties.author = "LegalAIZ.it"
        package.core_properties.comments = (
            "Paquete documental controlado, validado para la Macrofase C. "
            "No constituye decisión administrativa ni sustituye revisión profesional."
        )
        package.save(package_path)
        self._assert_clean_docx(package_path)
        hashes[self.PACKAGE_NAME] = self._hash_file(package_path)
        return evaluation, documents, hashes
