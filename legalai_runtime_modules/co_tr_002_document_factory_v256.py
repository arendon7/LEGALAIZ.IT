from __future__ import annotations

import hashlib
import json
import re
import secrets
from datetime import date, datetime, timezone
from pathlib import Path
import os
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SENTINEL_RE = re.compile(
    r"(?:\[\[|\]\]|\{\{|\}\}|\bNULL\b|\bundefined\b|\bN/?A\b|<[^>]+_PENDIENTE>)",
    re.IGNORECASE,
)
GEN_RE = re.compile(r"COTR002-[A-F0-9]{12}")



def _runtime_root(project_root: Path) -> Path:
    raw = os.environ.get("LEGAL_RUNTIME_DIR", "").strip()
    path = Path(raw).expanduser() if raw else Path(project_root) / "runtime"
    if not path.is_absolute():
        path = Path(project_root) / path
    return path.resolve()

class DocumentGenerationError(ValueError):
    pass


class CoTr002DocumentFactoryV256:
    VERSION = "2.56"
    PRODUCT_ID = "CO-TR-002"

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.product_dir = self.root / "app" / "assets" / "advanced-legal-library" / self.PRODUCT_ID
        self.output_dir = _runtime_root(self.root) / "generated" / "co-tr-002"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.document_specs = self._load_json("DOCUMENTOS_V256.json")
        self.blocks = self._load_json("BLOQUES_V256.json")
        self.sources = self._load_json("FUENTES_V256.json")
        self.source_verification = self._load_json("SOURCE_VERIFICATION_V256.json")

    def _load_json(self, name: str) -> dict[str, Any]:
        path = self.product_dir / name
        if not path.is_file():
            raise FileNotFoundError(f"No se encontró el activo requerido: {path}")
        return json.loads(path.read_text(encoding="utf-8"))

    @staticmethod
    def _now() -> str:
        return datetime.now(timezone.utc).isoformat()

    @staticmethod
    def _hash_file(path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    @staticmethod
    def _hash_obj(value: Any) -> str:
        raw = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        return hashlib.sha256(raw).hexdigest()

    @staticmethod
    def get_value(data: dict[str, Any], path: str, default: Any = None) -> Any:
        value: Any = data
        for part in path.split("."):
            if not isinstance(value, dict) or part not in value:
                return default
            value = value[part]
        return value

    @staticmethod
    def _human(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "Sí" if value else "No"
        if isinstance(value, (list, tuple)):
            return "; ".join(CoTr002DocumentFactoryV256._human(item) for item in value if item not in (None, ""))
        return str(value).strip()

    @staticmethod
    def _safe_filename(value: str) -> str:
        value = re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("._")
        return value or "documento"

    @staticmethod
    def _add_page_field(paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Página ")
        run.font.size = Pt(8)
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    @staticmethod
    def _mark_table_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)

    def _configure_document(self, doc: Document, title: str, generation_id: str, draft: bool = True) -> None:
        props = doc.core_properties
        props.title = title
        props.subject = f"{self.PRODUCT_ID} - documento jurídico controlado"
        props.author = "LegalAIZ.it"
        props.last_modified_by = "LegalAIZ.it"
        props.keywords = f"LegalAIZ.it, {self.PRODUCT_ID}, Colombia, tránsito, SAST"
        props.comments = "Generado por la Fábrica Documental LegalAIZ.it. Requiere revisión jurídica y QA."
        section = doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.12

        for name, size, color in (
            ("Title", 17, RGBColor(13, 19, 36)),
            ("Heading 1", 13, RGBColor(13, 19, 36)),
            ("Heading 2", 11.5, RGBColor(37, 99, 235)),
        ):
            style = styles[name]
            style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
            style.font.size = Pt(size)
            style.font.color.rgb = color

        if "LegalAIZ Notice" not in styles:
            notice = styles.add_style("LegalAIZ Notice", WD_STYLE_TYPE.PARAGRAPH)
            notice.font.name = "Aptos"
            notice.font.size = Pt(9)
            notice.font.color.rgb = RGBColor(70, 70, 70)
            notice.paragraph_format.space_before = Pt(4)
            notice.paragraph_format.space_after = Pt(8)

        header = section.header
        p = header.paragraphs[0]
        p.text = f"LegalAIZ.it · {self.PRODUCT_ID} · v{self.VERSION} · {generation_id}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Documento generado con tecnología. Requiere validación de hechos, anexos y estrategia antes de radicación."
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor(100, 100, 100)
        self._add_page_field(footer.add_paragraph())

        title_p = doc.add_paragraph(style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.add_run(title.upper()).bold = True
        meta = doc.add_paragraph(style="LegalAIZ Notice")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Producto {self.PRODUCT_ID} · Versión {self.VERSION} · Colombia")
        if draft:
            draft_p = doc.add_paragraph(style="LegalAIZ Notice")
            draft_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = draft_p.add_run("BORRADOR CONTROLADO — NO RADICAR SIN REVISIÓN")
            run.bold = True
            run.font.color.rgb = RGBColor(160, 70, 20)

    def _add_control_notice(self, doc: Document, evaluation: dict[str, Any]) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        self._mark_table_header(table.rows[0])
        cell = table.cell(0, 0)
        self._shade_cell(cell, "F7F5F1")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.add_run("CONTROL PREVIO. ").bold = True
        p.add_run(
            "Este producto no promete anulación, devolución ni archivo automático. "
            "La consecuencia de una falla de notificación depende de la etapa, la prueba disponible y la afectación efectiva del derecho de defensa."
        )
        if evaluation.get("review_requirements"):
            p2 = cell.add_paragraph()
            p2.add_run("Revisión profesional requerida: ").bold = True
            p2.add_run("; ".join(evaluation["review_requirements"]))
        if evaluation.get("release_blockers"):
            p3 = cell.add_paragraph()
            p3.add_run("Bloqueos de publicación: ").bold = True
            p3.add_run("; ".join(evaluation["release_blockers"]))

    def _add_address(self, doc: Document, answers: dict[str, Any]) -> None:
        city = self._human(self.get_value(answers, "filing.city") or self.get_value(answers, "identity.city"))
        filing_date = self._human(self.get_value(answers, "filing.date") or date.today().isoformat())
        p = doc.add_paragraph()
        p.add_run(f"{city}, {filing_date}").bold = True
        doc.add_paragraph()
        authority = self._human(self.get_value(answers, "authority.name"))
        ap = doc.add_paragraph()
        ap.add_run("Señores\n").bold = True
        ap.add_run(authority.upper()).bold = True
        authority_city = self._human(self.get_value(answers, "authority.city"))
        if authority_city:
            ap.add_run(f"\n{authority_city}")
        authority_email = self._human(self.get_value(answers, "authority.email"))
        if authority_email:
            ap.add_run(f"\nCorreo: {authority_email}")

    def _add_subject(self, doc: Document, title: str, answers: dict[str, Any]) -> None:
        p = doc.add_paragraph()
        p.add_run("ASUNTO: ").bold = True
        p.add_run(title)
        comparendo = self._human(self.get_value(answers, "infraction.comparendo_number"))
        plate = self._human(self.get_value(answers, "infraction.plate"))
        p.add_run(f" — Comparendo {comparendo} — Placa {plate}")

    def _facts(self, answers: dict[str, Any]) -> list[str]:
        facts = [
            f"La persona solicitante es {self._human(self.get_value(answers, 'identity.full_name'))}, identificada con {self._human(self.get_value(answers, 'identity.document_type') or 'documento')} número {self._human(self.get_value(answers, 'identity.document_number'))}.",
            f"Se registra la orden de comparendo número {self._human(self.get_value(answers, 'infraction.comparendo_number'))}, asociada a la placa {self._human(self.get_value(answers, 'infraction.plate'))} y a hechos del {self._human(self.get_value(answers, 'infraction.date'))}.",
            f"El estado informado de notificación es: {self._human(self.get_value(answers, 'infraction.notice_status'))}.",
        ]
        validation_date = self._human(self.get_value(answers, "infraction.validation_date"))
        if validation_date:
            facts.append(f"La fecha de validación reportada o conocida es {validation_date}.")
        notice_date = self._human(self.get_value(answers, "infraction.notice_date"))
        if notice_date:
            facts.append(f"La fecha de notificación o entrega informada es {notice_date}.")
        address_used = self._human(self.get_value(answers, "infraction.address_used"))
        if address_used:
            facts.append(f"La dirección utilizada por la autoridad fue: {address_used}.")
        narrative = self._human(self.get_value(answers, "case.narrative"))
        if narrative:
            facts.append(narrative)
        return facts

    def _add_numbered(self, doc: Document, heading: str, items: Iterable[str]) -> None:
        doc.add_heading(heading, level=1)
        for index, item in enumerate(items, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.add_run(f"{index}. ").bold = True
            p.add_run(str(item))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _evidence(self, answers: dict[str, Any]) -> list[str]:
        evidence = self.get_value(answers, "evidence.items", []) or []
        if not isinstance(evidence, list):
            evidence = [evidence]
        clean = [self._human(item) for item in evidence if self._human(item)]
        defaults = [
            "Copia de la orden de comparendo o consulta disponible.",
            "Consulta o captura del estado en SIMIT/RUNT, cuando corresponda.",
            "Soportes de dirección y datos registrados para notificación.",
        ]
        return list(dict.fromkeys(clean + defaults))

    def _legal_basis(self) -> list[str]:
        return [
            "Ley 1843 de 2017, especialmente su regulación del procedimiento y la notificación en detecciones mediante ayudas tecnológicas.",
            "Ley 1437 de 2011 y Ley 1755 de 2015, en cuanto a petición, acceso al expediente, notificaciones y revocación directa.",
            "Código Nacional de Tránsito y reglamentación compilatoria aplicable, sujetos a verificación de vigencia para la fecha de radicación.",
            "Jurisprudencia constitucional sobre responsabilidad personal, culpabilidad y debido proceso administrativo sancionador.",
        ]

    def _requests_for(self, doc_id: str, answers: dict[str, Any]) -> list[str]:
        common_record = [
            "Entregar copia íntegra, legible y cronológica del expediente administrativo, incluidos comparendo, soportes técnicos, validación, guías, constancias, avisos, actos y recursos.",
            "Informar las fechas, canales, direcciones y resultados de cada intento de notificación, con sus soportes verificables.",
            "Indicar el estado actual del procedimiento y los mecanismos disponibles para ejercer defensa y contradicción.",
        ]
        mapping = {
            "traffic_record_request": common_record,
            "traffic_notice_claim": [
                "Verificar integralmente la regularidad de la validación y de la notificación de la orden de comparendo.",
                "Reconocer, si se acredita falta o indebida notificación, los efectos procedentes respecto de las oportunidades de comparecencia y defensa, sin prejuzgar sobre la responsabilidad.",
                "Abstenerse de tratar la mera condición de propietario como prueba automática de responsabilidad sancionatoria.",
            ] + common_record,
            "traffic_hearing_request": [
                "Programar o habilitar la actuación de audiencia que jurídicamente corresponda para ejercer defensa.",
                "Decretar, incorporar y permitir controvertir las pruebas solicitadas y las que integren el expediente.",
                "Comunicar oportunamente fecha, hora, modalidad y canal de acceso a la actuación.",
            ],
            "traffic_revocation_request": [
                "Examinar de manera expresa y motivada la procedencia de la revocación directa frente al acto identificado.",
                "Valorar la regularidad de la vinculación, notificación, prueba, imputación personal y oportunidad real de defensa.",
                "Adoptar la decisión que corresponda y comunicarla con indicación clara de sus efectos y vías posteriores.",
            ],
            "traffic_registry_correction": [
                "Verificar la correspondencia entre el expediente, el acto vigente y la información reportada en SIMIT y/o RUNT.",
                "Corregir, actualizar o retirar el registro cuando exista soporte administrativo suficiente para ello.",
                "Informar la fecha de envío, aceptación y reflejo de la novedad en cada sistema.",
            ],
            "traffic_reiteration": [
                "Responder de fondo, de manera completa y congruente, la solicitud previamente radicada.",
                "Entregar los documentos e información pendientes o explicar de forma motivada la limitación legal aplicable.",
                "Informar el funcionario responsable y el estado interno del trámite.",
            ],
        }
        return mapping.get(doc_id, common_record)

    def _add_signature(self, doc: Document, answers: dict[str, Any]) -> None:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Atentamente,").bold = True
        doc.add_paragraph()
        name = self._human(self.get_value(answers, "identity.full_name"))
        dtype = self._human(self.get_value(answers, "identity.document_type") or "C.C.")
        number = self._human(self.get_value(answers, "identity.document_number"))
        email = self._human(self.get_value(answers, "identity.email"))
        address = self._human(self.get_value(answers, "identity.address"))
        sp = doc.add_paragraph()
        sp.add_run(name.upper()).bold = True
        sp.add_run(f"\n{dtype} {number}")
        sp.add_run(f"\nCorreo: {email}")
        if address:
            sp.add_run(f"\nDirección: {address}")

    def _add_annexes(self, doc: Document, answers: dict[str, Any]) -> None:
        self._add_numbered(doc, "ANEXOS", self._evidence(answers))

    def _add_sources(self, doc: Document) -> None:
        doc.add_heading("REFERENCIAS JURÍDICAS DE CONTROL", level=1)
        for item in self.sources.get("sources", []):
            p = doc.add_paragraph(style="LegalAIZ Notice")
            p.add_run(f"{item['id']}: ").bold = True
            p.add_run(item["title"])
            p.add_run(f" — Estado de control: {str(item.get('status', 'verificar')).replace('_', ' ')}.")
        p = doc.add_paragraph(style="LegalAIZ Notice")
        p.add_run("Advertencia: ").bold = True
        verified_at = self.source_verification.get("verified_at", "sin fecha")
        p.add_run(f"Verificación interna de fuentes: {verified_at}. La vigencia, modificación y aplicabilidad debe revalidarse antes de radicar.")

    def _render_standard(self, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any], generation_id: str, target: Path) -> None:
        doc = Document()
        self._configure_document(doc, spec["title"], generation_id)
        self._add_control_notice(doc, evaluation)
        self._add_address(doc, answers)
        self._add_subject(doc, spec["subject"], answers)
        self._add_numbered(doc, "HECHOS", self._facts(answers))
        self._add_numbered(doc, "SOLICITUDES", self._requests_for(spec["id"], answers))
        self._add_numbered(doc, "PRUEBAS Y DOCUMENTOS SOLICITADOS", self._evidence(answers))
        self._add_numbered(doc, "FUNDAMENTOS DE CONTROL", self._legal_basis())
        self._add_annexes(doc, answers)
        self._add_sources(doc)
        self._add_signature(doc, answers)
        self._save_checked(doc, target)

    def _render_escalation(self, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any], generation_id: str, target: Path) -> None:
        doc = Document()
        self._configure_document(doc, spec["title"], generation_id)
        self._add_control_notice(doc, evaluation)
        doc.add_heading("1. Estado del caso", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Campo"
        table.rows[0].cells[1].text = "Resultado"
        self._mark_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            self._shade_cell(cell, "0D1324")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
        rows = [
            ("Estado", evaluation.get("status")),
            ("Riesgo", evaluation.get("risk")),
            ("Etapa", self.get_value(answers, "procedure.stage")),
            ("Comparendo", self.get_value(answers, "infraction.comparendo_number")),
            ("Acto sancionatorio", self.get_value(answers, "procedure.sanction_resolution") or "No informado"),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = self._human(label)
            cells[1].text = self._human(value)

        self._add_numbered(
            doc,
            "2. Alertas que impiden una actuación automática definitiva",
            evaluation.get("release_blockers") or ["No se identificaron bloqueos automáticos, pero debe verificarse el expediente."],
        )
        self._add_numbered(
            doc,
            "3. Materias para revisión profesional",
            evaluation.get("review_requirements") or ["Validación final de hechos, pruebas, competencia y términos."],
        )
        self._add_numbered(
            doc,
            "4. Ruta sugerida",
            [
                "Conservar copia íntegra de la consulta, radicación y anexos.",
                "Obtener el expediente completo antes de definir una teoría definitiva del caso.",
                "Revisar términos, recursos, actos sancionatorios, cobro coactivo y eventuales medios de control.",
                "No efectuar pagos, acuerdos, desistimientos o reconocimientos sin comprender sus efectos jurídicos.",
                "Escalar inmediatamente a profesional cuando exista embargo, mandamiento de pago, proceso judicial o término próximo a vencer.",
            ],
        )
        self._add_sources(doc)
        self._save_checked(doc, target)

    def _render_traceability(self, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any], generation_id: str, target: Path) -> None:
        doc = Document()
        self._configure_document(doc, spec["title"], generation_id)
        self._add_control_notice(doc, evaluation)
        rows = [
            ("Identificador", generation_id),
            ("Producto", self.PRODUCT_ID),
            ("Versión", self.VERSION),
            ("Estado", evaluation.get("status")),
            ("Riesgo", evaluation.get("risk")),
            ("Revisión profesional", evaluation.get("professional_review_required")),
            ("Documentos seleccionados", evaluation.get("documents")),
            ("Bloques seleccionados", evaluation.get("blocks")),
            ("Bloqueos de publicación", evaluation.get("release_blockers")),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Elemento"
        table.rows[0].cells[1].text = "Valor"
        self._mark_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            self._shade_cell(cell, "0D1324")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = self._human(value)
        self._add_sources(doc)
        self._save_checked(doc, target)

    def _render_one(self, spec: dict[str, Any], answers: dict[str, Any], evaluation: dict[str, Any], generation_id: str, target: Path) -> None:
        kind = spec.get("kind", "standard")
        if kind == "escalation":
            self._render_escalation(spec, answers, evaluation, generation_id, target)
        elif kind == "traceability":
            self._render_traceability(spec, answers, evaluation, generation_id, target)
        else:
            self._render_standard(spec, answers, evaluation, generation_id, target)

    def _save_checked(self, doc: Document, target: Path) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        temp = target.with_suffix(".docx.tmp")
        doc.save(temp)
        data = temp.read_bytes()
        if not data.startswith(b"PK"):
            temp.unlink(missing_ok=True)
            raise DocumentGenerationError("El archivo DOCX generado no es válido.")
        try:
            with ZipFile(temp) as zf:
                if "word/document.xml" not in zf.namelist() or zf.testzip() is not None:
                    raise DocumentGenerationError("La estructura OOXML del documento es inválida.")
        except Exception as exc:
            temp.unlink(missing_ok=True)
            if isinstance(exc, DocumentGenerationError):
                raise
            raise DocumentGenerationError(f"No fue posible validar la estructura DOCX: {exc}") from exc
        temp.replace(target)
        self.assert_no_sentinels(target)

    def assert_no_sentinels(self, docx_path: Path) -> None:
        with ZipFile(docx_path) as zf:
            texts = []
            for name in zf.namelist():
                if name.endswith(".xml"):
                    texts.append(zf.read(name).decode("utf-8", errors="ignore"))
        joined = "\n".join(texts)
        match = SENTINEL_RE.search(joined)
        if match:
            raise DocumentGenerationError(f"Se detectó marcador o valor centinela en {docx_path.name}: {match.group(0)}")

    def _validate_document_inputs(self, spec: dict[str, Any], answers: dict[str, Any]) -> list[str]:
        missing = []
        for path in spec.get("required", []):
            if self.get_value(answers, path) in (None, "", []):
                missing.append(path)
        return missing

    def render_documents(self, answers: dict[str, Any], target_dir: Path, generation_id: str | None = None):
        evaluation = self.evaluator.evaluate(answers)
        document_answers = self.evaluator.to_document_answers(answers) if hasattr(self.evaluator, "to_document_answers") else answers
        if evaluation.get("missing_fields"):
            paths = ", ".join(item.get("path", "") for item in evaluation["missing_fields"])
            raise DocumentGenerationError(f"Faltan datos esenciales: {paths}")
        generation_id = generation_id or f"COTR002-{secrets.token_hex(6).upper()}"
        if not GEN_RE.fullmatch(generation_id):
            raise DocumentGenerationError("Identificador de generación inválido.")
        target_dir = Path(target_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        selected = set(evaluation.get("documents") or [])
        specs = [spec for spec in self.document_specs["documents"] if spec["id"] in selected]
        missing_by_doc: dict[str, list[str]] = {}
        for spec in specs:
            missing = self._validate_document_inputs(spec, document_answers)
            if missing:
                missing_by_doc[spec["id"]] = missing
        if missing_by_doc:
            raise DocumentGenerationError(f"Variables incompletas por documento: {missing_by_doc}")

        documents: list[str] = []
        hashes: dict[str, str] = {}
        for index, spec in enumerate(specs, start=1):
            filename = f"{index:02d}_{self._safe_filename(spec['filename'])}.docx"
            path = target_dir / filename
            self._render_one(spec, document_answers, evaluation, generation_id, path)
            documents.append(filename)
            hashes[filename] = self._hash_file(path)

        consolidated = target_dir / "00_PAQUETE_DOCUMENTAL_CO-TR-002_V256.docx"
        self._render_consolidated(specs, document_answers, evaluation, generation_id, consolidated)
        documents.insert(0, consolidated.name)
        hashes[consolidated.name] = self._hash_file(consolidated)
        return evaluation, documents, hashes

    def _render_consolidated(self, specs: list[dict[str, Any]], answers: dict[str, Any], evaluation: dict[str, Any], generation_id: str, target: Path) -> None:
        doc = Document()
        self._configure_document(doc, "Paquete documental — Fotomulta no notificada", generation_id)
        self._add_control_notice(doc, evaluation)
        doc.add_heading("ÍNDICE DE ENTREGABLES", level=1)
        for idx, spec in enumerate(specs, start=1):
            doc.add_paragraph(f"{idx}. {spec['title']}")
        doc.add_page_break()
        for idx, spec in enumerate(specs, start=1):
            if idx > 1:
                doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_heading(spec["title"].upper(), level=1)
            if spec.get("kind") == "traceability":
                rows = [
                    ("Identificador", generation_id),
                    ("Estado", evaluation.get("status")),
                    ("Riesgo", evaluation.get("risk")),
                    ("Documentos", evaluation.get("documents")),
                    ("Bloques", evaluation.get("blocks")),
                ]
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = "Elemento"
                table.rows[0].cells[1].text = "Valor"
                self._mark_table_header(table.rows[0])
                for cell in table.rows[0].cells:
                    self._shade_cell(cell, "0D1324")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                for label, value in rows:
                    cells = table.add_row().cells
                    cells[0].text = label
                    cells[1].text = self._human(value)
            elif spec.get("kind") == "escalation":
                self._add_numbered(doc, "Alertas", evaluation.get("release_blockers") or ["Validación profesional previa a la radicación."])
                self._add_numbered(doc, "Ruta", [
                    "Obtener expediente completo.",
                    "Controlar términos y etapa procesal.",
                    "Escalar a especialista ante sanción, cobro, embargo o proceso judicial.",
                ])
            else:
                self._add_address(doc, answers)
                self._add_subject(doc, spec["subject"], answers)
                self._add_numbered(doc, "HECHOS", self._facts(answers))
                self._add_numbered(doc, "SOLICITUDES", self._requests_for(spec["id"], answers))
                self._add_numbered(doc, "FUNDAMENTOS DE CONTROL", self._legal_basis())
                self._add_annexes(doc, answers)
                self._add_signature(doc, answers)
        self._add_sources(doc)
        self._save_checked(doc, target)

    def generate(self, answers: dict[str, Any], actor: dict[str, Any] | None = None) -> dict[str, Any]:
        generation_id = f"COTR002-{secrets.token_hex(6).upper()}"
        folder = self.output_dir / generation_id
        documents_dir = folder / "documents" / "revision-0001"
        evaluation, documents, hashes = self.render_documents(answers, documents_dir, generation_id=generation_id)
        package = self.output_dir / f"{generation_id}.zip"
        manifest = {
            "generation_id": generation_id,
            "product_id": self.PRODUCT_ID,
            "version": self.VERSION,
            "created_at": self._now(),
            "created_by": actor or {"id": "system", "role": "system"},
            "status": evaluation.get("status"),
            "risk": evaluation.get("risk"),
            "selected_blocks": evaluation.get("blocks", []),
            "documents": documents,
            "hashes": hashes,
            "document_folder": "documents/revision-0001",
            "review_requirements": evaluation.get("review_requirements", []),
            "release_blockers": evaluation.get("release_blockers", []),
            "answers_hash": self._hash_obj(answers),
            "evaluation_hash": self._hash_obj(evaluation),
            "source_verification_date": self.source_verification.get("verified_at"),
            "source_verification_status": self.source_verification.get("status"),
            "workflow_status": "draft_unregistered",
            "released": False,
            "package_filename": package.name,
            "package_sha256": None,
        }
        folder.mkdir(parents=True, exist_ok=True)
        manifest_path = folder / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self._build_package(folder, package)
        manifest["package_sha256"] = self._hash_file(package)
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return {
            "generation_id": generation_id,
            "folder": str(folder),
            "package": str(package),
            "manifest": manifest,
            "evaluation": evaluation,
        }

    @staticmethod
    def _build_package(folder: Path, target: Path) -> Path:
        temp = target.with_suffix(".zip.tmp")
        with ZipFile(temp, "w", ZIP_DEFLATED) as zf:
            for path in sorted(folder.rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        temp.replace(target)
        return target
