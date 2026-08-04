from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from co_em_003_document_factory_v243 import CoEm003DocumentFactoryV243, UNRESOLVED_PATTERN


GEN_RE = re.compile(r"COEM004-[A-F0-9]{12}")


class CoEm004DocumentFactoryV246(CoEm003DocumentFactoryV243):
    """Fábrica DOCX profunda para confidencialidad, secretos, datos y PI."""

    VERSION = "2.46"

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.output_dir = self.root / "data" / "generated" / "co-em-004-v246"
        self.output_dir.mkdir(parents=True, exist_ok=True)


    @staticmethod
    def _set_doc_styles(doc: Document):
        """Estilos propios de CO-EM-004; evita heredar el pie de CO-EM-003."""
        sec = doc.sections[0]
        sec.top_margin = Cm(2.3)
        sec.bottom_margin = Cm(2.1)
        sec.left_margin = Cm(2.6)
        sec.right_margin = Cm(2.6)
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.2)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for style_name in ("Title", "Heading 1", "Heading 2"):
            style = doc.styles[style_name]
            style.font.name = "Aptos Display"
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Heading 1"].font.size = Pt(12)
        doc.styles["Heading 2"].font.size = Pt(11)
        footer = sec.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("LegalAIZ.it | CO-EM-004 | Documento sujeto a revisión jurídica y QA | Página ")
        run.font.size = Pt(8)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer._p.append(fld)

    @staticmethod
    def _table(doc, headers, rows, widths=None):
        """Tabla jurídica estable: encabezado repetible, filas indivisibles y tipografía compacta."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr = table.rows[0]
        trPr = hdr._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        trPr.append(tblHeader)
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        for cell, label in zip(hdr.cells, headers):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(label))
            r.bold = True
            r.font.size = Pt(8.2)
        for row in rows:
            row_obj = table.add_row()
            row_pr = row_obj._tr.get_or_add_trPr()
            row_pr.append(OxmlElement("w:cantSplit"))
            for cell, value in zip(row_obj.cells, row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(CoEm003DocumentFactoryV243._plain(value))
                r.font.size = Pt(8.2)
        return table

    @staticmethod
    def _fragment(value, fallback="", lower_initial=False):
        """Normaliza texto dinámico para insertarlo dentro de una oración."""
        text = CoEm003DocumentFactoryV243._plain(value, fallback).strip()
        text = re.sub(r"[\s\.;,:]+$", "", text)
        if lower_initial and text and text[0].isupper() and not text[:2].isupper():
            text = text[0].lower() + text[1:]
        return text

    @staticmethod
    def _role_label(value):
        labels = {
            "responsible": "responsable del tratamiento",
            "controller": "responsable del tratamiento",
            "processor": "encargado del tratamiento",
            "joint_controller": "corresponsable del tratamiento",
            "recipient": "receptor autorizado",
            "independent": "responsable independiente",
            "independent_controller": "responsable independiente",
            "none": "sin rol de tratamiento aplicable",
            "to_define": "rol pendiente de definición",
        }
        return labels.get(str(value or "").strip(), CoEm004DocumentFactoryV246._fragment(value, "rol definido en el anexo"))

    @staticmethod
    def _allocation_label(value):
        labels = {
            "none": "la ausencia de asignación o transferencia de derechos sobre resultados",
            "license": "una licencia delimitada",
            "assignment": "una cesión patrimonial delimitada",
            "shared": "un esquema de titularidad compartida definido por activo y contribución",
            "joint_ownership": "un esquema de cotitularidad definido por activo",
            "party_a": "la titularidad de la Parte A en el alcance expresamente pactado",
            "party_b": "la titularidad de la Parte B en el alcance expresamente pactado",
            "case_by_case": "una definición individual por activo o entregable",
            "retained_by_creator": "la conservación de titularidad por su creador, con la licencia pactada",
            "to_define": "una distribución pendiente de definición antes de la aprobación final",
        }
        return labels.get(str(value or "").strip(), CoEm004DocumentFactoryV246._fragment(value, "la distribución descrita en el anexo de propiedad intelectual"))

    @staticmethod
    def _level_label(value):
        labels = {
            "critical": "crítico",
            "enhanced": "reforzado",
            "high": "alto",
            "standard": "estándar",
            "medium": "medio",
            "basic": "básico",
            "low": "básico",
        }
        return labels.get(str(value or "").strip(), CoEm004DocumentFactoryV246._fragment(value, "proporcional al riesgo"))

    @staticmethod
    def _years_label(value):
        try:
            number = int(value)
            return "1 año" if number == 1 else f"{number} años"
        except (TypeError, ValueError):
            return CoEm004DocumentFactoryV246._fragment(value, "el plazo definido")

    @staticmethod
    def _security_description(controls):
        controls = controls if isinstance(controls, dict) else {}
        parts = [f"nivel de protección {CoEm004DocumentFactoryV246._level_label(controls.get('level'))}"]
        labels = (("technical", "medidas técnicas"), ("organizational", "medidas organizacionales"), ("physical", "medidas físicas"))
        for key, label in labels:
            value = CoEm004DocumentFactoryV246._fragment(controls.get(key))
            if value:
                parts.append(f"{label}: {value}")
        return "; ".join(parts)

    @staticmethod
    def _clause(doc, number, heading, body):
        ordinals = [
            "PRIMERA", "SEGUNDA", "TERCERA", "CUARTA", "QUINTA", "SEXTA", "SÉPTIMA", "OCTAVA", "NOVENA", "DÉCIMA",
            "DÉCIMA PRIMERA", "DÉCIMA SEGUNDA", "DÉCIMA TERCERA", "DÉCIMA CUARTA", "DÉCIMA QUINTA", "DÉCIMA SEXTA",
            "DÉCIMA SÉPTIMA", "DÉCIMA OCTAVA", "DÉCIMA NOVENA", "VIGÉSIMA", "VIGÉSIMA PRIMERA", "VIGÉSIMA SEGUNDA",
            "VIGÉSIMA TERCERA", "VIGÉSIMA CUARTA", "VIGÉSIMA QUINTA", "VIGÉSIMA SEXTA", "VIGÉSIMA SÉPTIMA", "VIGÉSIMA OCTAVA",
            "VIGÉSIMA NOVENA", "TRIGÉSIMA", "TRIGÉSIMA PRIMERA", "TRIGÉSIMA SEGUNDA", "TRIGÉSIMA TERCERA",
            "TRIGÉSIMA CUARTA", "TRIGÉSIMA QUINTA", "TRIGÉSIMA SEXTA", "TRIGÉSIMA SÉPTIMA", "TRIGÉSIMA OCTAVA",
            "TRIGÉSIMA NOVENA", "CUADRAGÉSIMA"
        ]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.keep_together = False
        ordinal = ordinals[number - 1] if 0 < number <= len(ordinals) else str(number)
        r = p.add_run(f"CLÁUSULA {ordinal}. {heading.upper()}: ")
        r.bold = True
        p.add_run(str(body).strip())
        return p

    def _party_em4(self, answers, prefix, fallback):
        identification = self._as_dict(self._v(answers, f"{prefix}.identification"))
        signatory = self._as_dict(self._v(answers, f"{prefix}.signatory"))
        name = self._plain(identification.get("name"), fallback)
        id_number = self._plain(identification.get("id_number"), "identificación registrada")
        address = self._plain(identification.get("address"), "domicilio registrado")
        email = self._plain(identification.get("email"), "correo registrado")
        person_type = identification.get("type")
        sign_name = self._plain(signatory.get("name"), name)
        sign_id = self._plain(signatory.get("id_number"), "identificación registrada")
        capacity = self._plain(signatory.get("capacity"), "representante autorizado")
        authority = self._plain(signatory.get("authority_source"))
        if person_type == "natural_person":
            text = f"{name}, persona natural identificada con {id_number}, domiciliada en {address}"
        else:
            text = (
                f"{name}, persona jurídica identificada con NIT {id_number}, con domicilio en {address}, "
                f"representada para este acto por {sign_name}, identificado(a) con {sign_id}, quien actúa en calidad de {capacity}"
            )
            if authority:
                text += f", con fundamento en {authority}"
        return {
            "name": name,
            "id": id_number,
            "address": address,
            "email": email,
            "signatory": sign_name,
            "signatory_id": sign_id,
            "capacity": capacity,
            "authority": authority,
            "text": text,
        }

    def _agreement_configuration(self, answers):
        agreement_type = self._v(answers, "agreement.type")
        reciprocal = self._v(answers, "agreement.reciprocal")
        if agreement_type == "unilateral":
            return (
                "El acuerdo es unilateral. La Parte A actuará como PARTE REVELADORA y la Parte B como PARTE RECEPTORA, "
                "sin perjuicio de que una revelación excepcional de la Parte B requiera instrumento adicional o constancia expresa."
            )
        if agreement_type == "multilateral":
            return (
                "El acuerdo tiene configuración multilateral. Cada participante será PARTE REVELADORA respecto de la información que comunique "
                "y PARTE RECEPTORA respecto de la que reciba; las adhesiones deberán identificar al participante, su representante y la fecha de vinculación."
            )
        return (
            "El acuerdo es mutuo y recíproco." if reciprocal is not False else
            "El acuerdo fue denominado mutuo, pero las obligaciones se aplicarán únicamente en el alcance expresamente atribuido a cada parte."
        ) + " Cada parte será PARTE REVELADORA respecto de su información y PARTE RECEPTORA respecto de la información de la otra."

    def _ordinary_term(self, answers):
        term = self._as_dict(self._v(answers, "term_remedies"))
        agreement_years = term.get("agreement_years")
        ordinary_years = term.get("ordinary_confidentiality_years")
        trade_rule = term.get("trade_secret_rule")
        parts = []
        if agreement_years not in (None, ""):
            parts.append(f"El acuerdo tendrá una vigencia inicial de {self._years_label(agreement_years)}")
        else:
            parts.append("El acuerdo tendrá la vigencia indicada en el expediente")
        if ordinary_years not in (None, ""):
            parts.append(f"las obligaciones ordinarias de reserva sobrevivirán durante {self._years_label(ordinary_years)} después de su terminación")
        if trade_rule == "while_secret":
            parts.append("los secretos empresariales se protegerán mientras conserven esa calidad y subsistan medidas razonables de reserva")
        elif trade_rule == "fixed":
            parts.append("la reserva de secretos empresariales se sujetará al plazo fijo específicamente documentado, sin desconocer la protección legal aplicable")
        return "; ".join(parts) + "."

    def _contract(self, answers, evaluation, target):
        a = self._party_em4(answers, "party_a", "PARTE A")
        b = self._party_em4(answers, "party_b", "PARTE B")
        purpose = self._fragment(self._v(answers, "agreement.purpose"), "la finalidad autorizada registrada en el expediente", lower_initial=True)
        reference = self._fragment(self._v(answers, "agreement.reference"), "la relación o proyecto identificado en el expediente")
        categories = self._fragment(self._v(answers, "information.categories"), "las categorías delimitadas en el inventario de información")
        formats = self._fragment(self._v(answers, "information.formats_sources"), "medios escritos, orales, visuales, digitales y demostraciones autorizadas", lower_initial=True)
        exclusions = self._fragment(self._v(answers, "information.exclusions"), "las exclusiones legales y contractuales descritas en este acuerdo")
        recipients = self._fragment(self._v(answers, "access.authorized_recipients"), "las personas expresamente autorizadas", lower_initial=True)
        representatives = self._fragment(self._v(answers, "access.representatives"), "asesores y colaboradores sujetos a deberes equivalentes", lower_initial=True)
        need = self._fragment(self._v(answers, "access.need_to_know"), "acceso estrictamente limitado por necesidad de conocer", lower_initial=True)
        permitted = self._fragment(self._v(answers, "access.permitted_use"), purpose, lower_initial=True)
        compelled = self._fragment(self._v(answers, "access.compelled_disclosure"), "notificación previa cuando sea legalmente posible y revelación mínima necesaria", lower_initial=True)
        controls = self._as_dict(self._v(answers, "security.controls"))
        security_text = self._security_description(controls)
        personal = self._v(answers, "data.personal") is True
        roles = self._as_dict(self._v(answers, "data.roles"))
        lifecycle = self._fragment(self._v(answers, "data.lifecycle"), "el ciclo de vida definido en el anexo de datos", lower_initial=True)
        crossborder = self._v(answers, "data.crossborder")
        allocation = self._allocation_label(self._v(answers, "ip.results_allocation"))
        preexisting = self._fragment(self._v(answers, "ip.preexisting_materials"), "los materiales preexistentes identificados por cada parte")
        oss_raw = self._v(answers, "ip.oss_third_party")
        has_oss = bool(str(oss_raw or "").strip())
        oss = self._fragment(oss_raw, "los componentes de terceros u OSS inventariados")
        oss_clause = (
            f"Los componentes de terceros u OSS se administrarán conforme al inventario ANX-EM4-OSS-001, incluyendo {oss}. "
            "Cada componente deberá identificar titular, versión, licencia, avisos, obligaciones de atribución, copyleft, distribución de código fuente, restricciones y vulnerabilidades conocidas. "
            "Ninguna parte podrá prometer exclusividad sobre un activo sujeto a derechos de terceros."
            if has_oss else
            "No se prevé incorporar componentes de terceros u OSS dentro de la Finalidad informada. Si posteriormente se propone su uso, deberá identificarse y aprobarse previamente en el inventario ANX-EM4-OSS-001, "
            "verificando titular, versión, licencia, avisos, obligaciones de atribución, copyleft, distribución de código fuente, restricciones y vulnerabilidades conocidas. Ninguna parte podrá prometer exclusividad sobre un activo sujeto a derechos de terceros."
        )
        source_rule = self._fragment(self._v(answers, "ip.source_code_reverse_engineering"), "las restricciones y permisos expresamente documentados", lower_initial=True)
        ai_used = self._v(answers, "ai.used") is True
        ai_terms = self._fragment(self._v(answers, "ai.training_outputs"), "sin entrenamiento ni retención con información protegida, salvo autorización expresa", lower_initial=True)
        incident = self._fragment(self._v(answers, "security.incident_protocol"), "notificación sin demora injustificada, contención y preservación de evidencia", lower_initial=True)
        close = self._as_dict(self._v(answers, "closure_confirmation"))
        return_destroy = self._fragment(close.get("return_destroy"), "devolución o eliminación segura al finalizar", lower_initial=True)
        retained = self._fragment(close.get("retained_copies"), "conservación limitada a obligaciones legales, respaldo inalterable o defensa de derechos", lower_initial=True).replace("backups", "copias de respaldo").replace("sujetos al", "sujetas al")
        dispute = close.get("dispute_mechanism")
        dispute_text = {
            "ordinary_courts": "las autoridades judiciales competentes de Colombia",
            "negotiation_conciliation": "negociación directa y, si no se alcanza un acuerdo, a conciliación antes de acudir a la jurisdicción competente",
            "arbitration": "un tribunal de arbitramento conforme a las reglas y sede que LAS PARTES definan en el expediente",
            "to_define": "el mecanismo que LAS PARTES definan antes de la aprobación final",
        }.get(dispute, "el mecanismo de solución de controversias registrado en el expediente")
        term = self._as_dict(self._v(answers, "term_remedies"))
        liability = self._fragment(term.get("penalty_or_liability"), "responsabilidad por daños directos, probados y causalmente vinculados al incumplimiento", lower_initial=True)

        doc = Document()
        self._set_doc_styles(doc)
        self._title(doc, "Acuerdo de confidencialidad, uso restringido y propiedad intelectual")
        appearance = (
            f"Entre los suscritos, de una parte, {a['text']}, quien para efectos del presente acuerdo se denominará PARTE A; "
            f"y de otra parte, {b['text']}, quien se denominará PARTE B; conjuntamente LAS PARTES, se celebra el presente acuerdo "
            "de confidencialidad, uso restringido, seguridad de la información y, cuando corresponda, propiedad intelectual y tratamiento de datos personales, "
            "regido por la legislación colombiana y por las siguientes estipulaciones."
        )
        self._paragraph(doc, appearance)

        self._section(doc, "Consideraciones")
        considerations = [
            ("PRIMERA", f"Que LAS PARTES mantienen o proyectan mantener una relación vinculada con {reference} y que, para evaluarla, ejecutarla o cerrarla, pueden revelar información de valor jurídico, técnico, científico, comercial, financiero, operativo o estratégico."),
            ("SEGUNDA", f"Que la revelación se limitará a {purpose}, de modo que ninguna entrega autoriza usos distintos, competencia desleal, explotación autónoma, registro, publicación o transferencia a terceros."),
            ("TERCERA", "Que la protección contractual no convierte en secreto toda información indiscriminadamente, sino que exige categorías identificables, exclusiones verificables, acceso restringido y medidas razonables según la sensibilidad y el contexto."),
            ("CUARTA", "Que los derechos patrimoniales, los derechos morales, los materiales preexistentes, los componentes de terceros, los datos personales y los secretos empresariales obedecen a regímenes diferentes y deben mantenerse jurídicamente diferenciados."),
            ("QUINTA", "Que los anexos, inventarios, protocolos y actas identificados en el expediente forman parte del acuerdo únicamente cuando hayan sido generados, aprobados y vinculados a la revisión vigente."),
        ]
        for label, body in considerations:
            self._paragraph(doc, body, f"{label}: ")

        self._section(doc, "Cláusulas")
        clauses = [
            ("Configuración y roles", self._agreement_configuration(answers)),
            ("Definiciones", "Para este acuerdo, Información Confidencial es la información no pública revelada directa o indirectamente, cualquiera sea su soporte, que esté incluida en las categorías pactadas o que una persona razonable deba reconocer como reservada por su naturaleza, contexto, valor o forma de revelación. Parte Reveladora es quien comunica o permite el acceso; Parte Receptora es quien recibe; Representantes son las personas autorizadas que requieren acceso para la Finalidad; Incidente es cualquier pérdida, acceso, uso, alteración, divulgación o indisponibilidad no autorizada."),
            ("Finalidad autorizada", f"La Información Confidencial solo podrá recibirse, examinarse, reproducirse y utilizarse para {purpose}. La Finalidad se limita a {reference}. Cualquier uso adicional, comercialización, benchmarking, entrenamiento de sistemas, publicación, solicitud de registro o transferencia requerirá autorización previa, específica y verificable de la Parte Reveladora."),
            ("Información protegida y categorías", f"La protección comprende {categories}. Podrá revelarse mediante {formats}. El inventario ANX-EM4-INFO-001 deberá permitir identificar categorías, custodios, repositorios, nivel de sensibilidad, personas autorizadas, período de conservación y evento de devolución o eliminación."),
            ("Revelaciones orales, visuales y demostraciones", "Las revelaciones orales, visuales, audiovisuales, demostraciones, recorridos, muestras y accesos a sistemas se protegerán conforme a la regla seleccionada en el expediente. Cuando se exija confirmación escrita, esta deberá describir suficientemente la materia revelada sin reproducir secretos innecesariamente; la falta de confirmación no elimina la protección cuando el carácter reservado sea objetivamente evidente y la conducta de las partes demuestre el deber de reserva."),
            ("Identificación y marcado", "El marcado facilita la clasificación, pero no sustituye la diligencia. Si se utiliza un sistema híbrido, quedará protegida tanto la información marcada como aquella que razonablemente deba reconocerse como confidencial. Si se pacta protección exclusiva de lo marcado, la Parte Reveladora asumirá el deber de identificación y se aplicarán las excepciones legales y de buena fe correspondientes."),
            ("Exclusiones", f"No se considerará confidencial la información que la Parte Receptora demuestre documentalmente que: (i) era pública sin infracción; (ii) estaba legítimamente en su poder sin deber de reserva; (iii) fue recibida de tercero facultado para revelarla; (iv) fue desarrollada independientemente sin utilizar la información protegida; o (v) fue autorizada por escrito para divulgación. Se aplicarán además las exclusiones específicas siguientes: {exclusions}. La combinación de elementos públicos podrá conservar carácter reservado cuando su selección, disposición o integración no sea generalmente conocida."),
            ("Titularidad y ausencia de licencia implícita", "La revelación no transfiere propiedad, licencia, exclusividad, derecho de explotación, expectativa de negocio ni autorización para solicitar registros. Cada Parte Reveladora conserva sus derechos sobre la información, soportes y activos. Solo se concede una autorización limitada, revocable y no sublicenciable para ejecutar la Finalidad durante la vigencia aplicable."),
            ("Obligaciones de la Parte Receptora", "La Parte Receptora deberá proteger la información al menos con el mismo cuidado que emplea para sus activos reservados de importancia equivalente y nunca con menos de una diligencia razonable; evitará usos no autorizados; limitará copias y descargas; mantendrá trazabilidad; reportará incidentes; impedirá accesos por credenciales compartidas; cumplirá las instrucciones legítimas; y cooperará razonablemente en la contención, investigación y cierre."),
            ("Uso permitido, copias y reproducciones", f"Los usos autorizados son: {permitted}. Solo podrán realizarse copias, extractos, respaldos, transformaciones o reproducciones cuando sean necesarios para la Finalidad, estén protegidos con controles equivalentes y conserven leyendas o metadatos de clasificación cuando resulte aplicable. No se permitirá la extracción masiva, descompilación, análisis competitivo o reproducción para un proyecto distinto."),
            ("Acceso por necesidad de conocer", f"El acceso se limitará conforme a {need}. Las personas autorizadas son {recipients}. La Parte Receptora implementará alta, modificación y retiro de accesos, mínimo privilegio, segregación de funciones y revisión periódica; el acceso técnico posible no se considerará autorización jurídica."),
            ("Representantes, asesores, afiliadas y subcontratistas", f"Podrán acceder únicamente {representatives}, siempre que requieran la información para la Finalidad, conozcan su carácter reservado y estén obligados por deberes de protección no menos exigentes. La Parte Receptora responderá por su selección, instrucciones y control, sin perjuicio de la responsabilidad directa del tercero. No se autoriza una cadena ilimitada de subreceptores."),
            ("Medidas de seguridad", f"La Parte Receptora aplicará medidas técnicas, humanas, administrativas y físicas proporcionales a la naturaleza y riesgo de la información. Como mínimo: {security_text}. Los secretos empresariales, credenciales, datos sensibles, código fuente y resultados no publicados deberán contar con controles reforzados, registros de acceso, cifrado cuando corresponda y procedimientos de recuperación y continuidad."),
            ("Divulgación obligatoria", f"Ante requerimiento judicial, administrativo o legal, la Parte Receptora seguirá este procedimiento: {compelled}. Revelará únicamente la porción exigible, solicitará tratamiento reservado cuando proceda y conservará evidencia del requerimiento y de la respuesta. La cláusula no impide denunciar irregularidades, ejercer defensa ni acudir a autoridades competentes."),
            ("Secretos empresariales", "Cuando la información reúna los presupuestos de secreto empresarial, su protección no dependerá únicamente del plazo contractual. La Parte titular deberá adoptar medidas razonables para mantenerla secreta, y la Parte Receptora se abstendrá de adquirirla, usarla o divulgarla de manera contraria a las prácticas comerciales leales. La pérdida del carácter secreto deberá demostrarse respecto de la información concreta, no presumirse por el simple paso del tiempo."),
            ("Datos personales", (f"LAS PARTES tratarán datos personales conforme a los roles e instrucciones del ANX-EM4-DATA-001. La Parte A actuará como {self._role_label(roles.get('party_a_role'))} y la Parte B como {self._role_label(roles.get('party_b_role'))}. Las finalidades e instrucciones serán {self._fragment(roles.get('instructions'), purpose, lower_initial=True)}. El tratamiento se limitará a datos adecuados, pertinentes y necesarios, con acceso restringido, seguridad, confidencialidad, atención de derechos, gestión de incidentes y {lifecycle}." if personal else "Las categorías informadas no requieren tratamiento de datos personales por cuenta de la otra parte. Si durante la ejecución surge dicho tratamiento, deberá suspenderse hasta definir roles, finalidades, instrucciones, medidas y ciclo de vida en un anexo específico.")),
            ("Transferencias y servicios transfronterizos", ("La transferencia o transmisión internacional, el acceso remoto desde otro país o el uso de nube externa requerirá identificar jurisdicción, proveedor, ubicación, subencargados, base jurídica, instrucciones, medidas y mecanismo de retorno o supresión. La aprobación deberá preceder el acceso." if crossborder in {"yes", "unknown"} else "No se autoriza transferencia, transmisión o acceso internacional no previsto. Un cambio de proveedor, región o subencargado que implique tratamiento transfronterizo requerirá evaluación y aprobación previa.")),
            ("Propiedad intelectual sobre resultados", f"La distribución de derechos sobre resultados se regirá por {allocation}. Ninguna cláusula se interpretará como cesión universal de creaciones futuras indeterminadas. Las transferencias o licencias deberán identificar obras o categorías determinables, modalidades de explotación, territorio, duración, contraprestación cuando corresponda, entregables y limitaciones. Las ideas, métodos no protegibles y conocimientos generales se diferenciarán de las obras, invenciones, diseños y secretos."),
            ("Materiales preexistentes", f"Cada parte conserva la titularidad de sus materiales, herramientas, metodologías, bibliotecas, marcas, datos y conocimientos preexistentes, incluidos: {preexisting}. Su incorporación a un resultado deberá declararse y solo otorgará la licencia estrictamente necesaria para utilizar el entregable conforme a la Finalidad, salvo pacto escrito distinto."),
            ("Derechos morales y autoría", "Los derechos morales pertenecen al autor y no se ceden ni renuncian. Las autorizaciones de transformación o adaptación no podrán interpretarse como autorización para deformar la obra o afectar el honor o reputación del autor. Se registrarán autores, contribuciones, créditos, versiones y autorizaciones de modificación cuando sea material."),
            ("Código fuente, ingeniería inversa e interoperabilidad", f"El acceso a código fuente, binarios, APIs, modelos, esquemas, configuraciones o ambientes se regirá por las siguientes condiciones: {source_rule}. No se autoriza descompilar, eludir controles, extraer secretos, realizar pruebas intrusivas o replicar funciones, salvo permiso escrito o excepción legal aplicable. Las pruebas autorizadas deberán delimitar alcance, ambiente, horario, datos y reporte responsable de hallazgos."),
            ("Software de código abierto y componentes de terceros", oss_clause),
            ("Inteligencia artificial", (f"El uso de IA se limitará a sistemas autorizados y a estas condiciones: {ai_terms}. No se cargarán secretos, datos personales, credenciales, código restringido ni obras de terceros en servicios públicos o no aprobados. Los resultados deberán someterse a revisión humana, trazabilidad, verificación de exactitud, seguridad, sesgos, licencias y titularidad. La IA no adoptará decisiones finales que afecten derechos sin intervención humana competente." if ai_used else "No se autoriza utilizar la Información Confidencial como entrada, contexto, memoria, dataset, material de ajuste, entrenamiento, evaluación o recuperación en sistemas de inteligencia artificial. Cualquier uso posterior requerirá anexo específico, proveedor aprobado, revisión humana y reglas de retención y salida.")),
            ("Incidentes y preservación de evidencia", f"Ante pérdida, divulgación, acceso, uso, alteración o indisponibilidad no autorizada se aplicará: {incident}. La Parte afectada deberá contener el evento, preservar registros, evitar destrucción de evidencia, documentar decisiones, evaluar titulares o terceros afectados, coordinar comunicaciones y ejecutar acciones correctivas. La notificación inicial no constituye admisión automática de responsabilidad."),
            ("Vigencia y supervivencia", self._ordinary_term(answers) + " La terminación de la relación principal no extingue obligaciones de devolución, conservación restringida, datos personales, propiedad intelectual, incidentes, auditoría, responsabilidad o solución de controversias que por su naturaleza deban sobrevivir."),
            ("Devolución y destrucción", f"Al terminar la Finalidad o cuando lo solicite legítimamente la Parte Reveladora, la Parte Receptora aplicará {return_destroy}. Deberá devolver soportes, cerrar accesos, eliminar copias de trabajo, deshabilitar sincronizaciones y documentar las excepciones. La destrucción deberá ser técnicamente razonable y proporcional al soporte."),
            ("Copias retenidas y conservación probatoria", f"Podrán conservarse únicamente las siguientes copias autorizadas: {retained}. Estas permanecerán aisladas, con acceso limitado y sujetas a las obligaciones de reserva. Las copias de respaldo inalterables podrán eliminarse en el ciclo ordinario, siempre que no sean restauradas salvo necesidad legítima y se mantengan protegidas."),
            ("Conocimiento residual", "No se reconoce una licencia general sobre información recordada por las personas. El conocimiento general, experiencia y habilidades profesionales podrán utilizarse, pero no los detalles, combinaciones, datos, diseños, código, listas, modelos, estrategias, secretos o reproducciones sustanciales de la Información Confidencial. La carga de demostrar desarrollo independiente corresponderá a quien invoque la excepción."),
            ("Responsabilidad, cláusula penal y remedios", f"El régimen aplicable será: {liability}. La responsabilidad requerirá incumplimiento, daño, causalidad y prueba, salvo los eventos de responsabilidad objetiva previstos por ley. Cualquier cláusula penal deberá ser determinada o determinable, respetar las reglas imperativas y aclarar si opera como estimación anticipada, apremio o sin perjuicio de daños adicionales. Se mantienen los derechos a solicitar cesación, recuperación de soportes, preservación de evidencia y medidas cautelares cuando procedan."),
            ("Ausencia de no competencia general", "El acuerdo protege información y activos, pero no impide de manera general ejercer una profesión, trabajar, competir lícitamente o desarrollar soluciones independientes. Las restricciones sobre solicitud de clientes, uso de secretos, explotación de resultados o conflictos de interés deberán ser específicas, necesarias, proporcionales y estar vinculadas a un interés legítimo, sin convertirse en prohibición general postcontractual."),
            ("Declaraciones y ausencia de garantía implícita", "Cada Parte Reveladora declara que tiene facultad para comunicar la información en el alcance previsto. Salvo garantía expresa, la información se entrega para evaluación y podrá ser preliminar; la Parte Receptora deberá realizar su propio análisis antes de adoptar decisiones. Esta limitación no cubre fraude, dolo, ocultamiento deliberado ni incumplimiento de garantías expresas."),
            ("Fuerza mayor y deber de mitigación", "La fuerza mayor podrá justificar demoras en obligaciones operativas, pero no autoriza uso o divulgación no permitidos. La parte afectada deberá notificar, mitigar, aplicar continuidad y restablecer controles. La pérdida de capacidad tecnológica no elimina el deber de proteger, devolver o preservar la información."),
            ("Notificaciones", f"Las comunicaciones se enviarán a {a['email']} para la Parte A y {b['email']} para la Parte B, o a los contactos que se actualicen por medio verificable. Los incidentes utilizarán además los canales de emergencia previstos en PRO-EM4-INC-001. La comunicación electrónica será admisible sin perjuicio de formalidades especiales exigidas por ley."),
            ("Ley aplicable y controversias", f"El acuerdo se regirá por las leyes de la República de Colombia. Las controversias se intentarán resolver de buena fe. Si persisten, se someterán a {dispute_text}. Las medidas urgentes de protección de información, datos, secretos o propiedad intelectual podrán solicitarse ante la autoridad competente sin esperar el agotamiento de una etapa que haga ineficaz la protección."),
            ("Integridad, modificaciones, cesión y prelación", "El acuerdo, sus anexos y actas vigentes contienen el entendimiento sobre las materias reguladas. Toda modificación material deberá constar en una nueva revisión trazable y reiniciar las aprobaciones. Ninguna parte podrá ceder el acuerdo o transferir información a un sucesor sin verificar capacidad y protección equivalentes, salvo reorganización autorizada. En caso de contradicción prevalecerán las normas imperativas, luego el acuerdo, los otrosíes, anexos específicos y protocolos operativos."),
            ("Firma y aceptación electrónica", "LAS PARTES declaran haber leído el acuerdo, comprender su alcance y haber tenido oportunidad de solicitar aclaraciones. Podrá suscribirse mediante firma manuscrita, digital o electrónica confiable que permita identificar al firmante, vincularlo con el documento, preservar integridad y conservar evidencia de fecha, hora, método, versión y aceptación. Cada parte recibirá una copia íntegra."),
        ]
        for i, (heading, body) in enumerate(clauses, 1):
            self._clause(doc, i, heading, body)

        self._section(doc, "Documentos relacionados")
        names = [x["id"] if isinstance(x, dict) else str(x) for x in evaluation.get("documents", [])]
        self._bullets(doc, names)
        self._paragraph(doc, "Solo forman parte del expediente los documentos efectivamente generados, aprobados y vinculados a la revisión vigente. Los formularios de entrega y cierre deberán completarse cuando ocurra el evento correspondiente.")

        doc.add_page_break()
        self._section(doc, "Firmas")
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_props = table.rows[0]._tr.get_or_add_trPr()
        hdr = OxmlElement("w:tblHeader"); hdr.set(qn("w:val"), "true"); header_props.append(hdr)
        values = [
            ("PARTE A", "PARTE B"),
            (a["name"], b["name"]),
            (a["signatory"], b["signatory"]),
            (a["capacity"], b["capacity"]),
            ("Firma: ______________________________", "Firma: ______________________________"),
        ]
        for r_idx, row in enumerate(table.rows):
            for cell, value in zip(row.cells, values[r_idx]):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(value); run.bold = r_idx == 0
        doc.core_properties.title = "Acuerdo de confidencialidad, uso restringido y propiedad intelectual"
        doc.core_properties.subject = f"CO-EM-004 v{self.VERSION}"
        doc.core_properties.comments = "Generado por LegalAIZ.it. Requiere aprobación jurídica y QA antes de liberación."
        doc.save(target)

    def _annex_header(self, doc, answers, doc_id, title):
        a = self._party_em4(answers, "party_a", "PARTE A")
        b = self._party_em4(answers, "party_b", "PARTE B")
        self._title(doc, title)
        self._paragraph(doc, self._ensure_period(
            f"Documento {doc_id}, vinculado al acuerdo de confidencialidad y propiedad intelectual celebrado entre {a['name']} y {b['name']}"
        ) + " Corresponde a la revisión vigente del expediente y no podrá aplicarse separadamente a otra relación.")
        return a, b

    def _info_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc); self._compact_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-INFO-001", "Inventario de información y matriz de acceso")
        self._section(doc, "Finalidad y categorías")
        self._paragraph(doc, self._ensure_period(self._fragment(self._v(answers, "agreement.purpose"))), "FINALIDAD: ")
        self._paragraph(doc, self._ensure_period(self._fragment(self._v(answers, "information.categories"))), "CATEGORÍAS: ")
        rows = []
        categories = self._as_list(self._v(answers, "information.categories")) or ["Información descrita en el expediente"]
        level = self._level_label(self._v(answers, "security.controls.level"))
        recipients = self._fragment(self._v(answers, "access.authorized_recipients"), "Personas autorizadas")
        for i, category in enumerate(categories, 1):
            rows.append([i, self._fragment(category), level, recipients, "Repositorio autorizado", "Según ciclo de vida"])
        headers = ["No.", "Categoría", "Nivel", "Destinatarios", "Repositorio", "Conservación"]
        if len(rows) > 4:
            self._table(doc, headers, rows[:4])
            doc.add_page_break()
            self._section(doc, "Inventario de información — continuación")
            self._table(doc, headers, rows[4:])
        else:
            self._table(doc, headers, rows)
        self._section(doc, "Canales y fuentes")
        self._paragraph(doc, self._plain(self._v(answers, "information.formats_sources")))
        self._section(doc, "Reglas de acceso")
        self._subparagraph(doc, "Necesidad de conocer", self._plain(self._v(answers, "access.need_to_know")))
        self._subparagraph(doc, "Representantes", self._plain(self._v(answers, "access.representatives")))
        self._subparagraph(doc, "Usos y copias", self._plain(self._v(answers, "access.permitted_use")))
        self._subparagraph(doc, "Exclusiones", self._plain(self._v(answers, "information.exclusions")))
        self._section(doc, "Control de entrega")
        self._table(doc, ["Fecha", "Elemento / lote", "Parte reveladora", "Parte receptora", "Medio", "Observaciones"], [["", "", "", "", "", ""] for _ in range(4)])
        doc.save(target)

    def _relationship_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-REL-001", "Anexo especializado según la relación")
        context = self._v(answers, "relationship.context")
        labels = {
            "preliminary_negotiation": "Negociación preliminar",
            "commercial_supplier": "Proveedor o relación comercial",
            "independent_services": "Servicios independientes",
            "employment_collaboration": "Relación laboral o colaboración",
            "software_technology": "Software, SaaS o tecnología",
            "creative_content": "Contenidos y servicios creativos",
            "research_development": "Investigación y desarrollo",
            "strategic_alliance": "Alianza estratégica",
            "investment_due_diligence": "Inversión o debida diligencia",
            "data_processing": "Tratamiento de datos y outsourcing",
        }
        self._paragraph(doc, labels.get(context, "Relación definida en el expediente"), "CONTEXTO: ")
        clauses = {
            "preliminary_negotiation": [
                ("Ausencia de obligación de contratar", "La revelación permite evaluar una operación, pero no obliga a concluirla, otorgar exclusividad, financiarla ni continuar negociaciones. Cada parte asumirá sus costos salvo pacto escrito."),
                ("Contactos y no elusión", "El acceso a contactos, oportunidades o estructuras no autoriza su utilización fuera de la finalidad. Cualquier restricción deberá proteger una oportunidad identificable y no impedir competencia lícita general."),
            ],
            "commercial_supplier": [
                ("Información comercial", "Listas de precios, márgenes, clientes, proveedores, términos, volúmenes, previsiones y estrategias se usarán únicamente para la relación aprobada."),
                ("Personal y subcontratistas", "La parte receptora controlará el acceso de su personal y terceros y conservará evidencia de instrucciones y cierre."),
            ],
            "independent_services": [
                ("Coordinación independiente", "Los accesos, reportes y controles de seguridad no crean por sí solos subordinación laboral; tampoco autorizan usar información para otros clientes."),
                ("Entregables", "Los resultados, archivos fuente, materiales preexistentes, datos y herramientas se clasificarán antes de su entrega."),
            ],
            "employment_collaboration": [
                ("Compatibilidad laboral", "El anexo complementa, pero no sustituye, las obligaciones laborales ni puede impedir denuncias, defensa de derechos, actividad sindical o movilidad profesional lícita."),
                ("Información de terceros", "La persona colaboradora no deberá aportar secretos o materiales de anteriores empleadores o clientes."),
            ],
            "software_technology": [
                ("Ambientes y credenciales", "Se separarán producción, pruebas y desarrollo; los secretos y credenciales se manejarán con mínimo privilegio, MFA y repositorios autorizados."),
                ("Seguridad y vulnerabilidades", "La investigación de seguridad solo se realizará dentro del alcance autorizado, con preservación de evidencia y divulgación responsable."),
            ],
            "creative_content": [
                ("Materiales y créditos", "Se identificarán autores, materiales previos, licencias, releases, tipografías, música, imágenes, voces y autorizaciones de terceros."),
                ("Portafolio", "La inclusión en portafolio o premios requerirá autorización y no podrá revelar campañas, resultados o lanzamientos no públicos."),
            ],
            "research_development": [
                ("Cuadernos y resultados", "Se mantendrán registros de experimentos, contribuciones, muestras, datos, versiones y decisiones para soportar autoría, inventiva y reproducibilidad."),
                ("Publicaciones", "Abstracts, ponencias, tesis, artículos y divulgaciones requerirán revisión previa para proteger secretos, patentes, datos y derechos de terceros."),
            ],
            "strategic_alliance": [
                ("Gobierno de la alianza", "La información compartida no crea sociedad, mandato o representación. Cada parte solo podrá comprometer a la otra con autorización expresa."),
                ("Oportunidades", "Las oportunidades conjuntas se identificarán por escrito; el acuerdo no otorga exclusividad general ni derecho sobre negocios independientes."),
            ],
            "investment_due_diligence": [
                ("Data room", "El acceso se limitará por usuario, documento y período; se prohibirá descarga o impresión cuando no sea necesaria y se registrarán consultas y copias."),
                ("Información prospectiva", "Proyecciones y estimaciones se entregan para análisis y no constituyen garantía salvo declaración expresa."),
            ],
            "data_processing": [
                ("Instrucciones", "El encargado tratará datos solo por instrucciones documentadas del responsable, sin determinar fines incompatibles."),
                ("Subencargados", "Todo subencargado requerirá autorización y obligaciones equivalentes, con identificación de ubicación y servicios."),
            ],
        }
        selected = clauses.get(context, [("Alcance", "Las obligaciones especializadas se limitarán a la relación y finalidad descritas en el expediente.")])
        for i, (h, b) in enumerate(selected, 1): self._clause(doc, i, h, b)
        self._clause(doc, len(selected)+1, "Prelación", "Este anexo complementa el acuerdo principal. Si existe contradicción, prevalecerá la regla específica más protectora que sea compatible con la ley y la finalidad aprobada.")
        doc.save(target)

    def _ip_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-IP-001", "Anexo de propiedad intelectual y materiales preexistentes")
        allocation = self._allocation_label(self._v(answers, "ip.results_allocation"))
        pre = self._fragment(self._v(answers, "ip.preexisting_materials"), "ninguno informado")
        source = self._fragment(self._v(answers, "ip.source_code_reverse_engineering"), "no aplica o según autorización específica", lower_initial=True)
        clauses = [
            ("Clasificación", "Se distinguirán obras, software, bases de datos protegibles, invenciones, diseños, secretos empresariales, información no protegible, activos preexistentes y componentes de terceros."),
            ("Resultados", f"La distribución seleccionada es: {allocation}. Toda cesión o licencia deberá delimitar activos determinables, modalidades de explotación, territorio, plazo, entregables y contraprestación cuando corresponda."),
            ("Materiales preexistentes", f"Se identifican: {pre}. Su titular conserva los derechos y concede únicamente la licencia necesaria para la finalidad, salvo pacto distinto."),
            ("Derechos morales", "Los autores conservan los derechos morales inalienables e irrenunciables. Las autorizaciones de modificación no permiten afectar paternidad, integridad, honor o reputación."),
            ("Autoría e inventiva", "Se conservarán registros de contribuciones, fechas, versiones, cuadernos, commits, inventores, coautores y decisiones. El aporte de ideas o recursos por sí solo no atribuye autoría."),
            ("Código fuente y pruebas", f"El acceso, la ingeniería inversa, la interoperabilidad y las pruebas se regirán por las siguientes condiciones: {source}. No se autoriza la elusión de controles, la extracción de secretos ni el uso fuera del ambiente aprobado."),
            ("Terceros", "No se incorporarán obras, datos, marcas, código o materiales de terceros sin licencia compatible. Las restricciones y avisos deberán acompañar el entregable."),
            ("Entrega y continuidad", "Cuando corresponda se entregarán archivos fuente, documentación, dependencias, instrucciones de despliegue, formatos editables y materiales necesarios para el uso legítimo del resultado."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.add_page_break()
        self._section(doc, "Inventario inicial de activos y materiales")
        self._paragraph(doc, "Registre cada activo preexistente, resultado, componente de tercero o material sujeto a licencia antes de su incorporación. La omisión no modifica por sí sola la titularidad ni autoriza usos distintos de la finalidad aprobada.")
        self._table(doc, ["Activo", "Titular", "Clasificación", "Licencia / cesión", "Restricciones", "Evidencia"], [["", "", "", "", "", ""] for _ in range(8)])
        self._section(doc, "Aprobación del inventario")
        self._table(doc, ["Parte", "Nombre y calidad", "Fecha", "Firma / evidencia"], [["Parte A", "", "", ""], ["Parte B", "", "", ""]])
        doc.save(target)

    def _data_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-DATA-001", "Anexo de tratamiento de datos personales y transferencias")
        roles = self._as_dict(self._v(answers, "data.roles"))
        clauses = [
            ("Roles", f"Parte A: {self._role_label(roles.get('party_a_role'))}. Parte B: {self._role_label(roles.get('party_b_role'))}. La denominación deberá corresponder a las decisiones reales de cada parte sobre fines y medios."),
            ("Instrucciones y finalidades", self._ensure_period(self._fragment(roles.get("instructions"), self._v(answers, "agreement.purpose")))),
            ("Datos y titulares", "Se identificarán categorías de datos, titulares, origen, operaciones, sistemas, destinatarios y datos sensibles. No se tratarán datos adicionales por conveniencia técnica."),
            ("Deberes del responsable", "El responsable asegurará base jurídica, información al titular, calidad, instrucciones, atención de derechos y entrega lícita al encargado."),
            ("Deberes del encargado", "El encargado tratará los datos solo por instrucciones documentadas, mantendrá confidencialidad, seguridad, registros, cooperación, restricción de accesos y devolución o eliminación."),
            ("Subencargados", "Los subencargados requerirán autorización, identificación previa y obligaciones equivalentes. La parte que los vincule conservará responsabilidad por selección y control."),
            ("Seguridad", self._ensure_period(self._security_description(self._as_dict(self._v(answers, "security.controls"))))),
            ("Incidentes", self._ensure_period(self._fragment(self._v(answers, "security.incident_protocol")))),
            ("Transferencias", "Cualquier transferencia, transmisión o acceso transfronterizo deberá documentar país, proveedor, ubicación, garantías, subencargados, base jurídica y mecanismo de retorno o supresión."),
            ("Ciclo de vida", self._ensure_period(self._fragment(self._v(answers, "data.lifecycle"), "Conservación durante la finalidad y eliminación o devolución al finalizar, salvo obligación legal documentada"))),
            ("Auditoría", "La parte receptora suministrará evidencia razonable de cumplimiento sin revelar secretos de terceros ni comprometer la seguridad. Las auditorías serán proporcionales, coordinadas y trazables."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        doc.add_page_break()
        self._section(doc, "Registro de operaciones de tratamiento")
        self._paragraph(doc, "Complete una fila por categoría y operación. El registro deberá coincidir con las instrucciones documentadas, los sistemas autorizados, los destinatarios y el ciclo de vida aprobado.")
        self._table(doc, ["Categoría", "Titulares", "Finalidad", "Operación", "Repositorio", "Conservación", "Destinatario"], [["", "", "", "", "", "", ""] for _ in range(8)])
        self._section(doc, "Responsables de validación")
        self._table(doc, ["Rol", "Nombre", "Fecha", "Firma / evidencia"], [["Parte responsable", "", "", ""], ["Parte encargada o receptora", "", "", ""]])
        doc.save(target)

    def _ai_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-AI-001", "Anexo de uso controlado de inteligencia artificial")
        terms = self._ensure_period(self._fragment(self._v(answers, "ai.training_outputs"), "Sin retención ni entrenamiento con información protegida"))
        clauses = [
            ("Sistemas autorizados", "Solo se utilizarán proveedores, cuentas, modelos, regiones y configuraciones aprobados. Deberán revisarse términos, retención, entrenamiento, seguridad, subprocesadores y ubicación."),
            ("Entradas", "No se ingresarán secretos, datos personales, credenciales, código restringido, información privilegiada ni materiales de terceros sin autorización y controles suficientes."),
            ("Entrenamiento y retención", terms),
            ("Resultados", "Los resultados deberán ser revisados por una persona competente respecto de exactitud, alucinaciones, sesgo, seguridad, atribución, licencias, datos y adecuación al propósito."),
            ("Trazabilidad", "Cuando el uso sea material se registrarán herramienta, versión, finalidad, fecha, usuario, revisión humana y cambios, evitando duplicar información confidencial innecesaria."),
            ("Propiedad intelectual", "El uso de IA no garantiza titularidad, exclusividad ni ausencia de infracción. Se respetarán derechos sobre entradas, datasets, modelos, estilos, marcas, voces e imágenes."),
            ("Decisiones", "La IA no adoptará decisiones finales que afecten derechos o comprometan a LAS PARTES sin revisión humana autorizada."),
            ("Prohibiciones", "Quedan prohibidos la carga en sistemas públicos no autorizados, la evasión de controles, la suplantación, el entrenamiento oculto, la clonación no autorizada y el uso ilícito de contenidos."),
            ("Incidentes", "Se reportarán fugas, exposición de datos, respuestas inseguras, uso indebido, sesgos materiales o cambios de términos. El sistema podrá suspenderse hasta completar evaluación."),
        ]
        for i, (h, b) in enumerate(clauses, 1): self._clause(doc, i, h, b)
        self._section(doc, "Registro de herramientas")
        self._table(doc, ["Herramienta", "Proveedor", "Versión", "Finalidad", "Datos permitidos", "Retención", "Revisor"], [["", "", "", "", "", "", ""] for _ in range(4)])
        doc.save(target)

    def _oss_annex(self, answers, target):
        doc = Document(); self._set_doc_styles(doc); self._compact_styles(doc)
        self._annex_header(doc, answers, "ANX-EM4-OSS-001", "Inventario de OSS y componentes de terceros")
        self._paragraph(doc, self._ensure_period(self._fragment(self._v(answers, "ip.oss_third_party"))), "INFORMACIÓN INICIAL: ")
        self._table(doc, ["Componente", "Versión", "Titular / fuente", "Licencia", "Uso", "Obligaciones", "Vulnerabilidades"], [["", "", "", "", "", "", ""] for _ in range(6)])
        self._section(doc, "Controles")
        self._bullets(doc, [
            "Verificar compatibilidad de la licencia con el modelo de distribución y los demás componentes.",
            "Conservar avisos, atribuciones, textos de licencia y ofertas de código fuente cuando correspondan.",
            "No presentar como exclusivo un componente sujeto a derechos de terceros.",
            "Registrar procedencia, versión, hash y cambios locales.",
            "Evaluar vulnerabilidades, mantenimiento y sustitución antes de liberación.",
        ])
        doc.save(target)

    def _incident_protocol(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "PRO-EM4-INC-001", "Protocolo de incidentes y preservación de evidencia")
        self._paragraph(doc, self._ensure_period(self._fragment(self._v(answers, "security.incident_protocol"))), "REGLA PARTICULAR: ")
        steps = [
            ("Detección y reporte", "Registrar fecha, fuente, sistema, persona que detecta, información potencialmente afectada y canal de escalamiento."),
            ("Contención", "Limitar propagación, suspender accesos comprometidos, preservar continuidad y evitar alteraciones innecesarias."),
            ("Preservación", "Conservar logs, imágenes, archivos, hashes, comunicaciones y cadena de custodia; no borrar ni modificar evidencia relevante."),
            ("Evaluación", "Determinar alcance, categorías, titulares, secretos, terceros, jurisdicciones, impacto, causa y obligaciones de notificación."),
            ("Comunicación", "Coordinar comunicaciones internas, a la otra parte, titulares, proveedores, aseguradores y autoridades cuando corresponda, evitando especulación."),
            ("Erradicación y recuperación", "Corregir vulnerabilidad, restaurar servicios, rotar credenciales, verificar integridad y monitorear recurrencia."),
            ("Cierre", "Documentar causa raíz, decisiones, costos, acciones correctivas, lecciones y evidencia de implementación."),
        ]
        for i, (h, b) in enumerate(steps, 1): self._clause(doc, i, h, b)
        self._section(doc, "Registro inicial")
        self._table(doc, ["Campo", "Detalle"], [[x, ""] for x in ["Fecha y hora", "Reportante", "Sistema", "Información afectada", "Accesos", "Medidas inmediatas", "Evidencia preservada", "Responsable del incidente"]])
        doc.save(target)

    def _disclosure_act(self, answers, target):
        doc = Document(); self._set_doc_styles(doc); self._compact_styles(doc)
        self._annex_header(doc, answers, "ACT-EM4-DISC-001", "Acta de entrega y recepción de información")
        self._table(doc, ["No.", "Fecha", "Información / soporte", "Clasificación", "Revelador", "Receptor", "Medio", "Hash / evidencia"], [[i, "", "", "", "", "", "", ""] for i in range(1, 7)])
        self._paragraph(doc, "La recepción acredita la entrega descrita, no la exactitud, suficiencia o titularidad de todo el contenido. La Parte Receptora confirma que conoce la finalidad, clasificación, custodios y controles aplicables.")
        doc.save(target)

    def _closure_act(self, answers, target):
        doc = Document(); self._set_doc_styles(doc)
        self._annex_header(doc, answers, "ACT-EM4-CLOSE-001", "Acta de devolución, eliminación y cierre")
        self._table(doc, ["Componente", "Acción", "Evidencia", "Responsable", "Fecha", "Salvedad"], [
            ["Documentos y soportes físicos", "", "", "", "", ""],
            ["Archivos y copias de trabajo", "", "", "", "", ""],
            ["Repositorios y cuentas", "", "", "", "", ""],
            ["Credenciales, llaves y secretos", "", "", "", "", ""],
            ["Datos personales", "", "", "", "", ""],
            ["Copias de respaldo inalterables", "", "", "", "", ""],
            ["Materiales preexistentes", "", "", "", "", ""],
            ["OSS y componentes de terceros", "", "", "", "", ""],
            ["Incidentes o investigaciones abiertas", "", "", "", "", ""],
        ])
        self._paragraph(doc, "La firma acredita únicamente las acciones y salvedades registradas. No constituye paz y salvo general, renuncia de derechos, aceptación de daños no probados ni extinción de obligaciones que deban sobrevivir.")
        doc.save(target)

    def _annex(self, doc_id, answers, target):
        dispatch = {
            "ANX-EM4-INFO-001": self._info_annex,
            "ANX-EM4-REL-001": self._relationship_annex,
            "ANX-EM4-IP-001": self._ip_annex,
            "ANX-EM4-DATA-001": self._data_annex,
            "ANX-EM4-AI-001": self._ai_annex,
            "ANX-EM4-OSS-001": self._oss_annex,
            "PRO-EM4-INC-001": self._incident_protocol,
            "ACT-EM4-DISC-001": self._disclosure_act,
            "ACT-EM4-CLOSE-001": self._closure_act,
        }
        if doc_id not in dispatch:
            raise ValueError(f"Documento no soportado: {doc_id}")
        dispatch[doc_id](answers, target)

    def render_documents(self, answers, target_folder):
        evaluation = self.evaluator.evaluate(answers)
        if evaluation.get("blocked"):
            messages = "; ".join(x["message"] for x in evaluation.get("findings", []) if x.get("severity") == "blocker")
            raise ValueError("El expediente contiene bloqueos jurídicos: " + messages)
        if evaluation.get("missing_fields"):
            missing = ", ".join(x["label"] for x in evaluation["missing_fields"])
            raise ValueError("Faltan datos esenciales: " + missing)
        target_folder = Path(target_folder)
        target_folder.mkdir(parents=True, exist_ok=False)
        contract = target_folder / "CO-EM-004_Acuerdo_Confidencialidad_PI.docx"
        self._contract(answers, evaluation, contract)
        generated = [{"id": "DOC-EM4-NDA-001", "filename": contract.name}]
        for doc_id in evaluation.get("documents", []):
            if doc_id == "DOC-EM4-NDA-001":
                continue
            target = target_folder / f"{doc_id}.docx"
            self._annex(doc_id, answers, target)
            generated.append({"id": doc_id, "filename": target.name})
        hashes, unresolved = {}, []
        for item in generated:
            path = target_folder / item["filename"]
            text = self._extract_text(path)
            if UNRESOLVED_PATTERN.search(text):
                unresolved.append(item["id"])
            hashes[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()
        if unresolved:
            raise ValueError("Se detectaron variables o valores centinela sin resolver: " + ", ".join(unresolved))
        return evaluation, generated, hashes

    def generate(self, answers, actor=None):
        generation_id = "COEM004-" + uuid.uuid4().hex[:12].upper()
        folder = self.output_dir / generation_id
        folder.mkdir(parents=True, exist_ok=False)
        documents_dir = folder / "documents" / "revision-0001"
        evaluation, generated, hashes = self.render_documents(answers, documents_dir)
        manifest = {
            "generation_id": generation_id,
            "product_id": "CO-EM-004",
            "version": self.VERSION,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "actor": {"id": (actor or {}).get("id"), "role": (actor or {}).get("role")},
            "readiness": evaluation.get("readiness"),
            "requires_professional_review": evaluation.get("professional_review_required", False),
            "review_requirements": evaluation.get("review_requirements", []),
            "documents": generated,
            "document_folder": "documents/revision-0001",
            "hashes": hashes,
            "selected_blocks": evaluation.get("blocks", []),
            "unresolved_variables": 0,
            "status": "draft_generated",
            "current_revision": 1,
            "legal_approval": {"status": "pending"},
            "qa_approval": {"status": "pending"},
            "released": False,
        }
        (folder / "answers.json").write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self._package(folder, self.output_dir / f"{generation_id}.zip")
        manifest["package_filename"] = package.name
        manifest["package_sha256"] = hashlib.sha256(package.read_bytes()).hexdigest()
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return manifest

    @staticmethod
    def _package(folder, target):
        with ZipFile(target, "w", ZIP_DEFLATED) as zf:
            for path in sorted(Path(folder).rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        return target

    def package_path(self, generation_id):
        if not GEN_RE.fullmatch(generation_id or ""):
            return None
        path = self.output_dir / f"{generation_id}.zip"
        return path if path.is_file() else None
