from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt


UNRESOLVED_PATTERN = re.compile(r"\{\{[^{}]+\}\}|\b(?:NULL|undefined)\b", re.I)


class CoLa002DocumentFactoryV239:
    """Generate a deterministic DOCX package for the canonical CO-LA-002 flow.

    This factory is additive: it consumes the v2.38 evaluation result and does not
    replace the mature v2.24 generator. It refuses blocked or incomplete cases.
    """

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.output_dir = self.root / "data" / "generated" / "co-la-002-v239"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _v(data: dict, path: str, default=""):
        cur = data
        for part in path.split("."):
            if not isinstance(cur, dict) or part not in cur:
                return default
            cur = cur[part]
        return cur if cur is not None else default

    @staticmethod
    def _money(value):
        try:
            return "$" + f"{int(float(value)):,}".replace(",", ".")
        except (TypeError, ValueError):
            return str(value or "")


    @staticmethod
    def _date_es(value):
        if not value:
            return ""
        try:
            from datetime import date
            parsed = date.fromisoformat(str(value)[:10])
            months = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
            return f"{parsed.day} de {months[parsed.month-1]} de {parsed.year}"
        except (TypeError, ValueError):
            return str(value)

    @staticmethod
    def _label(value, mapping):
        return mapping.get(value, str(value or ""))

    @staticmethod
    def _set_doc_styles(doc: Document):
        sec = doc.sections[0]
        sec.top_margin = Cm(2.4)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.6)
        sec.right_margin = Cm(2.6)
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for name in ("Title", "Heading 1", "Heading 2"):
            style = doc.styles[name]
            style.font.name = "Aptos Display"
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Heading 1"].font.size = Pt(12)
        doc.styles["Heading 2"].font.size = Pt(11)

    @staticmethod
    def _title(doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(14)
        return p

    @staticmethod
    def _section(doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11.5)
        return p

    @staticmethod
    def _clause(doc: Document, ordinal: int, heading: str, body: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"CLÁUSULA {ordinal}. {heading.upper()}: ")
        r.bold = True
        p.add_run(body.strip())
        return p

    @staticmethod
    def _bullet_list(doc: Document, items):
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(str(item))

    def _identity(self, answers):
        employer_type = self._v(answers, "employer.type")
        if employer_type == "legal_person":
            employer = self._v(answers, "employer.legalName")
        else:
            employer = self._v(answers, "employer.naturalPersonFullName") or self._v(answers, "employer.legalName")
        return {
            "employer": employer,
            "employer_id": self._v(answers, "employer.identificationNumber"),
            "signatory": self._v(answers, "employerSignatory.fullName"),
            "signatory_capacity": self._v(answers, "employerSignatory.positionOrCapacity"),
            "worker": self._v(answers, "worker.fullName"),
            "worker_id": self._v(answers, "worker.identificationNumber"),
            "role": self._v(answers, "role.jobTitle"),
        }

    def _contract(self, answers: dict, evaluation: dict, target: Path):
        ident = self._identity(answers)
        doc = Document()
        self._set_doc_styles(doc)
        self._title(doc, "Contrato individual de trabajo a término indefinido")

        appearance = (
            f"Entre {ident['employer']}, identificado(a) con {ident['employer_id']}, "
            f"representado(a) para este acto por {ident['signatory'] or 'su titular'}, "
            f"quien actúa como {ident['signatory_capacity'] or 'empleador'}, en adelante EL EMPLEADOR; "
            f"y {ident['worker']}, identificado(a) con {ident['worker_id']}, en adelante EL TRABAJADOR, "
            "se celebra el presente contrato individual de trabajo a término indefinido."
        )
        p = doc.add_paragraph(appearance)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self._section(doc, "Consideraciones")
        considerations = [
            f"EL EMPLEADOR requiere la prestación personal de servicios en el cargo de {ident['role']}.",
            "LAS PARTES desean regular una relación laboral a término indefinido bajo la legislación colombiana.",
            "Los anexos y autorizaciones solo producirán los efectos expresamente identificados y legalmente procedentes.",
        ]
        ordinals = ["PRIMERA", "SEGUNDA", "TERCERA"]
        for o, text in zip(ordinals, considerations):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"{o}: "); r.bold = True; p.add_run(text)

        self._section(doc, "Cláusulas")
        clauses = []
        clauses.append(("Objeto", f"EL TRABAJADOR prestará personalmente sus servicios a EL EMPLEADOR en el cargo de {ident['role']}, bajo subordinación laboral legítima y a cambio de la remuneración pactada."))
        purpose = self._v(answers, "role.purpose", "cumplir las funciones propias del cargo")
        clauses.append(("Cargo y finalidad", f"EL TRABAJADOR desempeñará el cargo de {ident['role']}. Su propósito principal consiste en {purpose}."))

        functions = self._v(answers, "role.essentialFunctions", [])
        if isinstance(functions, str):
            functions = [x.strip() for x in functions.split("\n") if x.strip()]
        placement = self._v(answers, "role.functionsPlacement")
        if placement == "full_in_contract" and functions:
            clauses.append(("Funciones", "EL TRABAJADOR desarrollará las funciones que se relacionan a continuación:"))
        else:
            clauses.append(("Funciones", "Las funciones específicas constan en el Anexo de perfil y funciones vigente, que forma parte integral del expediente contractual."))

        workplace = self._v(answers, "work.mainWorkplace", "el lugar definido por LAS PARTES")
        modality = self._v(answers, "work.modality", "onsite")
        modality_label = {
            "onsite": "presencial", "telework_hybrid": "teletrabajo híbrido",
            "telework_autonomous": "teletrabajo autónomo", "telework_mobile": "teletrabajo móvil",
            "remote_work": "trabajo remoto"
        }.get(modality, modality)
        clauses.append(("Lugar y modalidad", f"El servicio se prestará principalmente en {workplace}, bajo modalidad {modality_label}, con sujeción a los anexos aplicables."))
        clauses.append(("Duración e inicio", f"El contrato es a término indefinido y la prestación efectiva inicia el {self._date_es(self._v(answers, 'work.actualStartDate'))}."))

        weekly = self._v(answers, "schedule.weeklyHours")
        schedule_type = self._v(answers, "schedule.type")
        schedule_label = self._label(schedule_type, {"fixed": "fijo", "flexible": "flexible", "rotating": "turnos rotativos", "special": "ciclo especial"})
        clauses.append(("Jornada", f"La jornada ordinaria será de {weekly} horas semanales, bajo esquema {schedule_label}, sin perjuicio de los límites, descansos y recargos vigentes."))
        clauses.append(("Desconexión laboral", "Fuera de la jornada y de disponibilidades expresamente programadas, EL TRABAJADOR no estará obligado a atender comunicaciones laborales, salvo las excepciones legalmente procedentes."))

        salary = self._money(self._v(answers, "compensation.baseSalary"))
        salary_type = self._v(answers, "compensation.salaryType", "ordinary")
        salary_type_label = self._label(salary_type, {"ordinary": "ordinario", "integral": "integral"})
        clauses.append(("Salario", f"EL EMPLEADOR pagará un salario mensual {salary_type_label} de {salary} M/CTE, con la periodicidad y mediante el medio acordados. Los conceptos que legalmente deban reconocerse separadamente no se entienden incluidos."))
        clauses.append(("Prestaciones y seguridad social", "EL TRABAJADOR tendrá derecho a prestaciones sociales, vacaciones, descansos, licencias, recargos y afiliaciones al Sistema de Seguridad Social Integral conforme a la ley."))
        clauses.append(("Obligaciones del empleador", "EL EMPLEADOR pagará oportunamente, proporcionará medios adecuados, protegerá la dignidad y la intimidad, implementará el SG-SST y respetará jornada, descansos, debido proceso y protección de datos."))
        clauses.append(("Obligaciones del trabajador", "EL TRABAJADOR ejecutará diligentemente sus funciones, observará instrucciones legítimas, protegerá la información, utilizará adecuadamente los activos y reportará incidentes, riesgos y novedades relevantes."))
        clauses.append(("Seguridad y salud en el trabajo", "LAS PARTES cumplirán las obligaciones que les correspondan dentro del SG-SST. La firma del contrato o de anexos no sustituye capacitación, controles, aptitud ni permisos de trabajo."))
        clauses.append(("Confidencialidad", "EL TRABAJADOR protegerá la información reservada conocida por razón de sus funciones. Esta obligación no impedirá revelaciones exigidas por ley, defensa legítima de derechos o denuncias ante autoridades competentes."))
        clauses.append(("Datos personales", "EL EMPLEADOR tratará los datos personales para finalidades legítimas de la relación laboral. Las autorizaciones de imagen, biometría y geolocalización serán independientes cuando resulten aplicables."))
        clauses.append(("Seguridad de la información", "EL TRABAJADOR utilizará sistemas autorizados, protegerá credenciales, respetará el mínimo privilegio y reportará incidentes de seguridad."))
        clauses.append(("Propiedad intelectual", "Los resultados protegibles creados en cumplimiento directo de las funciones o de encargos específicos se regirán por la ley y los pactos aplicables, preservando derechos morales, materiales preexistentes, licencias de terceros y conocimiento general."))
        clauses.append(("Equipos y activos", "Los activos entregados se documentarán en el acta correspondiente. El desgaste normal no genera responsabilidad y ningún valor de referencia constituye deuda o descuento automático."))
        clauses.append(("Procedimiento disciplinario", "Toda sanción disciplinaria respetará información clara de los hechos, acceso a pruebas, defensa, contradicción, imparcialidad, proporcionalidad y decisión motivada."))
        clauses.append(("Terminación", "El contrato podrá terminar por las causas y procedimientos previstos en la ley. Las protecciones especiales y autorizaciones previas deberán verificarse antes de cualquier decisión."))
        clauses.append(("Entrega del cargo", "Al terminar o producirse una transición legítima, EL TRABAJADOR realizará una entrega razonable de asuntos, documentos, activos y accesos institucionales, sin incluir contraseñas en texto plano ni renunciar a derechos."))
        clauses.append(("Integridad y modificaciones", "El contrato y sus anexos vigentes contienen las condiciones aplicables. Toda modificación material deberá constar en una nueva revisión trazable y no podrá desconocer derechos mínimos."))

        for i, (heading, body) in enumerate(clauses, 1):
            self._clause(doc, i, heading, body)
            if heading == "Funciones" and placement == "full_in_contract" and functions:
                self._bullet_list(doc, functions)

        self._section(doc, "Documentos relacionados")
        for doc_id in evaluation["documents"]:
            if doc_id != "DOC-LA-CONTRACT-001":
                doc.add_paragraph(doc_id, style="List Bullet")

        self._section(doc, "Firmas")
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        values = [
            ("EL EMPLEADOR", "EL TRABAJADOR"),
            (ident["signatory"] or ident["employer"], ident["worker"]),
            ("Firma: ____________________", "Firma: ____________________"),
        ]
        for row, vals in zip(table.rows, values):
            for cell, val in zip(row.cells, vals):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val); r.bold = vals == values[0]

        doc.core_properties.title = "Contrato individual de trabajo a término indefinido"
        doc.core_properties.subject = "CO-LA-002 v2.39"
        doc.core_properties.comments = "Generado por LegalAIZ.it. Requiere revisión jurídica y QA antes de aprobación."
        doc.save(target)

    def _annex(self, doc_id: str, answers: dict, target: Path):
        ident = self._identity(answers)
        doc = Document(); self._set_doc_styles(doc)
        titles = {
            "ANX-LA-FUN-001": "Anexo de perfil y funciones",
            "ANX-LA-MOD-001": "Anexo de modalidad no presencial",
            "ANX-LA-VAR-001": "Anexo de remuneración variable",
            "ACT-LA-EQP-001": "Acta de entrega de equipos y herramientas",
            "ANX-LA-DIS-001": "Anexo de disponibilidad",
            "ANX-LA-RIE-001": "Anexo de riesgos especiales",
            "AUT-LA-IMG-001": "Autorización para uso de imagen y voz",
            "AUT-LA-BIO-001": "Autorización para tratamiento de datos biométricos",
            "AUT-LA-GEO-001": "Información y autorización de geolocalización",
            "ACT-LA-ENT-001": "Acta de entrega del cargo",
            "ANX-LA-VPR-001": "Acuerdo de uso de vehículo propio",
            "COM-LA-TCH-001": "Comunicación de habilitación temporal de trabajo en casa",
        }
        self._title(doc, titles.get(doc_id, doc_id))
        p = doc.add_paragraph(
            f"Documento relacionado con el contrato celebrado entre {ident['employer']} y {ident['worker']}, "
            f"para el cargo de {ident['role']}. Código: {doc_id}."
        ); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if doc_id == "ANX-LA-FUN-001":
            self._section(doc, "Propósito del cargo")
            doc.add_paragraph(self._v(answers, "role.purpose", "Propósito pendiente de validación."))
            self._section(doc, "Funciones esenciales")
            funcs = self._v(answers, "role.essentialFunctions", []) or []
            if isinstance(funcs, str): funcs = [x for x in funcs.split("\n") if x.strip()]
            self._bullet_list(doc, funcs or ["Funciones definidas en el perfil ocupacional seleccionado y pendientes de validación final."])
        elif doc_id == "ANX-LA-MOD-001":
            self._clause(doc, 1, "Modalidad", f"La modalidad acordada es {self._v(answers, 'work.modality')}. El lugar remoto autorizado es {self._v(answers, 'remoteWork.authorizedLocation', 'pendiente de confirmación')}. ")
            self._clause(doc, 2, "Desconexión y privacidad", "La modalidad no autoriza disponibilidad continua, vigilancia audiovisual permanente ni ingreso al domicilio sin procedimiento válido.")
        elif doc_id == "ANX-LA-VAR-001":
            self._clause(doc, 1, "Esquema variable", self._v(answers, "variableCompensation.summary", "Pendiente de completar hecho generador, base, fórmula, fuente, período y pago."))
        elif doc_id == "ACT-LA-EQP-001":
            self._section(doc, "Activos")
            items = self._v(answers, "assets.items", []) or []
            if isinstance(items, str): items = [items]
            self._bullet_list(doc, items or [self._v(answers, "assets.summary", "Activos pendientes de inventario.")])
            doc.add_paragraph("El valor de referencia no constituye deuda automática ni autoriza descuento directo.")
        elif doc_id == "ANX-LA-DIS-001":
            self._clause(doc, 1, "Eventos de disponibilidad", self._v(answers, "availability.eventsSummary", "Eventos pendientes de delimitación."))
            self._clause(doc, 2, "Límites", "La disponibilidad no es permanente, no se presume por residencia y debe respetar descansos, licencias, incapacidades y vacaciones.")
        elif doc_id == "ANX-LA-RIE-001":
            self._clause(doc, 1, "Riesgos y controles", self._v(answers, "specialConditions.riskSummary", "La identificación detallada deberá ser completada y aprobada por SST antes de publicación."))
        elif doc_id.startswith("AUT-LA-"):
            self._clause(doc, 1, "Carácter independiente", "Esta autorización es facultativa, específica y separada del contrato. La negativa no autoriza represalias ni afecta derechos mínimos.")
            table = doc.add_table(rows=2, cols=1); table.style = "Table Grid"
            table.cell(0,0).text = "☐ AUTORIZO"
            table.cell(1,0).text = "☐ NO AUTORIZO"
        elif doc_id == "ANX-LA-VPR-001":
            self._clause(doc, 1, "Vehículo propio", "El vehículo, sus documentos, seguros, gastos, mantenimiento, accidentes y terminación del acuerdo deberán identificarse expresamente. No se trasladarán automáticamente todos los costos al trabajador.")
        elif doc_id == "ACT-LA-ENT-001":
            self._clause(doc, 1, "Entrega", "Se relacionarán asuntos pendientes, archivos, activos, accesos institucionales y observaciones. No se consignarán contraseñas en texto plano y la firma no constituye paz y salvo general.")
        elif doc_id == "COM-LA-TCH-001":
            self._clause(doc, 1, "Habilitación temporal", "La habilitación temporal no modifica la modalidad base y finaliza al desaparecer la causa o cumplirse el plazo comunicado.")

        self._section(doc, "Firmas")
        doc.add_paragraph(f"EL EMPLEADOR: {ident['signatory'] or ident['employer']}    Firma: ____________________")
        doc.add_paragraph(f"EL TRABAJADOR: {ident['worker']}    Firma: ____________________")
        doc.core_properties.subject = f"{doc_id} - CO-LA-002 v2.39"
        doc.save(target)

    @staticmethod
    def _extract_text(docx_path: Path):
        doc = Document(docx_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)

    def generate(self, answers: dict, actor: dict | None = None):
        evaluation = self.evaluator.evaluate(answers)
        if evaluation["blocked"]:
            raise ValueError("El expediente contiene bloqueos jurídicos y no puede generar documentos.")
        if evaluation["missing_fields"]:
            missing = ", ".join(x["label"] for x in evaluation["missing_fields"])
            raise ValueError(f"Faltan datos esenciales: {missing}.")

        generation_id = "COLA002-" + uuid.uuid4().hex[:12].upper()
        folder = self.output_dir / generation_id
        folder.mkdir(parents=True, exist_ok=False)
        contract = folder / "CO-LA-002_Contrato_Indefinido.docx"
        self._contract(answers, evaluation, contract)

        generated = [{"id": "DOC-LA-CONTRACT-001", "filename": contract.name}]
        for doc_id in evaluation["documents"]:
            if doc_id == "DOC-LA-CONTRACT-001":
                continue
            target = folder / f"{doc_id}.docx"
            self._annex(doc_id, answers, target)
            generated.append({"id": doc_id, "filename": target.name})

        unresolved = []
        hashes = {}
        for item in generated:
            path = folder / item["filename"]
            text = self._extract_text(path)
            if UNRESOLVED_PATTERN.search(text):
                unresolved.append(item["id"])
            hashes[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()
        if unresolved:
            raise ValueError("Se detectaron variables o valores centinela sin resolver: " + ", ".join(unresolved))

        manifest = {
            "generation_id": generation_id,
            "product_id": "CO-LA-002",
            "version": "2.39",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "actor": {"id": (actor or {}).get("id"), "role": (actor or {}).get("role")},
            "readiness": evaluation["readiness"],
            "requires_professional_review": bool(evaluation["review_requirements"] or evaluation["warnings"]),
            "review_requirements": evaluation["review_requirements"],
            "documents": generated,
            "hashes": hashes,
            "selected_blocks": evaluation["blocks"],
            "unresolved_variables": 0,
            "status": "draft_generated",
            "legal_approval": "pending",
            "qa_approval": "pending",
        }
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self.output_dir / f"{generation_id}.zip"
        with ZipFile(package, "w", ZIP_DEFLATED) as zf:
            for path in sorted(folder.iterdir()):
                zf.write(path, arcname=path.name)
        manifest["package_filename"] = package.name
        manifest["package_sha256"] = hashlib.sha256(package.read_bytes()).hexdigest()
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return manifest

    def package_path(self, generation_id: str):
        if not re.fullmatch(r"COLA002-[A-F0-9]{12}", generation_id or ""):
            return None
        path = self.output_dir / f"{generation_id}.zip"
        return path if path.is_file() else None
