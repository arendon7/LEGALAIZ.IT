from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from co_em_004_document_factory_v246 import CoEm004DocumentFactoryV246
from co_em_003_document_factory_v243 import UNRESOLVED_PATTERN


GEN_RE = re.compile(r"COAR001-[A-F0-9]{12}")


class CoAr001DocumentFactoryV249(CoEm004DocumentFactoryV246):
    """Fábrica DOCX para arrendamiento de vivienda urbana."""

    VERSION = "2.49"

    def __init__(self, root: Path, evaluator):
        self.root = Path(root)
        self.evaluator = evaluator
        self.output_dir = self.root / "data" / "generated" / "co-ar-001-v249"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _set_doc_styles(doc: Document):
        sec = doc.sections[0]
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.55)
        sec.right_margin = Cm(2.55)
        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.1)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for style_name in ("Title", "Heading 1", "Heading 2"):
            style = doc.styles[style_name]
            style.font.name = "Aptos Display"
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Heading 1"].font.size = Pt(12)
        doc.styles["Heading 2"].font.size = Pt(11)
        footer = sec.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("LegalAIZ.it | CO-AR-001 | Documento sujeto a revisión jurídica y QA | Página ")
        run.font.size = Pt(8)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer._p.append(fld)

    ENUM_LABELS = {
        "natural": "persona natural", "legal": "persona jurídica", "owner": "propietario acreditado",
        "agent": "apoderado o mandatario", "usufructuary": "usufructuario autorizado",
        "other_authorized": "otro autorizado", "pending": "soporte pendiente", "absent": "sin soporte",
        "individual": "individual", "joint": "mancomunado", "shared": "vivienda compartida", "pension": "pensión",
        "apartment": "apartamento", "house": "casa", "room": "habitación", "partial": "parte de inmueble", "other": "otro",
        "residential": "vivienda exclusivamente", "residential_remote": "vivienda y trabajo remoto sin atención al público",
        "residential_professional": "vivienda y actividad profesional sin transformación del uso",
        "mixed_or_commercial": "uso mixto o comercial", "none": "ninguna", "sublease_requested": "subarriendo o cesión solicitada",
        "tourism": "hospedaje turístico o por días", "platform": "plataforma de alojamiento", "unknown": "por verificar",
        "fit": "apto y funcional", "minor_defects": "defectos menores identificados",
        "material_defects": "defectos materiales pendientes", "unsafe": "riesgo grave de sanidad o seguridad",
        "detailed": "inventario detallado", "photos": "inventario fotográfico", "video": "video de entrega",
        "combined": "inventario, fotografías y video", "monthly": "mensual",
        "legal_ipc": "hasta el IPC legalmente aplicable después de doce meses",
        "lower_than_ipc": "porcentaje inferior al IPC", "fixed_or_other": "fórmula fija u otra",
        "tenant": "arrendatario", "landlord": "arrendador", "included": "incluida en el canon",
        "special_agreement": "acuerdo especial sujeto a revisión", "special": "distribución especial",
        "yes": "sí", "no": "no", "to_define": "por definir",
        "codebtor": "codeudor solidario", "guarantor": "fiador", "policy": "póliza de arrendamiento",
    }

    @classmethod
    def _plain(cls, value, default=""):
        if isinstance(value, str) and value.strip() in cls.ENUM_LABELS:
            return cls.ENUM_LABELS[value.strip()]
        return super(CoAr001DocumentFactoryV249, cls)._plain(value, default)


    @classmethod
    def _fragment(cls, value, fallback="", lower_initial=False):
        """Normaliza contenido dinámico insertado dentro de una oración."""
        text = cls._plain(value, fallback).strip()
        text = re.sub(r"[\s\.;,:]+$", "", text)
        if lower_initial and text and text[0].isupper() and not text[:2].isupper():
            text = text[0].lower() + text[1:]
        return text

    @classmethod
    def _renewal_label(cls, value):
        labels = {
            True: "prórroga conforme a las reglas legales aplicables",
            False: "prórroga únicamente cuando exista acuerdo válido o así lo determine la ley",
            "yes": "prórroga conforme a las reglas legales aplicables",
            "no": "prórroga únicamente cuando exista acuerdo válido o así lo determine la ley",
            "automatic": "prórroga automática en los eventos y condiciones permitidos por la ley",
            "legal": "prórroga conforme a las reglas legales aplicables",
        }
        return labels.get(value, cls._fragment(value, "las reglas legales aplicables"))

    @classmethod
    def _sublease_label(cls, value):
        labels = {
            "none": "no se contempla cesión, subarriendo ni uso turístico",
            "sublease_requested": "se ha solicitado evaluar una cesión o subarriendo, sujeto a autorización expresa",
            "tourism": "se ha informado una intención de alojamiento turístico, sujeta a revisión jurídica y regulatoria",
            "platform": "se ha informado una intención de uso en plataformas de alojamiento, sujeta a revisión jurídica y regulatoria",
            "unknown": "la situación debe ser verificada antes de aprobar el contrato",
        }
        return labels.get(value, cls._fragment(value, "no se autoriza salvo consentimiento escrito y cumplimiento del régimen aplicable"))

    @classmethod
    def _display(cls, value):
        if isinstance(value, str):
            return cls.ENUM_LABELS.get(value.strip(), value.strip())
        if isinstance(value, bool):
            return "Sí" if value else "No"
        if isinstance(value, list):
            return "; ".join(cls._display(x) for x in value)
        if isinstance(value, dict):
            return "; ".join(f"{cls._key_label(k)}: {cls._display(v)}" for k, v in value.items() if v not in (None, "", [], {}))
        if value is None:
            return ""
        return str(value)

    @classmethod
    def _table(cls, doc, headers, rows, widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr = table.rows[0]
        tr_pr = hdr._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); tr_pr.append(th); tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, label in zip(hdr.cells, headers):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(label)); run.bold = True; run.font.size = Pt(8.2)
        for row in rows:
            ro = table.add_row(); ro._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for cell, value in zip(ro.cells, row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.LEFT; para.paragraph_format.space_after = Pt(0)
                run = para.add_run(cls._display(value)); run.font.size = Pt(8.2)
        return table

    @staticmethod
    def _keep_table(table):
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:cantSplit"))
        return table

    @staticmethod
    def _signatures(doc, landlord, tenant, extra=None):
        rows = [
            ["EL ARRENDADOR", "EL ARRENDATARIO"],
            [landlord["name"], tenant["name"]],
            [landlord["id"], tenant["id"]],
            ["Firma: ______________________________", "Firma: ______________________________"],
        ]
        if extra:
            rows.extend(extra)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row_values in enumerate(rows):
            row_obj = table.add_row()
            if row_index == 0:
                tr_pr = row_obj._tr.get_or_add_trPr()
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)
            cells = row_obj.cells
            for i, value in enumerate(row_values):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(value))
                if row_index == 0:
                    r.bold = True
                r.font.size = Pt(9)
        CoAr001DocumentFactoryV249._keep_table(table)
        return table

    def _party_ar(self, answers, prefix, fallback):
        ident = self._as_dict(self._v(answers, f"{prefix}.identification"))
        signatory = self._as_dict(self._v(answers, f"{prefix}.signatory"))
        name = self._plain(ident.get("name"), fallback)
        ident_no = self._plain(ident.get("id_number"), "identificación registrada")
        address = self._plain(ident.get("address"), "domicilio registrado")
        email = self._plain(ident.get("email"), "correo contractual registrado")
        phone = self._plain(ident.get("phone"), "teléfono registrado")
        ptype = ident.get("type")
        sign_name = self._plain(signatory.get("name"), name)
        sign_id = self._plain(signatory.get("id_number"), ident_no)
        capacity = self._plain(signatory.get("capacity"), "persona autorizada")
        authority = self._plain(signatory.get("authority_source"))
        if ptype in {"legal", "legal_person", "company"}:
            text = (
                f"{name}, persona jurídica identificada con NIT {ident_no}, con domicilio en {address}, "
                f"representada para este acto por {sign_name}, identificado(a) con {sign_id}, quien actúa como {capacity}"
            )
            if authority:
                text += f", con fundamento en {authority}"
        else:
            text = f"{name}, persona natural identificada con {ident_no}, domiciliada en {address}"
        return {
            "name": name, "id": ident_no, "address": address, "email": email, "phone": phone,
            "signatory": sign_name, "signatory_id": sign_id, "capacity": capacity,
            "authority": authority, "text": text,
        }

    def _property(self, answers):
        ident = self._as_dict(self._v(answers, "property.identification"))
        included = self._as_dict(self._v(answers, "property.included_units"))
        return {
            "address": self._plain(ident.get("address"), "dirección registrada"),
            "municipality": self._plain(ident.get("municipality"), "municipio registrado"),
            "registration": self._plain(ident.get("registration"), "matrícula inmobiliaria registrada"),
            "cadastral": self._plain(ident.get("cadastral_id"), "no informada"),
            "type": self._plain(self._v(answers, "property.type"), "vivienda urbana"),
            "private_area": self._plain(included.get("private_area"), "unidad privada descrita"),
            "parking": self._plain(included.get("parking"), "no incluido"),
            "storage": self._plain(included.get("storage"), "no incluido"),
            "other": self._plain(included.get("other"), "sin otros bienes o usos incluidos"),
        }

    @staticmethod
    def _configuration_label(value):
        return {
            "individual": "individual",
            "joint": "con pluralidad de arrendatarios",
            "shared": "con ocupación compartida",
            "pension": "de pensión",
        }.get(str(value or ""), "individual")

    @staticmethod
    def _adjustment_label(value):
        return {
            "legal_ipc": "reajuste legal hasta el IPC del año calendario anterior, una vez transcurridos doce meses bajo el mismo precio",
            "none": "sin reajuste automático, sin perjuicio de la ley y de un acuerdo posterior válido",
            "fixed_or_other": "regla especial sujeta al límite legal y a revisión jurídica previa",
        }.get(str(value or ""), "reajuste conforme a la Ley 820 de 2003")

    def _contract(self, answers, evaluation, target):
        doc = Document(); self._set_doc_styles(doc)
        landlord = self._party_ar(answers, "landlord", "EL ARRENDADOR")
        tenant = self._party_ar(answers, "tenant", "EL ARRENDATARIO")
        prop = self._property(answers)
        rent = self._money(self._v(answers, "rent.amount"), "COP")
        payment = self._as_dict(self._v(answers, "rent.payment"))
        values = self._as_dict(self._v(answers, "rent.values"))
        term = self._as_dict(self._v(answers, "term.rules"))
        maintenance = self._as_dict(self._v(answers, "maintenance.rules"))
        notifications = self._as_dict(self._v(answers, "notifications.channels"))
        utilities = self._as_dict(self._v(answers, "charges.utilities"))
        admin = self._as_dict(self._v(answers, "charges.administration"))
        extra = self._as_dict(self._v(answers, "charges.additional_services"))
        guarantee = self._as_dict(self._v(answers, "guarantee.details"))
        ph = self._as_dict(self._v(answers, "property.ph_details"))
        screening = self._as_dict(self._v(answers, "data.screening"))
        pending_repairs = self._as_dict(self._v(answers, "condition.repairs_pending"))

        self._title(doc, "CONTRATO DE ARRENDAMIENTO DE VIVIENDA URBANA")
        self._paragraph(doc, (
            f"Entre los suscritos, a saber: de una parte, {landlord['text']}, quien en adelante se denominará EL ARRENDADOR; "
            f"y de otra, {tenant['text']}, quien en adelante se denominará EL ARRENDATARIO; se celebra el presente contrato "
            "de arrendamiento de vivienda urbana, regido por la Ley 820 de 2003, las normas civiles compatibles y las cláusulas siguientes."
        ))
        self._section(doc, "CONSIDERACIONES")
        considerations = [
            f"Que EL ARRENDADOR declara contar con título, representación o autorización suficiente para conceder el goce del inmueble ubicado en {prop['address']}, {prop['municipality']}, identificado con matrícula inmobiliaria {prop['registration']}.",
            "Que EL ARRENDATARIO manifiesta haber recibido información sobre la identificación, destinación, estado material, servicios, cargas y documentos del inmueble, sin que esta manifestación sustituya el inventario ni la obligación de entrega en condiciones de habitabilidad.",
            f"Que LAS PARTES acuerdan una modalidad de arrendamiento {self._configuration_label(self._v(answers, 'lease.configuration'))}, destinada exclusivamente a vivienda urbana ordinaria.",
            "Que los anexos, actas, comunicaciones y autorizaciones expresamente identificados forman parte del expediente contractual en su versión vigente y no podrán emplearse para desconocer normas imperativas.",
        ]
        for index, text in enumerate(considerations, 1):
            self._subparagraph(doc, ["PRIMERA", "SEGUNDA", "TERCERA", "CUARTA"][index-1], text)
        self._section(doc, "CLÁUSULAS")

        clauses = []
        clauses.append(("Objeto", f"EL ARRENDADOR concede a EL ARRENDATARIO el goce del inmueble descrito en este contrato para destinarlo exclusivamente a vivienda urbana, y EL ARRENDATARIO se obliga a pagar el canon y cumplir las obligaciones pactadas. El contrato no autoriza explotación turística, comercial, industrial, hotelera ni por plataformas, salvo instrumento distinto jurídicamente procedente."))
        clauses.append(("Identificación del inmueble", f"El inmueble se ubica en {prop['address']}, {prop['municipality']}, corresponde a {prop['type']}, se identifica con matrícula inmobiliaria {prop['registration']} y referencia catastral {prop['cadastral']}. La información detallada, linderos funcionales, unidad privada y soportes se incorporan en el Anexo de identificación del inmueble."))
        clauses.append(("Bienes, anexos y usos incluidos", f"Se incluyen la unidad o área {prop['private_area']}; parqueadero: {prop['parking']}; depósito o cuarto útil: {prop['storage']}; otros bienes o usos: {prop['other']}. Ningún bien, zona común o servicio se entenderá incluido por tolerancia, publicidad o acceso material si no aparece identificado en el expediente."))
        clauses.append(("Destino autorizado", f"El inmueble se destinará a {self._plain(self._v(answers, 'use.destination'), 'residencia permanente')}. EL ARRENDATARIO deberá observar las normas urbanísticas, sanitarias, de convivencia y de propiedad horizontal aplicables, y abstenerse de cambiar la destinación sin autorización escrita y sin verificar su legalidad."))
        clauses.append(("Ocupantes autorizados", f"Podrán residir las siguientes personas: {self._plain(self._v(answers, 'occupants.authorized'), 'las identificadas en el anexo de ocupantes')}. La presencia de ocupantes no los convierte automáticamente en arrendatarios ni libera a quienes hayan firmado el contrato; los cambios materiales deberán informarse y documentarse cuando afecten seguridad, convivencia, capacidad o servicios."))
        clauses.append(("Propiedad horizontal y convivencia", (
            f"El inmueble {'se encuentra' if self._v(answers, 'property.horizontal') is True else 'no se ha informado que se encuentre'} sometido a propiedad horizontal. "
            f"Cuando aplique, EL ARRENDATARIO observará la parte normativa del reglamento, manuales y decisiones oponibles que le hayan sido entregados, especialmente sobre uso, ruido, mascotas, mudanzas, residuos, parqueaderos y zonas comunes. "
            "EL ARRENDADOR conservará a su cargo las obligaciones propias del propietario que no hayan sido trasladadas válidamente."
        )))
        clauses.append(("Entrega material y habitabilidad", f"La entrega se proyecta para el {self._date_es(self._v(answers, 'delivery.date'))}. EL ARRENDADOR entregará el inmueble en estado de servir para la habitación, con servicios esenciales y condiciones razonables de seguridad, salubridad y funcionamiento. El acta de entrega registrará llaves, medidores, soportes, observaciones y reparaciones pendientes."))
        clauses.append(("Inventario y evidencia de estado", f"El inventario se elaborará mediante el método {self._plain(self._v(answers, 'delivery.inventory_method'), 'detallado con evidencia')}. Deberá individualizar ambientes, acabados, equipos, muebles, medidores y defectos visibles, con fotografías o soportes fechados cuando sea posible. Los valores de referencia no constituyen deudas automáticas."))
        clauses.append(("Defectos y reparaciones pendientes", f"Se registran como observaciones o reparaciones pendientes: {self._fragment(self._v(answers, 'condition.defects'), 'ninguna distinta de las consignadas en el inventario')}. Responsable y plazo: {self._fragment(pending_repairs.get('responsible'), 'según acta de entrega')}. La recepción con observaciones no implica renuncia a exigir condiciones de habitabilidad ni aceptación de vicios ocultos."))
        clauses.append(("Mantenimiento y reparaciones locativas", f"EL ARRENDATARIO asumirá las reparaciones locativas derivadas del uso ordinario o de hechos imputables, conforme a la ley y a la regla pactada: {self._fragment(maintenance.get('routine'), 'mantenimiento ordinario razonable')}. EL ARRENDADOR atenderá las reparaciones necesarias no imputables al arrendatario y las estructurales o de habitabilidad que legalmente le correspondan."))
        clauses.append(("Mejoras y adecuaciones", f"Las mejoras, instalaciones o modificaciones se sujetarán a la siguiente regla: {self._fragment(maintenance.get('improvements'), 'requieren acuerdo previo')}. El procedimiento de autorización será: {self._fragment(maintenance.get('authorization'), 'solicitud escrita con alcance, costo y reversibilidad')}. Ninguna autorización se presumirá por silencio, visita o conocimiento informal."))
        clauses.append(("Canon de arrendamiento", f"EL ARRENDATARIO pagará un canon mensual de {rent}. LAS PARTES declaran haber considerado el valor comercial soportado de {self._money(values.get('commercial_value'), 'COP')} y el avalúo catastral informado de {self._money(values.get('cadastral_value'), 'COP')}, con fuente o fecha {self._plain(values.get('source_date'), 'registrada en el expediente')}. El canon deberá mantenerse dentro de los límites legales."))
        clauses.append(("Forma y oportunidad de pago", f"El canon se pagará con periodicidad {self._plain(payment.get('frequency'), 'mensual')}, a más tardar el día {self._plain(payment.get('day'), 'acordado')} de cada período, mediante {self._plain(payment.get('method'), 'el medio autorizado')}, siguiendo estas instrucciones: {self._plain(payment.get('account'), 'las registradas por EL ARRENDADOR')}. Todo pago deberá contar con soporte verificable."))
        clauses.append(("Reajuste", f"El canon se sujetará a {self._adjustment_label(self._v(answers, 'rent.adjustment'))}. El reajuste solo será exigible después de doce meses de ejecución bajo el mismo precio, dentro del máximo legal y mediante comunicación del monto y fecha de aplicación por un canal contractual oponible."))
        clauses.append(("Administración", f"La administración ordinaria informada es {self._money(admin.get('ordinary_amount'), 'COP')}. Responsable de la ordinaria: {self._plain(admin.get('ordinary_responsible'), 'según la distribución pactada')}; responsable de las expensas extraordinarias: {self._plain(admin.get('extraordinary_responsible'), 'EL ARRENDADOR salvo pacto válido distinto')}. Las cuotas, sanciones y soportes deberán discriminarse y no podrán cobrarse dos veces."))
        clauses.append(("Servicios públicos y conectividad", f"El responsable principal será {self._plain(utilities.get('responsible'), 'EL ARRENDATARIO')}. Distribución: {self._plain(utilities.get('distribution'), 'según medidores y facturas')}. Internet: {self._plain(utilities.get('internet'), 'según contratación aplicable')}. Gas u otros: {self._plain(utilities.get('gas_or_other'), 'según disponibilidad y facturación')}. Las partes conservarán soportes de lectura, facturación y pago."))
        clauses.append(("Servicios, cosas o usos adicionales", (
            f"{'Se pactan' if extra.get('exists') is True else 'No se pactan'} servicios, cosas o usos adicionales. "
            f"Descripción: {self._plain(extra.get('description'), 'no aplica')}; valor mensual: {self._money(extra.get('value'), 'COP')}. "
            "Cuando existan, deberán ser identificables, separables y mantenerse dentro del límite legal frente al canon del inmueble."
        )))
        clauses.append(("Garantías permitidas", f"La garantía seleccionada es {self._plain(self._v(answers, 'guarantee.type'), 'ninguna')}. Garante o aseguradora: {self._plain(guarantee.get('party'), 'no aplica')}; identificación o póliza: {self._plain(guarantee.get('id_number'), 'no aplica')}; obligaciones cubiertas: {self._plain(guarantee.get('scope'), 'según anexo')}; vigencia: {self._plain(guarantee.get('validity'), 'según anexo')}. No se exigirá depósito en dinero ni caución real prohibida a favor del arrendador."))
        clauses.append(("Duración", f"El contrato tendrá una duración inicial de {self._plain(self._v(answers, 'term.duration_months'), 'doce')} meses, contados desde la entrega material o la fecha de inicio registrada. La duración no autoriza cobros por períodos no causados ni impide las formas de terminación previstas en la ley."))
        clauses.append(("Prórroga", f"La prórroga se regirá por la siguiente previsión: {self._renewal_label(term.get('automatic_extension'))}. La continuidad en la ocupación no modifica por sí sola el inventario, la distribución de cargas, las garantías ni las obligaciones de conservación, salvo acuerdo escrito."))
        clauses.append(("Obligaciones del arrendador", "EL ARRENDADOR deberá entregar el inmueble y los servicios incluidos en condiciones de uso, mantener el goce pacífico, realizar las reparaciones a su cargo, entregar copia del contrato y documentos aplicables, informar datos de pago y notificación, respetar la intimidad y atender oportunamente riesgos de habitabilidad. También deberá cumplir las obligaciones que le correspondan como propietario frente a la copropiedad y los servicios públicos."))
        clauses.append(("Obligaciones del arrendatario", "EL ARRENDATARIO deberá pagar oportunamente, cuidar el inmueble y bienes incluidos, usarlo conforme a su destinación, asumir reparaciones locativas imputables, pagar los servicios y cargas a su cargo, cumplir reglas oponibles de convivencia, permitir accesos justificados bajo procedimiento razonable, informar daños o riesgos y restituir el inmueble al finalizar."))
        clauses.append(("Conservación, daños y responsabilidad", "Cada parte responderá por los daños que le sean imputables conforme a la ley, la evidencia y el nexo causal. El desgaste normal, la vetustez, los vicios no atribuibles y los eventos de fuerza mayor no constituirán deuda automática del arrendatario. Antes de cualquier cobro se deberá comparar el inventario inicial y final, escuchar observaciones y descontar depreciación y mejoras autorizadas cuando proceda."))
        clauses.append(("Mascotas y zonas compartidas", f"Mascotas: {self._fragment(self._v(answers, 'pets.conditions'), 'no se informaron condiciones especiales')}. Zonas o instalaciones compartidas: {self._fragment(self._v(answers, 'use.shared_areas'), 'no informadas')}. Las restricciones deberán ser proporcionadas, compatibles con la ley y el reglamento aplicable, y orientadas a convivencia, seguridad, salubridad y reparación de daños comprobados."))
        clauses.append(("Cesión, subarriendo y turismo", f"Según la información suministrada, {self._sublease_label(self._v(answers, 'use.sublease_tourism'))}. EL ARRENDATARIO no podrá ceder el contrato, subarrendar ni entregar el goce total o parcial, ni destinar el inmueble a alojamiento turístico o plataformas, sin autorización escrita y sin cumplir el régimen aplicable."))
        clauses.append(("Acceso e inspecciones", "EL ARRENDADOR podrá solicitar acceso para reparaciones, inspecciones justificadas, lectura de medidores, avalúos, venta o restitución, mediante aviso razonable y coordinación previa, salvo emergencia real. No se autorizan ingresos arbitrarios, vigilancia permanente, retención de bienes ni perturbación del goce pacífico."))
        clauses.append(("Tratamiento de datos personales", f"Los datos se tratarán para {self._plain(screening.get('personal_data'), 'celebración, ejecución, seguridad y cierre contractual')}. El estudio de riesgo o crédito será {'aplicable' if screening.get('credit_study') is True else 'no aplicable'}, deberá contar con finalidad, información y autorización cuando corresponda. Los documentos sensibles se limitarán a lo necesario, con acceso restringido y conservación proporcional."))
        clauses.append(("Incumplimiento y subsanación", "Ante un incumplimiento subsanable, la parte afectada comunicará los hechos, soportes y medida requerida, otorgando un plazo razonable cuando la naturaleza del incumplimiento lo permita. La recepción de pagos parciales, visitas, comunicaciones o tolerancias no constituye renuncia general ni autoriza cobros, sanciones o terminaciones diferentes de las permitidas por la ley."))
        clauses.append(("Terminación por el arrendador", "EL ARRENDADOR podrá solicitar la terminación por las causales, procedimientos, preavisos, consignaciones e indemnizaciones legalmente aplicables. La cláusula contractual no sustituye la verificación del supuesto concreto ni permite vías de hecho, suspensión arbitraria de servicios, ingreso no autorizado o retención de bienes."))
        clauses.append(("Terminación por el arrendatario", "EL ARRENDATARIO podrá terminar por mutuo acuerdo, vencimiento o por las causales y procedimientos legales aplicables, incluyendo incumplimientos relevantes del arrendador. Cuando se exija preaviso, indemnización o entrega, deberá cumplirse por el medio y oportunidad que correspondan, dejando evidencia."))
        clauses.append(("Preavisos y terminaciones especiales", f"El preaviso previsto es de {self._plain(term.get('notice_days'), 'noventa')} días. Condiciones especiales informadas: {self._plain(term.get('special_termination'), 'ninguna adicional')}. Todo preaviso deberá identificar contrato, causal o fundamento, fecha efectiva, obligaciones económicas, mecanismo de entrega y soportes exigibles."))
        clauses.append(("Restitución del inmueble", "Al terminar, EL ARRENDATARIO restituirá el inmueble, llaves, controles, bienes y documentos mediante acta. LAS PARTES verificarán inventario, medidores, servicios, administración, daños y reparaciones. La entrega de llaves no equivale por sí sola a paz y salvo, y la negativa injustificada a recibir deberá documentarse y gestionarse por mecanismos legales."))
        clauses.append(("Liquidación y cierre", "El cierre discriminará cánones, servicios, administración, reparaciones comprobadas, saldos a favor y soportes. No se efectuarán compensaciones automáticas por valores discutidos ni se condicionará la recepción del inmueble a aceptar deudas no verificadas. Las garantías se gestionarán conforme a su naturaleza y vigencia."))
        clauses.append(("Notificaciones", f"Canal del arrendador: {self._plain(notifications.get('landlord_channel'), landlord['email'])}. Canal del arrendatario: {self._plain(notifications.get('tenant_channel'), tenant['email'])}. Direcciones físicas: {self._plain(notifications.get('physical_addresses'), landlord['address'] + ' / ' + tenant['address'])}. Las comunicaciones especiales se realizarán por los medios exigidos legalmente o pactados de forma oponible."))
        clauses.append(("Solución de controversias", "LAS PARTES procurarán resolver directamente las diferencias mediante comunicación documentada. Si no es posible, podrán acudir a conciliación y a la jurisdicción competente. Ninguna cláusula impedirá solicitar medidas urgentes, ejercer derechos de vivienda o consumidor cuando sean aplicables, ni acudir a autoridades administrativas o judiciales."))
        clauses.append(("Integridad, anexos y modificaciones", "El contrato, sus anexos vigentes, actas, comunicaciones y autorizaciones expresamente relacionadas contienen el acuerdo. Cualquier modificación material deberá constar por escrito, identificar la versión y no podrá desconocer normas imperativas. La ineficacia de una disposición no afectará las restantes, que se interpretarán de buena fe y conforme a la finalidad residencial."))
        clauses.append(("Prelación y documentos relacionados", "En caso de contradicción prevalecerán las normas imperativas, el contrato y sus otrosíes sobre anexos operativos, inventarios, comunicaciones y manuales, sin perjuicio de la regla más específica y favorable cuando resulte aplicable. Solo integrarán el expediente los documentos efectivamente generados, identificados y aceptados en la revisión vigente."))
        clauses.append(("Firma, lectura y copia", f"LAS PARTES declaran haber revisado el contrato y sus documentos aplicables, haber contado con oportunidad de formular observaciones y recibir una copia íntegra. Se firma en {prop['municipality']}, en la fecha registrada por el mecanismo de aceptación. La firma electrónica será válida cuando permita identificar a la persona y conservar evidencia de integridad y aceptación."))

        for number, (heading, body) in enumerate(clauses, 1):
            self._clause(doc, number, heading, body)

        doc.add_page_break()
        self._section(doc, "DOCUMENTOS RELACIONADOS")
        names = {x["id"]: x["name"] for x in self.evaluator.documents}
        rows = [[doc_id, names.get(doc_id, doc_id)] for doc_id in evaluation.get("documents", [])]
        self._table(doc, ["Código", "Documento"], rows)
        self._section(doc, "FIRMAS")
        self._signatures(doc, landlord, tenant)
        doc.save(target)

    def _annex_header(self, doc, answers, doc_id, title):
        self._set_doc_styles(doc); self._title(doc, title)
        landlord = self._party_ar(answers, "landlord", "EL ARRENDADOR")
        tenant = self._party_ar(answers, "tenant", "EL ARRENDATARIO")
        prop = self._property(answers)
        self._paragraph(doc, f"Documento {doc_id}, asociado al contrato entre {landlord['name']} y {tenant['name']} respecto del inmueble ubicado en {prop['address']}, {prop['municipality']}. Su contenido debe leerse con el contrato y la revisión vigente del expediente.")
        return landlord, tenant, prop

    def _property_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-PROPERTY-001", "ANEXO DE IDENTIFICACIÓN DEL INMUEBLE Y BIENES INCLUIDOS")
        rows = [
            ["Dirección", prop["address"]], ["Municipio y departamento", prop["municipality"]],
            ["Tipo", prop["type"]], ["Matrícula inmobiliaria", prop["registration"]],
            ["Identificación catastral", prop["cadastral"]], ["Unidad o área privada", prop["private_area"]],
            ["Parqueadero", prop["parking"]], ["Depósito", prop["storage"]], ["Otros", prop["other"]],
            ["Situación jurídica informada", self._plain(self._v(answers, "property.dispute_status"), "sin controversias informadas")],
        ]
        self._table(doc, ["Dato", "Descripción"], rows)
        self._paragraph(doc, "La identificación deberá contrastarse con los soportes disponibles. Este anexo no sustituye certificado de tradición, título, reglamento de propiedad horizontal, licencias ni verificaciones físicas o registrales.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _inventory_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-INVENTORY-001", "INVENTARIO Y EVIDENCIA DE ESTADO")
        self._paragraph(doc, f"Método seleccionado: {self._fragment(self._v(answers, 'delivery.inventory_method'), 'inventario detallado con evidencia')}. Defectos informados: {self._fragment(self._v(answers, 'condition.defects'), 'ninguno adicional')}.")
        rows = [[x, "", "", "", ""] for x in ["Acceso y cerraduras", "Sala y comedor", "Cocina", "Habitaciones", "Baños", "Lavandería", "Pisos, muros y techos", "Ventanas y vidrios", "Red eléctrica", "Red hidráulica y sanitaria", "Gas", "Medidores", "Equipos", "Otros"]]
        self._table(doc, ["Ambiente o elemento", "Descripción", "Estado inicial", "Evidencia", "Observaciones"], rows)
        self._paragraph(doc, "El inventario debe permitir comparación razonable al cierre. El desgaste normal, la vetustez y los daños no imputables deberán diferenciarse de los daños atribuibles.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _delivery_act(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ACT-AR-DELIVERY-001", "ACTA DE ENTREGA MATERIAL")
        self._table(doc, ["Aspecto", "Registro"], [
            ["Fecha y hora", self._date_es(self._v(answers, "delivery.date"))], ["Llaves y controles", ""], ["Lectura de medidores", ""], ["Servicios activos", ""], ["Estado de habitabilidad", self._plain(self._v(answers, "condition.habitability"))], ["Reparaciones pendientes", self._plain(self._v(answers, "condition.repairs_pending"))], ["Documentos entregados", "Contrato, anexos, inventario y reglamento cuando aplique"], ["Salvedades", ""],
        ])
        self._paragraph(doc, "La firma acredita la entrega material y las salvedades registradas; no implica renuncia frente a vicios ocultos, condiciones de habitabilidad o derechos imperativos.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _ph_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-PH-001", "PROPIEDAD HORIZONTAL Y REGLAS DE CONVIVENCIA")
        ph = self._as_dict(self._v(answers, "property.ph_details"))
        self._table(doc, ["Dato", "Contenido"], [["Copropiedad", ph.get("name")], ["Administrador", ph.get("administrator")], ["Entrega de parte normativa", ph.get("rules_delivery")], ["Administración ordinaria", self._money(ph.get("ordinary_fee"), "COP")]])
        self._bullets(doc, ["Uso residencial y respeto por la destinación de los bienes.", "Cumplimiento de reglas oponibles sobre ruido, residuos, mudanzas, parqueaderos, mascotas y zonas comunes.", "Prohibición de imponer al arrendatario obligaciones del propietario que no hayan sido trasladadas válidamente.", "Trazabilidad de sanciones, requerimientos y comunicaciones de la administración."])
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _services_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-SERVICES-001", "SERVICIOS PÚBLICOS, ADMINISTRACIÓN Y USOS ADICIONALES")
        utilities = self._as_dict(self._v(answers, "charges.utilities")); admin = self._as_dict(self._v(answers, "charges.administration")); extra = self._as_dict(self._v(answers, "charges.additional_services"))
        rows = [
            ["Administración ordinaria", self._money(admin.get("ordinary_amount"), "COP"), admin.get("ordinary_responsible"), "Factura o cuenta de cobro"],
            ["Expensas extraordinarias", "Según liquidación", admin.get("extraordinary_responsible"), "Decisión y soporte de copropiedad"],
            ["Servicios públicos", utilities.get("distribution"), utilities.get("responsible"), "Facturas y lecturas"],
            ["Internet", utilities.get("internet"), "Según contratación", "Factura"],
            ["Gas u otros", utilities.get("gas_or_other"), "Según contratación", "Factura"],
            ["Servicios o usos adicionales", extra.get("description"), self._money(extra.get("value"), "COP"), "Soporte discriminado"],
        ]
        self._table(doc, ["Concepto", "Regla o valor", "Responsable", "Evidencia"], rows)
        self._paragraph(doc, f"Denuncia del contrato ante empresas de servicios: {self._plain(self._v(answers, 'utilities.denunciation'), 'según decisión de las partes y procedimiento aplicable')}. Cuando se aplique, deberá utilizarse el formato y garantía aceptados por la empresa prestadora.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _guarantee_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-GUARANTEE-001", "GARANTÍA PERSONAL O PÓLIZA")
        details = self._as_dict(self._v(answers, "guarantee.details"))
        self._table(doc, ["Elemento", "Contenido"], [["Tipo", self._v(answers, "guarantee.type")], ["Garante o aseguradora", details.get("party")], ["Identificación o póliza", details.get("id_number")], ["Obligaciones cubiertas", details.get("scope")], ["Vigencia y renovaciones", details.get("validity")]])
        self._paragraph(doc, "La garantía se interpretará de manera delimitada, no autoriza cobros automáticos y deberá respetar las defensas, requisitos de reclamación, vigencia y soportes aplicables. No constituye depósito en dinero ni caución real a favor del arrendador.")
        self._section(doc, "ACEPTACIÓN DEL GARANTE O ASEGURADORA")
        self._table(doc, ["Nombre", "Identificación", "Firma", "Fecha"], [[details.get("party"), details.get("id_number"), "", ""]])
        doc.save(target)

    def _furnished_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-FURNISHED-001", "MUEBLES Y BIENES DE VALOR")
        rows = [["", "", "", "", "", ""] for _ in range(12)]
        self._table(doc, ["Bien", "Marca o referencia", "Serial", "Estado", "Evidencia", "Valor de referencia"], rows)
        self._paragraph(doc, "El valor de referencia facilita identificación y seguro, pero no prueba por sí solo responsabilidad ni cuantía de un daño. Al cierre deberán considerarse estado inicial, depreciación, desgaste normal, reparabilidad y evidencia.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _occupants_annex(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ANX-AR-OCCUPANTS-001", "OCUPANTES, MASCOTAS Y ZONAS COMPARTIDAS")
        additional = self._as_dict(self._v(answers, "tenant.additional"))
        self._table(doc, ["Aspecto", "Regla"], [["Arrendatarios adicionales", additional.get("names")], ["Solidaridad", additional.get("solidarity")], ["Ocupantes autorizados", self._v(answers, "occupants.authorized")], ["Mascotas", self._v(answers, "pets.conditions")], ["Zonas compartidas", self._v(answers, "use.shared_areas")]])
        self._paragraph(doc, "Las reglas de convivencia deberán ser razonables, no discriminatorias y compatibles con el reglamento aplicable. Las obligaciones económicas y la solidaridad solo vinculan a quienes las hayan aceptado válidamente.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _data_authorization(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "AUT-AR-DATA-001", "INFORMACIÓN Y AUTORIZACIÓN PARA ESTUDIO Y TRATAMIENTO DE DATOS")
        screening = self._as_dict(self._v(answers, "data.screening"))
        self._paragraph(doc, f"Finalidades informadas: {self._plain(screening.get('personal_data'), 'celebración, ejecución y cierre del contrato')}. Estudio de riesgo o crédito: {'sí' if screening.get('credit_study') is True else 'no'}. Documentos sensibles o especialmente delicados: {'sí' if screening.get('sensitive_documents') is True else 'no'}.")
        self._bullets(doc, ["Conocer, actualizar, rectificar y consultar los datos.", "Solicitar prueba de la autorización cuando proceda.", "Presentar consultas o reclamos.", "Ser informado sobre el uso y la finalidad.", "Revocar la autorización o solicitar supresión cuando legalmente proceda.", "No aportar datos sensibles salvo necesidad, información clara y decisión libre cuando corresponda."])
        self._section(doc, "DECISIÓN")
        self._table(doc, ["Opción", "Marcación"], [["Autorizo el tratamiento para las finalidades informadas", "[  ]"], ["No autorizo finalidades facultativas y solicito alternativa", "[  ]"]])
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _adjustment_communication(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "COM-AR-ADJUSTMENT-001", "COMUNICACIÓN DE REAJUSTE DEL CANON")
        self._table(doc, ["Dato", "Contenido"], [["Canon anterior", ""], ["Fecha desde la cual rige", ""], ["IPC de referencia", ""], ["Porcentaje aplicado", ""], ["Nuevo canon", ""], ["Verificación del límite del artículo 18", ""], ["Medio y fecha de notificación", ""]])
        self._paragraph(doc, "El reajuste deberá producirse después de doce meses de ejecución bajo el mismo precio, no superar el IPC del año calendario anterior ni el límite legal del canon, y ser comunicado de forma oponible.")
        doc.save(target)

    def _termination_notice(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "NOT-AR-TERMINATION-001", "MODELO DE PREAVISO O TERMINACIÓN")
        term = self._as_dict(self._v(answers, "term.rules"))
        self._table(doc, ["Elemento", "Contenido"], [["Parte que comunica", ""], ["Destinatario", ""], ["Fundamento o causal", term.get("special_termination")], ["Fecha efectiva", ""], ["Preaviso", str(term.get("notice_days") or "") + " días"], ["Indemnización o consignación aplicable", ""], ["Entrega y restitución", ""], ["Canal y evidencia", ""]])
        self._paragraph(doc, "Este modelo exige adecuación al supuesto concreto. No debe enviarse sin verificar causal, oportunidad, indemnización, consignación, autoridad, forma de notificación y demás requisitos aplicables.")
        doc.save(target)

    def _restitution_act(self, answers, target):
        doc = Document(); landlord, tenant, prop = self._annex_header(doc, answers, "ACT-AR-RESTITUTION-001", "ACTA DE RESTITUCIÓN, LIQUIDACIÓN Y CIERRE")
        self._table(doc, ["Aspecto", "Estado final", "Evidencia", "Salvedad"], [[x, "", "", ""] for x in ["Fecha y hora", "Llaves y controles", "Medidores", "Servicios públicos", "Administración", "Inventario", "Muebles y equipos", "Reparaciones", "Cánones y reajustes", "Garantía o póliza", "Correspondencia y documentos", "Saldo final"]])
        self._paragraph(doc, "La firma acredita los hechos y salvedades consignados. No constituye paz y salvo general, renuncia de derechos ni reconocimiento automático de daños o deudas no verificadas.")
        self._signatures(doc, landlord, tenant); doc.save(target)

    def _annex(self, doc_id, answers, target):
        dispatch = {
            "ANX-AR-PROPERTY-001": self._property_annex,
            "ANX-AR-INVENTORY-001": self._inventory_annex,
            "ACT-AR-DELIVERY-001": self._delivery_act,
            "ANX-AR-PH-001": self._ph_annex,
            "ANX-AR-SERVICES-001": self._services_annex,
            "ANX-AR-GUARANTEE-001": self._guarantee_annex,
            "ANX-AR-FURNISHED-001": self._furnished_annex,
            "ANX-AR-OCCUPANTS-001": self._occupants_annex,
            "AUT-AR-DATA-001": self._data_authorization,
            "COM-AR-ADJUSTMENT-001": self._adjustment_communication,
            "NOT-AR-TERMINATION-001": self._termination_notice,
            "ACT-AR-RESTITUTION-001": self._restitution_act,
        }
        if doc_id not in dispatch:
            raise ValueError(f"Documento no soportado: {doc_id}")
        dispatch[doc_id](answers, target)

    def render_documents(self, answers, target_folder):
        evaluation = self.evaluator.evaluate(answers)
        if evaluation.get("blocked"):
            messages = "; ".join(x["message"] for x in evaluation.get("findings", []) if x.get("severity") == "blocker")
            raise ValueError("El expediente contiene bloqueos jurídicos: " + messages)
        essential = [x for x in evaluation.get("missing_fields", []) if x.get("step_id") in {"parties", "property", "use", "economics", "term", "documents"}]
        if essential:
            raise ValueError("Faltan datos esenciales: " + ", ".join(x["label"] for x in essential))
        target_folder = Path(target_folder); target_folder.mkdir(parents=True, exist_ok=False)
        contract = target_folder / "CO-AR-001_Contrato_Arrendamiento_Vivienda_Urbana.docx"
        self._contract(answers, evaluation, contract)
        generated = [{"id": "DOC-AR-CONTRACT-001", "filename": contract.name}]
        for doc_id in evaluation.get("documents", []):
            if doc_id == "DOC-AR-CONTRACT-001":
                continue
            target = target_folder / f"{doc_id}.docx"
            self._annex(doc_id, answers, target)
            generated.append({"id": doc_id, "filename": target.name})
        hashes, unresolved = {}, []
        for item in generated:
            path = target_folder / item["filename"]
            text = self._extract_text(path)
            if UNRESOLVED_PATTERN.search(text):
                unresolved.append(item["id"])
            hashes[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()
        if unresolved:
            raise ValueError("Se detectaron variables o valores centinela sin resolver: " + ", ".join(unresolved))
        return evaluation, generated, hashes

    def generate(self, answers, actor=None):
        generation_id = "COAR001-" + uuid.uuid4().hex[:12].upper()
        folder = self.output_dir / generation_id; folder.mkdir(parents=True, exist_ok=False)
        documents_dir = folder / "documents" / "revision-0001"
        evaluation, generated, hashes = self.render_documents(answers, documents_dir)
        manifest = {
            "generation_id": generation_id,
            "product_id": "CO-AR-001",
            "version": self.VERSION,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "actor": {"id": (actor or {}).get("id"), "role": (actor or {}).get("role")},
            "readiness": evaluation.get("readiness") or evaluation.get("status"),
            "requires_professional_review": evaluation.get("professional_review_required", False),
            "review_requirements": evaluation.get("review_requirements", []),
            "documents": generated,
            "document_folder": "documents/revision-0001",
            "hashes": hashes,
            "selected_blocks": evaluation.get("blocks", []),
            "unresolved_variables": 0,
            "status": "draft_generated",
            "current_revision": 1,
            "legal_approval": {"status": "pending"},
            "qa_approval": {"status": "pending"},
            "released": False,
        }
        (folder / "answers.json").write_text(json.dumps(answers, ensure_ascii=False, indent=2), encoding="utf-8")
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        package = self._package(folder, self.output_dir / f"{generation_id}.zip")
        manifest["package_filename"] = package.name
        manifest["package_sha256"] = hashlib.sha256(package.read_bytes()).hexdigest()
        (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return manifest

    @staticmethod
    def _package(folder, target):
        with ZipFile(target, "w", ZIP_DEFLATED) as zf:
            for path in sorted(Path(folder).rglob("*")):
                if path.is_file():
                    zf.write(path, arcname=str(path.relative_to(folder)))
        return target

    def package_path(self, generation_id):
        if not GEN_RE.fullmatch(generation_id or ""):
            return None
        path = self.output_dir / f"{generation_id}.zip"
        return path if path.is_file() else None
