from __future__ import annotations

"""Presentación M33.2 para diagnósticos, consultas e informes jurídico-analíticos.

Esta capa es exclusivamente editorial. No altera hechos, reglas, cálculos, fuentes,
clasificaciones, conclusiones ni recomendaciones. Su objetivo es que los documentos
analíticos de LegalAIZ.it tengan una gramática propia: Book Antiqua, marca discreta,
título institucional, cuerpo justificado, secciones jerarquizadas y tablas legibles.
"""

from functools import wraps
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from document_standard_v33 import audit_docx_legal_standard

FONT_NAME = "Book Antiqua"
BODY_PT = 11
TABLE_PT = 9.5
TITLE_PT = 13
SUBTITLE_PT = 10.5
PARAGRAPH_AFTER_PT = 6
SECTION_BEFORE_PT = 9
SECTION_AFTER_PT = 5
TITLE_AFTER_PT = 8
NAVY = "0D1324"
IVORY = "F7F5F1"
WHITE = "FFFFFF"
CHARCOAL = "1F1F1F"

ANALYTICAL_PRODUCT_CODES = frozenset({
    "CO-LA-001",
    "CO-CD-001",
    "CO-CD-003",
    "CO-CD-004",
    "CO-SA-001",
    "CO-TR-001",
    "CO-TR-002",
})

_ANALYTICAL_TITLE_TOKENS = (
    "diagnóstico",
    "informe",
    "consulta integral",
    "concepto jurídico",
    "evaluación jurídica",
    "análisis jurídico",
)

_EXCLUDED_ANALYTICAL_TOKENS = (
    "matriz",
    "calendario",
    "cronograma",
    "trazabilidad",
    "resumen",
    "estado de cuenta",
    "acuerdo de pago",
    "pagaré",
    "carta de instrucciones",
    "recibo",
    "certificado",
    "índice probatorio",
)

_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+")
_BULLET_RE = re.compile(r"^\s*[•·-]\s+")
_LEAD_LABEL_RE = re.compile(
    r"^(Clasificación (?:propuesta|preliminar)|Regla aplicable|Resultado preliminar|"
    r"Nivel de riesgo|Conclusión(?: preliminar)?|Saldo neto preliminar|Hallazgo|"
    r"Recomendación):\s*",
    re.IGNORECASE,
)


def _set_run_font(
    run,
    *,
    size_pt: float = BODY_PT,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color: str | None = None,
    all_caps: bool | None = None,
) -> None:
    run.font.name = FONT_NAME
    rfonts = run._element.get_or_add_rPr().rFonts
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if all_caps is not None:
        run.font.all_caps = all_caps


def _set_paragraph_font(paragraph, *, size_pt: float = BODY_PT, color: str = CHARCOAL) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt, color=color)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _is_analytical_document(product_code: str, title: str) -> bool:
    code = str(product_code or "").strip().upper()
    if code not in ANALYTICAL_PRODUCT_CODES:
        return False
    lowered = str(title or "").strip().casefold()
    if any(token in lowered for token in _EXCLUDED_ANALYTICAL_TOKENS):
        return False
    return any(token in lowered for token in _ANALYTICAL_TITLE_TOKENS)


def _title_paragraph(document: Document, title: str):
    wanted = str(title or "").strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == wanted:
            return paragraph
    return next((p for p in document.paragraphs if p.style and p.style.name == "Title"), None)


def _paragraph_after(document: Document, paragraph):
    if paragraph is None:
        return None
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            return next((p for p in document.paragraphs if p._p is sibling), None)
        sibling = sibling.getnext()
    return None


def _remove_body_branding(document: Document, title_paragraph) -> int:
    if title_paragraph is None:
        return 0
    removed = 0
    for paragraph in list(document.paragraphs):
        if paragraph._p is title_paragraph._p:
            break
        text = paragraph.text.strip().casefold()
        if not text or text in {"legalaiz.it", "más que respuestas, soluciones."}:
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def _normalize_styles(document: Document) -> None:
    for style_name in ("Normal", "Title", "Heading 1"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_NAME
        style.font.size = Pt(BODY_PT)
        rfonts = style._element.get_or_add_rPr().rFonts
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT_NAME)


def _format_title_and_subtitle(document: Document, title: str) -> None:
    title_p = _title_paragraph(document, title)
    if title_p is None:
        return
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(TITLE_AFTER_PT)
    title_p.paragraph_format.keep_with_next = True
    for run in title_p.runs:
        _set_run_font(
            run,
            size_pt=TITLE_PT,
            bold=True,
            italic=False,
            underline=False,
            color=NAVY,
            all_caps=True,
        )

    subtitle = _paragraph_after(document, title_p)
    if subtitle is None or not subtitle.text.strip():
        return
    if subtitle.style and subtitle.style.name.lower().startswith("heading"):
        return
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True
    for run in subtitle.runs:
        _set_run_font(
            run,
            size_pt=SUBTITLE_PT,
            bold=False,
            italic=True,
            underline=False,
            color=CHARCOAL,
            all_caps=False,
        )


def _format_headings(document: Document, title: str) -> int:
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == str(title or "").strip():
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if not style_name.lower().startswith("heading"):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.upper() in {"FIRMA", "FIRMAS"} else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(SECTION_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(SECTION_AFTER_PT)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            _set_run_font(
                run,
                size_pt=BODY_PT,
                bold=True,
                italic=False,
                underline=False,
                color=NAVY,
                all_caps=False,
            )
        changed += 1
    return changed


def _apply_lead_label_bold(paragraph) -> bool:
    text = paragraph.text
    match = _LEAD_LABEL_RE.match(text.strip())
    if not match or len(paragraph.runs) != 1:
        return False
    run = paragraph.runs[0]
    raw = run.text
    left_padding = len(raw) - len(raw.lstrip())
    label = raw[:left_padding] + match.group(0)
    remainder = raw[len(label):]
    run.text = label
    _set_run_font(run, size_pt=BODY_PT, bold=True, color=CHARCOAL)
    if remainder:
        tail = paragraph.add_run(remainder)
        _set_run_font(tail, size_pt=BODY_PT, bold=False, color=CHARCOAL)
    return True


def _format_body(document: Document, title: str) -> int:
    changed = 0
    title_p = _title_paragraph(document, title)
    subtitle = _paragraph_after(document, title_p)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == str(title or "").strip():
            continue
        if subtitle is not None and paragraph._p is subtitle._p:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(PARAGRAPH_AFTER_PT)
        paragraph.paragraph_format.line_spacing = 1.0
        if _NUMBERED_RE.match(text):
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
        elif _BULLET_RE.match(text):
            paragraph.paragraph_format.left_indent = Pt(20)
            paragraph.paragraph_format.first_line_indent = Pt(-10)
        _set_paragraph_font(paragraph, size_pt=BODY_PT)
        _apply_lead_label_bold(paragraph)
        changed += 1
    return changed


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _is_signature_table(table) -> bool:
    tbl_pr = table._tbl.tblPr
    desc = tbl_pr.find(qn("w:tblDescription")) if tbl_pr is not None else None
    return bool(desc is not None and desc.get(qn("w:val")) == "LegalAIZ-SignatureTable")


def _mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _format_tables(document: Document) -> int:
    changed = 0
    for table in document.tables:
        signature = _is_signature_table(table)
        for row_index, row in enumerate(table.rows):
            _prevent_row_split(row)
            if row_index == 0 and not signature:
                _mark_header_row(row)
            for cell_index, cell in enumerate(row.cells):
                if not signature:
                    fill = NAVY if row_index == 0 else (IVORY if cell_index == 0 else WHITE)
                    _set_cell_shading(cell, fill)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2 if signature else 1)
                    for run in paragraph.runs:
                        _set_run_font(
                            run,
                            size_pt=TABLE_PT,
                            bold=(bool(run.bold) or (not signature and (row_index == 0 or cell_index == 0))),
                            color=(WHITE if (not signature and row_index == 0) else CHARCOAL),
                            all_caps=False,
                        )
        changed += 1
    return changed


def apply_m33_2_analytical_format(path: str | Path, *, product_code: str, title: str) -> dict:
    target = Path(path)
    if not _is_analytical_document(product_code, title):
        return {"applied": False, "profile": "M33.2-analytical", "reason": "non_analytical_document"}

    document = Document(target)
    _normalize_styles(document)
    title_p = _title_paragraph(document, title)
    removed_branding = _remove_body_branding(document, title_p)
    headings = _format_headings(document, title)
    paragraphs = _format_body(document, title)
    _format_title_and_subtitle(document, title)
    tables = _format_tables(document)
    document.save(target)

    report = audit_docx_legal_standard(target)
    if not report.get("valid"):
        raise ValueError(f"Documento analítico M33.2 no supera auditoría documental: {report.get('findings')}")
    return {
        "applied": True,
        "profile": "M33.2-analytical",
        "font": FONT_NAME,
        "removed_body_branding": removed_branding,
        "formatted_headings": headings,
        "formatted_paragraphs": paragraphs,
        "formatted_tables": tables,
        "paragraph_after_pt": PARAGRAPH_AFTER_PT,
        "table_header_fill": NAVY,
    }


def install_m33_2_analytical_format_gate() -> bool:
    """Envuelve `build_docx` después de la capa procedimental y revalida el archivo final."""
    import docx_builder
    from legalai_platform.document_release_gate import enforce_document_release_gate, infer_product_code

    current = docx_builder.build_docx
    if getattr(current, "_legalaiz_m33_2_analytical", False):
        return True

    @wraps(current)
    def guarded_build_docx(*args, **kwargs):
        call_args = list(args)
        path = kwargs.get("path") or (call_args[0] if call_args else None)
        title = kwargs.get("title") or (call_args[1] if len(call_args) >= 2 else "")
        metadata = kwargs.get("metadata")
        if metadata is None and len(call_args) >= 4:
            metadata = call_args[3]
        product_code = str(kwargs.get("product_code") or infer_product_code(path or "document.docx", metadata) or "").upper()

        result = current(*args, **kwargs)
        final_path = Path(path or result)
        presentation = apply_m33_2_analytical_format(final_path, product_code=product_code, title=str(title or ""))
        if presentation.get("applied"):
            enforce_document_release_gate(final_path, expected_product=product_code or None, metadata=metadata)
        return result

    guarded_build_docx._legalaiz_m33_2_analytical = True
    guarded_build_docx._legalaiz_original = current
    docx_builder.build_docx = guarded_build_docx
    return True


__all__ = [
    "ANALYTICAL_PRODUCT_CODES",
    "apply_m33_2_analytical_format",
    "install_m33_2_analytical_format_gate",
]
