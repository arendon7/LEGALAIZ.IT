from __future__ import annotations

"""Formato jurídico de referencia M33.2 para instrumentos contractuales LegalAIZ.it.

Esta capa es exclusivamente editorial. No modifica reglas, cláusulas, variables,
fuentes ni decisiones jurídicas. Recompone el DOCX ya generado para que el
instrumento de aprobación use la gramática visual validada por el usuario:
Book Antiqua, título centrado, tabla de identificación compacta, cuerpo justificado,
rótulos de cláusula en línea y cuadro visible de firmas.
"""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FONT_NAME = "Book Antiqua"
BODY_PT = 11
TABLE_PT = 10.5
TITLE_PT = 11
TITLE_AFTER_PT = 12
TABLE_TO_BODY_PT = 10
CONTRACT_PRODUCTS = frozenset({"CO-LA-002", "CO-EM-003", "CO-AR-001", "CO-EM-004"})

_ORDINAL = (
    r"PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[ÉE]PTIMA|OCTAVA|NOVENA|D[ÉE]CIMA|"
    r"D[ÉE]CIMA\s+PRIMERA|D[ÉE]CIMA\s+SEGUNDA|D[ÉE]CIMA\s+TERCERA|D[ÉE]CIMA\s+CUARTA|"
    r"D[ÉE]CIMA\s+QUINTA|D[ÉE]CIMA\s+SEXTA|D[ÉE]CIMA\s+S[ÉE]PTIMA|D[ÉE]CIMA\s+OCTAVA|"
    r"D[ÉE]CIMA\s+NOVENA|VIG[ÉE]SIMA|VIG[ÉE]SIMA\s+PRIMERA|VIG[ÉE]SIMA\s+SEGUNDA|"
    r"VIG[ÉE]SIMA\s+TERCERA|VIG[ÉE]SIMA\s+CUARTA|VIG[ÉE]SIMA\s+QUINTA|"
    r"VIG[ÉE]SIMA\s+SEXTA|VIG[ÉE]SIMA\s+S[ÉE]PTIMA|VIG[ÉE]SIMA\s+OCTAVA|"
    r"VIG[ÉE]SIMA\s+NOVENA|TRIG[ÉE]SIMA|CUADRAG[ÉE]SIMA|QUINCUAG[ÉE]SIMA"
)
_CLAUSE_HEADING_RE = re.compile(rf"^({_ORDINAL})(?:\.|:)\s*.+", re.IGNORECASE)
_DEFINED_TERM_RE = re.compile(
    r"(EL EMPLEADOR|LA PERSONA TRABAJADORA|EL TRABAJADOR|LA TRABAJADORA|"
    r"EL CONTRATANTE|LA CONTRATANTE|EL CONTRATISTA|LA CONTRATISTA|LAS PARTES|"
    r"LA PARTE ARRENDADORA|EL ARRENDADOR|LA ARRENDADORA|LA PARTE ARRENDATARIA|"
    r"EL ARRENDATARIO|LA ARRENDATARIA|LA PARTE REVELADORA|LA PARTE RECEPTORA)",
    re.IGNORECASE,
)


def _set_run_font(run, *, size_pt: float = BODY_PT, bold: bool | None = None, underline: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline


def _set_paragraph_font(paragraph, *, size_pt: float = BODY_PT) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_table_row(row) -> None:
    tr = row._tr
    parent = tr.getparent()
    if parent is not None:
        parent.remove(tr)


def _set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_table_borders(table, *, color: str = "666666", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _is_signature_table(table) -> bool:
    tbl_pr = table._tbl.tblPr
    desc = tbl_pr.find(qn("w:tblDescription")) if tbl_pr is not None else None
    return bool(desc is not None and desc.get(qn("w:val")) == "LegalAIZ-SignatureTable")


def _set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def _first_table(document: Document):
    return document.tables[0] if document.tables else None


def _normalize_first_table(table) -> None:
    if not table:
        return
    if table.rows and len(table.rows[0].cells) >= 2:
        first = [cell.text.strip().casefold() for cell in table.rows[0].cells[:2]]
        if first == ["campo", "información"]:
            _remove_table_row(table.rows[0])
    _set_table_borders(table, color="666666", size="4")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_shading(cell, "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=TABLE_PT, bold=(index == 0 or bool(run.bold)))


def _rebuild_signature_table(table) -> None:
    if not table or not _is_signature_table(table):
        return
    parties: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            text = "\n".join(line.strip() for line in cell.text.splitlines() if line.strip())
            if text:
                parties.append(text)
    if not parties:
        return

    while len(table.rows) < len(parties):
        table.add_row()
    while len(table.rows) > len(parties):
        _remove_table_row(table.rows[-1])

    for row, party_text in zip(table.rows, parties):
        if len(row.cells) < 2:
            continue
        left, right = row.cells[0], row.cells[1]
        left.text = party_text
        right.text = "FIRMA:\n\n\n"
        for cell_index, cell in enumerate((left, right)):
            _set_cell_shading(cell, "FFFFFF")
            for p_index, paragraph in enumerate(cell.paragraphs):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        size_pt=TABLE_PT,
                        bold=(cell_index == 1 or p_index == 0 or bool(run.bold)),
                    )
    _set_table_borders(table, color="666666", size="4")


def _paragraph_after_table(document: Document, table):
    if table is None:
        return None
    sibling = table._tbl.getnext()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            for paragraph in document.paragraphs:
                if paragraph._p is sibling:
                    return paragraph
            return None
        sibling = sibling.getnext()
    return None


def _title_paragraph(document: Document, title: str):
    wanted = str(title or "").strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == wanted:
            return paragraph
    return next((p for p in document.paragraphs if p.style and p.style.name == "Title"), None)


def _remove_leading_branding(document: Document, title_paragraph) -> None:
    if title_paragraph is None:
        return
    for paragraph in list(document.paragraphs):
        if paragraph._p is title_paragraph._p:
            break
        text = paragraph.text.strip().casefold()
        if text in {"legalaiz.it", "más que respuestas, soluciones."} or not text:
            _remove_paragraph(paragraph)


def _remove_cover_artifacts(document: Document, title: str, approval_subtitle: str) -> None:
    title_p = _title_paragraph(document, title)
    if title_p is None:
        return
    _remove_leading_branding(document, title_p)
    seen_title = False
    for paragraph in list(document.paragraphs):
        if paragraph._p is title_p._p:
            seen_title = True
            continue
        if not seen_title:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if approval_subtitle and text == approval_subtitle.strip():
            _remove_paragraph(paragraph)
            continue
        if text == str(title or "").strip():
            _remove_paragraph(paragraph)
        break


def _format_title(document: Document, title: str) -> None:
    paragraph = _title_paragraph(document, title)
    if paragraph is None:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(TITLE_AFTER_PT)
    _set_keep_with_next(paragraph, True)
    for run in paragraph.runs:
        _set_run_font(run, size_pt=TITLE_PT, bold=True, underline=True)


def _merge_clause_headings(document: Document) -> int:
    merged = 0
    for heading in list(document.paragraphs):
        label = heading.text.strip()
        if not label or not _CLAUSE_HEADING_RE.match(label):
            continue
        if not (heading.style and heading.style.name.lower().startswith("heading")):
            continue
        sibling = heading._p.getnext()
        if sibling is None or sibling.tag != qn("w:p"):
            continue
        body = next((p for p in document.paragraphs if p._p is sibling), None)
        if body is None or not body.text.strip():
            continue
        original = body.text.strip()
        for run in list(body.runs):
            body._p.remove(run._r)
        punctuated = label if label.endswith(":") else f"{label}:"
        run_label = body.add_run(punctuated + " ")
        _set_run_font(run_label, size_pt=BODY_PT, bold=True)
        run_body = body.add_run(original)
        _set_run_font(run_body, size_pt=BODY_PT, bold=False)
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body.paragraph_format.space_before = Pt(4)
        body.paragraph_format.space_after = Pt(3)
        _remove_paragraph(heading)
        merged += 1
    return merged


def _selective_bold(paragraph) -> None:
    if not paragraph.text.strip():
        return
    for run in list(paragraph.runs):
        text = run.text
        if not text or not _DEFINED_TERM_RE.search(text):
            _set_run_font(run, size_pt=BODY_PT)
            continue
        parts = _DEFINED_TERM_RE.split(text)
        index = paragraph._p.index(run._r)
        paragraph._p.remove(run._r)
        for part in parts:
            if not part:
                continue
            new = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attr}"), FONT_NAME)
            rpr.append(fonts)
            if run.bold or _DEFINED_TERM_RE.fullmatch(part):
                rpr.append(OxmlElement("w:b"))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(BODY_PT * 2))
            rpr.append(sz)
            szcs = OxmlElement("w:szCs")
            szcs.set(qn("w:val"), str(BODY_PT * 2))
            rpr.append(szcs)
            new.append(rpr)
            node = OxmlElement("w:t")
            if part[:1].isspace() or part[-1:].isspace():
                node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            node.text = part
            new.append(node)
            paragraph._p.insert(index, new)
            index += 1


def _format_body(document: Document, title: str) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == str(title or "").strip():
            continue
        if text.upper() == "CONSIDERACIONES":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                _set_run_font(run, size_pt=BODY_PT, bold=True)
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(3)
        _set_paragraph_font(paragraph, size_pt=BODY_PT)
        _selective_bold(paragraph)


def _set_post_table_spacing(document: Document, table) -> None:
    paragraph = _paragraph_after_table(document, table)
    if paragraph is None:
        return
    if paragraph.text.strip():
        paragraph.paragraph_format.space_before = Pt(TABLE_TO_BODY_PT)
    else:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(TABLE_TO_BODY_PT)
        sibling = paragraph._p.getnext()
        if sibling is not None and sibling.tag == qn("w:p"):
            next_p = next((p for p in document.paragraphs if p._p is sibling), None)
            if next_p is not None:
                next_p.paragraph_format.space_before = Pt(0)


def _normalize_styles(document: Document) -> None:
    # `Heading1` era un style_id heredado. python-docx recomienda resolver por el
    # nombre canónico del estilo para evitar advertencias y futuras incompatibilidades.
    for style_name in ("Normal", "Title", "Heading 1"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_NAME
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_NAME)
        style.font.size = Pt(BODY_PT)


def apply_m33_2_reference_format(
    path: str | Path,
    *,
    product_code: str,
    title: str,
    approval_subtitle: str = "",
) -> dict:
    """Aplica el formato validado a un contrato approval_candidate existente."""
    target = Path(path)
    if str(product_code or "").strip() not in CONTRACT_PRODUCTS:
        return {"applied": False, "profile": "M33.2", "reason": "non_contract_family"}

    document = Document(target)
    _normalize_styles(document)
    _remove_cover_artifacts(document, title, approval_subtitle)
    _format_title(document, title)

    identification = _first_table(document)
    if identification is not None and not _is_signature_table(identification):
        _normalize_first_table(identification)
        _set_post_table_spacing(document, identification)

    merged = _merge_clause_headings(document)
    _format_body(document, title)

    signature_tables = [table for table in document.tables if _is_signature_table(table)]
    for table in signature_tables:
        _rebuild_signature_table(table)

    document.save(target)
    return {
        "applied": True,
        "profile": "M33.2",
        "font": FONT_NAME,
        "contract_product": product_code,
        "inline_clause_headings": merged,
        "signature_tables": len(signature_tables),
    }


__all__ = ["CONTRACT_PRODUCTS", "apply_m33_2_reference_format"]