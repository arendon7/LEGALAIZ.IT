from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = "0D1324"
BLUE = "2563EB"
GOLD = "C9A96E"
IVORY = "F7F5F1"
LIGHT = "E6E6E1"
CHARCOAL = "1F1F1F"
MID = "A8AEB8"
RED = "8B1E1E"

CLAUSE_RE = re.compile(r"^CLÁUSULA\s+(.+?)\.\s*(.+)$", re.IGNORECASE)


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _keep_together(paragraph) -> None:
    paragraph.paragraph_format.keep_together = True


def _set_repeat_table_header(table) -> None:
    if table.rows:
        _repeat_header(table.rows[0])


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [("Title", 18, NAVY), ("Heading 1", 13, NAVY), ("Heading 2", 11, BLUE)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def _configure_section(section) -> None:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.3)


def _header_footer(doc: Document, footer_text: str, library_version: str = "M5") -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("LegalAIZ.it")
        r.font.name = "Aptos Display"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = _rgb(NAVY)
        p.add_run(f"  |  Biblioteca de playbooks jurídicos profundos {library_version}").font.color.rgb = _rgb(MID)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(footer_text + "  |  Página ")
        r.font.name = "Aptos"
        r.font.size = Pt(8)
        r.font.color.rgb = _rgb(MID)
        _page_number(p)


def _add_logo(doc: Document, logo_path: Path | None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path and logo_path.exists():
        shape = p.add_run().add_picture(str(logo_path), width=Inches(2.0))
        shape._inline.docPr.set("descr", "Logotipo oficial de LegalAIZ.it")
        shape._inline.docPr.set("title", "LegalAIZ.it")
    else:
        r = p.add_run("LegalAIZ.it")
        r.font.name = "Aptos Display"
        r.font.size = Pt(20)
        r.font.bold = True
        r.font.color.rgb = _rgb(NAVY)


def _add_cover(doc: Document, title: str, subtitle: str, metadata: list[tuple[str, str]], logo_path: Path | None) -> None:
    _add_logo(doc, logo_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title.upper())
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = _rgb(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.italic = True
    r.font.color.rgb = _rgb(BLUE)

    if metadata:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(1.65)
        table.columns[1].width = Inches(4.8)
        table.cell(0, 0).text = "Campo"
        table.cell(0, 1).text = "Información"
        for a, b in metadata:
            cells = table.add_row().cells
            cells[0].text = str(a)
            cells[1].text = str(b)
        _style_table(table)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Uso profesional controlado - verificar hechos, anexos, riesgo y vigencia antes de firma.")
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = _rgb(RED)
    doc.add_page_break()


def _style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_repeat_table_header(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index == 0:
                _shade(cell, NAVY)
            elif row_index % 2 == 0:
                _shade(cell, IVORY)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = _rgb("FFFFFF")


def _add_body_paragraph(doc: Document, text: str, *, italic: bool = False, color: str = CHARCOAL,
                        keep_with_next: bool = False) -> None:
    for part in str(text).split("\n"):
        if not part.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        if keep_with_next:
            _keep_with_next(p)
        r = p.add_run(part.strip())
        r.font.name = "Aptos"
        r.font.size = Pt(10.5)
        r.font.italic = italic
        r.font.color.rgb = _rgb(color)


def _add_clause(doc: Document, heading: str, text: str) -> None:
    match = CLAUSE_RE.match(heading.strip())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    _keep_together(p)
    if match:
        ordinal, title = match.groups()
        r = p.add_run(f"{ordinal.upper()}: {title.upper()}. ")
    else:
        r = p.add_run(heading.upper() + ". ")
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = _rgb(NAVY)
    body = p.add_run(str(text).strip())
    body.font.name = "Aptos"
    body.font.size = Pt(10.5)
    body.font.color.rgb = _rgb(CHARCOAL)


def _add_section(doc: Document, sec: dict[str, Any]) -> None:
    if sec.get("page_break_before") and len(doc.paragraphs) > 1:
        doc.add_page_break()
    heading = str(sec.get("heading") or "").strip()
    kind = sec.get("_type") or sec.get("type") or "section"
    if kind == "signature":
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 1"]
        p.add_run("FIRMAS")
        table = doc.add_table(rows=2, cols=max(1, len(sec.get("parties") or [])))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _repeat_header(table.rows[0])
        parties = sec.get("parties") or []
        for idx, party in enumerate(parties):
            table.cell(0, idx).text = "\n\n\n"
            p2 = table.cell(1, idx).paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(str(party.get("name") or ""))
            r.bold = True
            r.font.size = Pt(9)
            p2.add_run("\n" + str(party.get("label") or "")).font.size = Pt(8)
        return

    is_clause = bool(CLAUSE_RE.match(heading))
    is_control = kind == "control" or "CONTROL DE PUBLICACIÓN" in heading.upper()
    text = str(sec.get("text") or "")
    if is_clause:
        _add_clause(doc, heading, text)
    else:
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 1"]
        p.paragraph_format.page_break_before = bool(sec.get("page_break_before"))
        r = p.add_run(heading)
        if is_control:
            r.font.color.rgb = _rgb(RED)
        if text:
            _add_body_paragraph(doc, text, italic=is_control, color=RED if is_control else CHARCOAL)

    bullets = sec.get("bullets") or []
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(str(item))
        r.font.name = "Aptos"
        r.font.size = Pt(9.5 if is_control else 10)
        r.font.italic = is_control
        r.font.color.rgb = _rgb(RED if is_control else CHARCOAL)
    rows = sec.get("table")
    if rows:
        normalized = [list(row) for row in rows]
        cols = max(len(row) for row in normalized)
        table = doc.add_table(rows=0, cols=cols)
        for row in normalized:
            cells = table.add_row().cells
            for idx in range(cols):
                cells[idx].text = str(row[idx] if idx < len(row) else "")
        _style_table(table)


def build_professional_docx(path: Path, *, title: str, subtitle: str,
                            metadata: list[tuple[str, str]], sections: list[dict[str, Any]],
                            logo_path: Path | None = None,
                            footer: str = "LegalAIZ.it - Más que respuestas, soluciones.",
                            library_version: str = "M5") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_styles(doc)
    _configure_section(doc.sections[0])
    _add_cover(doc, title, subtitle, metadata, logo_path)
    for sec_index, sec in enumerate(sections):
        # The cover already ends with a page break. Suppress a second break on
        # the first substantive section to avoid a blank page after the cover.
        if sec_index == 0 and sec.get("page_break_before"):
            sec = dict(sec)
            sec["page_break_before"] = False
        _add_section(doc, sec)
    _header_footer(doc, footer, library_version)
    props = doc.core_properties
    props.title = title
    props.subject = subtitle
    props.author = "LegalAIZ.it"
    props.last_modified_by = "LegalAIZ.it"
    props.keywords = f"LegalAIZ.it, Colombia, playbook jurídico, {library_version}"
    props.comments = f"Documento generado por la biblioteca de playbooks jurídicos profundos {library_version}."
    props.created = datetime.now(timezone.utc)
    props.modified = datetime.now(timezone.utc)
    doc.save(path)
    return path


def append_package_divider(doc: Document, *, title: str, subtitle: str, position: int, total: int) -> None:
    # A paragraph-level page break avoids an empty page when the preceding
    # document finishes exactly at a physical page boundary.
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(f"PARTE {position} DE {total}")
    r.font.name = "Aptos Display"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = _rgb(GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = _rgb(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = _rgb(BLUE)
    doc.add_page_break()


def build_consolidated_docx(path: Path, *, product_code: str, product_title: str,
                            documents: list[tuple[str, str, list[dict[str, Any]]]],
                            metadata: list[tuple[str, str]], logo_path: Path | None = None,
                            library_version: str = "M5") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_styles(doc)
    _configure_section(doc.sections[0])
    _add_cover(
        doc,
        f"{product_code} - PAQUETE JURÍDICO OPERATIVO PROFUNDO",
        product_title + " · diagnóstico, actuaciones, evidencia y controles integrados",
        metadata,
        logo_path,
    )
    index = doc.add_table(rows=1, cols=3)
    index.cell(0, 0).text = "Parte"
    index.cell(0, 1).text = "Documento"
    index.cell(0, 2).text = "Finalidad"
    for pos, (title, subtitle, _) in enumerate(documents, 1):
        cells = index.add_row().cells
        cells[0].text = str(pos)
        cells[1].text = title
        cells[2].text = subtitle
    _style_table(index)
    total = len(documents)
    for pos, (title, subtitle, sections) in enumerate(documents, 1):
        append_package_divider(doc, title=title, subtitle=subtitle, position=pos, total=total)
        for sec_index, sec in enumerate(sections):
            # The divider already forces the first content page. Suppress a second
            # page-break on the first section to avoid blank pages in packages.
            if sec_index == 0 and sec.get("page_break_before"):
                sec = dict(sec)
                sec["page_break_before"] = False
            _add_section(doc, sec)
    _header_footer(doc, f"LegalAIZ.it - Paquete jurídico operativo profundo {library_version}", library_version)
    props = doc.core_properties
    props.title = f"{product_code} - {product_title}"
    props.subject = f"Paquete jurídico operativo profundo {library_version}"
    props.author = "LegalAIZ.it"
    props.last_modified_by = "LegalAIZ.it"
    props.created = datetime.now(timezone.utc)
    props.modified = datetime.now(timezone.utc)
    doc.save(path)
    return path
