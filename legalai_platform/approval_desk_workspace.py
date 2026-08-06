from __future__ import annotations

"""Capa operativa y segura para la mesa documental M32.5.

El motor M32.4 conserva los registros inmutables. Esta capa enlaza cada expediente
de revisión con un documento y un caso de la aplicación, aplica alcance RBAC,
prepara vistas de trabajo y entrega únicamente la copia liberada cuyo hash sigue
coincidiendo con las dos aprobaciones.
"""

from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any, Callable
import json
import os
import re
import shutil
import subprocess

from docx import Document
from pypdf import PdfReader

import core_v11 as core
from legalai_platform.application_services import can_access_case, document_row, validate_upload
from legalai_platform.document_approval_desk import (
    ApprovalDeskError,
    DocumentApprovalDesk,
    ImmutableRecordError,
    PermissionDenied,
    ReleaseBlocked,
)


PROFESSIONAL_ROLES = frozenset({"specialist", "admin"})
M32_5_SCHEMA = "M32.5"


def _sha256_file(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_segment(value: str, field: str) -> str:
    text = str(value or "").strip()
    if not text or not re.fullmatch(r"[A-Za-z0-9._-]+", text):
        raise ApprovalDeskError(f"{field} contiene caracteres no permitidos.")
    return text


def _actor(user: dict[str, Any]) -> dict[str, str]:
    return {
        "id": str(user.get("id") or ""),
        "role": str(user.get("role") or ""),
        "name": str(user.get("name") or ""),
    }


def _require_professional(user: dict[str, Any]) -> None:
    if str(user.get("role") or "") not in PROFESSIONAL_ROLES:
        raise PermissionDenied("La Mesa Jurídica está reservada para especialistas y administración.")


class ApprovalDeskWorkspace:
    """Adaptador entre documentos de la aplicación y el repositorio M32.4."""

    def __init__(
        self,
        root: str | Path | None = None,
        *,
        db_factory: Callable[[], Any] | None = None,
        document_lookup: Callable[[dict[str, Any], str], dict[str, Any] | None] | None = None,
        access_check: Callable[[dict[str, Any], str], bool] | None = None,
        upload_validator: Callable[[str, bytes], Any] | None = None,
    ):
        self.root = Path(root or (core.RUNTIME / "approval-desk")).resolve()
        self.root.mkdir(parents=True, exist_ok=True)
        self.desk = DocumentApprovalDesk(self.root)
        self.db_factory = db_factory or core.db
        self.document_lookup = document_lookup or document_row
        self.access_check = access_check or can_access_case
        self.upload_validator = upload_validator or validate_upload

    @staticmethod
    def desk_case_id(document_id: str) -> str:
        return f"DSK-{_safe_segment(document_id, 'document_id')}"

    def _detail_unchecked(self, desk_case_id: str) -> dict[str, Any]:
        return self.desk.detail(_safe_segment(desk_case_id, "desk_case_id"))

    def _source_case_id(self, detail: dict[str, Any]) -> str:
        return str(detail.get("case", {}).get("source_generation_id") or "")

    def _authorize_case(self, user: dict[str, Any], desk_case_id: str) -> dict[str, Any]:
        _require_professional(user)
        detail = self._detail_unchecked(desk_case_id)
        source_case_id = self._source_case_id(detail)
        if not source_case_id or not self.access_check(user, source_case_id):
            raise PermissionDenied("El expediente documental no existe o no está dentro del alcance del usuario.")
        return detail

    @staticmethod
    def _current_revision(detail: dict[str, Any]) -> dict[str, Any] | None:
        revision_id = detail.get("case", {}).get("current_revision_id")
        return next((row for row in detail.get("revisions", []) if row.get("revision_id") == revision_id), None)

    @staticmethod
    def _workflow_state(detail: dict[str, Any]) -> str:
        case = detail.get("case", {})
        if case.get("status") == "released":
            return "released"
        current = ApprovalDeskWorkspace._current_revision(detail)
        if not current:
            return "draft"
        findings = current.get("findings") or []
        if any(row.get("state") == "open" and row.get("severity") in {"blocking", "major"} for row in findings):
            return "changes_required"
        approvals = current.get("approvals") or {}
        legal = approvals.get("legal") or {}
        qa = approvals.get("qa") or {}
        if legal.get("decision") == "reject" or qa.get("decision") == "reject":
            return "rejected"
        if legal.get("decision") != "approve":
            return "legal_pending"
        if qa.get("decision") != "approve":
            return "qa_pending"
        if any(row.get("state") == "open" for row in findings):
            return "findings_pending"
        return "ready_to_release"

    def _summary_row(self, detail: dict[str, Any]) -> dict[str, Any]:
        case = detail["case"]
        current = self._current_revision(detail)
        findings = current.get("findings", []) if current else []
        approvals = current.get("approvals", {}) if current else {}
        return {
            "schema_version": M32_5_SCHEMA,
            "desk_case_id": case["case_id"],
            "source_case_id": case.get("source_generation_id"),
            "document_id": case.get("document_id"),
            "product_code": case.get("product_code"),
            "title": case.get("title"),
            "status": self._workflow_state(detail),
            "revision_count": case.get("revision_count", 0),
            "current_revision_id": case.get("current_revision_id"),
            "current_sha256": current.get("sha256") if current else None,
            "legal_decision": (approvals.get("legal") or {}).get("decision"),
            "qa_decision": (approvals.get("qa") or {}).get("decision"),
            "open_findings": sum(row.get("state") == "open" for row in findings),
            "blocking_findings": sum(
                row.get("state") == "open" and row.get("severity") in {"blocking", "major"}
                for row in findings
            ),
            "release_id": case.get("release_id"),
            "audit_valid": bool(detail.get("audit", {}).get("valid")),
            "created_at": case.get("created_at"),
        }

    def list_for_user(self, user: dict[str, Any], status: str | None = None) -> dict[str, Any]:
        _require_professional(user)
        rows: list[dict[str, Any]] = []
        for manifest in sorted(self.root.glob("*/case.json")):
            try:
                detail = self._detail_unchecked(manifest.parent.name)
                source_case_id = self._source_case_id(detail)
                if not source_case_id or not self.access_check(user, source_case_id):
                    continue
                row = self._summary_row(detail)
                if status and row["status"] != status:
                    continue
                rows.append(row)
            except (ApprovalDeskError, OSError, ValueError, json.JSONDecodeError):
                continue
        rows.sort(key=lambda row: (row["status"] == "released", str(row.get("created_at") or "")), reverse=False)
        states = [row["status"] for row in rows]
        return {
            "schema_version": M32_5_SCHEMA,
            "cases": rows,
            "metrics": {
                "total": len(rows),
                "legal_pending": states.count("legal_pending"),
                "qa_pending": states.count("qa_pending"),
                "changes_required": states.count("changes_required") + states.count("rejected"),
                "ready_to_release": states.count("ready_to_release"),
                "released": states.count("released"),
                "open_findings": sum(row["open_findings"] for row in rows),
                "invalid_audit_chains": sum(not row["audit_valid"] for row in rows),
            },
            "notice": "La aprobación de la plantilla no aprueba el documento concreto. Cada decisión se vincula a la revisión y al SHA-256 visibles en esta mesa.",
        }

    def bootstrap(self, user: dict[str, Any], *, limit: int = 100) -> dict[str, Any]:
        if user.get("role") != "admin":
            raise PermissionDenied("Solo administración puede preparar inicialmente la bandeja documental.")
        limit_value = max(1, min(int(limit or 100), 500))
        con = self.db_factory()
        try:
            rows = [dict(row) for row in con.execute(
                """SELECT d.id,d.case_id,d.product_code,d.name,d.kind,d.mime_type,d.file_path
                   FROM documents d JOIN cases c ON c.id=d.case_id
                   WHERE d.file_path IS NOT NULL AND d.file_path!=''
                   ORDER BY d.updated_at DESC LIMIT ?""",
                (limit_value,),
            ).fetchall()]
        finally:
            con.close()
        created: list[str] = []
        skipped: list[dict[str, str]] = []
        for row in rows:
            source = Path(str(row.get("file_path") or ""))
            if source.suffix.casefold() != ".docx" or not source.is_file() or row.get("kind") == "audit":
                skipped.append({"document_id": row["id"], "reason": "No es un DOCX revisable disponible."})
                continue
            desk_case_id = self.desk_case_id(row["id"])
            if (self.root / desk_case_id / "case.json").is_file():
                skipped.append({"document_id": row["id"], "reason": "Ya existe en la mesa."})
                continue
            self.desk.create_case(
                case_id=desk_case_id,
                product_code=row["product_code"],
                document_id=row["id"],
                title=row.get("name") or row["id"],
                source_generation_id=row["case_id"],
                actor=_actor(user),
            )
            self.desk.add_revision(
                case_id=desk_case_id,
                source_file=source,
                actor=_actor(user),
                note="Revisión inicial registrada desde el documento vigente de LegalAIZ.it.",
            )
            created.append(desk_case_id)
        return {"schema_version": M32_5_SCHEMA, "created": created, "created_count": len(created), "skipped": skipped}

    def detail(self, user: dict[str, Any], desk_case_id: str) -> dict[str, Any]:
        detail = self._authorize_case(user, desk_case_id)
        current = self._current_revision(detail)
        source_case_id = self._source_case_id(detail)
        detail["schema_version"] = M32_5_SCHEMA
        detail["workflow_status"] = self._workflow_state(detail)
        detail["source_case_id"] = source_case_id
        detail["capabilities"] = {
            "add_revision": user.get("role") in PROFESSIONAL_ROLES,
            "add_finding": user.get("role") in PROFESSIONAL_ROLES,
            "resolve_finding": user.get("role") in PROFESSIONAL_ROLES,
            "legal_approve": user.get("role") == "specialist",
            "qa_approve": user.get("role") == "admin",
            "release": user.get("role") == "admin",
            "download_released": bool(detail.get("release")),
        }
        if current:
            current["preview_url"] = f"/api/m32/approval-desk/cases/{desk_case_id}/revisions/{current['revision_id']}/preview"
        return detail

    def register_current_document(self, user: dict[str, Any], desk_case_id: str, note: str) -> dict[str, Any]:
        detail = self._authorize_case(user, desk_case_id)
        document_id = detail["case"]["document_id"]
        row = self.document_lookup(user, document_id)
        if not row or not row.get("file_path"):
            raise ApprovalDeskError("El documento fuente no existe o no está dentro del alcance del usuario.")
        source = Path(row["file_path"])
        if not source.is_file() or source.suffix.casefold() != ".docx":
            raise ApprovalDeskError("La versión vigente no es un DOCX disponible.")
        current = detail["case"].get("current_revision_id")
        current_manifest = self._current_revision(detail)
        if current_manifest and _sha256_file(source) == current_manifest.get("sha256"):
            raise ImmutableRecordError("El documento vigente ya corresponde al hash de la revisión actual.")
        return self.desk.add_revision(
            case_id=desk_case_id,
            source_file=source,
            actor=_actor(user),
            note=str(note or "Nueva revisión registrada desde el documento vigente.").strip(),
            parent_revision_id=current,
        )

    def upload_revision(self, user: dict[str, Any], desk_case_id: str, filename: str, data: bytes, note: str) -> dict[str, Any]:
        detail = self._authorize_case(user, desk_case_id)
        if len(data) > core.MAX_UPLOAD:
            raise ApprovalDeskError("El archivo supera el límite de 10 MB.")
        detected, digest, security_status = self.upload_validator(filename, data)
        if detected != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise ApprovalDeskError("La revisión debe ser un DOCX válido.")
        original_name = core.safe_filename(Path(filename).name)
        if not original_name.casefold().endswith(".docx"):
            raise ApprovalDeskError("El nombre de la revisión debe terminar en .docx.")
        with TemporaryDirectory(prefix="legalaiz-m325-") as temporary:
            target = Path(temporary) / original_name
            target.write_bytes(data)
            revision = self.desk.add_revision(
                case_id=desk_case_id,
                source_file=target,
                actor=_actor(user),
                note=str(note or "Revisión DOCX cargada desde la Mesa Jurídica.").strip(),
                parent_revision_id=detail["case"].get("current_revision_id"),
            )
        revision["upload_sha256"] = digest
        revision["security_status"] = security_status
        revision["original_filename"] = Path(filename).name
        return revision

    def add_finding(self, user: dict[str, Any], desk_case_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._authorize_case(user, desk_case_id)
        return self.desk.add_finding(
            case_id=desk_case_id,
            revision_id=str(payload.get("revision_id") or ""),
            actor=_actor(user),
            severity=str(payload.get("severity") or ""),
            description=str(payload.get("description") or ""),
            page=payload.get("page"),
            clause=payload.get("clause"),
            block_id=payload.get("block_id"),
        )

    def resolve_finding(self, user: dict[str, Any], desk_case_id: str, finding_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._authorize_case(user, desk_case_id)
        return self.desk.resolve_finding(
            case_id=desk_case_id,
            finding_id=finding_id,
            actor=_actor(user),
            resolution=str(payload.get("resolution") or ""),
            state=str(payload.get("state") or "resolved"),
        )

    def approve(self, user: dict[str, Any], desk_case_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._authorize_case(user, desk_case_id)
        return self.desk.approve(
            case_id=desk_case_id,
            revision_id=str(payload.get("revision_id") or ""),
            approval_type=str(payload.get("approval_type") or ""),
            decision=str(payload.get("decision") or ""),
            actor=_actor(user),
            comment=str(payload.get("comment") or ""),
            expected_sha256=str(payload.get("expected_sha256") or ""),
        )

    def compare(self, user: dict[str, Any], desk_case_id: str, from_revision_id: str, to_revision_id: str) -> dict[str, Any]:
        self._authorize_case(user, desk_case_id)
        return self.desk.compare(
            case_id=desk_case_id,
            from_revision_id=_safe_segment(from_revision_id, "from_revision_id"),
            to_revision_id=_safe_segment(to_revision_id, "to_revision_id"),
        )

    def release(self, user: dict[str, Any], desk_case_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._authorize_case(user, desk_case_id)
        return self.desk.release(
            case_id=desk_case_id,
            revision_id=str(payload.get("revision_id") or ""),
            actor=_actor(user),
            expected_sha256=str(payload.get("expected_sha256") or ""),
        )

    def _revision_file(self, detail: dict[str, Any], revision_id: str) -> tuple[dict[str, Any], Path]:
        revision = next((row for row in detail.get("revisions", []) if row.get("revision_id") == revision_id), None)
        if not revision:
            raise ApprovalDeskError("Revisión documental no encontrada.")
        target = self.root / detail["case"]["case_id"] / "revisions" / revision_id / revision["stored_filename"]
        target = target.resolve()
        expected_root = (self.root / detail["case"]["case_id"] / "revisions" / revision_id).resolve()
        if target.parent != expected_root or not target.is_file() or _sha256_file(target) != revision.get("sha256"):
            raise ReleaseBlocked("El archivo de revisión no coincide con su manifiesto inmutable.")
        return revision, target

    def preview(self, user: dict[str, Any], desk_case_id: str, revision_id: str) -> dict[str, Any]:
        detail = self._authorize_case(user, desk_case_id)
        revision_id = _safe_segment(revision_id, "revision_id")
        revision, source = self._revision_file(detail, revision_id)
        preview_root = self.root / desk_case_id / "previews" / revision_id
        preview_root.mkdir(parents=True, exist_ok=True)
        pdf_path = preview_root / "document.pdf"
        renderer = shutil.which("libreoffice") or shutil.which("soffice")
        render_error = None
        if renderer and not pdf_path.is_file():
            try:
                with TemporaryDirectory(prefix="legalaiz-render-") as profile:
                    environment = dict(os.environ)
                    environment["HOME"] = profile
                    completed = subprocess.run(
                        [renderer, "--headless", "--convert-to", "pdf", "--outdir", str(preview_root), str(source)],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        timeout=90,
                        check=False,
                        env=environment,
                    )
                generated = preview_root / f"{source.stem}.pdf"
                if generated.is_file() and generated != pdf_path:
                    generated.replace(pdf_path)
                if completed.returncode != 0 or not pdf_path.is_file():
                    render_error = (completed.stderr or completed.stdout).decode("utf-8", errors="replace")[-1000:]
            except (OSError, subprocess.SubprocessError) as exc:
                render_error = str(exc)
        if pdf_path.is_file():
            reader = PdfReader(str(pdf_path))
            pages = []
            for index, page in enumerate(reader.pages, 1):
                pages.append({"page": index, "text": (page.extract_text() or "").strip()[:30000]})
            return {
                "schema_version": M32_5_SCHEMA,
                "case_id": desk_case_id,
                "revision_id": revision_id,
                "sha256": revision["sha256"],
                "rendered": True,
                "rendering_engine": "LibreOffice",
                "page_count": len(pages),
                "pages": pages,
                "pdf_url": f"/api/m32/approval-desk/cases/{desk_case_id}/revisions/{revision_id}/preview.pdf",
                "warning": "La vista PDF es una representación de revisión. La liberación continúa ligada al DOCX y SHA-256 aprobados.",
            }
        document = Document(source)
        sections = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table_index, table in enumerate(document.tables, 1):
            sections.append(f"[TABLA {table_index}]")
            sections.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return {
            "schema_version": M32_5_SCHEMA,
            "case_id": desk_case_id,
            "revision_id": revision_id,
            "sha256": revision["sha256"],
            "rendered": False,
            "rendering_engine": None,
            "page_count": None,
            "pages": [],
            "structural_preview": sections[:800],
            "warning": "El motor de paginación no está instalado. Se muestra una vista estructural que no acredita saltos de página ni composición final.",
            "render_error": render_error,
        }

    def preview_pdf_path(self, user: dict[str, Any], desk_case_id: str, revision_id: str) -> Path:
        preview = self.preview(user, desk_case_id, revision_id)
        if not preview.get("rendered"):
            raise ApprovalDeskError("La vista PDF no está disponible en este entorno.")
        target = (self.root / desk_case_id / "previews" / revision_id / "document.pdf").resolve()
        expected = (self.root / desk_case_id / "previews" / revision_id).resolve()
        if target.parent != expected or not target.is_file():
            raise ApprovalDeskError("La vista PDF no está disponible.")
        return target

    def released_path(self, user: dict[str, Any], desk_case_id: str) -> tuple[Path, dict[str, Any]]:
        detail = self._detail_unchecked(desk_case_id)
        source_case_id = self._source_case_id(detail)
        if not source_case_id or not self.access_check(user, source_case_id):
            raise PermissionDenied("El documento liberado no existe o no está dentro del alcance del usuario.")
        release = detail.get("release")
        if not release or detail.get("case", {}).get("status") != "released":
            raise ReleaseBlocked("El documento todavía no cuenta con liberación jurídica y QA vigente.")
        release_id = _safe_segment(release["release_id"], "release_id")
        filename = Path(str(release.get("filename") or "")).name
        target = (self.root / desk_case_id / "releases" / release_id / filename).resolve()
        expected_root = (self.root / desk_case_id / "releases" / release_id).resolve()
        if target.parent != expected_root or not target.is_file() or _sha256_file(target) != release.get("sha256"):
            raise ReleaseBlocked("La copia liberada no conserva el hash aprobado.")
        return target, release

    def audit(self, user: dict[str, Any], desk_case_id: str) -> dict[str, Any]:
        detail = self._authorize_case(user, desk_case_id)
        return {"schema_version": M32_5_SCHEMA, "case_id": desk_case_id, "audit": detail["audit"]}


__all__ = [
    "ApprovalDeskWorkspace",
    "ApprovalDeskError",
    "ImmutableRecordError",
    "PermissionDenied",
    "ReleaseBlocked",
    "M32_5_SCHEMA",
]
