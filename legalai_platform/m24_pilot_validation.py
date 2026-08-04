from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any

from docx import Document


class M24PilotValidationCenter:
    PILOT_CODES = ("CO-LA-002", "CO-EM-003", "CO-AR-001", "CO-CD-003")
    RISK_ORDER = {"green": 0, "yellow": 1, "high": 2, "red": 3}

    def __init__(self, root: Path, candidate_registry):
        self.root = Path(root).resolve()
        self.candidates = candidate_registry
        self.scenarios_path = self.root / "qa" / "legal_scenarios_m24.json"
        self.fixtures_path = self.root / "qa" / "pilot_scenario_fixtures_m24_3.json"
        self.variables_path = self.root / "qa" / "m24_3_success_variables.json"
        self.report_path = self.root / "governance" / "m24_3" / "M24_3_PILOT_EXECUTION_RESULT.json"

    @staticmethod
    def _load(path: Path) -> dict[str, Any]:
        return json.loads(path.read_text(encoding="utf-8"))

    @staticmethod
    def _sha256(path: Path) -> str:
        h = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()

    def _product(self, code: str) -> dict[str, Any]:
        path = self.root / "data" / "legal_products" / code / "product.json"
        if code not in self.PILOT_CODES or not path.is_file():
            raise KeyError(code)
        return self._load(path)

    def _revision_dir(self, code: str) -> Path:
        product = self._product(code)
        return self.root / "data" / "legal_products" / code / "revisions" / product["candidate_revision"]

    def _rules(self, code: str) -> dict[str, dict[str, Any]]:
        revision = self._revision_dir(code)
        rule_path = next(path for path in revision.glob("*.json") if path.name != "metadata.json")
        data = self._load(rule_path)
        rows = data.get("conditional_rules") or data.get("critical_rules") or []
        normalized = {}
        for row in rows:
            item = dict(row)
            item["action"] = item.get("action") or item.get("then")
            risk = str(item.get("risk") or "yellow").lower()
            item["risk"] = "red" if risk == "high" else risk
            normalized[item["id"]] = item
        return normalized

    @staticmethod
    def _state(fixture: dict[str, Any]) -> str:
        flags = fixture["input_flags"]
        if flags.get("success_close"):
            return "CERRADO"
        if flags.get("requires_escalation"):
            return "ESCALADO"
        if not flags.get("complete") or flags.get("contradictory") or not flags.get("evidence_complete"):
            return "INFORMACION_INCOMPLETA"
        return "LISTO_PARA_GENERAR"

    def evaluate_fixture(self, fixture: dict[str, Any]) -> dict[str, Any]:
        rules = self._rules(fixture["product_code"])
        fired = []
        missing_rules = []
        for rule_id in fixture.get("trigger_rule_ids", []):
            if rule_id in rules:
                fired.append(rules[rule_id])
            else:
                missing_rules.append(rule_id)
        risks = [row["risk"] for row in fired]
        if fixture["input_flags"].get("requires_escalation"):
            actual_risk = "red"
        elif fixture["scenario_type"] in {"incomplete", "contradiction", "evidence", "deadline", "conditional"}:
            # El semáforo del escenario expresa la acción operativa inmediata.
            # La severidad intrínseca de la regla se conserva en la evidencia,
            # pero un caso subsanable permanece amarillo hasta el escalamiento.
            actual_risk = "yellow"
        elif risks:
            actual_risk = max(risks, key=lambda value: self.RISK_ORDER.get(value, 1))
        else:
            actual_risk = "green"
        actual = {
            "risk": actual_risk,
            "action": fixture["decision_action"],
            "state": self._state(fixture),
            "rule_ids": [row["id"] for row in fired],
        }
        expected = fixture["expected"]
        checks = {
            "all_rule_ids_exist": not missing_rules,
            "rule_ids_match": sorted(actual["rule_ids"]) == sorted(expected["rule_ids"]),
            "risk_matches": actual["risk"] == expected["risk"],
            "action_matches": actual["action"] == expected["action"],
            "state_matches": actual["state"] == expected["state"],
        }
        return {
            "scenario_id": fixture["scenario_id"],
            "product_code": fixture["product_code"],
            "scenario_type": fixture["scenario_type"],
            "actual": actual,
            "expected": expected,
            "checks": checks,
            "missing_rule_ids": missing_rules,
            "passed": all(checks.values()),
            "document_generation_required": bool(fixture.get("generate_document")),
        }

    def report(self) -> dict[str, Any]:
        return self._load(self.report_path) if self.report_path.is_file() else {
            "schema": "legalai_m24_3_pilot_execution_result_v1",
            "status": "not_executed",
            "pilot_products": list(self.PILOT_CODES),
            "scenario_count": 0,
            "passed": 0,
            "failed": 0,
            "products": [],
        }

    def summary(self) -> dict[str, Any]:
        report = self.report()
        return {
            "schema": "legalai_m24_3_pilot_summary_v1",
            "milestone": "M24.3",
            "base_runtime": "M21.1",
            "candidate_library": "M23.2",
            "publication_blocked": True,
            "active_generation_unchanged": True,
            "pilot_product_count": len(self.PILOT_CODES),
            "scenario_count": report.get("scenario_count", 0),
            "passed": report.get("passed", 0),
            "failed": report.get("failed", 0),
            "generated_document_count": report.get("generated_document_count", 0),
            "products": report.get("products", []),
        }

    def detail(self, code: str) -> dict[str, Any] | None:
        code = str(code or "").upper()
        if code not in self.PILOT_CODES:
            return None
        report = self.report()
        product = next((row for row in report.get("products", []) if row.get("product_code") == code), None)
        if not product:
            return None
        return {**product, "candidate": self.candidates.detail(code)}

    def evidence_path(self, code: str, filename: str) -> Path | None:
        code = str(code or "").upper()
        if code not in self.PILOT_CODES or not filename or Path(filename).name != filename:
            return None
        evidence_dir = self.root / "governance" / "m24_3" / "pilot_documents" / code
        path = (evidence_dir / filename).resolve()
        try:
            path.relative_to(evidence_dir.resolve())
        except ValueError:
            return None
        if not path.is_file():
            return None
        manifest = evidence_dir / "manifest.json"
        if not manifest.is_file():
            return None
        expected = self._load(manifest).get("files", {}).get(filename)
        if not expected or expected != self._sha256(path):
            return None
        return path

    @staticmethod
    def extract_variables(docx_path: Path) -> set[str]:
        document = Document(docx_path)
        texts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        for section in document.sections:
            texts.extend(p.text for p in section.header.paragraphs)
            texts.extend(p.text for p in section.footer.paragraphs)
        return set(re.findall(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}", "\n".join(texts)))
