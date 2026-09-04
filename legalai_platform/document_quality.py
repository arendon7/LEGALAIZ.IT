from __future__ import annotations

import hashlib
import posixpath
import re
from collections import Counter
from pathlib import Path
from zipfile import BadZipFile, ZipFile
import xml.etree.ElementTree as ET

from docx import Document


REQUIRED_OOXML_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
    "word/styles.xml",
}
RELATIONSHIP_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
UNRESOLVED_PATTERNS = (
    re.compile(r"\{\{[^{}]+\}\}"),
    re.compile(r"\$\{[^{}]+\}"),
    re.compile(r"\[\[[^\[\]]+\]\]"),
    re.compile(r"<<[^<>]+>>"),
    re.compile(r"\b(?:NULL|undefined|NaN)\b", re.I),
)
SOFT_SENTINELS = (
    re.compile(r"\bN/?A\b", re.I),
    re.compile(r"_{5,}"),
)


def _relationship_target(rels_path: str, target: str) -> str:
    """Resolve an internal OOXML relationship target to a package part."""
    target = str(target or "").split("#", 1)[0]
    if target.startswith("/"):
        return posixpath.normpath(target.lstrip("/"))
    if rels_path == "_rels/.rels":
        base = ""
    else:
        owner = rels_path.replace("/_rels/", "/")
        if owner.endswith(".rels"):
            owner = owner[:-5]
        base = posixpath.dirname(owner)
    return posixpath.normpath(posixpath.join(base, target)).lstrip("./")


def _content_type_errors(package: ZipFile, names: set[str]) -> list[str]:
    """Validate OPC content-type declarations that Word may otherwise repair.

    A technically readable ZIP can still be repaired by Microsoft Word when its
    [Content_Types].xml contains duplicate declarations, dangling overrides or
    package parts without a matching content type. Those repairs are especially
    undesirable for a document factory because the approved bytes would no longer
    be the bytes the user ultimately edits.
    """
    path = "[Content_Types].xml"
    if path not in names:
        return []  # the required-part check reports the canonical error

    errors: list[str] = []
    try:
        root = ET.fromstring(package.read(path))
    except ET.ParseError:
        return []  # the generic XML parser reports the parse error once

    expected_tag = f"{{{CONTENT_TYPES_NS}}}Types"
    if root.tag != expected_tag:
        errors.append("[Content_Types].xml no usa el elemento Types/namespace OPC esperado.")
        return errors

    defaults: dict[str, str] = {}
    overrides: dict[str, str] = {}
    for node in list(root):
        if node.tag == f"{{{CONTENT_TYPES_NS}}}Default":
            extension = str(node.attrib.get("Extension") or "").strip().lower()
            content_type = str(node.attrib.get("ContentType") or "").strip()
            if not extension or not content_type:
                errors.append("Existe una declaración Default incompleta en [Content_Types].xml.")
                continue
            if extension in defaults:
                errors.append(f"Declaración Default duplicada para la extensión .{extension}.")
            else:
                defaults[extension] = content_type
        elif node.tag == f"{{{CONTENT_TYPES_NS}}}Override":
            raw_part = str(node.attrib.get("PartName") or "").strip()
            content_type = str(node.attrib.get("ContentType") or "").strip()
            if not raw_part.startswith("/") or not content_type:
                errors.append("Existe una declaración Override incompleta o con PartName inválido.")
                continue
            part = posixpath.normpath(raw_part.lstrip("/"))
            if not part or part.startswith("../"):
                errors.append(f"Override inseguro o inválido en [Content_Types].xml: {raw_part}.")
                continue
            if part in overrides:
                errors.append(f"Declaración Override duplicada para /{part}.")
            else:
                overrides[part] = content_type
            if part not in names:
                errors.append(f"Override huérfano en [Content_Types].xml: no existe {part}.")

    for name in sorted(names):
        if name == path or name.endswith("/"):
            continue
        extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if name not in overrides and extension not in defaults:
            errors.append(f"La parte OOXML {name} no tiene tipo de contenido declarado.")
    return errors


def _document_text(document: Document) -> str:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part is not None)


def _duplicate_paragraphs(document: Document) -> list[str]:
    normalized: list[str] = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text or "").strip().casefold()
        if len(text) >= 80:
            normalized.append(text)
    return [text[:160] for text, count in Counter(normalized).items() if count > 1]


def validate_docx(path: str | Path, expected_product: str | None = None) -> dict:
    """Validate a DOCX as an OOXML package and as an editable Word document.

    Errors block delivery. Warnings identify issues that require human review but
    may be legitimate in an editable draft, such as signature lines or N/A.
    """
    file_path = Path(path)
    errors: list[str] = []
    warnings: list[str] = []
    metrics = {
        "bytes": file_path.stat().st_size if file_path.is_file() else 0,
        "package_parts": 0,
        "paragraphs": 0,
        "tables": 0,
        "characters": 0,
    }

    if not file_path.is_file():
        errors.append("El archivo DOCX no existe.")
        return {"valid": False, "errors": errors, "warnings": warnings, "metrics": metrics, "sha256": None}
    if file_path.stat().st_size < 1_000:
        errors.append("El archivo DOCX está vacío o es anormalmente pequeño.")

    names: set[str] = set()
    try:
        with ZipFile(file_path) as package:
            corrupt_part = package.testzip()
            if corrupt_part:
                errors.append(f"La parte OOXML {corrupt_part} no supera la comprobación CRC.")

            package_names = package.namelist()
            duplicate_parts = sorted(name for name, count in Counter(package_names).items() if count > 1)
            if duplicate_parts:
                errors.append(
                    "El paquete DOCX contiene partes OOXML duplicadas: " + ", ".join(duplicate_parts[:20])
                )
            names = set(package_names)
            metrics["package_parts"] = len(names)
            missing = sorted(REQUIRED_OOXML_PARTS - names)
            if missing:
                errors.append("Faltan partes OOXML obligatorias: " + ", ".join(missing))

            errors.extend(_content_type_errors(package, names))

            for name in sorted(names):
                if not (name.endswith(".xml") or name.endswith(".rels")):
                    continue
                try:
                    root = ET.fromstring(package.read(name))
                except ET.ParseError as exc:
                    errors.append(f"XML inválido en {name}: {exc}.")
                    continue
                if not name.endswith(".rels"):
                    continue
                for rel in root.findall(f"{{{RELATIONSHIP_NS}}}Relationship"):
                    if str(rel.attrib.get("TargetMode", "")).casefold() == "external":
                        continue
                    target = rel.attrib.get("Target", "")
                    resolved = _relationship_target(name, target)
                    if not resolved or resolved.startswith("../"):
                        errors.append(f"Relación interna insegura o inválida en {name}: {target}.")
                    elif resolved not in names:
                        errors.append(f"Relación rota en {name}: no existe {resolved}.")
    except BadZipFile:
        errors.append("El archivo no es un paquete DOCX/ZIP válido.")
    except OSError as exc:
        errors.append(f"No fue posible leer el archivo DOCX: {exc}.")

    document = None
    text = ""
    if not errors or names:
        try:
            document = Document(file_path)
            text = _document_text(document)
            metrics["paragraphs"] = len(document.paragraphs)
            metrics["tables"] = len(document.tables)
            metrics["characters"] = len(text.strip())
        except Exception as exc:  # python-docx exposes several package exceptions
            errors.append(f"python-docx no pudo abrir el documento: {type(exc).__name__}: {exc}.")

    if document is not None:
        if len(text.strip()) < 80:
            errors.append("El documento no contiene texto jurídico suficiente.")
        for pattern in UNRESOLVED_PATTERNS:
            matches = sorted(set(pattern.findall(text)))
            if matches:
                errors.append("Se detectaron variables o valores centinela sin resolver: " + ", ".join(map(str, matches[:10])))
        for pattern in SOFT_SENTINELS:
            if pattern.search(text):
                warnings.append(f"Se detectó un marcador editable que requiere revisión humana: {pattern.pattern}.")
        if expected_product:
            metadata = " ".join(
                str(value or "")
                for value in (
                    document.core_properties.title,
                    document.core_properties.subject,
                    document.core_properties.comments,
                )
            )
            if expected_product.casefold() not in (text + "\n" + metadata).casefold():
                errors.append(f"El documento no conserva el identificador esperado {expected_product}.")
        duplicates = _duplicate_paragraphs(document)
        if duplicates:
            warnings.append(f"Se detectaron {len(duplicates)} párrafos extensos duplicados; deben revisarse antes de aprobar.")

    sha256 = hashlib.sha256(file_path.read_bytes()).hexdigest() if file_path.is_file() else None
    return {
        "valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "metrics": metrics,
        "sha256": sha256,
    }


def assert_docx_quality(path: str | Path, expected_product: str | None = None) -> dict:
    report = validate_docx(path, expected_product=expected_product)
    if not report["valid"]:
        raise ValueError("Control de calidad DOCX fallido: " + " | ".join(report["errors"]))
    return report
