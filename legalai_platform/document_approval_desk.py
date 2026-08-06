from __future__ import annotations

"""Mesa documental inmutable con aprobación dual ligada al hash exacto.

La aprobación de una plantilla canónica y la aprobación de un documento generado
son controles distintos. Este módulo gobierna el segundo nivel: revisiones del
archivo concreto, hallazgos por página/cláusula, decisiones independientes del
especialista jurídico y QA, y liberación fail-closed del mismo SHA-256 aprobado.
"""

from dataclasses import dataclass
from datetime import datetime
from difflib import unified_diff
from hashlib import sha256
import json
import os
from pathlib import Path
import re
import shutil
from threading import RLock
from typing import Any
from uuid import uuid4
from zoneinfo import ZoneInfo

from docx import Document


LEGAL_ROLES = frozenset({"specialist", "legal_specialist", "abogado_especialista"})
QA_ROLES = frozenset({"qa", "admin", "qa_specialist"})
REVIEW_ROLES = LEGAL_ROLES | QA_ROLES
FINDING_SEVERITIES = frozenset({"blocking", "major", "minor", "observation"})
FINDING_STATES = frozenset({"open", "resolved", "dismissed"})
APPROVAL_TYPES = frozenset({"legal", "qa"})
DECISIONS = frozenset({"approve", "reject"})
_LOCK = RLock()


class ApprovalDeskError(RuntimeError):
    pass


class PermissionDenied(ApprovalDeskError):
    pass


class ImmutableRecordError(ApprovalDeskError):
    pass


class ReleaseBlocked(ApprovalDeskError):
    pass


@dataclass(frozen=True)
class Actor:
    actor_id: str
    role: str
    display_name: str = ""

    @classmethod
    def from_value(cls, value: dict[str, Any] | "Actor") -> "Actor":
        if isinstance(value, cls):
            return value
        actor_id = str(value.get("id") or value.get("actor_id") or "").strip()
        role = str(value.get("role") or "").strip().casefold()
        if not actor_id or not role:
            raise PermissionDenied("El actor debe incluir identificación y rol.")
        return cls(actor_id=actor_id, role=role, display_name=str(value.get("name") or "").strip())


def _now() -> str:
    return datetime.now(ZoneInfo("America/Bogota")).isoformat(timespec="seconds")


def _canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def _sha256_file(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _atomic_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + f".{uuid4().hex}.tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(temporary, path)


def _read_json(path: Path, default: Any = None) -> Any:
    if not path.is_file():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def _safe_segment(value: str, field: str) -> str:
    text = str(value or "").strip()
    if not text or not re.fullmatch(r"[A-Za-z0-9._-]+", text):
        raise ApprovalDeskError(f"{field} contiene caracteres no permitidos.")
    return text


def _docx_lines(path: Path) -> list[str]:
    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(document.tables, 1):
        lines.append(f"[TABLA {table_index}]")
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return lines


class DocumentApprovalDesk:
    """Repositorio append-only de expedientes documentales revisables."""

    SCHEMA_VERSION = "M32.4"

    def __init__(self, root: str | Path):
        self.root = Path(root).resolve()
        self.root.mkdir(parents=True, exist_ok=True)

    def _case_dir(self, case_id: str) -> Path:
        return self.root / _safe_segment(case_id, "case_id")

    def _case_path(self, case_id: str) -> Path:
        return self._case_dir(case_id) / "case.json"

    def _events_path(self, case_id: str) -> Path:
        return self._case_dir(case_id) / "events.jsonl"

    def _revision_dir(self, case_id: str, revision_id: str) -> Path:
        return self._case_dir(case_id) / "revisions" / _safe_segment(revision_id, "revision_id")

    def _revision_manifest_path(self, case_id: str, revision_id: str) -> Path:
        return self._revision_dir(case_id, revision_id) / "revision.json"

    def _load_case(self, case_id: str) -> dict[str, Any]:
        payload = _read_json(self._case_path(case_id))
        if not isinstance(payload, dict):
            raise ApprovalDeskError("Expediente documental no encontrado.")
        return payload

    def _load_revision(self, case_id: str, revision_id: str) -> dict[str, Any]:
        payload = _read_json(self._revision_manifest_path(case_id, revision_id))
        if not isinstance(payload, dict):
            raise ApprovalDeskError("Revisión documental no encontrada.")
        return payload

    def _append_event(self, case_id: str, event_type: str, actor: Actor, payload: dict[str, Any]) -> dict[str, Any]:
        path = self._events_path(case_id)
        previous_hash = "0" * 64
        sequence = 1
        if path.is_file():
            lines = [line for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
            if lines:
                previous = json.loads(lines[-1])
                previous_hash = previous["event_hash"]
                sequence = int(previous["sequence"]) + 1
        event = {
            "schema_version": self.SCHEMA_VERSION,
            "sequence": sequence,
            "event_id": f"EVT-{sequence:06d}",
            "event_type": event_type,
            "created_at": _now(),
            "actor": {
                "id": actor.actor_id,
                "role": actor.role,
                "name": actor.display_name,
            },
            "payload": payload,
            "previous_hash": previous_hash,
        }
        event["event_hash"] = sha256(_canonical_json(event).encode("utf-8")).hexdigest()
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("a", encoding="utf-8") as stream:
            stream.write(json.dumps(event, ensure_ascii=False, separators=(",", ":")) + "\n")
            stream.flush()
            os.fsync(stream.fileno())
        return event

    def create_case(
        self,
        *,
        case_id: str,
        product_code: str,
        document_id: str,
        title: str,
        actor: dict[str, Any] | Actor,
        source_generation_id: str | None = None,
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        case_dir = self._case_dir(case_id)
        with _LOCK:
            if case_dir.exists():
                raise ImmutableRecordError("El expediente ya existe y no puede sobrescribirse.")
            case_dir.mkdir(parents=True)
            payload = {
                "schema_version": self.SCHEMA_VERSION,
                "case_id": _safe_segment(case_id, "case_id"),
                "product_code": _safe_segment(product_code, "product_code").upper(),
                "document_id": _safe_segment(document_id, "document_id"),
                "title": str(title).strip(),
                "source_generation_id": str(source_generation_id or "").strip() or None,
                "created_at": _now(),
                "created_by": actor_value.actor_id,
                "current_revision_id": None,
                "revision_count": 0,
                "release_id": None,
                "status": "draft",
            }
            _atomic_json(self._case_path(case_id), payload)
            self._append_event(case_id, "case.created", actor_value, payload)
            return payload

    def add_revision(
        self,
        *,
        case_id: str,
        source_file: str | Path,
        actor: dict[str, Any] | Actor,
        note: str,
        parent_revision_id: str | None = None,
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        source = Path(source_file).resolve()
        if source.suffix.casefold() != ".docx" or not source.is_file():
            raise ApprovalDeskError("La revisión debe provenir de un archivo DOCX existente.")
        Document(source)  # apertura obligatoria antes de persistir
        with _LOCK:
            case = self._load_case(case_id)
            current = case.get("current_revision_id")
            if current and parent_revision_id != current:
                raise ImmutableRecordError(
                    "La nueva revisión debe declarar como padre la revisión vigente para evitar bifurcaciones silenciosas."
                )
            if not current and parent_revision_id:
                raise ApprovalDeskError("La primera revisión no admite revisión padre.")
            revision_number = int(case.get("revision_count") or 0) + 1
            revision_id = f"REV-{revision_number:04d}"
            revision_dir = self._revision_dir(case_id, revision_id)
            if revision_dir.exists():
                raise ImmutableRecordError("La revisión calculada ya existe.")
            revision_dir.mkdir(parents=True)
            target = revision_dir / "document.docx"
            shutil.copy2(source, target)
            digest = _sha256_file(target)
            manifest = {
                "schema_version": self.SCHEMA_VERSION,
                "case_id": case_id,
                "revision_id": revision_id,
                "revision_number": revision_number,
                "parent_revision_id": current,
                "filename": source.name,
                "stored_filename": target.name,
                "sha256": digest,
                "size_bytes": target.stat().st_size,
                "created_at": _now(),
                "created_by": actor_value.actor_id,
                "note": str(note or "").strip(),
                "approval_state": {"legal": "pending", "qa": "pending"},
                "release_candidate": False,
            }
            _atomic_json(self._revision_manifest_path(case_id, revision_id), manifest)
            case.update({
                "current_revision_id": revision_id,
                "revision_count": revision_number,
                "release_id": None,
                "status": "in_review",
            })
            _atomic_json(self._case_path(case_id), case)
            self._append_event(case_id, "revision.created", actor_value, manifest)
            return manifest

    def add_finding(
        self,
        *,
        case_id: str,
        revision_id: str,
        actor: dict[str, Any] | Actor,
        severity: str,
        description: str,
        page: int | None = None,
        clause: str | None = None,
        block_id: str | None = None,
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        if actor_value.role not in REVIEW_ROLES:
            raise PermissionDenied("Solo especialista jurídico o QA pueden registrar hallazgos.")
        severity_value = str(severity).casefold()
        if severity_value not in FINDING_SEVERITIES:
            raise ApprovalDeskError("Severidad de hallazgo inválida.")
        revision = self._load_revision(case_id, revision_id)
        if revision_id != self._load_case(case_id).get("current_revision_id"):
            raise ImmutableRecordError("Los hallazgos nuevos solo pueden vincularse a la revisión vigente.")
        finding_id = f"FND-{uuid4().hex[:12].upper()}"
        finding = {
            "finding_id": finding_id,
            "case_id": case_id,
            "revision_id": revision_id,
            "revision_sha256": revision["sha256"],
            "severity": severity_value,
            "state": "open",
            "description": str(description).strip(),
            "locator": {
                "page": int(page) if page is not None else None,
                "clause": str(clause or "").strip() or None,
                "block_id": str(block_id or "").strip() or None,
            },
            "created_at": _now(),
            "created_by": actor_value.actor_id,
        }
        path = self._case_dir(case_id) / "findings" / f"{finding_id}.json"
        _atomic_json(path, finding)
        self._append_event(case_id, "finding.created", actor_value, finding)
        return finding

    def resolve_finding(
        self,
        *,
        case_id: str,
        finding_id: str,
        actor: dict[str, Any] | Actor,
        resolution: str,
        state: str = "resolved",
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        if actor_value.role not in REVIEW_ROLES:
            raise PermissionDenied("Solo especialista jurídico o QA pueden resolver hallazgos.")
        state_value = str(state).casefold()
        if state_value not in {"resolved", "dismissed"}:
            raise ApprovalDeskError("El cierre debe ser resolved o dismissed.")
        original_path = self._case_dir(case_id) / "findings" / f"{_safe_segment(finding_id, 'finding_id')}.json"
        original = _read_json(original_path)
        if not isinstance(original, dict):
            raise ApprovalDeskError("Hallazgo no encontrado.")
        closure_path = self._case_dir(case_id) / "finding-resolutions" / f"{finding_id}.json"
        if closure_path.exists():
            raise ImmutableRecordError("El hallazgo ya tiene una resolución inmutable.")
        closure = {
            "finding_id": finding_id,
            "case_id": case_id,
            "revision_id": original["revision_id"],
            "state": state_value,
            "resolution": str(resolution).strip(),
            "resolved_at": _now(),
            "resolved_by": actor_value.actor_id,
        }
        _atomic_json(closure_path, closure)
        self._append_event(case_id, "finding.closed", actor_value, closure)
        return closure

    def _finding_state(self, case_id: str, finding: dict[str, Any]) -> str:
        closure = _read_json(
            self._case_dir(case_id) / "finding-resolutions" / f"{finding['finding_id']}.json"
        )
        return str(closure.get("state")) if isinstance(closure, dict) else "open"

    def findings(self, case_id: str, revision_id: str | None = None) -> list[dict[str, Any]]:
        folder = self._case_dir(case_id) / "findings"
        values: list[dict[str, Any]] = []
        for path in sorted(folder.glob("*.json")) if folder.is_dir() else []:
            finding = _read_json(path)
            if revision_id and finding.get("revision_id") != revision_id:
                continue
            finding = dict(finding)
            finding["state"] = self._finding_state(case_id, finding)
            values.append(finding)
        return values

    def approve(
        self,
        *,
        case_id: str,
        revision_id: str,
        approval_type: str,
        decision: str,
        actor: dict[str, Any] | Actor,
        comment: str,
        expected_sha256: str,
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        approval_value = str(approval_type).casefold()
        decision_value = str(decision).casefold()
        if approval_value not in APPROVAL_TYPES or decision_value not in DECISIONS:
            raise ApprovalDeskError("Tipo o decisión de aprobación inválidos.")
        if approval_value == "legal" and actor_value.role not in LEGAL_ROLES:
            raise PermissionDenied("La aprobación jurídica requiere rol de especialista jurídico.")
        if approval_value == "qa" and actor_value.role not in QA_ROLES:
            raise PermissionDenied("La aprobación QA requiere rol QA o administrador autorizado.")
        case = self._load_case(case_id)
        if revision_id != case.get("current_revision_id"):
            raise ReleaseBlocked("No es posible aprobar una revisión histórica como revisión vigente.")
        revision = self._load_revision(case_id, revision_id)
        stored_file = self._revision_dir(case_id, revision_id) / revision["stored_filename"]
        actual_hash = _sha256_file(stored_file)
        expected = str(expected_sha256).casefold()
        if actual_hash != revision["sha256"] or expected != actual_hash:
            raise ReleaseBlocked("El hash esperado, el manifiesto y el archivo no coinciden.")
        open_blockers = [
            finding for finding in self.findings(case_id, revision_id)
            if finding["state"] == "open" and finding["severity"] in {"blocking", "major"}
        ]
        if decision_value == "approve" and open_blockers:
            raise ReleaseBlocked("Existen hallazgos bloqueantes o mayores sin resolver.")
        approvals_dir = self._revision_dir(case_id, revision_id) / "approvals"
        approval_path = approvals_dir / f"{approval_value}.json"
        if approval_path.exists():
            raise ImmutableRecordError("La decisión ya existe; una nueva decisión requiere nueva revisión.")
        other_type = "qa" if approval_value == "legal" else "legal"
        other = _read_json(approvals_dir / f"{other_type}.json")
        if isinstance(other, dict) and other.get("actor", {}).get("id") == actor_value.actor_id:
            raise ReleaseBlocked("La aprobación jurídica y QA deben pertenecer a personas distintas.")
        if approval_value == "qa" and decision_value == "approve":
            legal = _read_json(approvals_dir / "legal.json")
            if not isinstance(legal, dict) or legal.get("decision") != "approve":
                raise ReleaseBlocked("QA solo puede aprobar después de la aprobación jurídica del mismo hash.")
        approval = {
            "schema_version": self.SCHEMA_VERSION,
            "case_id": case_id,
            "revision_id": revision_id,
            "sha256": actual_hash,
            "approval_type": approval_value,
            "decision": decision_value,
            "comment": str(comment or "").strip(),
            "actor": {
                "id": actor_value.actor_id,
                "role": actor_value.role,
                "name": actor_value.display_name,
            },
            "created_at": _now(),
        }
        approval["record_hash"] = sha256(_canonical_json(approval).encode("utf-8")).hexdigest()
        _atomic_json(approval_path, approval)
        self._append_event(case_id, f"approval.{approval_value}.{decision_value}", actor_value, approval)
        return approval

    def compare(self, *, case_id: str, from_revision_id: str, to_revision_id: str) -> dict[str, Any]:
        source_revision = self._load_revision(case_id, from_revision_id)
        target_revision = self._load_revision(case_id, to_revision_id)
        source_file = self._revision_dir(case_id, from_revision_id) / source_revision["stored_filename"]
        target_file = self._revision_dir(case_id, to_revision_id) / target_revision["stored_filename"]
        source_lines = _docx_lines(source_file)
        target_lines = _docx_lines(target_file)
        diff = list(unified_diff(
            source_lines,
            target_lines,
            fromfile=from_revision_id,
            tofile=to_revision_id,
            lineterm="",
        ))
        return {
            "case_id": case_id,
            "from_revision_id": from_revision_id,
            "to_revision_id": to_revision_id,
            "from_sha256": source_revision["sha256"],
            "to_sha256": target_revision["sha256"],
            "changed": source_revision["sha256"] != target_revision["sha256"],
            "diff_lines": diff,
            "summary": {
                "added_lines": sum(line.startswith("+") and not line.startswith("+++") for line in diff),
                "removed_lines": sum(line.startswith("-") and not line.startswith("---") for line in diff),
            },
        }

    def release(
        self,
        *,
        case_id: str,
        revision_id: str,
        actor: dict[str, Any] | Actor,
        expected_sha256: str,
    ) -> dict[str, Any]:
        actor_value = Actor.from_value(actor)
        if actor_value.role not in QA_ROLES:
            raise PermissionDenied("La liberación requiere rol QA o administrador autorizado.")
        with _LOCK:
            case = self._load_case(case_id)
            if case.get("release_id"):
                raise ImmutableRecordError("El expediente ya tiene una liberación inmutable.")
            if revision_id != case.get("current_revision_id"):
                raise ReleaseBlocked("Solo la revisión vigente puede liberarse.")
            revision = self._load_revision(case_id, revision_id)
            stored_file = self._revision_dir(case_id, revision_id) / revision["stored_filename"]
            actual_hash = _sha256_file(stored_file)
            expected = str(expected_sha256).casefold()
            if actual_hash != revision["sha256"] or actual_hash != expected:
                raise ReleaseBlocked("El archivo cambió o el hash solicitado no coincide.")
            approvals_dir = self._revision_dir(case_id, revision_id) / "approvals"
            legal = _read_json(approvals_dir / "legal.json")
            qa = _read_json(approvals_dir / "qa.json")
            for label, approval in (("jurídica", legal), ("QA", qa)):
                if not isinstance(approval, dict) or approval.get("decision") != "approve":
                    raise ReleaseBlocked(f"Falta aprobación {label} válida.")
                if approval.get("sha256") != actual_hash or approval.get("revision_id") != revision_id:
                    raise ReleaseBlocked(f"La aprobación {label} no corresponde al archivo vigente.")
            if legal["actor"]["id"] == qa["actor"]["id"]:
                raise ReleaseBlocked("Las aprobaciones deben pertenecer a personas distintas.")
            open_findings = [finding for finding in self.findings(case_id, revision_id) if finding["state"] == "open"]
            if open_findings:
                raise ReleaseBlocked("Todos los hallazgos deben cerrarse antes de liberar.")
            release_id = f"REL-{uuid4().hex[:12].upper()}"
            release_dir = self._case_dir(case_id) / "releases" / release_id
            release_dir.mkdir(parents=True)
            released_document = release_dir / revision["filename"]
            shutil.copy2(stored_file, released_document)
            if _sha256_file(released_document) != actual_hash:
                raise ReleaseBlocked("La copia liberada no conserva el hash aprobado.")
            release = {
                "schema_version": self.SCHEMA_VERSION,
                "release_id": release_id,
                "case_id": case_id,
                "revision_id": revision_id,
                "sha256": actual_hash,
                "filename": released_document.name,
                "legal_approval_record_hash": legal["record_hash"],
                "qa_approval_record_hash": qa["record_hash"],
                "released_at": _now(),
                "released_by": actor_value.actor_id,
                "status": "released_exact_hash",
            }
            release["release_record_hash"] = sha256(_canonical_json(release).encode("utf-8")).hexdigest()
            _atomic_json(release_dir / "release.json", release)
            case.update({"release_id": release_id, "status": "released"})
            _atomic_json(self._case_path(case_id), case)
            self._append_event(case_id, "document.released", actor_value, release)
            return release

    def verify_audit_chain(self, case_id: str) -> dict[str, Any]:
        path = self._events_path(case_id)
        events = []
        if path.is_file():
            events = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
        previous_hash = "0" * 64
        errors: list[str] = []
        for expected_sequence, event in enumerate(events, 1):
            stored_hash = event.get("event_hash")
            unsigned = dict(event)
            unsigned.pop("event_hash", None)
            calculated = sha256(_canonical_json(unsigned).encode("utf-8")).hexdigest()
            if event.get("sequence") != expected_sequence:
                errors.append(f"Secuencia inválida en evento {expected_sequence}.")
            if event.get("previous_hash") != previous_hash:
                errors.append(f"Cadena previa inválida en evento {expected_sequence}.")
            if stored_hash != calculated:
                errors.append(f"Hash inválido en evento {expected_sequence}.")
            previous_hash = str(stored_hash or "")
        return {"valid": not errors, "events": len(events), "errors": errors, "last_hash": previous_hash}

    def detail(self, case_id: str) -> dict[str, Any]:
        case = self._load_case(case_id)
        revisions = []
        revisions_root = self._case_dir(case_id) / "revisions"
        for path in sorted(revisions_root.glob("*/revision.json")) if revisions_root.is_dir() else []:
            revision = _read_json(path)
            approvals = {}
            approvals_dir = path.parent / "approvals"
            for approval_type in APPROVAL_TYPES:
                value = _read_json(approvals_dir / f"{approval_type}.json")
                approvals[approval_type] = value
            revision["approvals"] = approvals
            revision["findings"] = self.findings(case_id, revision["revision_id"])
            revisions.append(revision)
        release = None
        if case.get("release_id"):
            release = _read_json(self._case_dir(case_id) / "releases" / case["release_id"] / "release.json")
        return {
            "case": case,
            "revisions": revisions,
            "release": release,
            "audit": self.verify_audit_chain(case_id),
        }
