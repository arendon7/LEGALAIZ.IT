from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


_ORDINAL = (
    r"(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|"
    r"DÉCIMA\s+PRIMERA|DÉCIMA\s+SEGUNDA|DÉCIMA\s+TERCERA|DÉCIMA\s+CUARTA|"
    r"DÉCIMA\s+QUINTA|DÉCIMA\s+SEXTA|DÉCIMA\s+SÉPTIMA|DÉCIMA\s+OCTAVA|DÉCIMA\s+NOVENA|"
    r"VIGÉSIMA(?:\s+(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA))?|"
    r"TRIGÉSIMA(?:\s+(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA))?|"
    r"CUADRAGÉSIMA(?:\s+(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA))?|\d+)"
)
CLAUSE_PATTERN = re.compile(rf"^(?:CLÁUSULA\s+)?{_ORDINAL}(?:\.|:|\b)", re.I)
SIGNATURE_PATTERN = re.compile(r"\bFIRMA(?:S)?\b|\bEL EMPLEADOR\b|\bEL TRABAJADOR\b|\bLAS PARTES\b", re.I)


def _centimeters(length) -> float | None:
    if length is None:
        return None
    try:
        return round(float(length.cm), 3)
    except (AttributeError, TypeError, ValueError):
        return None


def _has_page_field(document: Document) -> bool:
    for section in document.sections:
        xml = section.footer._element.xml
        if "PAGE" in xml or "w:pgNumType" in xml:
            return True
    return False


def _table_header_repeats(table) -> bool:
    if not table.rows:
        return False
    row_properties = table.rows[0]._tr.trPr
    return bool(row_properties is not None and row_properties.find(qn("w:tblHeader")) is not None)


def _row_cannot_split(row) -> bool:
    row_properties = row._tr.trPr
    return bool(row_properties is not None and row_properties.find(qn("w:cantSplit")) is not None)


def _paragraph_is_heading(paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False
    style_name = str(getattr(paragraph.style, "name", "") or "").casefold()
    if style_name.startswith("heading") or style_name in {"title", "título"}:
        return True
    visible_runs = [run for run in paragraph.runs if (run.text or "").strip()]
    return bool(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        and text == text.upper()
        and len(text) <= 120
        and visible_runs
        and all(run.bold for run in visible_runs)
    )


def validate_visual_structure(path: str | Path, expected_product: str | None = None) -> dict:
    """Perform a deterministic visual preflight over a DOCX structure.

    This is not a substitute for page-by-page human inspection after rendering.
    Errors identify layout configurations that are predictably unusable. Warnings
    identify risks that must be reviewed in the rendered PDF/PNG evidence.
    """
    file_path = Path(path)
    errors: list[str] = []
    warnings: list[str] = []
    metrics = {
        "sections": 0,
        "headings": 0,
        "clauses": 0,
        "tables": 0,
        "tables_with_repeating_header": 0,
        "rows_protected_from_split": 0,
        "long_paragraphs": 0,
        "consecutive_empty_paragraphs": 0,
        "has_page_field": False,
        "has_signature_area": False,
        "page_layouts": [],
    }

    if not file_path.is_file():
        return {
            "valid": False,
            "errors": ["El archivo DOCX no existe para el preflight visual."],
            "warnings": [],
            "metrics": metrics,
            "requires_human_visual_review": True,
        }

    try:
        document = Document(file_path)
    except Exception as exc:
        return {
            "valid": False,
            "errors": [f"No fue posible abrir el DOCX para el preflight visual: {type(exc).__name__}: {exc}."],
            "warnings": [],
            "metrics": metrics,
            "requires_human_visual_review": True,
        }

    metrics["sections"] = len(document.sections)
    if not document.sections:
        errors.append("El documento no contiene una sección de página válida.")

    for index, section in enumerate(document.sections, 1):
        width = _centimeters(section.page_width)
        height = _centimeters(section.page_height)
        top = _centimeters(section.top_margin)
        bottom = _centimeters(section.bottom_margin)
        left = _centimeters(section.left_margin)
        right = _centimeters(section.right_margin)
        layout = {
            "section": index,
            "page_width_cm": width,
            "page_height_cm": height,
            "top_margin_cm": top,
            "bottom_margin_cm": bottom,
            "left_margin_cm": left,
            "right_margin_cm": right,
        }
        metrics["page_layouts"].append(layout)
        if width is None or height is None or width < 15 or height < 20:
            errors.append(f"La sección {index} tiene un tamaño de página inválido o anormal.")
        margins = [value for value in (top, bottom, left, right) if value is not None]
        if len(margins) != 4 or any(value < 1.2 or value > 4.0 for value in margins):
            warnings.append(f"La sección {index} usa márgenes fuera del rango recomendado de 1,2 a 4,0 cm.")
        if width and left is not None and right is not None and width - left - right < 10:
            errors.append(f"La sección {index} deja un ancho imprimible inferior a 10 cm.")
        if height and top is not None and bottom is not None and height - top - bottom < 15:
            errors.append(f"La sección {index} deja una altura imprimible inferior a 15 cm.")

    empty_run = 0
    text_parts: list[str] = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text or "").strip()
        if not text:
            empty_run += 1
            if empty_run >= 3:
                metrics["consecutive_empty_paragraphs"] += 1
            continue
        empty_run = 0
        text_parts.append(text)
        if _paragraph_is_heading(paragraph):
            metrics["headings"] += 1
            keep_with_next = paragraph.paragraph_format.keep_with_next
            if keep_with_next is False:
                warnings.append(f"El encabezado «{text[:80]}» permite quedar huérfano al final de página.")
        if CLAUSE_PATTERN.match(text):
            metrics["clauses"] += 1
            if not paragraph.runs or not paragraph.runs[0].bold:
                warnings.append(f"La cláusula «{text[:80]}» no inicia con ordinal y título en negrita.")
        if len(text) > 2_000:
            metrics["long_paragraphs"] += 1
            warnings.append(f"Existe un párrafo de {len(text)} caracteres con riesgo de lectura y paginación.")

    for table in document.tables:
        metrics["tables"] += 1
        if len(table.columns) > 7:
            warnings.append(f"Una tabla contiene {len(table.columns)} columnas y puede desbordarse en formato vertical.")
        if _table_header_repeats(table):
            metrics["tables_with_repeating_header"] += 1
        elif len(table.rows) >= 4:
            warnings.append("Una tabla de cuatro o más filas no tiene encabezado repetible configurado.")
        for row in table.rows:
            if _row_cannot_split(row):
                metrics["rows_protected_from_split"] += 1

    full_text = "\n".join(text_parts)
    metrics["has_signature_area"] = bool(SIGNATURE_PATTERN.search(full_text))
    metrics["has_page_field"] = _has_page_field(document)

    if metrics["headings"] == 0:
        warnings.append("No se detectaron encabezados jurídicos diferenciados para navegación visual.")
    if expected_product and expected_product in {"CO-EM-003", "CO-EM-004", "CO-AR-001", "CO-LA-002"} and metrics["clauses"] < 5:
        warnings.append(f"El documento {expected_product} contiene menos de cinco cláusulas detectables; debe verificarse que no sea un anexo breve.")
    if not metrics["has_signature_area"]:
        warnings.append("No se detectó una zona de firmas; puede ser legítimo en informes o comunicaciones, pero requiere revisión.")
    if not metrics["has_page_field"]:
        warnings.append("No se detectó un campo de número de página en el pie de página.")
    if metrics["consecutive_empty_paragraphs"]:
        warnings.append("Se detectaron tres o más párrafos vacíos consecutivos, con riesgo de saltos visuales innecesarios.")

    return {
        "valid": not errors,
        "errors": errors,
        "warnings": list(dict.fromkeys(warnings)),
        "metrics": metrics,
        "requires_human_visual_review": True,
        "review_statement": "El preflight estructural no sustituye la inspección humana de todas las páginas renderizadas.",
    }


def assert_visual_structure(path: str | Path, expected_product: str | None = None) -> dict:
    report = validate_visual_structure(path, expected_product=expected_product)
    if not report["valid"]:
        raise ValueError("Preflight visual DOCX fallido: " + " | ".join(report["errors"]))
    return report