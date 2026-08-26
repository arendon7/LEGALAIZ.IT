from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any
import re
import unicodedata

from docx import Document


STANDARD = "M38.7-CONTRACT-COMPOSITION-COVERAGE"


class ContractCompositionCoverageError(RuntimeError):
    pass


# Estas políticas son deliberadamente de cobertura estructural, no de suficiencia
# jurídica. Exigen que un contrato siga teniendo una composición materialmente
# profunda y que las familias esenciales sobrevivan al render final.
POLICIES: dict[str, dict[str, Any]] = {
    "CO-EM-003": {
        "min_clauses": 18,
        "min_text_chars": 6500,
        "families": {
            "object": ("objeto",),
            "scope": ("alcance", "exclusiones"),
            "deliverables": ("entregables", "aceptacion"),
            "independence": ("autonomia", "laboralidad", "subordinacion"),
            "economic": ("honorarios", "facturacion", "pago"),
            "risk_liability": ("responsabilidad", "riesgos", "indemn"),
            "termination": ("terminacion", "cierre"),
            "disputes": ("controversias", "ley aplicable", "domicilio"),
        },
    },
    "CO-EM-004": {
        "min_clauses": 14,
        "min_text_chars": 5500,
        "families": {
            "confidentiality": ("confidencial",),
            "trade_secret": ("secreto empresarial",),
            "purpose_use": ("finalidad", "uso autorizado", "proposito"),
            "exclusions": ("exclusiones", "no constituye informacion confidencial"),
            "access_security": ("acceso", "seguridad", "incidente"),
            "term": ("vigencia", "duracion", "plazo"),
            "return_deletion": ("devolucion", "eliminacion", "destruccion"),
            "remedies_disputes": ("responsabilidad", "medidas", "controversias"),
        },
    },
    "CO-AR-001": {
        "min_clauses": 16,
        "min_text_chars": 6000,
        "families": {
            "property_object": ("inmueble", "objeto"),
            "term": ("duracion", "plazo"),
            "rent_payment": ("canon", "pago"),
            "delivery_inventory": ("entrega", "inventario"),
            "party_obligations": ("obligaciones", "arrendador", "arrendatario"),
            "maintenance_repairs": ("reparaciones", "mantenimiento"),
            "utilities_charges": ("servicios publicos", "administracion"),
            "termination_restitution": ("terminacion", "restitucion"),
        },
    },
    "CO-LA-002": {
        "min_clauses": 16,
        "min_text_chars": 6000,
        "families": {
            "role_functions": ("cargo", "funciones", "objeto"),
            "term": ("termino indefinido", "duracion"),
            "salary": ("salario", "remuneracion"),
            "working_time": ("jornada", "horas semanales"),
            "workplace_mode": ("lugar de trabajo", "modalidad", "trabajo remoto"),
            "duties": ("obligaciones", "deberes"),
            "social_security": ("seguridad social", "riesgos laborales", "sg-sst"),
            "termination": ("terminacion", "justa causa"),
        },
    },
}


def _normalize(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.casefold()
    return re.sub(r"\s+", " ", text).strip()


def _section_text(section: dict[str, Any]) -> str:
    values: list[str] = []
    for key in ("heading", "text", "notes"):
        if section.get(key):
            values.append(str(section[key]))
    for key in ("paragraphs", "bullets", "numbered"):
        if isinstance(section.get(key), list):
            values.extend(str(item) for item in section[key] if item not in (None, ""))
    if isinstance(section.get("table"), list):
        for row in section["table"]:
            if isinstance(row, (list, tuple)):
                values.extend(str(cell) for cell in row if cell not in (None, ""))
    if isinstance(section.get("parties"), list):
        for party in section["parties"]:
            if isinstance(party, dict):
                values.extend(str(value) for value in party.values() if value not in (None, ""))
    return "\n".join(values)


def composition_text(sections: list[dict[str, Any]]) -> str:
    return _normalize("\n".join(_section_text(section) for section in sections or []))


def rendered_docx_text(path: str | Path) -> str:
    document = Document(Path(path))
    values: list[str] = []
    values.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    for section in document.sections:
        for area in (section.header, section.footer):
            values.extend(paragraph.text for paragraph in area.paragraphs if paragraph.text)
            for table in area.tables:
                for row in table.rows:
                    for cell in row.cells:
                        values.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return _normalize("\n".join(values))


def _match_family(text: str, aliases: tuple[str, ...]) -> str | None:
    for alias in aliases:
        normalized = _normalize(alias)
        if normalized and normalized in text:
            return alias
    return None


def _public_clause_count(sections: list[dict[str, Any]]) -> int:
    count = 0
    for section in sections or []:
        if section.get("_type") == "control":
            continue
        heading = _normalize(section.get("heading"))
        if section.get("_type") == "clause" or section.get("clause_number") or re.match(
            r"^(primera|segunda|tercera|cuarta|quinta|sexta|septima|octava|novena|decima|vigesima|trigesima|cuadragesima|quincuagesima)",
            heading,
        ):
            count += 1
    return count


def assess_contractual_coverage(
    *,
    product_code: str,
    sections: list[dict[str, Any]],
    rendered_text: str,
) -> dict[str, Any]:
    code = str(product_code or "").strip().upper()
    policy = POLICIES.get(code)
    if not policy:
        return {
            "standard": STANDARD,
            "applicable": False,
            "product_code": code,
            "passed": True,
            "legal_sufficiency_claimed": False,
        }

    source_sections = deepcopy(list(sections or []))
    source_text = composition_text(source_sections)
    final_text = _normalize(rendered_text)
    clause_count = _public_clause_count(source_sections)
    signature_in_composition = any(section.get("_type") == "signature" for section in source_sections)
    signature_in_render = "firmas" in final_text

    families: dict[str, dict[str, Any]] = {}
    blockers: list[str] = []
    for family, aliases in policy["families"].items():
        composition_hit = _match_family(source_text, tuple(aliases))
        rendered_hit = _match_family(final_text, tuple(aliases))
        passed = bool(composition_hit and rendered_hit)
        families[family] = {
            "passed": passed,
            "composition_match": composition_hit,
            "rendered_match": rendered_hit,
        }
        if not passed:
            blockers.append(f"family_missing:{family}")

    if clause_count < int(policy["min_clauses"]):
        blockers.append(f"clause_count:{clause_count}<{policy['min_clauses']}")
    if len(final_text) < int(policy["min_text_chars"]):
        blockers.append(f"rendered_text_chars:{len(final_text)}<{policy['min_text_chars']}")
    if not signature_in_composition:
        blockers.append("signature_missing:composition")
    if not signature_in_render:
        blockers.append("signature_missing:rendered")

    return {
        "standard": STANDARD,
        "applicable": True,
        "product_code": code,
        "passed": not blockers,
        "public_clause_count": clause_count,
        "minimum_public_clauses": int(policy["min_clauses"]),
        "rendered_text_chars": len(final_text),
        "minimum_rendered_text_chars": int(policy["min_text_chars"]),
        "signature_in_composition": signature_in_composition,
        "signature_in_render": signature_in_render,
        "families": families,
        "blockers": blockers,
        "legal_sufficiency_claimed": False,
        "review_rule": (
            "La cobertura estructural evita regresiones hacia instrumentos básicos, pero no acredita suficiencia jurídica, "
            "vigencia normativa ni adecuación al caso concreto; la liberación conserva revisión jurídica y QA independientes."
        ),
    }


def assess_contractual_docx(
    path: str | Path,
    *,
    product_code: str,
    sections: list[dict[str, Any]],
) -> dict[str, Any]:
    return assess_contractual_coverage(
        product_code=product_code,
        sections=sections,
        rendered_text=rendered_docx_text(path),
    )


def assert_contractual_docx_coverage(
    path: str | Path,
    *,
    product_code: str,
    sections: list[dict[str, Any]],
) -> dict[str, Any]:
    report = assess_contractual_docx(path, product_code=product_code, sections=sections)
    if report.get("applicable") and not report.get("passed"):
        raise ContractCompositionCoverageError(
            f"{product_code}: cobertura contractual M38.7 insuficiente: {report.get('blockers')}"
        )
    return report


__all__ = [
    "ContractCompositionCoverageError",
    "POLICIES",
    "STANDARD",
    "assess_contractual_coverage",
    "assess_contractual_docx",
    "assert_contractual_docx_coverage",
    "composition_text",
    "rendered_docx_text",
]
