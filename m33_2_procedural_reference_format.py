from __future__ import annotations

"""Presentación M33.2 para peticiones, reclamaciones y escritos procedimentales.

Esta capa es exclusivamente editorial. No cambia hechos, reglas, cálculos, fuentes,
solicitudes ni conclusiones jurídicas. Convierte los escritos formales generados por
LegalAIZ.it en instrumentos con gramática documental tradicional: Book Antiqua,
marca discreta en encabezado, título jurídico centrado, secciones claras, cuerpo
justificado y separación estable entre párrafos sin insertar líneas vacías.
"""

from functools import wraps
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from document_standard_v33 import audit_docx_legal_standard

FONT_NAME = "Book Antiqua"
BODY_PT = 11
TABLE_PT = 10
TITLE_PT = 11
SUBTITLE_PT = 10.5
PARAGRAPH_AFTER_PT = 4
NUMBERED_AFTER_PT = 2
SECTION_BEFORE_PT = 4
SECTION_AFTER_PT = 3
SIGNATURE_BEFORE_PT = 2
SIGNATURE_AFTER_PT = 1
TITLE_AFTER_PT = 7
SUBTITLE_AFTER_PT = 7

FORMAL_PRODUCT_CODES = frozenset({
    "CO-LA-001",
    "CO-CD-001",
    "CO-CD-003",
    "CO-CD-004",
    "CO-SA-001",
    "CO-TR-001",
    "CO-TR-002",
})

_FORMAL_TITLE_TOKENS = (
    "reclamación",
    "reclamo",
    "petición",
    "solicitud",
    "reiteración",
    "revocación",
    "corrección",
    "requerimiento",
    "cobro",
    "carta",
    "comunicación",
    "aviso",
    "supersalud",
    "escalamiento",
    "terminación por no entrega",
    "no entrega",
    "retiro condicionado",
)

_EXCLUDED_INSTRUMENT_TOKENS = (
    "diagnóstico",
    "matriz",
    "calendario",
    "cronograma",
    "informe",
    "trazabilidad",
    "resumen",
    "estado de cuenta",
    "acuerdo de pago",
    "pagaré",
    "carta de instrucciones",
    "recibo",
    "certificado de cierre",
    "índice probatorio",
)

_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+")


def _set_run_font(run, *, size_pt: float = BODY_PT, bold: bool | None = None, italic: bool | None = None, underline: bool | None = None) -> None:
    run.font.name = FONT_NAME
    rfonts = run._element.get_or_add_rPr().rFonts
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def _set_paragraph_font(paragraph, *, size_pt: float = BODY_PT) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _has_explicit_page_break(paragraph) -> bool:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    page_break_before = p_pr.find(qn("w:pageBreakBefore"))
    return bool(
        page_break_before is not None
        and page_break_before.get(qn("w:val"), "1") not in {"0", "false"}
    )


def _remove_top_level_blank_paragraphs(document: Document) -> int:
    """Elimina separadores vacíos del cuerpo; el espaciado queda en propiedades.

    No toca párrafos que materializan un salto de página explícito ni párrafos dentro
    de tablas, de modo que firmas y estructuras tabulares conservan su semántica.
    """
    removed = 0
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() or _has_explicit_page_break(paragraph):
            continue
        _remove_paragraph(paragraph)
        removed += 1
    return removed


def _is_formal_writing(product_code: str, title: str) -> bool:
    code = str(product_code or "").strip().upper()
    if code not in FORMAL_PRODUCT_CODES:
        return False
    lowered = str(title or "").strip().casefold()
    if any(token in lowered for token in _EXCLUDED_INSTRUMENT_TOKENS):
        return False
    return any(token in lowered for token in _FORMAL_TITLE_TOKENS)


def _title_paragraph(document: Document, title: str):
    wanted = str(title or "").strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == wanted:
            return paragraph
    return next((p for p in document.paragraphs if p.style and p.style.name == "Title"), None)


def _remove_body_branding(document: Document, title_paragraph) -> int:
    """Deja la marca en el encabezado y elimina la pseudoportada del cuerpo."""
    if title_paragraph is None:
        return 0
    removed = 0
    for paragraph in list(document.paragraphs):
        if paragraph._p is title_paragraph._p:
            break
        text = paragraph.text.strip().casefold()
        if not text or text in {"legalaiz.it", "más que respuestas, soluciones."}:
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def _paragraph_after(document: Document, paragraph):
    if paragraph is None:
        return None
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            return next((p for p in document.paragraphs if p._p is sibling), None)
        sibling = sibling.getnext()
    return None


def _format_title_and_subtitle(document: Document, title: str) -> None:
    title_p = _title_paragraph(document, title)
    if title_p is None:
        return
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(TITLE_AFTER_PT)
    title_p.paragraph_format.keep_with_next = True
    for run in title_p.runs:
        _set_run_font(run, size_pt=TITLE_PT, bold=True, italic=False, underline=True)

    subtitle = _paragraph_after(document, title_p)
    if subtitle is None or not subtitle.text.strip():
        return
    if subtitle.style and subtitle.style.name.lower().startswith("heading"):
        return
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(SUBTITLE_AFTER_PT)
    subtitle.paragraph_format.keep_with_next = True
    for run in subtitle.runs:
        _set_run_font(run, size_pt=SUBTITLE_PT, bold=False, italic=True, underline=False)


def _is_control_heading(text: str) -> bool:
    return "control de uso" in str(text or "").casefold()


def _format_headings(document: Document, title: str) -> int:
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == str(title or "").strip():
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if not style_name.lower().startswith("heading"):
            continue
        signature = text.upper() in {"FIRMA", "FIRMAS"}
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if signature else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(SIGNATURE_BEFORE_PT if signature else SECTION_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(SIGNATURE_AFTER_PT if signature else SECTION_AFTER_PT)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            _set_run_font(run, size_pt=BODY_PT, bold=True, italic=_is_control_heading(text), underline=False)
        changed += 1
    return changed


def _format_body(document: Document, title: str) -> int:
    changed = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text == str(title or "").strip():
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        numbered = bool(_NUMBERED_RE.match(text))
        paragraph.paragraph_format.space_after = Pt(NUMBERED_AFTER_PT if numbered else PARAGRAPH_AFTER_PT)
        paragraph.paragraph_format.line_spacing = 1.0
        if numbered:
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
        _set_paragraph_font(paragraph, size_pt=BODY_PT)
        changed += 1
    return changed


def _set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _is_signature_table(table) -> bool:
    tbl_pr = table._tbl.tblPr
    desc = tbl_pr.find(qn("w:tblDescription")) if tbl_pr is not None else None
    return bool(desc is not None and desc.get(qn("w:val")) == "LegalAIZ-SignatureTable")


def _format_tables(document: Document) -> int:
    changed = 0
    for table in document.tables:
        signature = _is_signature_table(table)
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if not signature:
                    _set_cell_shading(cell, "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        _set_run_font(
                            run,
                            size_pt=TABLE_PT,
                            bold=(bool(run.bold) or (not signature and (row_index == 0 or cell_index == 0))),
                        )
        changed += 1
    return changed


def _normalize_styles(document: Document) -> None:
    # `Heading1` era un style_id histórico y python-docx advierte que su lookup está
    # deprecado. Los documentos M33 usan el nombre canónico `Heading 1`.
    for style_name in ("Normal", "Title", "Heading 1"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_NAME
        rfonts = style._element.get_or_add_rPr().rFonts
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), FONT_NAME)
        style.font.size = Pt(BODY_PT)


def apply_m33_2_procedural_format(path: str | Path, *, product_code: str, title: str) -> dict:
    target = Path(path)
    if not _is_formal_writing(product_code, title):
        return {"applied": False, "profile": "M33.2-procedural", "reason": "non_formal_writing"}

    document = Document(target)
    _normalize_styles(document)
    title_p = _title_paragraph(document, title)
    removed_branding = _remove_body_branding(document, title_p)
    removed_blank_paragraphs = _remove_top_level_blank_paragraphs(document)
    headings = _format_headings(document, title)
    paragraphs = _format_body(document, title)
    _format_title_and_subtitle(document, title)
    tables = _format_tables(document)
    document.save(target)

    report = audit_docx_legal_standard(target)
    if not report.get("valid"):
        raise ValueError(f"Escrito M33.2 no supera auditoría documental: {report.get('findings')}")
    return {
        "applied": True,
        "profile": "M33.2-procedural",
        "font": FONT_NAME,
        "removed_body_branding": removed_branding,
        "removed_blank_paragraphs": removed_blank_paragraphs,
        "formatted_headings": headings,
        "formatted_paragraphs": paragraphs,
        "formatted_tables": tables,
        "paragraph_after_pt": PARAGRAPH_AFTER_PT,
        "numbered_after_pt": NUMBERED_AFTER_PT,
    }


def install_m33_2_procedural_format_gate() -> bool:
    """Envuelve `build_docx` después de la compuerta M32.3 y revalida el DOCX final."""
    import docx_builder
    from legalai_platform.document_release_gate import enforce_document_release_gate, infer_product_code

    current = docx_builder.build_docx
    if getattr(current, "_legalaiz_m33_2_procedural", False):
        return True

    @wraps(current)
    def guarded_build_docx(*args, **kwargs):
        call_args = list(args)
        path = kwargs.get("path") or (call_args[0] if call_args else None)
        title = kwargs.get("title") or (call_args[1] if len(call_args) >= 2 else "")
        metadata = kwargs.get("metadata")
        if metadata is None and len(call_args) >= 4:
            metadata = call_args[3]
        product_code = str(kwargs.get("product_code") or infer_product_code(path or "document.docx", metadata) or "").upper()

        result = current(*args, **kwargs)
        final_path = Path(path or result)
        presentation = apply_m33_2_procedural_format(final_path, product_code=product_code, title=str(title or ""))
        if presentation.get("applied"):
            enforce_document_release_gate(final_path, expected_product=product_code or None, metadata=metadata)
        return result

    guarded_build_docx._legalaiz_m33_2_procedural = True
    guarded_build_docx._legalaiz_original = current
    docx_builder.build_docx = guarded_build_docx
    return True


__all__ = [
    "FORMAL_PRODUCT_CODES",
    "apply_m33_2_procedural_format",
    "install_m33_2_procedural_format_gate",
]
