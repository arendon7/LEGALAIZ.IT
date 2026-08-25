from __future__ import annotations

"""Presentación M33.2 para matrices, calendarios y cronogramas jurídico-operativos.

La capa es exclusivamente editorial: conserva hechos, reglas, cálculos, fuentes,
clasificaciones, conclusiones y contenido sustantivo. Optimiza lectura operativa,
control temporal, trazabilidad y consulta de evidencia mediante jerarquía compacta,
tablas estables y orientación horizontal solo cuando una matriz de cinco o más
columnas necesita ancho adicional.
"""

from functools import wraps
from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from document_standard_v33 import audit_docx_legal_standard

FONT_NAME = "Book Antiqua"
BODY_PT = 11
TABLE_PT = 9
TITLE_PT = 12.5
SUBTITLE_PT = 10
PARAGRAPH_AFTER_PT = 5
SECTION_BEFORE_PT = 8
SECTION_AFTER_PT = 4
TITLE_AFTER_PT = 7
NAVY = "0D1324"
IVORY = "F7F5F1"
WHITE = "FFFFFF"
CHARCOAL = "1F1F1F"

OPERATIONAL_PRODUCT_CODES = frozenset({"CO-LA-001","CO-CD-001","CO-CD-003","CO-CD-004","CO-SA-001","CO-TR-001","CO-TR-002"})
_OPERATIONAL_TITLE_TOKENS = (
    "matriz",
    "calendario",
    "cronograma",
    "índice probatorio",
    "hoja de trazabilidad",
)
_EXCLUDED_OPERATIONAL_TOKENS = ("diagnóstico", "informe", "consulta integral", "reclamación", "petición", "pagaré", "acuerdo de pago")

def _target(product_code, title):
    if str(product_code or "").upper() not in OPERATIONAL_PRODUCT_CODES:
        return False
    t=str(title or '').casefold()
    if any(x in t for x in _EXCLUDED_OPERATIONAL_TOKENS):
        return False
    return any(x in t for x in _OPERATIONAL_TITLE_TOKENS)

def _set_run(run, size=BODY_PT, bold=None, italic=None, color=CHARCOAL):
    run.font.name=FONT_NAME
    rp=run._element.get_or_add_rPr().rFonts
    for a in ("ascii","hAnsi","eastAsia","cs"): rp.set(qn(f"w:{a}"), FONT_NAME)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    run.font.color.rgb=RGBColor.from_string(color)

def _remove_p(p):
    parent=p._element.getparent()
    if parent is not None: parent.remove(p._element)

def _title_p(doc,title):
    for p in doc.paragraphs:
        if p.text.strip()==str(title).strip(): return p
    for p in doc.paragraphs:
        if p.style and p.style.name=="Title": return p
    return None

def _next_p(doc,p):
    if p is None: return None
    sib=p._p.getnext()
    while sib is not None:
        if sib.tag==qn("w:p"):
            return next((x for x in doc.paragraphs if x._p is sib),None)
        sib=sib.getnext()
    return None

def _remove_branding(doc,titlep):
    n=0
    for p in list(doc.paragraphs):
        if titlep is not None and p._p is titlep._p: break
        if not p.text.strip() or p.text.strip().casefold() in {"legalaiz.it","más que respuestas, soluciones."}:
            _remove_p(p); n+=1
    return n

def _normalize(doc):
    for name in ("Normal","Title","Heading 1"):
        try: s=doc.styles[name]
        except KeyError: continue
        s.font.name=FONT_NAME; s.font.size=Pt(BODY_PT)
        rf=s._element.get_or_add_rPr().rFonts
        for a in ("ascii","hAnsi","eastAsia","cs"): rf.set(qn(f"w:{a}"),FONT_NAME)

def _landscape_needed(doc, title=""):
    return any(len(t.columns) >= 5 for t in doc.tables)

def _set_orientation(doc, landscape):
    if not landscape: return
    for sec in doc.sections:
        sec.orientation=WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

def _format_title(doc,title):
    tp=_title_p(doc,title)
    if not tp: return
    tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before=Pt(0); tp.paragraph_format.space_after=Pt(TITLE_AFTER_PT)
    tp.paragraph_format.keep_with_next=True
    for r in tp.runs: _set_run(r,TITLE_PT,bold=True,italic=False,color=NAVY)
    sp=_next_p(doc,tp)
    if sp and sp.text.strip() and not (sp.style and sp.style.name.lower().startswith("heading")):
        sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(10)
        sp.paragraph_format.keep_with_next=True
        for r in sp.runs: _set_run(r,SUBTITLE_PT,bold=False,italic=True,color=CHARCOAL)

def _format_paras(doc,title):
    tp=_title_p(doc,title); sp=_next_p(doc,tp)
    for p in doc.paragraphs:
        text=p.text.strip()
        if not text or (tp and p._p is tp._p) or (sp and p._p is sp._p): continue
        style=p.style.name if p.style else ""
        if style.lower().startswith("heading"):
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before=Pt(SECTION_BEFORE_PT)
            p.paragraph_format.space_after=Pt(SECTION_AFTER_PT)
            p.paragraph_format.keep_with_next=True
            for r in p.runs:_set_run(r,BODY_PT,bold=True,italic=False,color=NAVY)
            continue
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after=Pt(PARAGRAPH_AFTER_PT)
        p.paragraph_format.line_spacing=1.0
        if re.match(r"^\s*\d+[.)]\s+", text):
            p.paragraph_format.left_indent=Pt(22); p.paragraph_format.first_line_indent=Pt(-11)
        for r in p.runs:_set_run(r,BODY_PT,color=CHARCOAL)

def _shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr()
    for old in pr.findall(qn("w:shd")): pr.remove(old)
    x=OxmlElement("w:shd"); x.set(qn("w:fill"),fill); pr.append(x)

def _repeat(row):
    pr=row._tr.get_or_add_trPr()
    if pr.find(qn("w:tblHeader")) is None:
        x=OxmlElement("w:tblHeader"); x.set(qn("w:val"),"true"); pr.append(x)

def _cant(row):
    pr=row._tr.get_or_add_trPr()
    if pr.find(qn("w:cantSplit")) is None: pr.append(OxmlElement("w:cantSplit"))

def _table_layout_fixed(table):
    pr=table._tbl.tblPr
    old=pr.find(qn("w:tblLayout"))
    if old is None:
        old=OxmlElement("w:tblLayout"); pr.append(old)
    old.set(qn("w:type"),"fixed")

def _set_widths(table, usable_twips):
    n=len(table.columns)
    hdr=[c.text.strip().casefold() for c in table.rows[0].cells] if table.rows else []
    if n==5:
        props=[0.10,0.22,0.26,0.14,0.28]
    elif n==4 and hdr[:2] == ["hito", "regla"]:
        props=[0.16,0.29,0.22,0.33]
    elif n==4 and hdr and hdr[0] in {"cuota","no.","n°"}:
        props=[0.12,0.32,0.28,0.28]
    elif n==4 and hdr and hdr[0] == "id":
        props=[0.12,0.20,0.42,0.26]
    elif n==4:
        props=[0.18,0.28,0.28,0.26]
    elif n==3:
        props=[0.22,0.33,0.45]
    elif n==2:
        props=[0.34,0.66]
    else:
        props=[1/n]*n
    widths=[int(usable_twips*p) for p in props]
    widths[-1]+=usable_twips-sum(widths)
    pr=table._tbl.tblPr
    tw=pr.find(qn("w:tblW"))
    if tw is None: tw=OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"),str(usable_twips)); tw.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        g=OxmlElement("w:gridCol"); g.set(qn("w:w"),str(w)); grid.append(g)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcpr=cell._tc.get_or_add_tcPr()
            tcw=tcpr.find(qn("w:tcW"))
            if tcw is None: tcw=OxmlElement("w:tcW"); tcpr.append(tcw)
            tcw.set(qn("w:w"),str(widths[i])); tcw.set(qn("w:type"),"dxa")
    return widths

def _format_tables(doc):
    sec=doc.sections[0]
    usable=int(sec.page_width.twips-sec.left_margin.twips-sec.right_margin.twips)
    centered_headers={"id","fecha","estado","cuota","sensibilidad","nivel","resultado"}
    for table in doc.tables:
        if not table.rows: continue
        _table_layout_fixed(table); _set_widths(table,usable)
        for ri,row in enumerate(table.rows):
            _cant(row)
            if ri==0:_repeat(row)
            for ci,cell in enumerate(row.cells):
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                fill=NAVY if ri==0 else (IVORY if ci==0 else WHITE)
                _shade(cell,fill)
                hdr=table.rows[0].cells[ci].text.strip().casefold()
                for p in cell.paragraphs:
                    p.paragraph_format.space_before=Pt(0)
                    p.paragraph_format.space_after=Pt(1)
                    p.paragraph_format.line_spacing=1.0
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER if hdr in centered_headers else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        _set_run(r,TABLE_PT,bold=(ri==0 or ci==0),color=(WHITE if ri==0 else CHARCOAL))

def apply_m33_2_operational_format(path: str | Path, *, product_code: str, title: str) -> dict:
    if not _target(product_code, title):
        return {"applied": False, "profile": "M33.2-operational", "reason": "non_operational_document"}
    doc=Document(path)
    _normalize(doc)
    tp=_title_p(doc,title); rem=_remove_branding(doc,tp)
    landscape=_landscape_needed(doc,title)
    _set_orientation(doc,landscape)
    _format_title(doc,title); _format_paras(doc,title); _format_tables(doc)
    doc.save(path)
    report = audit_docx_legal_standard(path)
    if not report.get("valid"):
        raise ValueError(
            f"Documento operativo M33.2 no supera auditoría documental: {report.get('findings')}"
        )
    return {
        "applied": True,
        "profile": "M33.2-operational",
        "font": FONT_NAME,
        "landscape": landscape,
        "removed_body_branding": rem,
        "paragraph_after_pt": PARAGRAPH_AFTER_PT,
        "table_header_fill": NAVY,
        "table_font_pt": TABLE_PT,
    }


def install_m33_2_operational_format_gate() -> bool:
    """Envuelve `build_docx` tras las capas procedimental y analítica y revalida el archivo final."""
    import docx_builder
    from legalai_platform.document_release_gate import enforce_document_release_gate, infer_product_code

    current = docx_builder.build_docx
    if getattr(current, "_legalaiz_m33_2_operational", False):
        return True

    @wraps(current)
    def guarded_build_docx(*args, **kwargs):
        call_args = list(args)
        path = kwargs.get("path") or (call_args[0] if call_args else None)
        title = kwargs.get("title") or (call_args[1] if len(call_args) >= 2 else "")
        metadata = kwargs.get("metadata")
        if metadata is None and len(call_args) >= 4:
            metadata = call_args[3]
        product_code = str(
            kwargs.get("product_code")
            or infer_product_code(path or "document.docx", metadata)
            or ""
        ).upper()

        result = current(*args, **kwargs)
        final_path = Path(path or result)
        presentation = apply_m33_2_operational_format(
            final_path,
            product_code=product_code,
            title=str(title or ""),
        )
        if presentation.get("applied"):
            enforce_document_release_gate(
                final_path,
                expected_product=product_code or None,
                metadata=metadata,
            )
        return result

    guarded_build_docx._legalaiz_m33_2_operational = True
    guarded_build_docx._legalaiz_original = current
    docx_builder.build_docx = guarded_build_docx
    return True


__all__ = [
    "OPERATIONAL_PRODUCT_CODES",
    "apply_m33_2_operational_format",
    "install_m33_2_operational_format_gate",
]
