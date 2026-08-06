from __future__ import annotations

from argparse import ArgumentParser
from datetime import datetime, timedelta
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile
from zoneinfo import ZoneInfo
import json
import shutil
import sqlite3

from docx import Document

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def write_docx(path: Path, code: str) -> None:
    document = Document()
    title = document.add_heading(f"DOCUMENTO SINTÉTICO {code}", 0)
    title.alignment = 1
    document.add_paragraph("Evidencia técnica M32.6. No contiene datos personales ni está destinada a firma.")
    document.add_heading("PRIMERA: OBJETO DEMOSTRATIVO", level=1)
    document.add_paragraph("Validar asignación, SLA, alertas, actividad y expediente de aprobación por documento.")
    document.save(path)


def main() -> None:
    parser = ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    output = Path(args.output).resolve()
    output.mkdir(parents=True, exist_ok=True)
    fixed_now = datetime(2026, 8, 6, 9, 0, tzinfo=BOGOTA)

    with TemporaryDirectory(prefix="legalaiz-m326-") as temporary:
        root = Path(temporary)
        database = root / "m326.db"

        def db_factory():
            con = sqlite3.connect(database)
            con.row_factory = sqlite3.Row
            return con

        con = db_factory()
        con.executescript(
            """
            CREATE TABLE users(id TEXT PRIMARY KEY,name TEXT,role TEXT,specialty TEXT,active INTEGER DEFAULT 1);
            CREATE TABLE cases(id TEXT PRIMARY KEY,owner_id TEXT,specialist_id TEXT);
            CREATE TABLE documents(id TEXT PRIMARY KEY,case_id TEXT,product_code TEXT,name TEXT,kind TEXT,mime_type TEXT,file_path TEXT,updated_at TEXT);
            """
        )
        con.executemany("INSERT INTO users VALUES(?,?,?,?,?)", [
            ("qa-demo", "QA demostrativo", "admin", "QA documental", 1),
            ("legal-demo", "Especialista demostrativo", "specialist", "Portafolio jurídico", 1),
            ("client-demo", "Cliente demostrativo", "client", None, 1),
        ])
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            source = root / f"{code}_sintetico.docx"
            write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,NULL)", (f"CASE-{index:02d}", "client-demo"))
            con.execute("INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)", (
                f"DOC-{index:02d}", f"CASE-{index:02d}", code, source.name, "main", DOCX_MIME, str(source), fixed_now.isoformat(),
            ))
        con.commit(); con.close()

        def access_check(user, case_id):
            con = db_factory(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
            if not row: return False
            if user["role"] == "admin": return True
            if user["role"] == "client": return row["owner_id"] == user["id"]
            return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

        def document_lookup(user, document_id):
            con = db_factory(); row = con.execute("SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?", (document_id,)).fetchone(); con.close()
            return dict(row) if row and access_check(user, row["case_id"]) else None

        def upload_validator(filename, data):
            return DOCX_MIME, sha256(data).hexdigest(), "clean:synthetic-m326"

        admin = {"id":"qa-demo","role":"admin","name":"QA demostrativo"}
        legal = {"id":"legal-demo","role":"specialist","name":"Especialista demostrativo"}
        workspace = ApprovalDeskWorkspace(
            root / "approval-desk", db_factory=db_factory, document_lookup=document_lookup,
            access_check=access_check, upload_validator=upload_validator,
        )
        operations = ApprovalDeskOperations(
            root / "approval-desk", workspace=workspace, db_factory=db_factory, now_factory=lambda: fixed_now,
        )

        sync = operations.sync_portfolio(admin)
        priorities = ["critical", "high", "normal", "low"]
        case_ids = [f"DSK-DOC-{index:02d}" for index in range(1, 12)]
        for index, case_id in enumerate(case_ids):
            operations.update_assignment(admin, case_id, "legal-demo", "qa-demo")
            operations.update_priority(admin, case_id, priorities[index % len(priorities)])
        operations.update_deadline(admin, case_ids[1], (fixed_now - timedelta(hours=3)).isoformat(), 24)
        operations.update_deadline(admin, case_ids[2], (fixed_now + timedelta(hours=2)).isoformat(), 24)
        operations.add_note(legal, case_ids[0], "Cotejo sintético iniciado; se verificaron variables y anexos.")

        first_detail = workspace.detail(legal, case_ids[0])
        current = first_detail["revisions"][0]
        workspace.approve(legal, case_ids[0], {
            "revision_id": current["revision_id"], "approval_type": "legal", "decision": "approve",
            "comment": "Decisión sintética para validar M32.6.", "expected_sha256": current["sha256"],
        })
        workspace.approve(admin, case_ids[0], {
            "revision_id": current["revision_id"], "approval_type": "qa", "decision": "approve",
            "comment": "Control QA sintético para validar M32.6.", "expected_sha256": current["sha256"],
        })
        workspace.release(admin, case_ids[0], {"revision_id": current["revision_id"], "expected_sha256": current["sha256"]})

        dossier_path, dossier_name = operations.export_dossier(admin, case_ids[0])
        dossier_copy = output / dossier_name
        shutil.copy2(dossier_path, dossier_copy)
        with ZipFile(dossier_copy) as archive:
            dossier = json.loads(archive.read("expediente_aprobacion.json"))
            entries = sorted(archive.namelist())

        portfolio = operations.portfolio(admin)
        states = {case_id: operations.state(admin, case_id) for case_id in case_ids}
        evidence = {
            "schema_version": "M32.6",
            "portfolio": portfolio["portfolio"],
            "metrics": portfolio["metrics"],
            "desk_cases": len(case_ids),
            "assigned_cases": sum(bool(item["operations"]["assigned_specialist"] and item["operations"]["assigned_qa"]) for item in states.values()),
            "valid_operations_chains": sum(bool(item["operations_audit"]["valid"]) for item in states.values()),
            "overdue_cases": sum(item["sla"]["status"] == "overdue" for item in states.values()),
            "at_risk_cases": sum(item["sla"]["status"] == "at_risk" for item in states.values()),
            "released_cases": portfolio["metrics"]["released"],
            "pending_professional_review": len(case_ids) - portfolio["metrics"]["released"],
            "dossier": {
                "filename": dossier_name,
                "sha256": sha256(dossier_copy.read_bytes()).hexdigest(),
                "entries": entries,
                "professional_approval_complete": dossier["professional_approval_complete"],
                "human_review_required": dossier["human_review_required"],
                "revision_sha256": dossier["current_revision"]["sha256"],
                "released_sha256": dossier["release"]["sha256"],
            },
            "real_legal_approval": False,
            "real_qa_approval": False,
            "declaration": "Toda decisión de esta evidencia es sintética y no constituye aprobación profesional de documentos para usuarios reales.",
            "sync_created": sync["bootstrap"]["created_count"],
        }
        (output / "m32-6-portfolio-evidence.json").write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
        (output / "m32-6-portfolio.json").write_text(json.dumps(portfolio, ensure_ascii=False, indent=2), encoding="utf-8")
        print(json.dumps(evidence, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
