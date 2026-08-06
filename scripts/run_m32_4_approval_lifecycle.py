#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
import shutil
import sys
from tempfile import TemporaryDirectory

from docx import Document

from legalai_platform.document_approval_desk import DocumentApprovalDesk, ReleaseBlocked


AUTHOR = {"id": "author-demo", "role": "author", "name": "Autor documental demo"}
LEGAL = {"id": "legal-demo", "role": "specialist", "name": "Especialista jurídico demo"}
QA = {"id": "qa-demo", "role": "qa", "name": "QA documental demo"}
ADMIN = {"id": "admin-demo", "role": "admin", "name": "Administrador demo"}


def make_docx(path: Path, body: str) -> None:
    document = Document()
    document.add_heading("CONTRATO DEMOSTRATIVO M32.4", level=0)
    document.add_paragraph("Documento sintético. No corresponde a un expediente real.")
    document.add_heading("CLÁUSULA PRIMERA: OBJETO", level=1)
    document.add_paragraph(body)
    document.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Ejecuta el ciclo integral de aprobación M32.4.")
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    output = args.output.resolve()
    if output.exists():
        shutil.rmtree(output)
    output.mkdir(parents=True)

    source_one = output / "revision-1-source.docx"
    source_two = output / "revision-2-source.docx"
    make_docx(source_one, "La parte contratista prestará servicios de diagnóstico documental.")
    make_docx(
        source_two,
        "La parte contratista prestará servicios de diagnóstico documental y entregará una matriz de control verificable.",
    )

    desk = DocumentApprovalDesk(output / "approval-desk")
    case = desk.create_case(
        case_id="M32-4-DEMO-001",
        product_code="CO-EM-003",
        document_id="DOC-EM-CONTRACT-001",
        title="Contrato demostrativo de servicios",
        actor=AUTHOR,
        source_generation_id="M32-4-GEN-DEMO",
    )
    revision_one = desk.add_revision(
        case_id=case["case_id"],
        source_file=source_one,
        actor=AUTHOR,
        note="Versión inicial sintética",
    )
    finding = desk.add_finding(
        case_id=case["case_id"],
        revision_id=revision_one["revision_id"],
        actor=LEGAL,
        severity="major",
        description="El objeto no identifica un entregable verificable.",
        page=1,
        clause="PRIMERA: OBJETO",
        block_id="CLAUSE-OBJECT",
    )
    blocked_before_resolution = False
    try:
        desk.approve(
            case_id=case["case_id"],
            revision_id=revision_one["revision_id"],
            approval_type="legal",
            decision="approve",
            actor=LEGAL,
            comment="Intento que debe bloquearse.",
            expected_sha256=revision_one["sha256"],
        )
    except ReleaseBlocked:
        blocked_before_resolution = True
    if not blocked_before_resolution:
        raise RuntimeError("La compuerta no bloqueó un hallazgo mayor abierto.")

    desk.resolve_finding(
        case_id=case["case_id"],
        finding_id=finding["finding_id"],
        actor=LEGAL,
        resolution="Se exige una nueva revisión con entregable verificable.",
    )
    revision_two = desk.add_revision(
        case_id=case["case_id"],
        source_file=source_two,
        actor=AUTHOR,
        note="Se incorporó la matriz de control como entregable.",
        parent_revision_id=revision_one["revision_id"],
    )
    comparison = desk.compare(
        case_id=case["case_id"],
        from_revision_id=revision_one["revision_id"],
        to_revision_id=revision_two["revision_id"],
    )
    desk.approve(
        case_id=case["case_id"],
        revision_id=revision_two["revision_id"],
        approval_type="legal",
        decision="approve",
        actor=LEGAL,
        comment="Texto jurídico demostrativo aprobado para QA.",
        expected_sha256=revision_two["sha256"],
    )
    desk.approve(
        case_id=case["case_id"],
        revision_id=revision_two["revision_id"],
        approval_type="qa",
        decision="approve",
        actor=QA,
        comment="Integridad, formato y trazabilidad demostrativos verificados.",
        expected_sha256=revision_two["sha256"],
    )
    release = desk.release(
        case_id=case["case_id"],
        revision_id=revision_two["revision_id"],
        actor=ADMIN,
        expected_sha256=revision_two["sha256"],
    )
    detail = desk.detail(case["case_id"])
    result = {
        "schema_version": "M32.4",
        "case_id": case["case_id"],
        "revision_count": len(detail["revisions"]),
        "current_revision_id": detail["case"]["current_revision_id"],
        "released_revision_id": release["revision_id"],
        "released_sha256": release["sha256"],
        "comparison": comparison["summary"],
        "blocked_before_resolution": blocked_before_resolution,
        "legal_actor": detail["revisions"][-1]["approvals"]["legal"]["actor"]["id"],
        "qa_actor": detail["revisions"][-1]["approvals"]["qa"]["actor"]["id"],
        "release_status": release["status"],
        "audit": detail["audit"],
        "real_legal_approval": False,
        "real_qa_approval": False,
        "declaration": (
            "Evidencia sintética de funcionamiento. Las decisiones demo no constituyen "
            "aprobación jurídica ni QA de documentos destinados a usuarios reales."
        ),
    }
    if result["revision_count"] != 2:
        raise RuntimeError("El ciclo no conservó dos revisiones inmutables.")
    if result["released_sha256"] != revision_two["sha256"]:
        raise RuntimeError("La liberación no corresponde al hash aprobado.")
    if result["legal_actor"] == result["qa_actor"]:
        raise RuntimeError("La separación de aprobadores fue vulnerada.")
    if not result["audit"]["valid"]:
        raise RuntimeError("La cadena de auditoría no es válida.")
    (output / "m32-4-lifecycle.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
