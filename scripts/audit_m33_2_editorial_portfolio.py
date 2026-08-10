#!/usr/bin/env python3
from __future__ import annotations

import argparse
from collections import Counter
import json
from pathlib import Path
import re
import sys
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXPECTED_M33_DOCS = 53
EXPECTED_PROFILE_COUNTS = {
    "M33.2-contract": 4,
    "M33.2-procedural": 20,
    "M33.2-operational": 10,
    "M33.2-analytical": 7,
    "M33.2-special-guide": 3,
    "M33.2-special-communication": 2,
    "M33.2-special-statement": 1,
    "M33.2-special-agreement": 1,
    "M33.2-special-note": 1,
    "M33.2-special-instructions": 1,
    "M33.2-special-receipt": 1,
    "M33.2-special-certificate": 1,
    "M33.2-special-authorization": 1,
}
EXPECTED_LANDSCAPE = {"wave3/CO-SA-001_health_evidence_M33_0.docx"}
SENTINEL_PATTERNS = (
    re.compile(r"\bNULL\b", re.IGNORECASE),
    re.compile(r"\bundefined\b", re.IGNORECASE),
    re.compile(r"\bN/A\b", re.IGNORECASE),
    re.compile(r"\{\{[^{}]{1,120}\}\}"),
    re.compile(r"\$\{[^{}]{1,120}\}"),
)


def _container_text(container) -> str:
    parts: list[str] = []
    for paragraph in container.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text.strip())
    return "\n".join(parts)


def _visible_text(document: Document) -> str:
    parts = [_container_text(document)]
    for section in document.sections:
        parts.append(_container_text(section.header))
        parts.append(_container_text(section.footer))
    return "\n".join(part for part in parts if part)


def _iter_runs(document: Document):
    for paragraph in document.paragraphs:
        yield from paragraph.runs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from paragraph.runs
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield from paragraph.runs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield from paragraph.runs


def _has_page_field(path: Path) -> bool:
    with ZipFile(path) as archive:
        footer_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    return bool(re.search(r"<w:instrText[^>]*>\s*PAGE\b", footer_xml, re.IGNORECASE))


def _manifest_records(samples: Path) -> list[tuple[str, dict]]:
    records: list[tuple[str, dict]] = []
    for manifest in sorted(samples.rglob("m33-*-samples.json")):
        data = json.loads(manifest.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            raise ValueError(f"{manifest}: manifiesto no es una lista")
        group = manifest.parent.name
        for item in data:
            if not isinstance(item, dict):
                raise ValueError(f"{manifest}: registro inválido")
            records.append((group, item))
    return records


def _profile_for(group: str, record: dict) -> str:
    if group == "contracts":
        return "M33.2-contract"
    return str(record.get("presentation_profile") or "")


def audit_portfolio(samples: Path) -> dict:
    samples = samples.resolve()
    findings: list[str] = []
    m33_docs = sorted(samples.rglob("*_M33_0.docx"))
    if len(m33_docs) != EXPECTED_M33_DOCS:
        findings.append(f"Se esperaban {EXPECTED_M33_DOCS} DOCX M33 y se encontraron {len(m33_docs)}")

    records = _manifest_records(samples)
    if len(records) != EXPECTED_M33_DOCS:
        findings.append(f"Se esperaban {EXPECTED_M33_DOCS} registros de manifiesto y se encontraron {len(records)}")

    record_by_rel: dict[str, tuple[str, dict]] = {}
    profiles: Counter[str] = Counter()
    for group, record in records:
        sample = str(record.get("sample") or "").strip()
        if not sample:
            findings.append(f"{group}: registro sin nombre de muestra")
            continue
        rel = f"{group}/{sample}"
        if rel in record_by_rel:
            findings.append(f"Muestra duplicada en manifiestos: {rel}")
        record_by_rel[rel] = (group, record)

        if record.get("document_standard") != "M33.0":
            findings.append(f"{rel}: document_standard distinto de M33.0")
        if record.get("released") is not False:
            findings.append(f"{rel}: released debe permanecer false")
        if record.get("legal_approval") != "pending":
            findings.append(f"{rel}: legal_approval debe permanecer pending")
        if record.get("qa_approval") != "pending":
            findings.append(f"{rel}: qa_approval debe permanecer pending")

        profile = _profile_for(group, record)
        if group != "contracts":
            if record.get("presentation_standard") != "M33.2":
                findings.append(f"{rel}: presentation_standard distinto de M33.2")
            if not profile or profile == "M33.2-base":
                findings.append(f"{rel}: documento sin familia editorial M33.2 explícita")
        profiles[profile] += 1

    if dict(sorted(profiles.items())) != dict(sorted(EXPECTED_PROFILE_COUNTS.items())):
        findings.append(
            "Distribución de perfiles M33.2 cambió: "
            + json.dumps(dict(sorted(profiles.items())), ensure_ascii=False, sort_keys=True)
        )

    landscape_docs: set[str] = set()
    document_summaries: list[dict] = []
    for path in m33_docs:
        rel = path.relative_to(samples).as_posix()
        if rel not in record_by_rel:
            findings.append(f"{rel}: DOCX sin registro de manifiesto")
        try:
            document = Document(path)
        except Exception as exc:
            findings.append(f"{rel}: DOCX no abre con python-docx: {exc}")
            continue

        titles = [p for p in document.paragraphs if p.style and p.style.name == "Title" and p.text.strip()]
        if len(titles) != 1:
            findings.append(f"{rel}: se esperaba un único párrafo Title y se encontraron {len(titles)}")
        elif titles[0].alignment != WD_ALIGN_PARAGRAPH.CENTER:
            findings.append(f"{rel}: título no está centrado")

        for section_index, section in enumerate(document.sections, start=1):
            margins = {
                "top": section.top_margin.cm if section.top_margin else None,
                "bottom": section.bottom_margin.cm if section.bottom_margin else None,
                "left": section.left_margin.cm if section.left_margin else None,
                "right": section.right_margin.cm if section.right_margin else None,
            }
            for name, value in margins.items():
                if value is None or abs(value - 2.5) > 0.06:
                    findings.append(f"{rel}: margen {name} de sección {section_index} = {value}, esperado 2.5 cm")
            if section.orientation == WD_ORIENT.LANDSCAPE:
                landscape_docs.add(rel)
            if not _container_text(section.header).strip():
                findings.append(f"{rel}: encabezado vacío en sección {section_index}")
            if not _container_text(section.footer).strip():
                findings.append(f"{rel}: pie vacío en sección {section_index}")

        if not _has_page_field(path):
            findings.append(f"{rel}: pie sin campo PAGE")

        for style_name in ("Normal", "Title", "Heading 1"):
            try:
                style = document.styles[style_name]
            except KeyError:
                findings.append(f"{rel}: falta estilo {style_name}")
                continue
            if (style.font.name or "").strip() != "Book Antiqua":
                findings.append(f"{rel}: estilo {style_name} usa {style.font.name!r} en vez de Book Antiqua")

        bad_fonts = sorted({
            str(run.font.name).strip()
            for run in _iter_runs(document)
            if run.text.strip() and run.font.name and str(run.font.name).strip() != "Book Antiqua"
        })
        if bad_fonts:
            findings.append(f"{rel}: fuentes directas ajenas a Book Antiqua: {bad_fonts}")

        visible = _visible_text(document)
        for pattern in SENTINEL_PATTERNS:
            if pattern.search(visible):
                findings.append(f"{rel}: marcador o centinela visible detectado por {pattern.pattern!r}")

        group, record = record_by_rel.get(rel, (path.parent.name, {}))
        document_summaries.append({
            "document": rel,
            "profile": _profile_for(group, record),
            "sections": len(document.sections),
            "landscape": rel in landscape_docs,
            "tables": len(document.tables),
        })

    missing_doc_records = sorted(set(record_by_rel) - {path.relative_to(samples).as_posix() for path in m33_docs})
    for rel in missing_doc_records:
        findings.append(f"{rel}: registro de manifiesto sin DOCX M33")

    if landscape_docs != EXPECTED_LANDSCAPE:
        findings.append(
            "Conjunto de documentos horizontales inesperado: "
            + json.dumps(sorted(landscape_docs), ensure_ascii=False)
        )

    report = {
        "schema": "legalaizit-m33-2-editorial-portfolio-audit-v1",
        "expected_documents": EXPECTED_M33_DOCS,
        "documents": len(m33_docs),
        "manifest_records": len(records),
        "profile_counts": dict(sorted(profiles.items())),
        "landscape_documents": sorted(landscape_docs),
        "findings": findings,
        "ok": not findings,
        "document_summaries": document_summaries,
    }
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Audita de forma transversal el portafolio editorial M33.2.")
    parser.add_argument("--samples", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    report = audit_portfolio(args.samples)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: report[k] for k in ("documents", "manifest_records", "profile_counts", "landscape_documents", "ok")}, ensure_ascii=False))
    if report["findings"]:
        for finding in report["findings"]:
            print(f"ERROR: {finding}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
