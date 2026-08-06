#!/usr/bin/env python3
from __future__ import annotations

import argparse
from datetime import datetime, timedelta
from hashlib import sha256
import json
from pathlib import Path
import shutil
import sqlite3
from tempfile import TemporaryDirectory

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace
from legalai_platform.approval_notification_center import ApprovalNotificationCenter
from legalai_platform.transactional_communications import TransactionalCommunications


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FIXED_NOW = datetime(2026, 8, 6, 9, 30, tzinfo=BOGOTA)


def connection(path: Path):
    con = sqlite3.connect(path)
    con.row_factory = sqlite3.Row
    return con


def write_docx(path: Path, code: str):
    document = Document()
    document.add_heading(f"DOCUMENTO SINTÉTICO {code}", 0)
    document.add_paragraph("Evidencia técnica M32.8 sin datos personales ni decisión jurídica real.")
    document.save(path)


def build(output: Path) -> dict:
    output.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as temporary:
        root = Path(temporary)
        database = root / "m328-evidence.db"

        def db():
            return connection(database)

        con = db()
        con.executescript(
            """
            CREATE TABLE users(id TEXT PRIMARY KEY,name TEXT NOT NULL,email TEXT,role TEXT NOT NULL,specialty TEXT,active INTEGER NOT NULL DEFAULT 1);
            CREATE TABLE cases(id TEXT PRIMARY KEY,owner_id TEXT,specialist_id TEXT);
            CREATE TABLE documents(id TEXT PRIMARY KEY,case_id TEXT NOT NULL,product_code TEXT NOT NULL,name TEXT NOT NULL,kind TEXT NOT NULL,mime_type TEXT NOT NULL,file_path TEXT,updated_at TEXT NOT NULL);
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana Administradora", "ana@example.test", "admin", "Operación", 1),
                ("USR-QA", "Quinn QA", "qa@example.test", "qa", "QA", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-LEGAL-2", "Carlos Legal", "carlos@example.test", "specialist", "Derecho público", 1),
                ("USR-CLIENT", "Cliente sintético", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = root / f"{code}_{index:02d}.docx"
            write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), FIXED_NOW.isoformat()),
            )
        con.commit(); con.close()

        admin = {"id":"USR-ADMIN","role":"admin","name":"Ana Administradora"}
        qa = {"id":"USR-QA","role":"qa","name":"Quinn QA"}
        legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}

        def access_check(user, case_id):
            con = db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
            if not row: return False
            if user["role"] in {"admin", "qa"}: return True
            if user["role"] == "client": return row["owner_id"] == user["id"]
            return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

        def lookup(user, document_id):
            con = db()
            row = con.execute("SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?", (document_id,)).fetchone()
            con.close()
            return dict(row) if row and access_check(user, row["case_id"]) else None

        workspace = ApprovalDeskWorkspace(
            root / "approval-desk",
            db_factory=db,
            document_lookup=lookup,
            access_check=access_check,
            upload_validator=lambda filename, data: (DOCX_MIME, sha256(data).hexdigest(), "clean:evidence"),
        )
        operations = ApprovalDeskOperations(
            root / "approval-desk",
            workspace=workspace,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        center = ApprovalNotificationCenter(
            root / "approval-desk",
            operations=operations,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        communications = TransactionalCommunications(
            root / "approval-desk",
            notification_center=center,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )

        synced = operations.sync_portfolio(admin)
        cases = operations.portfolio(admin)["cases"]
        for row in cases:
            operations.update_assignment(admin, row["desk_case_id"], "USR-LEGAL", "USR-ADMIN")
            operations.update_priority(admin, row["desk_case_id"], "normal")
        operations.update_deadline(admin, cases[0]["desk_case_id"], (FIXED_NOW - timedelta(hours=3)).isoformat(), 24)
        operations.update_deadline(admin, cases[1]["desk_case_id"], (FIXED_NOW + timedelta(hours=2)).isoformat(), 4)
        center.update_policy(admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
            "admin_escalation_after_hours": 0,
        })
        evaluation = center.evaluate(admin)

        created_template = communications.create_template_version(admin, {
            "template_id": "professional-alert",
            "name": "Alerta profesional M32.8",
            "subject": "LegalAIZ.it · {{product_code}} · {{title}}",
            "body": (
                "Hola {{recipient_name}}. Revise el expediente {{case_id}} del producto {{product_code}}. "
                "Fecha objetivo: {{due_at}}. Consulte el detalle únicamente dentro de la Mesa Jurídica."
            ),
        })
        communications.activate_template(qa, "professional-alert", created_template["template"]["version"])
        communications.update_policy(admin, {
            "sandbox_enabled": True,
            "max_attempts": 3,
            "initial_backoff_seconds": 60,
            "max_backoff_seconds": 3600,
            "batch_size": 100,
        })
        imported = communications.sync_outbox(admin)
        processed = communications.process(admin)
        queue = communications.queue(admin)
        accepted = next((item for item in queue["dispatches"] if item["status"] == "accepted_sandbox"), None)
        if accepted:
            communications.record_receipt(
                qa,
                accepted["dispatch_id"],
                provider_status="delivered",
                provider_event_id="EVIDENCE-SBX-DELIVERED-001",
                occurred_at=FIXED_NOW.isoformat(),
                synthetic=True,
            )
        dashboard = communications.dashboard(admin)
        personal = communications.queue(legal)
        final_queue = communications.queue(admin)

        events_copy = output / "m32-8-events.jsonl"
        source_events_copy = output / "m32-7-source-events.jsonl"
        shutil.copy2(root / "approval-desk" / "transactional-communications" / "events.jsonl", events_copy)
        shutil.copy2(root / "approval-desk" / "notification-center" / "events.jsonl", source_events_copy)

        evidence = {
            "schema_version": "M32.8",
            "generated_at": FIXED_NOW.isoformat(),
            "synthetic_evidence": True,
            "portfolio": synced["portfolio"],
            "m32_7_evaluation": {
                "evaluated_cases": evaluation["evaluated_cases"],
                "queued_messages": len(evaluation["queued_messages"]),
                "external_delivery_performed": evaluation["external_delivery_performed"],
            },
            "template": {
                "template_id": created_template["template"]["template_id"],
                "version": created_template["template"]["version"],
                "sha256": created_template["template"]["template_sha256"],
                "active_version": dashboard["templates"]["active_templates"]["professional-alert"],
                "independent_activation": True,
                "contains_document_content": False,
                "attachments_allowed": False,
            },
            "sync": imported,
            "processing": processed,
            "queue_metrics": final_queue["metrics"],
            "queue_statuses": sorted({item["status"] for item in final_queue["dispatches"]}),
            "personal_queue_total": personal["metrics"]["total"],
            "recipient_addresses_stored": final_queue["recipient_addresses_stored"],
            "contains_document_content": final_queue["contains_document_content"],
            "all_attachments_empty": all(item.get("attachments") == [] for item in final_queue["dispatches"]),
            "all_real_delivery_false": all(item.get("real_delivery") is False for item in final_queue["dispatches"]),
            "m32_8_chain": communications.verify_chain(),
            "m32_7_chain": center.verify_chain(),
            "real_external_delivery": False,
            "real_delivery_evidence": False,
            "real_legal_approval": False,
            "real_qa_approval": False,
            "files": {"events": events_copy.name, "source_events": source_events_copy.name},
            "notice": "Evidencia sintética: el proveedor sandbox no envía correos ni acredita recepción externa.",
        }
        evidence_path = output / "m32-8-transactional-communications-evidence.json"
        evidence_path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
        return evidence


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    evidence = build(Path(args.output).resolve())
    print(json.dumps({
        "coverage": evidence["portfolio"]["coverage_percent"],
        "imported": len(evidence["sync"]["imported_dispatches"]),
        "accepted_sandbox": len(evidence["processing"]["accepted_sandbox"]),
        "delivered_sandbox": evidence["queue_metrics"]["delivered_sandbox"],
        "chain_valid": evidence["m32_8_chain"]["valid"],
        "real_delivery": evidence["real_external_delivery"],
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
