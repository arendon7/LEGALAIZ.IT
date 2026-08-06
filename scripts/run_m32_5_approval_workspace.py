from __future__ import annotations

from argparse import ArgumentParser
from hashlib import sha256
from pathlib import Path
from tempfile import TemporaryDirectory
import json
import shutil
import sqlite3

from docx import Document
from docx.enum.text import WD_BREAK

from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace, ReleaseBlocked


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def write_document(path: Path, version: int) -> None:
    document = Document()
    title = document.add_heading("ACUERDO DEMOSTRATIVO DE SERVICIOS", 0)
    title.alignment = 1
    document.add_paragraph(
        "Documento sintético generado exclusivamente para validar la Mesa Jurídica M32.5. "
        "No contiene datos personales ni constituye un contrato destinado a firma."
    )
    document.add_heading("CLÁUSULA PRIMERA: OBJETO", level=1)
    object_text = (
        "El prestador ejecutará actividades de diagnóstico documental y presentará entregables trazables."
        if version == 1
        else "El prestador ejecutará actividades de diagnóstico documental, presentará entregables trazables y documentará cada supuesto utilizado."
    )
    document.add_paragraph(object_text)
    document.add_heading("CLÁUSULA SEGUNDA: CONTROL DE CAMBIOS", level=1)
    document.add_paragraph(
        "Toda modificación deberá registrarse como una nueva revisión inmutable y conservar el vínculo con su versión anterior."
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("ANEXO DE TRAZABILIDAD", level=1)
    document.add_paragraph(
        "La aprobación jurídica y el control QA recaen sobre el archivo concreto identificado por su SHA-256."
    )
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Estado"
    table.cell(1, 0).text = "Revisión humana"
    table.cell(1, 1).text = "Requerida"
    table.cell(2, 0).text = "Aprobación real"
    table.cell(2, 1).text = "No otorgada en esta evidencia"
    document.save(path)


def make_database(path: Path, document_path: Path):
    con = sqlite3.connect(path)
    con.row_factory = sqlite3.Row
    con.executescript(
        """
        CREATE TABLE cases(
          id TEXT PRIMARY KEY,
          owner_id TEXT,
          specialist_id TEXT
        );
        CREATE TABLE documents(
          id TEXT PRIMARY KEY,
          case_id TEXT NOT NULL,
          product_code TEXT NOT NULL,
          name TEXT NOT NULL,
          kind TEXT NOT NULL,
          mime_type TEXT NOT NULL,
          file_path TEXT,
          updated_at TEXT NOT NULL
        );
        """
    )
    con.execute("INSERT INTO cases VALUES(?,?,?)", ("M32-5-CASE-001", "client-demo", "legal-demo"))
    con.execute(
        "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
        (
            "M32-5-DOC-001",
            "M32-5-CASE-001",
            "CO-EM-003",
            "M32_5_Acuerdo_Demostrativo.docx",
            "main",
            DOCX_MIME,
            str(document_path),
            "2026-08-05T22:00:00-05:00",
        ),
    )
    con.commit(); con.close()


def main() -> None:
    parser = ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    output = Path(args.output).resolve()
    output.mkdir(parents=True, exist_ok=True)

    with TemporaryDirectory(prefix="legalaiz-m325-evidence-") as temporary:
        root = Path(temporary)
        database = root / "workspace.db"
        first = root / "M32_5_Acuerdo_Demostrativo_v1.docx"
        revised = root / "M32_5_Acuerdo_Demostrativo_v2.docx"
        write_document(first, 1)
        write_document(revised, 2)
        make_database(database, first)

        def db_factory():
            con = sqlite3.connect(database)
            con.row_factory = sqlite3.Row
            return con

        def access_check(user, case_id):
            con = db_factory()
            row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone()
            con.close()
            if not row:
                return False
            if user["role"] == "admin":
                return True
            if user["role"] == "client":
                return row["owner_id"] == user["id"]
            return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

        def document_lookup(user, document_id):
            con = db_factory()
            row = con.execute(
                "SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?",
                (document_id,),
            ).fetchone()
            con.close()
            if not row or not access_check(user, row["case_id"]):
                return None
            return dict(row)

        def upload_validator(filename, data):
            if not filename.casefold().endswith(".docx") or not data.startswith(b"PK"):
                raise ValueError("La evidencia requiere un DOCX válido.")
            return DOCX_MIME, sha256(data).hexdigest(), "clean:synthetic-evidence"

        admin = {"id": "qa-demo", "role": "admin", "name": "QA demostrativo"}
        legal = {"id": "legal-demo", "role": "specialist", "name": "Especialista demostrativo"}
        client = {"id": "client-demo", "role": "client", "name": "Cliente demostrativo"}
        workspace = ApprovalDeskWorkspace(
            root / "approval-desk",
            db_factory=db_factory,
            document_lookup=document_lookup,
            access_check=access_check,
            upload_validator=upload_validator,
        )

        bootstrap = workspace.bootstrap(admin)
        case_id = bootstrap["created"][0]
        first_detail = workspace.detail(legal, case_id)
        first_revision = first_detail["case"]["current_revision_id"]
        second = workspace.upload_revision(
            legal,
            case_id,
            revised.name,
            revised.read_bytes(),
            "Se añade el deber de documentar los supuestos utilizados.",
        )
        comparison = workspace.compare(legal, case_id, first_revision, second["revision_id"])
        preview = workspace.preview(legal, case_id, second["revision_id"])
        if not preview.get("rendered") or int(preview.get("page_count") or 0) < 2:
            raise SystemExit("El motor no produjo una vista paginada de al menos dos páginas.")

        finding = workspace.add_finding(
            legal,
            case_id,
            {
                "revision_id": second["revision_id"],
                "severity": "major",
                "description": "Confirmar que el anexo declara expresamente el carácter sintético de la evidencia.",
                "page": 2,
                "clause": "ANEXO DE TRAZABILIDAD",
                "block_id": "m32-5-anexo-control",
            },
        )
        blocked = False
        try:
            workspace.approve(
                legal,
                case_id,
                {
                    "revision_id": second["revision_id"],
                    "approval_type": "legal",
                    "decision": "approve",
                    "comment": "Intento que debe quedar bloqueado.",
                    "expected_sha256": second["sha256"],
                },
            )
        except ReleaseBlocked:
            blocked = True
        if not blocked:
            raise SystemExit("Un hallazgo mayor abierto no bloqueó la aprobación jurídica.")

        resolution = workspace.resolve_finding(
            legal,
            case_id,
            finding["finding_id"],
            {
                "resolution": "La declaración sintética figura en el encabezado y en el anexo de trazabilidad.",
                "state": "resolved",
            },
        )
        legal_approval = workspace.approve(
            legal,
            case_id,
            {
                "revision_id": second["revision_id"],
                "approval_type": "legal",
                "decision": "approve",
                "comment": "Decisión sintética para validar el flujo técnico.",
                "expected_sha256": second["sha256"],
            },
        )
        qa_approval = workspace.approve(
            admin,
            case_id,
            {
                "revision_id": second["revision_id"],
                "approval_type": "qa",
                "decision": "approve",
                "comment": "Control sintético de integridad, renderizado y trazabilidad.",
                "expected_sha256": second["sha256"],
            },
        )
        release = workspace.release(
            admin,
            case_id,
            {"revision_id": second["revision_id"], "expected_sha256": second["sha256"]},
        )
        released_path, released_manifest = workspace.released_path(client, case_id)
        final_detail = workspace.detail(admin, case_id)
        summary = workspace.list_for_user(admin)

        preview_pdf = workspace.preview_pdf_path(legal, case_id, second["revision_id"])
        released_copy = output / released_path.name
        preview_copy = output / "M32_5_Vista_Paginada.pdf"
        shutil.copy2(released_path, released_copy)
        shutil.copy2(preview_pdf, preview_copy)
        shutil.copy2(first, output / first.name)
        shutil.copy2(revised, output / revised.name)

        evidence = {
            "schema_version": "M32.5",
            "case_id": case_id,
            "source_case_id": "M32-5-CASE-001",
            "document_id": "M32-5-DOC-001",
            "revision_count": final_detail["case"]["revision_count"],
            "current_revision_id": final_detail["case"]["current_revision_id"],
            "current_sha256": second["sha256"],
            "rendered_preview": True,
            "rendering_engine": preview["rendering_engine"],
            "page_count": preview["page_count"],
            "comparison": comparison["summary"],
            "comparison_changed": comparison["changed"],
            "finding_id": finding["finding_id"],
            "finding_locator": finding["locator"],
            "approval_blocked_before_resolution": blocked,
            "finding_resolution_state": resolution["state"],
            "legal_actor": legal_approval["actor"]["id"],
            "qa_actor": qa_approval["actor"]["id"],
            "separation_of_functions": legal_approval["actor"]["id"] != qa_approval["actor"]["id"],
            "released_revision_id": release["revision_id"],
            "released_sha256": release["sha256"],
            "released_physical_sha256": sha256(released_copy.read_bytes()).hexdigest(),
            "released_manifest_sha256": released_manifest["sha256"],
            "workflow_status": final_detail["workflow_status"],
            "audit": final_detail["audit"],
            "summary_metrics": summary["metrics"],
            "real_legal_approval": False,
            "real_qa_approval": False,
            "declaration": (
                "Evidencia sintética de funcionamiento. Las decisiones demo no constituyen aprobación jurídica ni QA de documentos destinados a usuarios reales."
            ),
            "files": {
                "revision_1": first.name,
                "revision_2": revised.name,
                "preview_pdf": preview_copy.name,
                "released_docx": released_copy.name,
            },
        }
        if len({evidence["current_sha256"], evidence["released_sha256"], evidence["released_physical_sha256"], evidence["released_manifest_sha256"]}) != 1:
            raise SystemExit("La evidencia liberada no conserva un único SHA-256.")
        (output / "m32-5-workspace-evidence.json").write_text(
            json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        (output / "m32-5-case-detail.json").write_text(
            json.dumps(final_detail, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        (output / "m32-5-comparison.json").write_text(
            json.dumps(comparison, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(json.dumps(evidence, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
