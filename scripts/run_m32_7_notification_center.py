#!/usr/bin/env python3
from __future__ import annotations

import argparse
from datetime import datetime, timedelta
from hashlib import sha256
import json
from pathlib import Path
import shutil
import sqlite3
from tempfile import TemporaryDirectory

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace
from legalai_platform.approval_notification_center import ApprovalNotificationCenter


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FIXED_NOW = datetime(2026, 8, 6, 9, 0, tzinfo=BOGOTA)


def connection(path: Path):
    con = sqlite3.connect(path)
    con.row_factory = sqlite3.Row
    return con


def write_docx(path: Path, code: str):
    document = Document()
    document.add_heading(f"DOCUMENTO SINTÉTICO {code}", 0)
    document.add_paragraph("Evidencia técnica M32.7. No contiene información personal ni decisión jurídica real.")
    document.save(path)


def build(output: Path) -> dict:
    output.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as temporary:
        root = Path(temporary)
        database = root / "m327-evidence.db"

        def db():
            return connection(database)

        con = db()
        con.executescript(
            """
            CREATE TABLE users(id TEXT PRIMARY KEY,name TEXT NOT NULL,email TEXT,role TEXT NOT NULL,specialty TEXT,active INTEGER NOT NULL DEFAULT 1);
            CREATE TABLE cases(id TEXT PRIMARY KEY,owner_id TEXT,specialist_id TEXT);
            CREATE TABLE documents(id TEXT PRIMARY KEY,case_id TEXT NOT NULL,product_code TEXT NOT NULL,name TEXT NOT NULL,kind TEXT NOT NULL,mime_type TEXT NOT NULL,file_path TEXT,updated_at TEXT NOT NULL);
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana QA", "ana@example.test", "admin", "QA documental", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-LEGAL-2", "Carlos Legal", "carlos@example.test", "specialist", "Derecho público", 1),
                ("USR-CLIENT", "Cliente sintético", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = root / f"{code}_{index:02d}.docx"
            write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), FIXED_NOW.isoformat()),
            )
        con.commit(); con.close()

        admin = {"id":"USR-ADMIN","role":"admin","name":"Ana QA"}
        legal = {"id":"USR-LEGAL","role":"specialist","name":"María Jurídica"}

        def access_check(user, case_id):
            con = db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
            if not row: return False
            if user["role"] == "admin": return True
            if user["role"] == "client": return row["owner_id"] == user["id"]
            return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

        def lookup(user, document_id):
            con = db()
            row = con.execute("SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?", (document_id,)).fetchone()
            con.close()
            return dict(row) if row and access_check(user, row["case_id"]) else None

        workspace = ApprovalDeskWorkspace(
            root / "approval-desk",
            db_factory=db,
            document_lookup=lookup,
            access_check=access_check,
            upload_validator=lambda filename, data: (DOCX_MIME, sha256(data).hexdigest(), "clean:evidence"),
        )
        operations = ApprovalDeskOperations(
            root / "approval-desk",
            workspace=workspace,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        center = ApprovalNotificationCenter(
            root / "approval-desk",
            operations=operations,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )

        synced = operations.sync_portfolio(admin)
        cases = operations.portfolio(admin)["cases"]
        for row in cases:
            operations.update_assignment(admin, row["desk_case_id"], "USR-LEGAL", "USR-ADMIN")
            operations.update_priority(admin, row["desk_case_id"], "normal")
        operations.update_deadline(admin, cases[0]["desk_case_id"], (FIXED_NOW - timedelta(hours=3)).isoformat(), 24)
        operations.update_deadline(admin, cases[1]["desk_case_id"], (FIXED_NOW + timedelta(hours=2)).isoformat(), 4)

        center.update_calendar(admin, {
            "name": "Calendario sintético de operación",
            "weekdays": [0,1,2,3,4],
            "open_time": "08:00",
            "close_time": "17:00",
            "holidays": ["2026-08-07"],
        })
        center.update_policy(admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
            "admin_escalation_after_hours": 0,
        })
        schedule = center.schedule_case(
            admin,
            cases[2]["desk_case_id"],
            10,
            datetime(2026,8,6,16,0,tzinfo=BOGOTA).isoformat(),
        )
        evaluation = center.evaluate(admin)
        legal_inbox = center.inbox(legal)
        if legal_inbox["notifications"]:
            center.mark_read(legal, legal_inbox["notifications"][0]["notification_id"])
        legal_inbox = center.inbox(legal)
        all_inbox = center.inbox(admin, include_all=True, limit=500)
        dashboard = center.dashboard(admin)
        outbox = center.outbox(admin)
        audit = center.verify_chain()

        events_copy = output / "m32-7-events.jsonl"
        shutil.copy2(root / "approval-desk" / "notification-center" / "events.jsonl", events_copy)
        evidence = {
            "schema_version": "M32.7",
            "generated_at": FIXED_NOW.isoformat(),
            "synthetic_evidence": True,
            "portfolio": synced["portfolio"],
            "evaluation": evaluation,
            "calendar": dashboard["calendar"],
            "business_schedule": schedule["business_sla"],
            "personal_inbox_metrics": legal_inbox["metrics"],
            "all_inbox_metrics": all_inbox["metrics"],
            "workload": dashboard["workload"],
            "outbox_metrics": outbox["metrics"],
            "outbox_statuses": sorted({item["status"] for item in outbox["messages"]}),
            "outbox_contains_document_content": any(item["contains_document_content"] for item in outbox["messages"]),
            "outbox_recipient_addresses_stored": any(item.get("recipient_address_stored", True) for item in outbox["messages"]),
            "external_delivery_active": False,
            "external_delivery_performed": evaluation["external_delivery_performed"],
            "notification_chain": audit,
            "real_legal_approval": False,
            "real_qa_approval": False,
            "real_external_delivery": False,
            "statutory_deadline_calculation": False,
            "files": {"events": events_copy.name},
            "notice": "Evidencia sintética: la cola no envía correo y el calendario no representa términos legales oficiales.",
        }
        evidence_path = output / "m32-7-notification-center-evidence.json"
        evidence_path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
        return evidence


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    evidence = build(Path(args.output).resolve())
    print(json.dumps({
        "coverage": evidence["portfolio"]["coverage_percent"],
        "notifications": evidence["all_inbox_metrics"]["total"],
        "queued": evidence["outbox_metrics"]["queued"],
        "chain_valid": evidence["notification_chain"]["valid"],
        "external_delivery": evidence["real_external_delivery"],
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
