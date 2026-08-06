#!/usr/bin/env python3
from __future__ import annotations

import argparse
from datetime import datetime, timedelta
from hashlib import sha256
import json
from pathlib import Path
import shutil
import sqlite3
from tempfile import TemporaryDirectory

from docx import Document
from zoneinfo import ZoneInfo

from legalai_platform.approval_desk_operations import ApprovalDeskOperations, PORTFOLIO_CODES
from legalai_platform.approval_desk_workspace import ApprovalDeskWorkspace
from legalai_platform.approval_notification_center import ApprovalNotificationCenter
from legalai_platform.contact_governance import ContactGovernance, GovernedTransactionalCommunications


BOGOTA = ZoneInfo("America/Bogota")
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FIXED_NOW = datetime(2026, 8, 6, 11, 0, tzinfo=BOGOTA)


def connection(path: Path):
    con = sqlite3.connect(path)
    con.row_factory = sqlite3.Row
    return con


def write_docx(path: Path, code: str):
    document = Document()
    document.add_heading(f"DOCUMENTO SINTÉTICO {code}", 0)
    document.add_paragraph("Evidencia técnica M32.9 sin información personal real.")
    document.save(path)


def build(output: Path) -> dict:
    output.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as temporary:
        root = Path(temporary)
        database = root / "m329-evidence.db"

        def db():
            return connection(database)

        con = db()
        con.executescript(
            """
            CREATE TABLE users(id TEXT PRIMARY KEY,name TEXT NOT NULL,email TEXT,role TEXT NOT NULL,specialty TEXT,active INTEGER NOT NULL DEFAULT 1);
            CREATE TABLE cases(id TEXT PRIMARY KEY,owner_id TEXT,specialist_id TEXT);
            CREATE TABLE documents(id TEXT PRIMARY KEY,case_id TEXT NOT NULL,product_code TEXT NOT NULL,name TEXT NOT NULL,kind TEXT NOT NULL,mime_type TEXT NOT NULL,file_path TEXT,updated_at TEXT NOT NULL);
            """
        )
        con.executemany(
            "INSERT INTO users VALUES(?,?,?,?,?,?)",
            [
                ("USR-ADMIN", "Ana Administradora", "ana@example.test", "admin", "Operación", 1),
                ("USR-QA", "Quinn QA", "qa@example.test", "qa", "QA", 1),
                ("USR-LEGAL", "María Jurídica", "maria@example.test", "specialist", "Contratos", 1),
                ("USR-CLIENT", "Cliente sintético", "cliente@example.test", "client", None, 1),
            ],
        )
        for index, code in enumerate(PORTFOLIO_CODES, 1):
            case_id = f"CASE-{index:02d}"
            document_id = f"DOC-{index:02d}"
            source = root / f"{code}_{index:02d}.docx"
            write_docx(source, code)
            con.execute("INSERT INTO cases VALUES(?,?,?)", (case_id, "USR-CLIENT", None))
            con.execute(
                "INSERT INTO documents VALUES(?,?,?,?,?,?,?,?)",
                (document_id, case_id, code, source.name, "main", DOCX_MIME, str(source), FIXED_NOW.isoformat()),
            )
        con.commit(); con.close()

        admin = {"id":"USR-ADMIN","role":"admin","name":"Ana Administradora"}
        qa = {"id":"USR-QA","role":"qa","name":"Quinn QA"}

        def access_check(user, case_id):
            con = db(); row = con.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone(); con.close()
            if not row: return False
            if user["role"] in {"admin", "qa"}: return True
            if user["role"] == "client": return row["owner_id"] == user["id"]
            return user["role"] == "specialist" and row["specialist_id"] in (None, user["id"])

        def lookup(user, document_id):
            con = db()
            row = con.execute("SELECT d.*,c.owner_id,c.specialist_id FROM documents d JOIN cases c ON c.id=d.case_id WHERE d.id=?", (document_id,)).fetchone()
            con.close()
            return dict(row) if row and access_check(user, row["case_id"]) else None

        workspace = ApprovalDeskWorkspace(
            root / "approval-desk",
            db_factory=db,
            document_lookup=lookup,
            access_check=access_check,
            upload_validator=lambda filename, data: (DOCX_MIME, sha256(data).hexdigest(), "clean:evidence"),
        )
        operations = ApprovalDeskOperations(
            root / "approval-desk",
            workspace=workspace,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        center = ApprovalNotificationCenter(
            root / "approval-desk",
            operations=operations,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        governance = ContactGovernance(
            root / "approval-desk",
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
        )
        communications = GovernedTransactionalCommunications(
            root / "approval-desk",
            notification_center=center,
            db_factory=db,
            now_factory=lambda: FIXED_NOW,
            governance=governance,
        )

        synced = operations.sync_portfolio(admin)
        cases = operations.portfolio(admin)["cases"]
        for row in cases:
            operations.update_assignment(admin, row["desk_case_id"], "USR-LEGAL", "USR-ADMIN")
        operations.update_deadline(admin, cases[0]["desk_case_id"], (FIXED_NOW - timedelta(hours=3)).isoformat(), 24)
        center.update_policy(admin, {
            "external_email_enabled": True,
            "external_min_severity": "high",
            "repeat_critical_hours": 24,
        })
        evaluation = center.evaluate(admin)
        communications.sync_outbox(admin)

        notice = governance.create_notice_version(admin, {
            "notice_id": "contact-governance",
            "name": "Aviso sintético M32.9",
            "text": "Finalidades, canales, derechos, revocatoria y supresión para evidencia técnica M32.9.",
        })
        governance.activate_notice(qa, "contact-governance", notice["notice"]["version"])
        relationship = governance.record_relationship(admin, {
            "subject_id": "USR-CLIENT",
            "relationship_type": "client",
            "lawful_basis": "contract",
            "status": "active",
            "evidence_reference": "CONTRACT-EVIDENCE-001",
        })
        preference = governance.record_preference(admin, {
            "subject_id": "USR-CLIENT",
            "purpose": "commercial_marketing",
            "channel": "email",
            "state": "granted",
            "basis": "consent",
            "evidence_reference": "CONSENT-EVIDENCE-001",
            "reason": "Evidencia sintética",
        })
        marketing_allowed = governance.evaluate(
            admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
            scheduled_at=FIXED_NOW,
            context_reference="EVIDENCE-MARKETING-ALLOWED",
        )["decision"]
        marketing_sunday = governance.evaluate(
            admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
            scheduled_at=datetime(2026, 8, 9, 11, 0, tzinfo=BOGOTA),
            context_reference="EVIDENCE-MARKETING-SUNDAY",
        )["decision"]
        suppression = governance.add_suppression(qa, {
            "subject_id": "USR-CLIENT",
            "scope": "purpose_channel",
            "purpose": "commercial_marketing",
            "channel": "email",
            "reason": "Solicitud sintética verificada",
            "source": "verified_request",
        })
        marketing_suppressed = governance.evaluate(
            admin,
            subject_id="USR-CLIENT",
            purpose="commercial_marketing",
            channel="email",
            scheduled_at=FIXED_NOW,
            context_reference="EVIDENCE-MARKETING-SUPPRESSED",
        )["decision"]
        processing = communications.process(admin)
        dashboard = governance.dashboard(admin)
        queue = communications.queue(admin)

        governance_events = output / "m32-9-events.jsonl"
        communication_events = output / "m32-8-events.jsonl"
        notification_events = output / "m32-7-events.jsonl"
        shutil.copy2(root / "approval-desk" / "contact-governance" / "events.jsonl", governance_events)
        shutil.copy2(root / "approval-desk" / "transactional-communications" / "events.jsonl", communication_events)
        shutil.copy2(root / "approval-desk" / "notification-center" / "events.jsonl", notification_events)

        evidence = {
            "schema_version": "M32.9",
            "generated_at": FIXED_NOW.isoformat(),
            "synthetic_evidence": True,
            "portfolio": synced["portfolio"],
            "m32_7": {
                "evaluated_cases": evaluation["evaluated_cases"],
                "queued_messages": len(evaluation["queued_messages"]),
                "external_delivery_performed": evaluation["external_delivery_performed"],
            },
            "notice": {
                "version": notice["notice"]["version"],
                "sha256": notice["notice"]["notice_sha256"],
                "active_version": dashboard["notices"]["active_notices"]["contact-governance"],
                "independent_activation": True,
            },
            "relationship": {
                "status": relationship["relationship"]["status"],
                "evidence_reference_stored": relationship["relationship"]["evidence_reference_stored"],
                "evidence_sha256": relationship["relationship"]["evidence_sha256"],
            },
            "preference": {
                "state": preference["preference"]["state"],
                "basis": preference["preference"]["basis"],
                "notice_sha256": preference["preference"]["notice_sha256"],
                "evidence_reference_stored": preference["preference"]["evidence_reference_stored"],
                "evidence_sha256": preference["preference"]["evidence_sha256"],
            },
            "decisions": {
                "marketing_allowed": marketing_allowed,
                "marketing_sunday": marketing_sunday,
                "marketing_suppressed": marketing_suppressed,
            },
            "suppression": suppression["suppression"],
            "processing": processing,
            "queue_metrics": queue["metrics"],
            "governance_metrics": dashboard["metrics"],
            "m32_9_chain": governance.verify_chain(),
            "m32_8_chain": communications.verify_chain(),
            "m32_7_chain": center.verify_chain(),
            "real_external_delivery": False,
            "real_contact_performed": False,
            "legal_conclusion": False,
            "official_holiday_calendar": False,
            "files": {
                "governance_events": governance_events.name,
                "communication_events": communication_events.name,
                "notification_events": notification_events.name,
            },
            "notice_text": "Evidencia sintética; las decisiones no sustituyen la valoración jurídica ni acreditan contacto real.",
        }
        path = output / "m32-9-contact-governance-evidence.json"
        path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
        return evidence


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    evidence = build(Path(args.output).resolve())
    print(json.dumps({
        "coverage": evidence["portfolio"]["coverage_percent"],
        "marketing_allowed": evidence["decisions"]["marketing_allowed"]["allowed"],
        "sunday_blocked": not evidence["decisions"]["marketing_sunday"]["allowed"],
        "suppression_blocked": not evidence["decisions"]["marketing_suppressed"]["allowed"],
        "accepted_sandbox": len(evidence["processing"]["accepted_sandbox"]),
        "chain_valid": evidence["m32_9_chain"]["valid"],
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
